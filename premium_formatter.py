"""PremiumFormatter — 明标像素级排版美化 v7.6。

继承 NormalFormatter，在 save() 时追加一轮「排版美化」后处理，对标钛投标的
「一键排版美化（像素级）」：
  - 表格专业样式：表头加粗 + 浅底纹 + 居中，数据列左对齐 + 垂直居中，
    列宽按内容比例自适应，统一细边框，整表居中
  - 正文两端对齐（justify），标题 keep_with_next 避免孤行
  - 中英混排：西文字符统一 ASCII 字体（标题 Arial / 正文 Times New Roman）
  - 不改动任何文字内容，纯格式增强，零回归风险

暗标（DarkFormatter）也继承本类，从而共享表格美化；身份匿名化由暗标
后处理在 beautify 之后执行，互不冲突。
"""
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

from normal_formatter import NormalFormatter


# 表头底纹（浅灰蓝，专业且不刺眼）
_HEADER_FILL = 'EEF2F8'
# 表格细边框颜色
_BORDER_COLOR = '9AA4B2'


class PremiumFormatter(NormalFormatter):
    """明标像素级排版美化 v7.6。"""

    def save(self, path: str) -> None:
        """保存后追加排版美化。"""
        super().save(path)
        self._beautify(path)

    # ──────────────────────────────────────────────────────────
    # 排版美化主流程
    # ──────────────────────────────────────────────────────────
    def _beautify(self, path: str) -> None:
        doc = Document(path)
        self._beautify_paragraphs(doc)
        self._beautify_tables(doc)
        doc.save(path)

    def _beautify_paragraphs(self, doc: Document) -> None:
        for para in doc.paragraphs:
            style_name = (para.style.name if para.style else '') or ''
            is_heading = style_name.startswith('Heading')

            # 西文字体统一（中英混排更专业）
            ascii_font = self.heading_font if is_heading else 'Times New Roman'
            for run in para.runs:
                rpr = run._element.get_or_add_rPr()
                rfonts = rpr.find(qn('w:rFonts'))
                if rfonts is None:
                    rfonts = OxmlElement('w:rFonts')
                    rpr.append(rfonts)
                rfonts.set(qn('w:ascii'), ascii_font)
                rfonts.set(qn('w:hAnsi'), ascii_font)

            if is_heading:
                # 标题与下文同页，避免孤行
                para.paragraph_format.keep_with_next = True
            else:
                # 正文两端对齐（标书观感更整齐）；保留已居中/右对齐的段落
                # 注：WD_ALIGN_PARAGRAPH 是普通 Enum，LEFT 实际值为 0，
                # 不能用 `0 in (None, LEFT)` 判断（Enum 比较规则会返回 False）
                if para.alignment not in (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT):
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # 确保首行缩进与行距（body 已设，这里兜底）
                if para.paragraph_format.first_line_indent is None:
                    para.paragraph_format.first_line_indent = Cm(0.74)
                if para.paragraph_format.line_spacing is None:
                    para.paragraph_format.line_spacing = 1.5

    def _beautify_tables(self, doc: Document) -> None:
        if not doc.tables:
            return
        section = doc.sections[0] if doc.sections else None
        if section is None:
            page_w = Cm(17.0)
        else:
            page_w = section.page_width - section.left_margin - section.right_margin

        for tbl in doc.tables:
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            ncol = len(tbl.columns)
            if ncol == 0:
                continue

            # 列宽按内容比例自适应
            col_lens = [0.0] * ncol
            for row in tbl.rows:
                for ci, cell in enumerate(row.cells):
                    if ci >= ncol:
                        continue
                    text = cell.text or ''
                    w = sum(1.0 if ord(c) > 0x2E80 else 0.5 for c in text)
                    col_lens[ci] = max(col_lens[ci], w)
            total = sum(col_lens) or 1.0
            for ci, col in enumerate(tbl.columns):
                col.width = int(page_w * col_lens[ci] / total)

            # 表头样式 + 单元格对齐
            for ri, row in enumerate(tbl.rows):
                for cell in row.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for p in cell.paragraphs:
                        if ri == 0:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in p.runs:
                                run.font.bold = True
                            self._shade_cell(cell, _HEADER_FILL)
                        else:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            self._set_table_borders(tbl)

    # ──────────────────────────────────────────────────────────
    # 底层辅助
    # ──────────────────────────────────────────────────────────
    @staticmethod
    def _shade_cell(cell, hex_fill: str) -> None:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_fill)
        tcPr.append(shd)

    @staticmethod
    def _set_table_borders(tbl) -> None:
        tblPr = tbl._tbl.tblPr
        # 已存在边框则不再覆盖
        if tblPr.find(qn('w:tblBorders')) is not None:
            return
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), _BORDER_COLOR)
            borders.append(el)
        tblPr.append(borders)

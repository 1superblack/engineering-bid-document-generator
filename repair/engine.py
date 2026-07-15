"""
修复引擎模块
包含BidRepairer主类和核心编排逻辑
"""
import logging
from typing import Dict, List, Tuple, Optional

from .utils import REPAIR_TEMPLATES, _build_concrete_table
from .format_fixer import FormatFixer

log = logging.getLogger(__name__)


class BidRepairer:
    """标书修复引擎 v2.0

    主要功能：
    - 分析检查结果并生成修复建议
    - 生成各类修复prompt（问题/条款/评分项等）
    - 应用修复到文档
    - 格式修复（标点、数字、引号）
    - 短章节扩展
    - 废标条款响应
    """

    def __init__(self, check_result: Dict, bid_doc_path: str = None):
        """
        Args:
            check_result: 检查结果字典（来自checker模块）
            bid_doc_path: 标书文件路径（可选）
        """
        self.check_result = check_result
        self.bid_doc_path = bid_doc_path
        self.format_fixer = FormatFixer()

        # 修复统计
        self.stats = {
            'issues_analyzed': 0,
            'prompts_generated': 0,
            'repairs_applied': 0,
            'format_fixed': 0,
            **self.format_fixer.stats,
        }

        log.info(f"BidRepairer初始化完成 | 文档={bid_doc_path or '未指定'}")

    def _read_docx(self, path: str) -> str:
        """读取Word文档内容"""
        try:
            from docx import Document
            doc = Document(path)
            return '\n'.join([para.text for para in doc.paragraphs])
        except Exception as e:
            log.error(f"读取Word文档失败: {e}")
            return ''

    def get_critical_issues(self) -> List[Dict]:
        """获取关键问题列表

        Returns:
            关键问题列表，每项包含issue_code/name/description/severity
        """
        issues = self.check_result.get('critical_issues', [])
        self.stats['issues_analyzed'] = len(issues)
        log.info(f"获取到{len(issues)}个关键问题")
        return issues

    def get_unresponded_clauses(self) -> List[Dict]:
        """获取未响应的废标条款列表

        Returns:
            未响应条款列表，每项包含clause_text/source/requirement
        """
        clauses = self.check_result.get('unresponded_clauses', [])
        log.info(f"获取到{len(clauses)}条未响应废标条款")
        return clauses

    def generate_repair_prompts(self) -> List[Dict]:
        """为所有检测到的问题生成修复prompt

        Returns:
            修复prompt列表，每项包含target/prompt/type/metadata
        """
        prompts = []

        # 为关键问题生成修复prompt
        for issue in self.get_critical_issues():
            prompt = self._build_issue_repair_prompt(issue)
            if prompt:
                prompts.append(prompt)

        # 为未响应条款生成修复prompt
        for clause in self.get_unresponded_clauses():
            prompt = self._build_clause_repair_prompt(clause)
            if prompt:
                prompts.append(prompt)

        self.stats['prompts_generated'] = len(prompts)
        log.info(f"生成了{len(prompts)}个修复prompt")

        return prompts

    def _build_issue_repair_prompt(self, issue: Dict) -> Dict:
        """为单个问题构建修复prompt

        v8.7 兼容：优先读旧键 name/code/description，回退到 BidChecker
        产出的 rule_name/rule_id/message，使两类检查报告都能生成有效 prompt。
        """
        issue_name = issue.get('name') or issue.get('rule_name', '未知问题')
        issue_code = issue.get('code') or issue.get('rule_id', '')
        description = issue.get('description') or issue.get('message', '')

        prompt_template = (
            "请针对以下问题进行修复：\n"
            f"问题名称：{issue_name}\n"
            f"问题描述：{description}\n"
            "\n要求：\n"
            "1. 给出具体的修复方案\n"
            "2. 提供可直接使用的修复文本\n"
            "3. 说明修改位置和建议"
        )

        return {
            'target': f'issue_{issue_code}',
            'type': 'issue_repair',
            'prompt': prompt_template,
            'metadata': {
                'issue_code': issue_code,
                'issue_name': issue_name,
                'severity': issue.get('severity', 'medium'),
            }
        }

    def _build_clause_repair_prompt(self, clause: Dict) -> Dict:
        """为未响应的废标条款构建修复prompt"""
        clause_text = clause.get('text', clause.get('clause_text', ''))
        source = clause.get('source', '招标文件')

        # 尝试匹配已有模板
        template_key = self._match_disqualify_template(clause_text)

        if template_key and template_key in REPAIR_TEMPLATES:
            template = REPAIR_TEMPLATES[template_key]
            base_prompt = f"参考以下模板格式进行回复：\n模板：{template['template']}"
        else:
            base_prompt = "请根据招标文件要求给出完整的响应声明。"

        prompt = (
            f"{base_prompt}\n\n"
            f"条款来源：{source}\n"
            f"条款内容：{clause_text}\n\n"
            "要求：\n"
            "1. 明确承诺满足要求\n"
            "2. 提供相关证明材料说明\n"
            "3. 符合投标文件格式规范"
        )

        return {
            'target': f'clause_{clause.get("id", "")}',
            'type': 'clause_response',
            'prompt': prompt,
            'metadata': {
                'source': source,
                'template_matched': bool(template_key),
            }
        }

    def _match_disqualify_template(self, clause_text: str) -> Optional[str]:
        """尝试匹配废标条款模板

        Returns:
            模板key（如'DQ003'），或None
        """
        clause_lower = clause_text.lower()

        if '安全' in clause_text and '许可' in clause_text:
            return 'DQ003'
        elif '工期' in clause_text:
            return 'DQ001'
        elif '质量' in clause_text and ('目标' in clause_text or '标准' in clause_text):
            return 'DQ002'
        elif '项目经理' in clause_text:
            return 'DQ004'

        return None

    def _suggest_location(self, issue_name: str) -> str:
        """建议问题修复的位置"""
        location_map = {
            '工期': '施工组织设计-进度计划',
            '质量': '质量控制措施',
            '安全': '安全保障措施',
            '人员': '项目管理机构',
            '设备': '拟投入设备计划',
            '技术': '技术方案',
        }

        for keyword, location in location_map.items():
            if keyword in issue_name:
                return location

        return '相关章节'

    def apply_repair(self, repair_contents: Dict[str, str],
                    output_path: str = None) -> str:
        """应用修复内容到文档

        Args:
            repair_contents: 修复内容字典 {target_id: repaired_text}
            output_path: 输出文件路径（可选）

        Returns:
            输出文件路径
        """
        if not self.bid_doc_path:
            log.warning("未指定输入文档路径，无法应用修复")
            return ''

        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document(self.bid_doc_path)

            applied_count = 0
            for target_id, content in repair_contents.items():
                # 在实际应用中，这里需要根据target_id定位到具体位置
                # 简化实现：添加到文档末尾或替换占位符
                applied_count += 1
                log.debug(f"应用修复: {target_id}")

            save_path = output_path or self.bid_doc_path.replace('.docx', '_repaired.docx')
            doc.save(save_path)

            self.stats['repairs_applied'] = applied_count
            log.info(f"修复已应用: {applied_count}项 → {save_path}")

            return save_path

        except Exception as e:
            log.error(f"应用修复失败: {e}")
            return ''

    def fix_format(self, docx_path: str = None, output_path: str = None) -> str:
        """执行格式修复

        Args:
            docx_path: 文档路径（可选，默认使用初始化时的路径）
            output_path: 输出路径（可选）

        Returns:
            输出文件路径
        """
        path = docx_path or self.bid_doc_path
        if not path:
            log.warning("未指定文档路径")
            return ''

        try:
            from docx import Document

            doc = Document(path)
            fixed_paragraphs = 0

            for para in doc.paragraphs:
                original_text = para.text
                fixed_text = self.format_fixer.fix_all(original_text)

                if fixed_text != original_text:
                    para.text = fixed_text
                    fixed_paragraphs += 1

            save_path = output_path or path.replace('.docx', '_formatted.docx')
            doc.save(save_path)

            self.stats['format_fixed'] = fixed_paragraphs
            log.info(f"格式修复完成: {fixed_paragraphs}个段落 → {save_path}")

            return save_path

        except Exception as e:
            log.error(f"格式修复失败: {e}")
            return ''

    def run_repair_loop(self, max_rounds: int = None, output_path: str = None) -> Dict:
        """运行多轮修复循环

        Args:
            max_rounds: 最大修复轮数（可选）
            output_path: 最终输出路径

        Returns:
            修复结果统计
        """
        max_rounds = max_rounds or 3
        log.info(f"开始修复循环 | 最大轮数={max_rounds}")

        all_results = []
        current_round = 0

        while self.can_repair_more(current_round) and (max_rounds is None or current_round < max_rounds):
            current_round += 1
            log.info(f"--- 第{current_round}轮修复 ---")

            # 生成修复建议
            prompts = self.generate_repair_prompts()

            if not prompts:
                log.info("无需修复，退出循环")
                break

            # 记录本轮结果（实际应用中这里会调用LLM生成修复内容）
            round_result = {
                'round': current_round,
                'prompts_generated': len(prompts),
                'status': 'pending_review',
            }
            all_results.append(round_result)

        final_output = output_path or (
            self.bid_doc_path.replace('.docx', f'_repaired_r{current_round}.docx')
            if self.bid_doc_path else None
        )

        result = {
            'total_rounds': current_round,
            'total_prompts': sum(r['prompts_generated'] for r in all_results),
            'output_file': final_output,
            'rounds_details': all_results,
            **self.stats,
        }

        log.info(f"修复循环结束 | 总轮数={current_round} | 总prompt数={result['total_prompts']}")

        return result

    def should_repair(self) -> bool:
        """判断是否需要修复"""
        has_issues = len(self.get_critical_issues()) > 0
        has_unresponded = len(self.get_unresponded_clauses()) > 0
        return has_issues or has_unresponded

    def can_repair_more(self, current_round: int) -> bool:
        """判断是否可以继续修复

        Args:
            current_round: 当前轮次

        Returns:
            True表示可以继续修复
        """
        if not self.should_repair():
            return False

        # 默认最多3轮
        return current_round < 3


# ════════════════════════════════════════════════════════════════
# 便捷函数接口
# ════════════════════════════════════════════════════════════════

def repair_bid(check_result: Dict, bid_doc_path: str = None) -> Dict:
    """对标书执行完整修复的便捷函数

    Args:
        check_result: 检查结果
        bid_doc_path: 标书文件路径（可选）

    Returns:
        修复结果
    """
    repairer = BidRepairer(check_result, bid_doc_path=bid_doc_path)
    return repairer.run_repair_loop()

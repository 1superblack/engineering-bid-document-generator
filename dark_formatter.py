"""DarkFormatter — 暗标格式化引擎 v7.0。

继承 NormalFormatter，覆写需要过滤身份信息的方法。

暗标核心要求：
1. 不得出现投标人名称/标识 → 过滤身份信息
2. 不得有页眉页脚（含公司名） → 生成后清除
3. 全文格式统一（字体/字号/行距一致） → 强制覆盖
4. 不得有封面（或封面不含身份信息） → 简化封面
5. 不得有页码 → 不添加页码
"""
import re

from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from normal_formatter import NormalFormatter
from premium_formatter import PremiumFormatter


class DarkFormatter(PremiumFormatter):
    """暗标格式 v7.0（继承 PremiumFormatter，共享表格美化）。"""

    DEFAULT_FILTER_WORDS: list[str] = [
        '投标人', '投标单位', '本公司', '我公司', '我方',
        '我单位', '本企业', '本团体',
    ]

    # v7.5: 标题前导序号 → 匿名标记（防止"记号废标"）
    # 匹配：1.1.1 / 2.3 / 第X章 / 第X节 / 一二三… 、 等层级编号
    _HEADING_PREFIX = re.compile(
        r'^\s*(?:(?:\d+\.)+\d+|\d+|[一二三四五六七八九十百千零]+[、.．]?|'
        r'第[一二三四五六七八九十百千零0-9]+[章节目节部分])\s*'
    )

    def __init__(self, *args, is_dark_bid: bool = True, filter_words: list | None = None,
                 heading_marker: str = '●', **kwargs):
        super().__init__(*args, **kwargs)
        self.is_dark_bid = is_dark_bid
        self._filter_words = list(self.DEFAULT_FILTER_WORDS)
        if filter_words:
            self._filter_words.extend(filter_words)
        # 暗标标题匿名标记（可配置，如 '●' / '§' / '—'）
        self._heading_marker = heading_marker or '●'

    def set_filter_words(self, words: list[str]) -> None:
        """设置额外过滤词表。"""
        self._filter_words.extend(words)

    def _filter_text(self, text: str) -> str:
        """过滤身份信息。"""
        filtered = text
        for word in self._filter_words:
            if word in filtered:
                filtered = filtered.replace(word, '***')
        return filtered

    def body(self, *args) -> None:
        """过滤后再输出。"""
        text = ''.join(str(a) for a in args)
        filtered = self._filter_text(text)
        NormalFormatter.body(self, filtered)

    def body_list(self, items: list) -> None:
        """过滤后再输出。"""
        for i, text in enumerate(items, 1):
            self.body(f"{i}、{text}")

    def body_bold(self, text: str) -> None:
        """过滤后加粗输出。"""
        filtered = self._filter_text(text)
        NormalFormatter.body_bold(self, filtered)

    def table(self, headers: list, rows: list) -> None:
        """表格也要过滤（继承宽表格自动适配）。"""
        filtered_headers = [self._filter_text(h) for h in headers]
        filtered_rows = [[self._filter_text(str(v)) for v in row] for row in rows]
        NormalFormatter.table(self, filtered_headers, filtered_rows)

    def add_professional_table(
        self,
        table_num: str,
        table_title: str,
        headers: list,
        rows: list,
    ) -> None:
        """三线表也要过滤。"""
        filtered_num = self._filter_text(table_num)
        filtered_title = self._filter_text(table_title)
        filtered_headers = [self._filter_text(h) for h in headers]
        filtered_rows = [[self._filter_text(str(v)) for v in row] for row in rows]
        NormalFormatter.add_professional_table(
            self, filtered_num, filtered_title, filtered_headers, filtered_rows
        )

    def add_note_box(self, text: str, bg_color: str = 'F2F2F2', border_color: str = 'CCCCCC') -> None:
        """注意事项框也要过滤。"""
        filtered = self._filter_text(text)
        NormalFormatter.add_note_box(self, filtered, bg_color, border_color)

    def add_cover_page(self, project_info: dict) -> None:
        """暗标封面：只有项目名称，不含投标人信息。"""
        name = project_info.get('name', '本项目')
        bid_type = project_info.get('bid_type', 'construction')
        type_label = '施工组织设计' if bid_type == 'construction' else '服务方案'

        for _ in range(8):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name)
        r.font.name = self.heading_font
        r.font.size = Pt(26)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), self.heading_font)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(type_label)
        r.font.name = self.heading_font
        r.font.size = Pt(36)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), self.heading_font)

        self.doc.add_page_break()

    def add_page_numbers(self) -> None:
        """暗标模式不添加页码。"""
        pass

    def add_header_footer(self, project_name: str | None = None, chapter_name: str | None = None) -> None:
        """暗标模式不添加页眉页脚。"""
        pass

    def save(self, path: str) -> None:
        """保存前清除页眉页脚、统一格式。"""
        super().save(path)
        self._post_process_dark_bid(path)

    def _post_process_dark_bid(self, path: str) -> None:
        """暗标后处理 v7.0 - 修复格式覆盖问题。

        之前的问题：任何 < 12pt 的字号都被强制改为 14pt，导致标题字号层级被破坏
        修复策略：
        - 标题（Heading 1-4）保留原始字号层级
        - 正文确保 >= 14pt（四号字）
        - 清除所有非黑色
        - 统一页边距
        """
        from docx import Document
        from docx.shared import Pt, Cm

        doc = Document(path)

        # 1. 清除所有页眉页脚
        for section in doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            for para in header.paragraphs:
                para.clear()
            footer = section.footer
            footer.is_linked_to_previous = False
            for para in footer.paragraphs:
                para.clear()

        # 2. 强制格式统一（v5.0修复：区分标题和正文）
        for para in doc.paragraphs:
            is_heading = para.style and 'Heading' in str(para.style.name)

            # v7.5: 暗标标题序号匿名化——将前导层级编号重写为标记，
            # 规避评委依据固定序号判定"记号"导致废标（钛投标/链企同款能力）。
            if is_heading and para.runs:
                new_text = self._HEADING_PREFIX.sub(
                    f'{self._heading_marker} ', para.runs[0].text
                )
                if new_text != para.runs[0].text:
                    para.runs[0].text = new_text

            for run in para.runs:
                if is_heading:
                    # 标题：只统一颜色为黑色，保留字号层级
                    if run.font.color and run.font.color.rgb and run.font.color.rgb != (0, 0, 0):
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.name = self.heading_font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), self.heading_font)
                else:
                    # 正文：确保字号不小于14pt，字体统一
                    if run.font.size and run.font.size < Pt(14):
                        run.font.size = Pt(16)
                    elif run.font.size and run.font.size < Pt(16) and not run.font.bold:
                        run.font.size = Pt(16)
                    if run.font.color and run.font.color.rgb and run.font.color.rgb != (0, 0, 0):
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    if not run.font.bold:
                        run.font.name = self.body_font
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)

        doc.save(path)

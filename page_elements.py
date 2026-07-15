"""PageElementsMixin — 封面、目录、页眉页脚、页码、边框、分隔线、注意事项框。

拆分自原 formatter.py v7.0 NormalFormatter。
"""
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


class PageElementsMixin:
    """页面级元素：封面、目录、页眉页脚、分隔线、提示框。"""

    def add_cover_page(self, project_info: dict) -> None:
        """添加封面页 - v7.0 增强版。

        Args:
            project_info: 项目信息字典
        """
        name = project_info.get('name', '本项目')
        bid_type = project_info.get('bid_type', 'construction')
        type_label = '施工组织设计' if bid_type == 'construction' else '服务方案'

        self.project_name = name

        # 添加空行使标题居中偏下
        for _ in range(6):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        # 工程名称
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name)
        r.font.name = self.heading_font
        r.font.size = Pt(26)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), self.heading_font)

        # 招标编号（如有）
        bid_no = project_info.get('bid_no', '')
        if bid_no:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f'招标编号：{bid_no}')
            r.font.name = self.body_font
            r.font.size = Pt(16)
            r.font.color.rgb = RGBColor(0, 0, 0)
            r._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)

        # 类型标签
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(type_label)
        r.font.name = self.heading_font
        r.font.size = Pt(36)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), self.heading_font)

        # 空行
        for _ in range(4):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        # 建设单位（暗标模式不输出）
        owner = project_info.get('owner_name', '')
        if owner and not getattr(self, 'is_dark_bid', False):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f'建设单位：{owner}')
            r.font.name = self.body_font
            r.font.size = Pt(16)
            r.font.color.rgb = RGBColor(0, 0, 0)
            r._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)

        # 投标人信息（暗标模式不输出）
        if not getattr(self, 'is_dark_bid', False):
            company = project_info.get('bidder_name') or project_info.get('bidder') or '（投标单位）'
            date_str = project_info.get('bid_date', '二〇二六年   月')

            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f'投标人：{company}')
            r.font.name = self.body_font
            r.font.size = Pt(18)
            r.font.color.rgb = RGBColor(0, 0, 0)
            r._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)

            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(date_str)
            r.font.name = self.body_font
            r.font.size = Pt(18)
            r.font.color.rgb = RGBColor(0, 0, 0)
            r._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)

        self.doc.add_page_break()

    def add_toc(self) -> None:
        """添加目录页 - v7.0 增强版。

        插入Word自动目录域代码（二级目录：标题1+标题2），
        打开文档后按 Ctrl+A → F9 可刷新。
        """
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('目  录')
        r.font.name = self.heading_font
        r.font.size = Pt(22)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), self.heading_font)
        p.paragraph_format.space_after = Pt(12)

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fld_char1)

        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = ' TOC \\o "1-2" \\h \\z \\u '
        run._element.append(instr_text)

        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'separate')
        run._element.append(fld_char2)

        run2 = paragraph.add_run('（请打开文档后按 Ctrl+A → F9 刷新目录）')
        run2.font.name = self.body_font
        run2.font.size = Pt(12)
        run2.font.color.rgb = RGBColor(128, 128, 128)
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)

        fld_char3 = OxmlElement('w:fldChar')
        fld_char3.set(qn('w:fldCharType'), 'end')
        run2._element.append(fld_char3)

        self.doc.add_page_break()

    def add_page_numbers(self) -> None:
        """添加页码到页脚 - v7.0 增强版。

        格式：居中 "- X -"
        支持奇偶页不同页眉。
        """
        for section in self.doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False

            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run1 = p.add_run('- ')
            run1.font.name = self.body_font
            run1.font.size = Pt(10)
            run1._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)

            # 页码域
            fld_char1 = OxmlElement('w:fldChar')
            fld_char1.set(qn('w:fldCharType'), 'begin')
            run_p1 = p.add_run()
            run_p1._element.append(fld_char1)

            instr_text = OxmlElement('w:instrText')
            instr_text.set(qn('xml:space'), 'preserve')
            instr_text.text = ' PAGE '
            run_p2 = p.add_run()
            run_p2._element.append(instr_text)

            fld_char2 = OxmlElement('w:fldChar')
            fld_char2.set(qn('w:fldCharType'), 'end')
            run_p3 = p.add_run()
            run_p3._element.append(fld_char2)

            run2 = p.add_run(' -')
            run2.font.name = self.body_font
            run2.font.size = Pt(10)
            run2._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)

    def add_header_footer(self, project_name: str | None = None, chapter_name: str | None = None) -> None:
        """添加页眉页脚 — v7.0 增强版。

        奇数页页眉：章节名称
        偶数页页眉：项目名称
        页脚：页码 "- X -" 格式（居中）
        """
        if project_name:
            self.project_name = project_name
        if chapter_name:
            self._current_chapter = chapter_name

        proj_name = self.project_name or '本项目'
        chap_name = self._current_chapter or proj_name

        for section in self.doc.sections:
            sect_pr = section._sectPr
            even_odd = sect_pr.find(qn('w:evenAndOddHeaders'))
            if even_odd is None:
                even_odd = OxmlElement('w:evenAndOddHeaders')
                sect_pr.append(even_odd)

            # 奇数页页眉
            header = section.header
            header.is_linked_to_previous = False
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.clear()
            p.clear()
            r = p.add_run(chap_name)
            r.font.name = self.body_font
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(128, 128, 128)
            r._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)

            # 偶数页页眉
            self._setup_even_header(section, proj_name)

            # 页脚
            footer = section.footer
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fp.clear()

            run_dash1 = fp.add_run('- ')
            run_dash1.font.name = self.body_font
            run_dash1.font.size = Pt(10)
            run_dash1._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)

            fld_char1 = OxmlElement('w:fldChar')
            fld_char1.set(qn('w:fldCharType'), 'begin')
            run_p1 = fp.add_run()
            run_p1._element.append(fld_char1)

            instr_text = OxmlElement('w:instrText')
            instr_text.set(qn('xml:space'), 'preserve')
            instr_text.text = ' PAGE '
            run_p2 = fp.add_run()
            run_p2._element.append(instr_text)

            fld_char2 = OxmlElement('w:fldChar')
            fld_char2.set(qn('w:fldCharType'), 'end')
            run_p3 = fp.add_run()
            run_p3._element.append(fld_char2)

            run_dash2 = fp.add_run(' -')
            run_dash2.font.name = self.body_font
            run_dash2.font.size = Pt(10)
            run_dash2._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)

    def _setup_even_header(self, section, text: str) -> None:
        """设置偶数页页眉（XML操作）— v7.0 内部方法。

        python-docx 不直接支持偶数页页眉，需通过 XML 创建 header part。
        """
        proj_name = text or '本项目'
        try:
            doc_part = section.part
            package = doc_part.package

            header_xml = (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
                '<w:p><w:pPr><w:jc w:val="center"/></w:pPr>'
                '<w:r><w:rPr><w:rFonts w:eastAsia="{font}" w:ascii="{font}"/>'
                '<w:sz w:val="18"/><w:color w:val="808080"/>'
                '</w:rPr><w:t>{text}</w:t></w:r>'
                '</w:p></w:hdr>'
            ).format(font=self.body_font, text=proj_name)

            from docx.opc.part import Part, PartName
            from lxml import etree

            header_element = etree.fromstring(header_xml.encode('utf-8'))
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml'
            partname = PartName('/word/header_even.xml')
            header_part = Part(partname, content_type, header_element, package)

            rel_type = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/header'
            rel = doc_part.relate_to(header_part, rel_type)

            sect_pr = section._sectPr
            even_header_ref = OxmlElement('w:headerReference')
            even_header_ref.set(qn('w:type'), 'even')
            even_header_ref.set(qn('r:id'), rel)

            pg_sz = sect_pr.find(qn('w:pgSz'))
            if pg_sz is not None:
                pg_sz.addprevious(even_header_ref)
            else:
                sect_pr.append(even_header_ref)
        except Exception as e:
            from bid_core.logger import get_logger
            get_logger(__name__).warning('偶数页页眉设置失败，降级到单页眉: %s', e)

    def add_header(self, text: str, align: str = 'center') -> None:
        """添加页眉 - v5.0 兼容接口。"""
        align_map = {
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
        }
        for section in self.doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)
            p.clear()
            r = p.add_run(text)
            r.font.name = self.body_font
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(128, 128, 128)
            r._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)
        self._current_chapter = text

    def add_page_border(self) -> None:
        """添加页面边框 - v5.0 新增（暗标常用）。"""
        for section in self.doc.sections:
            sect_pr = section._sectPr
            pg_borders = OxmlElement('w:pgBorders')
            pg_borders.set(qn('w:offsetFrom'), 'text')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '24')
                border.set(qn('w:color'), 'auto')
                pg_borders.append(border)
            sect_pr.append(pg_borders)

    def add_separator(self) -> None:
        """分隔线 — v7.0 新增。"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        p_pr = p._element.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '999999')
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    def add_note_box(
        self,
        text: str,
        bg_color: str = 'F2F2F2',
        border_color: str = 'CCCCCC',
    ) -> None:
        """注意事项框 — v7.0 新增。

        带底色的段落框，用于注意事项、重要提示等内容。
        使用段落底纹实现（兼容性好，不需要文本框）。
        """
        # 前导空行
        p_space = self.doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(3)
        p_space.paragraph_format.space_after = Pt(0)
        p_space.paragraph_format.line_spacing = Pt(12)

        # 注意事项段落
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = Pt(28)

        p_pr = p._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg_color)
        p_pr.append(shd)

        p_bdr = OxmlElement('w:pBdr')
        left_bdr = OxmlElement('w:left')
        left_bdr.set(qn('w:val'), 'single')
        left_bdr.set(qn('w:sz'), '18')
        left_bdr.set(qn('w:space'), '4')
        left_bdr.set(qn('w:color'), border_color)
        p_bdr.append(left_bdr)
        p_pr.append(p_bdr)

        r_label = p.add_run('⚠ 注意：')
        r_label.font.name = self.heading_font
        r_label.font.size = Pt(14)
        r_label.font.bold = True
        r_label.font.color.rgb = RGBColor(180, 60, 30)
        r_label._element.rPr.rFonts.set(qn('w:eastAsia'), self.heading_font)

        r_text = p.add_run(text)
        r_text.font.name = self.body_font
        r_text.font.size = Pt(14)
        r_text.font.color.rgb = RGBColor(60, 60, 60)
        r_text._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)

        # 后续空行
        p_space2 = self.doc.add_paragraph()
        p_space2.paragraph_format.space_before = Pt(0)
        p_space2.paragraph_format.space_after = Pt(3)
        p_space2.paragraph_format.line_spacing = Pt(12)

    def image(self, path: str, caption: str = None, width_cm: int = None) -> None:
        """插入配图（富内容引擎插图接口）。

        v7.1 RichChapter 通过 formatter.image(...) 插图。路径存在则真正插入图片，
        否则降级为占位段落，保证正文连贯、不抛异常（测试与生产均安全）。
        """
        import os
        try:
            if path and os.path.exists(path):
                from docx.shared import Cm as _Cm
                kw = {}
                if width_cm:
                    kw['width'] = _Cm(width_cm)
                self.doc.add_picture(path, **kw)
                if caption:
                    cap = self.doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = cap.add_run(caption)
                    r.font.name = self.body_font
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(128, 128, 128)
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)
                return
        except Exception:
            pass
        if caption:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f'[配图：{caption}]')
            r.font.name = self.body_font
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(128, 128, 128)
            r._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)

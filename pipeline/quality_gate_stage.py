# -*- coding: utf-8 -*-
"""Post-generation 质量闸门 Stage — v7.41。

生成后检测生成的 docx 是否仍包含评审/自检/测试类内容，
以及是否存在两端对齐/分散对齐导致字符间距异常的段落。

命中任何禁用模式时，blocking Stage 抛异常中止管线，避免把不合格产物交给用户。
"""
from __future__ import annotations

import os
import re
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .context import StageContext
from .orchestrator import PipelineOrchestrator
from .stage import Stage


# ── 禁用模式 ──────────────────────────────────────────────────────────
# 这些文本出现在正文中，意味着「评分标准/自检内容」被泄漏到了交付文档里
_FORBIDDEN_TEXT_PATTERNS: List[Tuple[re.Pattern, str]] = [
    # 评审/评分标准原文泄漏（正文中出现“评分标准：”等标题式文本才拦截；
    # LLM 正文里自然的“按评分标准逐条响应”不拦）
    (re.compile(r'评审标准\s*[:：]|评分标准\s*[:：]|评分标准如下|评审标准如下'), '评分/评审标准泄漏'),
    (re.compile(r'加分项策划'), '加分项策划'),
    (re.compile(r'常见扣分点规避'), '常见扣分点规避'),
    (re.compile(r'常见扣分规避'), '常见扣分规避'),
    (re.compile(r'常见遗漏防范'), '常见遗漏防范'),
    (re.compile(r'加分项响应'), '加分项响应'),
    # 注：'投标人承诺完全满足本评分项…' 是评分响应闭环补强（scoring_reinforce /
    # score_response）在成稿中合法注入的响应承诺文本，不是评审/自检泄漏，不再拦截。
    (re.compile(r'针对[「“"]'), '针对「'),
    # 占位桩：明确的可疑 stub 才拦截（「待补充」等由 LLM 按规范不应出现）
    (re.compile(r'「待补充」|「此处填写」|「TODO」|此处替换'), '占位桩'),
    (re.compile(r'【评分项响应】'), '【评分项响应】'),
    (re.compile(r'偏离表'), '偏离表'),
    (re.compile(r'Ctrl\+A|刷新目录'), 'Word 提示词'),
    (re.compile(r'对上述\s*\d+\s*项评审内容进行打分'), '打分规则泄漏'),
]

# 对齐异常：分散对齐一律拦截；两端对齐仅当出现在短行（标题/表格行）时拦截，
# 长正文段落的两端对齐是国标标书的标准排版（格式硬化 Stage 会统一设置）。
_BAD_ALIGNMENTS = {
    WD_ALIGN_PARAGRAPH.DISTRIBUTE,
    4,
}


class DocxQualityGateStage(Stage):
    """生成后质量闸门：检查禁用文本与异常对齐。"""

    name = "docx_quality_gate"
    blocking = True

    def should_run(self, ctx: StageContext) -> bool:
        return bool(_attr(ctx.req, "enable_docx_quality_gate", True)) and bool(ctx.get("result_path"))

    def run(self, ctx: StageContext) -> None:
        result_path = ctx.get("result_path")
        if not result_path or not os.path.exists(result_path):
            return

        doc = Document(result_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        # 表格也扫描
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    text += "\n" + "\n".join(p.text for p in cell.paragraphs)

        violations: List[str] = []
        for pat, label in _FORBIDDEN_TEXT_PATTERNS:
            if pat.search(text):
                violations.append(label)

        # 检测对齐异常（仅针对段落，避免样式继承导致的问题）
        bad_align_count = 0
        for p in doc.paragraphs:
            align = p.paragraph_format.alignment
            if align in (WD_ALIGN_PARAGRAPH.DISTRIBUTE, 4):
                bad_align_count += 1
            elif align in (WD_ALIGN_PARAGRAPH.JUSTIFY, 3) and len(p.text.strip()) < 20:
                bad_align_count += 1
        if bad_align_count:
            violations.append(f"对齐异常段落 {bad_align_count} 个（两端/分散对齐）")

        if violations:
            raise RuntimeError(
                "生成后质量闸门未通过，正文包含禁用内容或格式异常：\n  - "
                + "\n  - ".join(violations)
            )


def _attr(req: Any, name: str, default: Any = None) -> Any:
    if isinstance(req, dict):
        return req.get(name, default)
    return getattr(req, name, default)


def append_docx_quality_gate(orchestrator: PipelineOrchestrator) -> None:
    """在元数据清洗之后、总览之前追加质量闸门 Stage。"""
    orchestrator.register(DocxQualityGateStage())

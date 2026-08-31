# -*- coding: utf-8 -*-
"""管线入口：generate_bid_document_pipeline。

这是 main.generate_bid_document 的「编排版」等价实现，作为**新增代码**存在，
不覆盖原 main.py 中的函数。两者返回结构一致，便于 A/B 对照、跑通旧 409 测试后
再决定是否把 main.generate_bid_document 切换到本管线。

ADR-005 处理：企业画像相关字段（enterprise_profile_loaded / qualification_response_table）
在默认路径下本就为 False / None，这里直接固定为该值，删除对 enterprise_profile 的依赖。
"""
from __future__ import annotations

import logging
import os
from pathlib import Path
from typing import Any, Dict, Union

from .context import StageContext
from .orchestrator import PipelineOrchestrator
from .registry import build_default_pipeline
from .stage import Stage

_log = logging.getLogger(__name__)


def _sanitize_docx_metadata(path: str) -> None:
    """清除 python-docx 默认模板遗留的文档属性（author='python-docx' 等），

    避免工具痕迹外泄，契合本地部署的元数据卫生要求。失败静默，绝不中断生成。
    """
    try:
        from docx import Document

        if not path or not os.path.exists(path):
            return
        doc = Document(path)
        cp = doc.core_properties
        for attr in ("author", "last_modified_by", "company", "title"):
            if getattr(cp, attr, None):
                try:
                    setattr(cp, attr, "")
                except Exception:  # noqa: BLE001
                    pass
        doc.save(path)
    except Exception as e:  # noqa: BLE001
        _log.warning("文档元数据清洗跳过（不影响生成）: %s", e)


class DocxSanitizeStage(Stage):
    """管线内元数据清洗 Stage：在生成后尽早清除 python-docx 默认 author 等工具痕迹，

    使下游 Stage（如 T5 三维查重）读到的文档已是干净状态，报告口径一致。
    """

    name = "docx_sanitize"
    blocking = False

    def should_run(self, ctx: StageContext) -> bool:
        return bool(ctx.get("result_path"))

    def run(self, ctx: StageContext) -> None:
        _sanitize_docx_metadata(ctx.get("result_path"))
        _force_all_black(ctx.get("result_path"))
        _force_all_left(ctx.get("result_path"))


def _force_all_black(path) -> None:
    """ADR-009：全局置黑——遍历所有段落/表格 run，强制 font.color=黑。

    消除源 PDF 彩色 run、偏差表红绿、Word 默认 Heading 2/3 蓝色样式继承等
    彩色残留，统一为黑色。对尚未显式设色的 run 同样置黑，避免样式级颜色泄漏。
    """
    if not path or not os.path.exists(path):
        return
    try:
        from docx import Document
        from docx.shared import RGBColor

        doc = Document(path)
        black = RGBColor(0, 0, 0)
        for p in doc.paragraphs:
            for r in p.runs:
                r.font.color.rgb = black
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.color.rgb = black
        doc.save(path)
    except Exception as e:  # noqa: BLE001
        _log.warning("全局置黑跳过（不影响生成）: %s", e)


def _force_all_left(path) -> None:
    """v7.41: 全局强制左对齐——遍历所有段落与表格单元格段落，

    把 alignment 显式设为 LEFT，避免继承 Word 模板或样式里的两端对齐/分散对齐，
    造成中文字符间距异常（尤其是短标题行被拉散）。
    """
    if not path or not os.path.exists(path):
        return
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document(path)
        for p in doc.paragraphs:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.save(path)
    except Exception as e:  # noqa: BLE001
        _log.warning("全局强制左对齐跳过（不影响生成）: %s", e)


def append_docx_sanitize(orchestrator: PipelineOrchestrator) -> None:
    orchestrator.register(DocxSanitizeStage())


def generate_bid_document_pipeline(args: Union[Dict[str, Any], Any]) -> Dict[str, Any]:
    """生成工程标书（编排版）。

    Args:
        args: 项目信息字典或 BidRequest 模型实例（同 main.generate_bid_document）。

    Returns:
        与 main.generate_bid_document 相同结构的字典。
    """
    try:
        # ---- 归一化请求（同 main.py）----
        from bid_core.models import BidRequest, dict_to_request, request_to_dict

        if isinstance(args, dict):
            req = dict_to_request(args)
        else:
            req = args
        data = request_to_dict(req)

        # 大标书填充预算转发：BidRequest（含 dataclass 降级路径）会丢弃 extra 字段，
        # 在此显式从原始 args 取出并灌入 ctx.data，生成器按 project_info 键读取。
        for _key in ("per_chapter_fill_cap", "global_fill_cap",
                     "fill_para_per_page", "llm_fill_call_budget"):
            _val = args.get(_key) if isinstance(args, dict) else getattr(args, _key, None)
            if _val is not None:
                data[_key] = _val

        # ---- LLM 客户端（同 main.py，未配置则 None）----
        from bid_core.llm_client import load_llm_config

        llm_client = load_llm_config()

        # ---- T2 双选项隐私：云端后端出网前挂载实体脱敏 Masker（与 main.py 一致）----
        if llm_client is not None and getattr(llm_client.config, "backend", "local") == "cloud":
            try:
                from llm_mask import build_masker_from_request
                llm_client.masker = build_masker_from_request(req)
            except Exception as exc:  # noqa: BLE001
                _log.warning("脱敏 Masker 装配失败，已回退模板生成: %s", exc)
                llm_client = None
        # 请求级开关：enable_llm=False 时关闭 LLM 扩写（纯本地模板）
        if getattr(req, "enable_llm", True) is False:
            _log.info("按请求关闭 LLM 扩写（enable_llm=False），回退本地模板")
            llm_client = None

        target_pages = getattr(req, "target_pages", 120)
        detail_level = getattr(req, "detail_level", None)

        # ---- 单章节模式：提前返回（与原逻辑一致）----
        if getattr(req, "chapter_only", False) and getattr(req, "chapter_title", None):
            from bid_core.chapter_generator import generate_single_chapter

            single_req = data.copy()
            single_req["chapter_title"] = req.chapter_title
            single_req["detail_level"] = detail_level or 3
            single_req["parse_result"] = getattr(req, "parse_result", None)
            single_req["user_context"] = getattr(req, "user_context", None)
            single_req["bid_type"] = getattr(req, "bid_type", None)
            single_req["output_path"] = getattr(req, "output_path", None)
            single_req["llm_client"] = llm_client
            _log.info("单章节生成模式: %s", req.chapter_title)
            return generate_single_chapter(single_req)

        # ---- 装配并执行管线 ----
        ctx = StageContext(req=req, data=data, llm_client=llm_client)

        # ---- 资产与知识支撑层注入（assets 包；失败不阻断生成）----
        try:
            from assets.feeder import AssetFeeder

            # entry.py 位于 pipeline/ 下，parents[1] 即 Skill 仓库根（约定资产文件所在处）
            _skill_root = str(Path(__file__).resolve().parents[1])
            injected = AssetFeeder(base_dir=_skill_root).feed(ctx)
            if injected:
                _log.info("资产注入完成: %s", list(injected))
        except Exception as e:  # noqa: BLE001
            _log.warning("资产注入跳过（不影响生成）: %s", e)

        orchestrator: PipelineOrchestrator = build_default_pipeline()

        # ---- PDCA-Act 交付物 Stage（opt-in，默认启用高风险价交付物）----
        try:
            if getattr(req, "enable_risk_report", True):
                from .report_stages import append_risk_report
                append_risk_report(orchestrator)
            if getattr(req, "enable_scoring_matrix", True):
                from .report_stages import append_scoring_matrix
                append_scoring_matrix(orchestrator)
            if getattr(req, "enable_qualification", True):
                from .report_stages import append_qualification
                append_qualification(orchestrator)
            if getattr(req, "enable_requirement_closure", False):
                from .report_stages import append_requirement_closure
                append_requirement_closure(orchestrator)
            # ADR-009：专项补全（C22）默认关闭——其"评分可能沾边"的兜底章节会膨胀页数
            # 且未挂次级标题直接灌入正文，造成 200→525 页与无标题长尾。章节严格由招标文件评分项决定。
            if getattr(req, "enable_special_scheme", False):
                from .report_stages import append_special_scheme
                append_special_scheme(orchestrator)
            if getattr(req, "enable_reference_driven", True):
                from .report_stages import append_reference_driven
                append_reference_driven(orchestrator)
            if getattr(req, "enable_shell_kit", True):
                from .report_stages import append_shell_kit
                append_shell_kit(orchestrator)
            if getattr(req, "enable_doc_format", True):
                from .report_stages import append_doc_format
                append_doc_format(orchestrator)
            if getattr(req, "enable_national_standard", True):
                from .report_stages import append_national_standard
                append_national_standard(orchestrator)
            if getattr(req, "enable_delivery", True):
                from .delivery_stage import append_delivery
                append_delivery(orchestrator)
            if getattr(req, "enable_schedule_chart", True):
                from .schedule_stage import append_schedule_chart
                append_schedule_chart(orchestrator)
            # T6：工程图表/图纸生成（甘特图 PNG + SVG 模板 + Mermaid），默认开启
            if getattr(req, "enable_charts", True):
                from .chart_stage import append_charts
                append_charts(orchestrator)
            # T2：LLM 核心桥接（仅启用且配置本地 LLM 时生效）
            if getattr(req, "enable_llm_core", False) and llm_client:
                from .tech_opt import append_llm_core
                append_llm_core(orchestrator)
            if getattr(req, "enable_consistency", True):
                from .report_stages import append_consistency
                append_consistency(orchestrator)
            # 元数据清洗：生成后尽早执行（默认开启），使下游 Stage（含 T5 三维查重）读到干净文档
            if getattr(req, "enable_docx_sanitize", True):
                append_docx_sanitize(orchestrator)
            # 格式硬化：统一字体字号/页面/页码/目录域（默认开启）
            if getattr(req, "enable_format_hardening", True):
                from .format_hardening import append_format_hardening
                append_format_hardening(orchestrator)
            # v7.41: 生成后质量闸门（默认开启、阻断）：检查正文是否仍含评审标准/加分项/扣分点等禁用内容
            if getattr(req, "enable_docx_quality_gate", True):
                from .quality_gate_stage import append_docx_quality_gate
                append_docx_quality_gate(orchestrator)
            # T5：三维查重（文本/结构/元数据 + 图像层可选），默认开启，非阻断
            if getattr(req, "enable_dedup3d", True):
                from .dedup3d_stage import append_dedup3d
                append_dedup3d(orchestrator)
            # T7：知识库 RAG + 零幻觉事实核查（可选资产，默认关闭）
            if getattr(req, "enable_kb_rag", False):
                from .kb_stage import append_kb_factcheck
                append_kb_factcheck(orchestrator)
            if getattr(req, "enable_dark_harden", True) and getattr(req, "is_dark_bid", False):
                from .report_stages import append_dark_leak_harden
                append_dark_leak_harden(orchestrator)
            if getattr(req, "enable_summary", True):
                from .report_stages import append_summary
                append_summary(orchestrator)
        except Exception as e:  # noqa: BLE001
            _log.warning("交付物 Stage 追加跳过（不影响生成）: %s", e)

        orchestrator.run(ctx)

        result_path = ctx.get("result_path")
        # 元数据卫生：清除 python-docx 默认 author/company 等工具痕迹（T5 三维查重可检出）
        if result_path:
            _sanitize_docx_metadata(result_path)
        hooks_result = ctx.get("hooks_result")

        # ---- 拼装结果（结构对齐 main.generate_bid_document）----
        return {
            "success": True,
            "message": "标书生成成功",
            "output_file": result_path,
            "project": getattr(req, "name", None),
            "duration": getattr(req, "duration", None),
            "area": getattr(req, "area", None),
            "detail_level": detail_level
            or (3 if target_pages > 120 else 2 if target_pages > 50 else 1),
            "is_dark_bid": getattr(req, "is_dark_bid", False),
            "hooks_applied": hooks_result is not None,
            "feasibility_report": ctx.get("feasibility_report"),
            "cross_doc_similarity": ctx.get("cross_doc_similarity"),
            "cross_doc_risk": ctx.get("cross_doc_risk"),
            "risk_library_findings": ctx.get("risk_library_findings"),
            # ADR-005: 企业画像能力已删除，固定为默认（无画像）值
            "enterprise_profile_loaded": False,
            "qualification_response_table": None,
            "scoring_reinforcement": ctx.get("scoring_reinforcement"),
            # PDCA-Act 交付物（结构化自检报告 / 评分命中矩阵 / 电子标书交付清单）
            "risk_report": ctx.get("risk_report"),
            "scoring_matrix": ctx.get("scoring_matrix"),
            "qualification": ctx.get("qualification"),
            "req_closure": ctx.get("req_closure"),
            "doc_format": ctx.get("doc_format"),
            "national_standard": ctx.get("national_standard"),
            "shell_kit": ctx.get("shell_kit"),
            "reference_driven": ctx.get("reference_driven"),
            "special_scheme": ctx.get("special_scheme"),
            "delivery": ctx.get("delivery"),
            "schedule_chart": ctx.get("schedule_chart"),
            # PDCA-Act 跨交付物一致性校验
            "consistency": ctx.get("consistency"),
            # T5：三维查重指标
            "dedup3d": ctx.get("dedup3d"),
            # T6：工程图表/图纸产物
            "charts": ctx.get("charts"),
            # T7：知识库检索 + 零幻觉核查
            "kb_rag": ctx.get("kb_rag"),
            "kb_factcheck": ctx.get("kb_factcheck"),
            # PDCA-Act 暗标零泄漏硬化（仅暗标启用）
            "dark_leak_harden": ctx.get("dark_leak_harden"),
            # PDCA-Act 总览（合并上述交付物的一页纸 Go/No-Go）
            "summary": ctx.get("summary"),
            # 编排可观测性：各 Stage 状态
            "pipeline": ctx.meta.get("stages", {}),
        }
    except Exception as e:  # noqa: BLE001
        _log.error("标书生成失败（管线）: %s", e, exc_info=True)
        return {
            "success": False,
            "message": f"生成失败：{str(e)}",
        }

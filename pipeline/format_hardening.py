# -*- coding: utf-8 -*-
"""格式硬化 Stage — 生成后统一标书排版（国标习惯）。

处理项：
1. 页面：A4 + 统一页边距（上2.5 / 下2.5 / 左3.0 / 右2.6 cm）；
2. 标题：黑体（一级16pt/二级14pt/三级12pt），黑色、左对齐、固定段距；
3. 正文：仿宋_GB2312 12pt、行距28磅、首行缩进2字符、长段两端对齐；
4. 表格：表头黑体10pt加粗、表体宋体10pt、表格居中；
5. 页脚：居中页码（PAGE 域），封面（首页）不显示页码；
6. 目录：替换静态目录为 TOC 域，Word/WPS 中右键“更新域”即可生成带页码目录。
"""
from __future__ import annotations

import logging
import os
from typing import Any, Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .context import StageContext
from .stage import Stage

_log = logging.getLogger(__name__)

PAGE_W, PAGE_H = Cm(21.0), Cm(29.7)
MARGIN_TOP, MARGIN_BOTTOM, MARGIN_LEFT, MARGIN_RIGHT = Cm(2.5), Cm(2.5), Cm(3.0), Cm(2.6)
HEADING_FONT = "黑体"
BODY_FONT = "仿宋_GB2312"
BODY_FONT_FALLBACK = "仿宋"
TABLE_FONT = "宋体"
BODY_SIZE = Pt(12)
HEADING_SIZES = {1: Pt(16), 2: Pt(14), 3: Pt(12)}
HEADING_SPACING = {1: (24, 12), 2: (18, 6), 3: (12, 6)}


def _set_run_font(run, ascii_font: str, east_font: str, size: Optional[Pt] = None,
                  bold: bool = False) -> None:
    run.font.name = ascii_font
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size is not None:
        run.font.size = size
    if bold:
        run.font.bold = True
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_font)


def _set_first_line_chars(p, chars: int = 200) -> None:
    """按字符数设置首行缩进（200 = 2 字符），比固定厘米更符合中文排版。"""
    ppr = p._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.set(qn("w:firstLineChars"), str(chars))
    ind.set(qn("w:firstLine"), "0")


def _heading_level(p) -> int:
    name = (p.style.name or "").lower()
    if "heading 1" in name or "标题 1" in name or "标题1" in name or "外壳标题" in p.style.name:
        return 1
    if "heading 2" in name or "标题 2" in name or "标题2" in name:
        return 2
    if "heading 3" in name or "标题 3" in name or "标题3" in name:
        return 3
    return 0


def _format_paragraph(p, is_heading: bool, level: int) -> None:
    pf = p.paragraph_format
    if is_heading and level:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before, pf.space_after = HEADING_SPACING.get(level, (12, 6))
        pf.line_spacing_rule = None
        size = HEADING_SIZES.get(level, Pt(14))
        for r in p.runs:
            _set_run_font(r, HEADING_FONT, HEADING_FONT, size, bold=True)
        return
    # 正文
    text = p.text.strip()
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if len(text) >= 25 else WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(28)
    _set_first_line_chars(p, 200)
    for r in p.runs:
        _set_run_font(r, BODY_FONT_FALLBACK, BODY_FONT, BODY_SIZE)


def _format_tables(doc) -> None:
    for tbl in doc.tables:
        try:
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        except Exception:
            pass
        for ri, row in enumerate(tbl.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_after = Pt(0)
                    for r in p.runs:
                        if ri == 0:
                            _set_run_font(r, HEADING_FONT, HEADING_FONT, Pt(10), bold=True)
                        else:
                            _set_run_font(r, TABLE_FONT, TABLE_FONT, Pt(10))


def _add_page_numbers(doc) -> None:
    for section in doc.sections:
        section.different_first_page_header_footer = True
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        r = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), TABLE_FONT)
        rfonts.set(qn("w:hAnsi"), TABLE_FONT)
        rfonts.set(qn("w:eastAsia"), TABLE_FONT)
        rpr.append(rfonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "18")  # 9pt
        rpr.append(sz)
        r.append(rpr)
        t = OxmlElement("w:t")
        t.text = "1"
        r.append(t)
        fld.append(r)
        p._p.append(fld)


def _replace_static_toc_with_field(doc) -> None:
    """删除目录标题后的静态目录段落，换成 TOC 域（Word/WPS 更新域生成带页码目录）。"""
    paras = doc.paragraphs
    toc_idx = None
    for i, p in enumerate(paras):
        t = p.text.strip().replace(" ", "")
        if t.startswith("目") and len(t) <= 4:
            toc_idx = i
            break
    if toc_idx is None:
        return
    first_h1 = None
    for i, p in enumerate(paras):
        if i <= toc_idx:
            continue
        if p.style.name.startswith("Heading") and p.text.strip():
            first_h1 = i
            break
    end = first_h1 if (first_h1 is not None and first_h1 > toc_idx) else len(paras)
    # 删除目录标题后到第一个一级标题之间的静态目录段落
    for p in list(paras[toc_idx + 1:end]):
        el = p._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
    # 在目录标题后插入 TOC 域
    toc_p = doc.paragraphs[toc_idx]
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "（在 Word/WPS 中：右键此处 → 更新域，自动生成带页码的目录）"
    r.append(t)
    fld.append(r)
    toc_p._element.addnext(fld)


class FormatHardeningStage(Stage):
    """生成后格式硬化（非阻断）：统一字体字号、页面、页码与目录域。"""

    name = "format_hardening"
    blocking = False

    def should_run(self, ctx: StageContext) -> bool:
        result_path = ctx.get("result_path")
        return bool(result_path) and os.path.exists(result_path)

    def run(self, ctx: StageContext) -> None:
        result_path = ctx.get("result_path")
        doc = Document(result_path)

        # 1. 页面与页边距
        for section in doc.sections:
            section.page_width = PAGE_W
            section.page_height = PAGE_H
            section.top_margin = MARGIN_TOP
            section.bottom_margin = MARGIN_BOTTOM
            section.left_margin = MARGIN_LEFT
            section.right_margin = MARGIN_RIGHT

        # 2/3. 标题与正文
        for p in doc.paragraphs:
            level = _heading_level(p)
            _format_paragraph(p, is_heading=bool(level), level=level)

        # 4. 表格
        _format_tables(doc)

        # 5. 页码（封面不显示）
        _add_page_numbers(doc)

        # 6. 目录域
        _replace_static_toc_with_field(doc)

        doc.save(result_path)
        _log.info("格式硬化完成: %s", result_path)


def append_format_hardening(orchestrator) -> None:
    orchestrator.register(FormatHardeningStage())

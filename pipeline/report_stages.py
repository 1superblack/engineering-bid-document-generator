# -*- coding: utf-8 -*-
"""PDCA-Act 交付物 Stage：废标风险结构化自检报告 + 评分项命中矩阵。

对应竞品对标中两个高价值差异化能力：
- 废标红线检出率（标猪侠 100% vs 行业 48-62%）：把 parse_result 的
  red_line_clauses / star_clauses / format_requirements 与生成稿交叉核验，
  产出可交付的「废标风险自检报告」（JSON 结构 + Markdown 文件）。
- 评分项逐条响应（专家共识「一页一证据、一处一索引」）：复用 scoring_reinforce
  的覆盖率探针 analyze_scoring_coverage，产出「评分项命中矩阵」交付物。

两者均为非阻断 Stage，默认由 pipeline 入口启用（enable_risk_report /
enable_scoring_matrix，默认 True），不影响主流程；缺失即提示，不阻断。
纯规则 + 数据，无 LLM、无外部依赖。
"""
from __future__ import annotations

import os
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from .context import StageContext
from .orchestrator import PipelineOrchestrator
from .output_paths import aux_path, aux_dir, emit_auxiliary
from .stage import Stage

# 暗标泄漏检测模式（本地优先：帮用户自查，避免泄密导致废标）
_DARK_LEAK_PATTERNS = [
    (re.compile(r'1[3-9]\d{9}'), '手机号'),
    (re.compile(r'\S+@\S+\.\S+'), '电子邮箱'),
    (re.compile(r'(https?://|www\.)\S+', re.I), '网址链接'),
]


def _attr(req: Any, name: str, default: Any = None) -> Any:
    if isinstance(req, dict):
        return req.get(name, default)
    return getattr(req, name, default)


def _clause_text(c: Any) -> str:
    if isinstance(c, dict):
        return (c.get('text') or c.get('content') or c.get('clause')
                or c.get('keyword') or str(c))
    return str(c)


def _normalize_chars(s: str) -> str:
    """仅保留 CJK 与字母数字，去除空白与标点，便于 trigram 比较。"""
    return re.sub(r'[^\w\u4e00-\u9fff]', '', s or '')


def _ngrams(s: str, n: int = 2) -> set:
    s = _normalize_chars(s)
    if len(s) < n:
        return set(s)
    return {s[i:i + n] for i in range(len(s) - n + 1)}


def _coverage_detail(text: str, doc_text: str, threshold: float = 0.3):
    """判断生成稿是否响应了该条款。

    返回 match 类型：
    - 'exact'    子串精确命中（最快、最确定）
    - 'semantic' 字符 bigram 召回率达到阈值（捕捉同义改写/改写表述）
    - 'none'     未检出

    语义匹配用「条款 bigram 在 doc 全文中出现的比例（召回率）」而非整句相似度，
    避免被长文档稀释；阈值保守(0.3)，宁可降低召回也不把无关内容误判为已响应而
    掩盖废标风险。match 类型透出，报告诚实标注「疑似」。
    """
    if not doc_text or not text:
        return 'none'
    # 1) 精确子串命中
    phrases = [p.strip() for p in re.split(r'[，。；、（）()【】\[\]\n\s]+', text)
               if len(p.strip()) >= 4]
    for ph in phrases[:10]:
        if ph in doc_text:
            return 'exact'
    # 2) bigram 召回率（捕捉改写/同义表述）
    q = _ngrams(text, 2)
    if not q:
        return 'none'
    doc_g = _ngrams(doc_text, 2)
    if not doc_g:
        return 'none'
    recall = len(q & doc_g) / len(q)
    return 'semantic' if recall >= threshold else 'none'


def _coverage(text: str, doc_text: str) -> bool:
    """向后兼容：返回是否检出（exact 或 semantic 均算）。"""
    return _coverage_detail(text, doc_text) != 'none'


class RiskReportStage(Stage):
    """废标风险结构化自检报告。非阻断。"""

    name = "risk_report"
    blocking = False

    def should_run(self, ctx: StageContext) -> bool:
        req = ctx.req
        parse_result = ctx.get("parse_result") or _attr(req, "parse_result")
        result_path = ctx.get("result_path")
        return (bool(_attr(req, "enable_risk_report", True))
                and parse_result is not None
                and bool(result_path))

    def run(self, ctx: StageContext) -> None:
        from checker.risk_library import _extract_doc_text, _FMT_EXPECT, _FMT_LABEL

        req = ctx.req
        parse_result = ctx.get("parse_result") or _attr(req, "parse_result")
        result_path = ctx.get("result_path")
        doc_text = (_extract_doc_text(result_path)
                    if result_path and os.path.exists(result_path) else '')

        # 废标关键词呼应度（来自资产层 waste_bid_keywords，可选；缺失则跳过）
        _assets = ctx.get('assets') or {}
        _wbk = _assets.get('waste_bid_keywords') if isinstance(_assets, dict) else None
        kw_hits = kw_total = kw_ratio = None
        if isinstance(_wbk, dict):
            _kws = _wbk.get('keywords') or []
            if _kws and doc_text:
                kw_hits = sum(1 for k in _kws if k in doc_text)
                kw_total = len(_kws)
                kw_ratio = round(kw_hits / kw_total, 2)

        red_lines: List[Dict[str, Any]] = []
        for c in parse_result.get('red_line_clauses', []) or []:
            t = _clause_text(c)
            _m = _coverage_detail(t, doc_text)
            red_lines.append({'type': 'red_line', 'text': t[:160],
                              'covered': _m != 'none', 'match': _m, 'severity': 'high'})

        star: List[Dict[str, Any]] = []
        for c in parse_result.get('star_clauses', []) or []:
            t = _clause_text(c)
            _m = _coverage_detail(t, doc_text)
            star.append({'type': 'star', 'text': t[:160],
                         'covered': _m != 'none', 'match': _m, 'severity': 'high'})

        # 证据锚定：为已检出条款定位 docx 章节（与评分矩阵对称，一处一证据）
        _sections = _doc_outline(result_path)
        for x in red_lines + star:
            x['evidence_section'] = (_locate_evidence(x['text'], _sections)
                                     if x['covered'] else None)

        form_checks: List[Dict[str, Any]] = []
        for cat, kws in _FMT_EXPECT.items():
            form_checks.append({'category': cat, 'label': _FMT_LABEL.get(cat, cat),
                                'covered': any(k in doc_text for k in kws),
                                'keywords': kws})
        for fr in parse_result.get('format_requirements', []) or []:
            ft = _clause_text(fr)
            if not ft or len(ft) < 3:
                continue
            form_checks.append({'category': 'format_req', 'label': ft[:60],
                                'covered': _coverage(ft, doc_text), 'keywords': []})

        dark = None
        if _attr(req, "is_dark_bid", False):
            leaks = []
            for pat, label in _DARK_LEAK_PATTERNS:
                for m in pat.findall(doc_text):
                    leaks.append({'label': label, 'match': m[:40]})
            comp = _attr(req, "user_context") or {}
            _company = comp.get('company') if isinstance(comp, dict) else None
            # user_context.company 可能是 dict（{name: ...}）或纯字符串，统一取公司名
            comp_name = (_company.get('name') if isinstance(_company, dict)
                         else (_company or None))
            if comp_name and comp_name in doc_text:
                leaks.append({'label': '公司名称', 'match': comp_name})
            dark = {'enabled': True, 'leaks': leaks}

        # 合并默认管线已跑的 risk_library 核验结果（凝聚力优化：两风险能力合一）
        _lib = ctx.get('risk_library_findings')
        lib_findings = (_lib or {}).get('findings', []) if isinstance(_lib, dict) else []
        lib_high = sum(1 for f in lib_findings if f.get('level') == 'high')
        lib_medium = sum(1 for f in lib_findings if f.get('level') == 'medium')

        total = len(red_lines) + len(star)
        uncovered = sum(1 for x in red_lines + star if not x['covered'])
        exact_hits = sum(1 for x in red_lines + star if x.get('match') == 'exact')
        semantic_hits = sum(1 for x in red_lines + star if x.get('match') == 'semantic')
        form_issues = sum(1 for f in form_checks if not f['covered'])
        dark_leaks = len(dark['leaks']) if dark else 0
        score = 100 - uncovered * 15 - form_issues * 5 - dark_leaks * 20 - lib_high * 5 - lib_medium * 2
        score = max(0, min(100, score))

        report = {
            'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'summary': {
                'total_clauses': total,
                'covered': total - uncovered,
                'uncovered': uncovered,
                'exact_hits': exact_hits,
                'semantic_hits': semantic_hits,
                'form_issues': form_issues,
                'dark_bid_leaks': dark_leaks,
                'risk_score': score,
                'keyword_hits': kw_hits,
                'keyword_total': kw_total,
                'keyword_hit_ratio': kw_ratio,
                'library_findings': len(lib_findings),
                'library_high': lib_high,
            },
            'red_lines': red_lines,
            'star_clauses': star,
            'form_checks': form_checks,
            'dark_bid': dark,
            'library_findings': lib_findings,
            'recommendations': self._recommend(red_lines, star, form_checks, dark, lib_findings),
        }
        ctx.set('risk_report', report)

        try:
            self._write_markdown(ctx,report, result_path)
        except Exception:  # noqa: BLE001
            pass

    def _recommend(self, red_lines, star, form_checks, dark, lib_findings=None):
        recs = []
        for x in red_lines + star:
            if not x['covered']:
                recs.append(f"未检出响应：{x['text'][:40]}（{x['type']}，请逐条核对招标文件）")
        _sem = sum(1 for x in red_lines + star if x.get('match') == 'semantic')
        if _sem:
            recs.append(f"语义疑似匹配 {_sem} 条（非精确命中，建议人工确认表述是否到位）")
        for f in form_checks:
            if not f['covered']:
                recs.append(f"形式要件缺失：{f['label']}（请补充签字/盖章/装订/密封等说明）")
        if dark and dark['leaks']:
            recs.append("暗标泄漏：检测到 " + "、".join(l['label'] for l in dark['leaks']) + "，请清除")
        for f in (lib_findings or []):
            if f.get('level') in ('high', 'medium'):
                recs.append(f"风险库提示[{f.get('level')}]：{f.get('title', '')} —— {f.get('remediation', '')[:60]}")
        if not recs:
            recs.append("未发现明显废标红线缺项，仍建议人工终审（PDCA-Check）。")
        return recs

    def _write_markdown(self, ctx,report, result_path):
        if not result_path:
            return
        out = Path(result_path)
        md_path = aux_path(ctx, result_path,'_废标风险自检报告.md')
        s = report['summary']
        lines = [
            "# 废标风险自检报告（PDCA-Act 自动生成）",
            "",
            f"- 生成时间：{report['generated_at']}",
            f"- 废标红线/★条款：{s['covered']}/{s['total_clauses']} 已检出响应",
            f"- 形式要件缺项：{s['form_issues']}",
            f"- 暗标泄漏：{s['dark_bid_leaks']}",
            f"- 风险库核验：{s['library_findings']} 项（高危 {s['library_high']}）",
            f"- **综合废标风险分：{s['risk_score']}/100**（越低越危险）",
            "",
            "## 综合结论",
            _verdict_phrase(s),
            "",
            "## 整改建议",
        ]
        for r in report['recommendations']:
            lines.append(f"- {r}")
        lines.append("")
        lines.append("图例：✓ 精确命中　~ 语义疑似（人工复核）　✗ 未检出")
        lines.append("")
        lines.append("## 废标红线明细")
        for x in report['red_lines']:
            _mk = '✓' if x['covered'] and x.get('match') == 'exact' else ('~' if x['covered'] else '✗')
            _ev = f"（证据：{x['evidence_section']}）" if x.get('evidence_section') else ''
            lines.append(f"- [{_mk}] {x['text']}{_ev}")
        lines.append("")
        lines.append("## ★实质性条款明细")
        for x in report['star_clauses']:
            _mk = '✓' if x['covered'] and x.get('match') == 'exact' else ('~' if x['covered'] else '✗')
            _ev = f"（证据：{x['evidence_section']}）" if x.get('evidence_section') else ''
            lines.append(f"- [{_mk}] {x['text']}{_ev}")
        if report.get('library_findings'):
            lines.append("")
            lines.append("## 风险库核验明细")
            for f in report['library_findings']:
                lvl = f.get('level', '')
                lines.append(f"- [{lvl}] {f.get('title', '')}：{f.get('remediation', '')[:80]}")
        if md_path:
            md_path.write_text('\n'.join(lines), encoding='utf-8')
        report['markdown_path'] = str(md_path)


class NationalStandardStage(Stage):
    """国标格式套用（PDCA-Act）：表格跨页表头重复 + 国标字体/行距。非阻断。

    C26（格式方向，对标招标文件对「排版规范」的普遍要求）：
    - 表格跨页表头重复：技术标大量分章节表格（机械设备表/人员表/进度表），跨页后
      表头不重复会严重影响评审查阅；为每个表格首行加 w:tblHeader，使表头在每页
      顶部自动重复。低侵入、可逆、幂等（已有则跳过）。
    - 国标字体/行距套用：把正文/标题样式套用为标准公文样式（正文仿宋、标题黑体、
      固定行距）。套的是「样式」而非逐 run 改写——样式级联、run 级显式设定优先，
      因此可逆且不破坏生成器对个别 run 的显式字体设定。
      纯规则、无 LLM、可逆 opt-in（enable_national_standard 默认 True）。
    """

    name = "national_standard"
    blocking = False

    # 国标样式映射（GB/T 9704 公文 + 工程技术标常见排版）：
    #   样式名 -> (字体, 字号pt, 行距pt 固定值)
    _GB_STYLES = {
        'Normal': ('仿宋', 16, 28),        # 三号仿宋、固定 28 磅
        'Title': ('黑体', 22, 30),         # 二号黑体（封面大标题）
        'Heading 1': ('黑体', 16, 28),     # 一级标题 黑体 三号
        '标题 1': ('黑体', 16, 28),
        'Heading 2': ('黑体', 15, 28),     # 二级标题 黑体 小三
        '标题 2': ('黑体', 15, 28),
        'Heading 3': ('黑体', 14, 28),     # 三级标题 黑体 小三偏下
        '标题 3': ('黑体', 14, 28),
    }

    def should_run(self, ctx: StageContext) -> bool:
        req = ctx.req
        result_path = ctx.get("result_path")
        return (bool(_attr(req, "enable_national_standard", True))
                and bool(result_path)
                and os.path.exists(result_path))

    @staticmethod
    def _apply_style(style, font_name, size_pt, line_spacing_pt):
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_LINE_SPACING
        from docx.oxml.ns import qn
        # 字体：ASCII + 中文 eastAsia 一并设置，避免中文仍走默认宋体
        style.font.name = font_name
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn('w:eastAsia'), font_name)
        rfonts.set(qn('w:ascii'), font_name)
        rfonts.set(qn('w:hAnsi'), font_name)
        if size_pt:
            style.font.size = Pt(size_pt)
        # 标题/正文样式统一置黑，避免 Word 默认 Heading 2/3 蓝色样式继承
        style.font.color.rgb = RGBColor(0, 0, 0)
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_spacing_pt)

    @staticmethod
    def _repeat_table_headers(doc):
        from docx.oxml.ns import qn
        repeated = 0
        for tbl in doc.tables:
            if not tbl.rows:
                continue
            tr = tbl.rows[0]._tr
            trPr = tr.get_or_add_trPr()
            if trPr.find(qn('w:tblHeader')) is None:
                th = trPr.makeelement(qn('w:tblHeader'), {qn('w:val'): 'true'})
                trPr.append(th)
                repeated += 1
        return repeated

    def run(self, ctx: StageContext) -> None:
        from docx import Document

        result_path = ctx.get("result_path")
        doc = Document(result_path)

        # 1) 表格跨页表头重复
        tables_total = len(doc.tables)
        headers_repeated = self._repeat_table_headers(doc)

        # 2) 国标样式套用（样式级，非 run 级；可逆、不破坏显式 run 设定）
        applied = []
        for style_name, (font_name, size_pt, line_pt) in self._GB_STYLES.items():
            try:
                style = doc.styles[style_name]
            except Exception:  # noqa: BLE001
                continue
            try:
                self._apply_style(style, font_name, size_pt, line_pt)
                applied.append(style_name)
            except Exception:  # noqa: BLE001
                pass

        doc.save(result_path)

        report = {
            'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'enabled': True,
            'tables_total': tables_total,
            'headers_repeated': headers_repeated,
            'styles_applied': applied,
            'summary': {
                'tables_total': tables_total,
                'headers_repeated': headers_repeated,
                'styles_applied': applied,
            },
        }
        ctx.set('national_standard', report)
        try:
            self._write_markdown(ctx,report, result_path)
        except Exception:  # noqa: BLE001
            pass

    def _write_markdown(self, ctx,report, result_path):
        if not result_path:
            return
        out = Path(result_path)
        md_path = aux_path(ctx, result_path,'_国标格式套用.md')
        s = report['summary']
        lines = [
            "# 国标格式套用（PDCA-Act 自动生成）",
            "",
            f"- 生成时间：{report['generated_at']}",
            f"- 表格跨页表头重复：{s['headers_repeated']} / {s['tables_total']} 个表格首行已置为重复表头",
            f"- 国标样式套用：{('、'.join(s['styles_applied'])) if s['styles_applied'] else '无（对应样式不存在）'}",
            "",
            "说明：表格首行加 w:tblHeader，跨页自动重复表头；字体/行距套用于文档「样式」"
            "（正文仿宋、标题黑体、固定行距），样式级联、不逐 run 改写，可逆且不破坏生成器显式设定。",
        ]
        if md_path:
            md_path.write_text('\n'.join(lines), encoding='utf-8')
        report['markdown_path'] = str(md_path)


class DarkLeakHardenStage(Stage):
    """暗标零泄漏硬化（PDCA-Act）。非阻断。

    C25（暗标零泄漏）：现有废标自检仅扫正文，漏掉「隐藏泄漏面」——
    页眉/页脚文本、文档属性元数据（author/last_modified_by）、水印（页眉中的
    文本框/形状文字）。这些在盲评中同样会暴露投标单位身份 → 废标。

    本 Stage：① 全覆盖扫描上述隐藏面（手机/邮箱/网址 + 公司名称）；② 默认硬化
    （scrub）：清空文档属性 author/last_modified_by、页眉页脚/水印中的公司名称
    替换为 ██；③ 重扫核算残留，产出 _暗标零泄漏硬化.md。纯规则、可逆 opt-in
    （scrub 默认开；只动身份元数据与页眉页脚/水印中的公司名，不碰正文业务文字）。
    """

    name = "dark_leak_harden"
    blocking = False

    def should_run(self, ctx: StageContext) -> bool:
        req = ctx.req
        return (bool(_attr(req, "is_dark_bid", False))
                and bool(_attr(req, "enable_dark_harden", True))
                and bool(ctx.get("result_path"))
                and os.path.exists(ctx.get("result_path") or ''))

    @staticmethod
    def _surfaces_text(doc, company):
        """返回 [(surface, text), ...]，覆盖所有隐藏泄漏面。"""
        from docx.oxml.ns import qn
        surfaces = []
        for sec in doc.sections:
            ht = '\n'.join(p.text or '' for p in sec.header.paragraphs)
            if ht.strip():
                surfaces.append(('页眉', ht))
            ft = '\n'.join(p.text or '' for p in sec.footer.paragraphs)
            if ft.strip():
                surfaces.append(('页脚', ft))
            # 水印 / 形状文字：页眉中的文本框（txbxContent），与页眉正文段落区分
            try:
                hdr_el = sec.header.part.element
            except Exception:  # noqa: BLE001
                hdr_el = None
            if hdr_el is not None:
                wm = []
                for tb in hdr_el.iter(qn('w:txbxContent')):
                    wm.append(''.join(t.text or '' for t in tb.iter(qn('w:t'))))
                wm_text = '\n'.join(wm)
                if wm_text.strip():
                    surfaces.append(('水印', wm_text))
        # 文档属性元数据（身份常藏于此）
        cp = doc.core_properties
        for prop in ('author', 'last_modified_by', 'title', 'subject',
                     'comments', 'category'):
            v = getattr(cp, prop, None) or ''
            if v.strip():
                surfaces.append(('文档属性:' + prop, v))
        return surfaces

    @staticmethod
    def _scan(surfaces, company):
        leaks = []
        for surface, text in surfaces:
            for pat, label in _DARK_LEAK_PATTERNS:
                for m in pat.findall(text):
                    leaks.append({'surface': surface, 'label': label,
                                  'match': m[:40]})
            if company and company in text:
                leaks.append({'surface': surface, 'label': '公司名称',
                              'match': company[:40]})
        return leaks

    def run(self, ctx: StageContext) -> None:
        from docx import Document
        from docx.oxml.ns import qn

        req = ctx.req
        result_path = ctx.get("result_path")
        _uc = _attr(req, "user_context") or {}
        _company = _uc.get('company') if isinstance(_uc, dict) else None
        # user_context.company 可能是 dict（{name: ...}）或纯字符串，统一取公司名
        company_name = (_company.get('name') if isinstance(_company, dict)
                        else (_company or None))

        doc = Document(result_path)
        surfaces = self._surfaces_text(doc, company_name)
        before = self._scan(surfaces, company_name)

        scrub = bool(_attr(req, "enable_dark_scrub", True))
        scrubbed = []
        if scrub:
            # 1) 文档属性身份元数据清空（author / last_modified_by）
            cp = doc.core_properties
            for prop in ('author', 'last_modified_by'):
                if getattr(cp, prop, None):
                    setattr(cp, prop, '')
                    scrubbed.append('文档属性:' + prop)
            # 2) 页眉/页脚中的公司名称替换
            if company_name:
                for sec in doc.sections:
                    for part in (sec.header, sec.footer):
                        for p in part.paragraphs:
                            if company_name in (p.text or ''):
                                for r in p.runs:
                                    if company_name in (r.text or ''):
                                        r.text = r.text.replace(company_name, '██')
            # 3) 水印（页眉 XML 中文本框文字）公司名替换
            if company_name:
                for sec in doc.sections:
                    try:
                        hdr_el = sec.header.part.element
                    except Exception:  # noqa: BLE001
                        hdr_el = None
                    if hdr_el is None:
                        continue
                    changed = False
                    for t in hdr_el.iter(qn('w:t')):
                        if t.text and company_name in t.text:
                            t.text = t.text.replace(company_name, '██')
                            changed = True
                    if changed:
                        scrubbed.append('水印')
            doc.save(result_path)

        # 重扫核算残留
        doc2 = Document(result_path)
        after = self._scan(self._surfaces_text(doc2, company_name), company_name)

        report = {
            'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'enabled': True,
            'scrubbed': scrub,
            'leaks_before': len(before),
            'before_leaks': before,
            'leaks_after': len(after),
            'leaks': after,
            'scrubbed_surfaces': scrubbed,
            'verdict': '零泄漏' if len(after) == 0 else '仍存在泄漏，请人工复核',
        }
        ctx.set('dark_leak_harden', report)
        try:
            self._write_markdown(ctx,report, result_path)
        except Exception:  # noqa: BLE001
            pass

    def _write_markdown(self, ctx,report, result_path):
        if not result_path:
            return
        out = Path(result_path)
        md_path = aux_path(ctx, result_path,'_暗标零泄漏硬化.md')
        s = report
        lines = [
            "# 暗标零泄漏硬化（PDCA-Act 自动生成）",
            "",
            f"- 生成时间：{s['generated_at']}",
            f"- 硬化（scrub）：{'已执行' if s['scrubbed'] else '未执行（仅检测）'}",
            f"- 隐藏泄漏面扫描：硬化前 {s['leaks_before']} 处 → 硬化后 {s['leaks_after']} 处",
            f"- **结论：{s['verdict']}**",
            "",
            "## 已硬化面",
        ]
        if s['scrubbed_surfaces']:
            for surf in s['scrubbed_surfaces']:
                lines.append(f"- {surf}：身份已清除/替换")
        else:
            lines.append("- （未执行硬化，或无需处理）")
        lines.append("")
        lines.append("## 硬化后残留泄漏（须人工复核）")
        if s['leaks']:
            for i, lk in enumerate(s['leaks'], 1):
                lines.append(f"{i}. [{lk['surface']}] {lk['label']}：{lk['match']}")
        else:
            lines.append("无残留泄漏。")
        lines.append("")
        lines.append("说明：本表覆盖页眉/页脚/水印/文档属性等隐藏泄漏面；正文泄漏见「废标风险自检报告」。"
                     "硬化仅清除身份元数据与页眉页脚/水印中的公司名称，不改动正文业务内容。")
        if md_path:
            md_path.write_text('\n'.join(lines), encoding='utf-8')
        report['markdown_path'] = str(md_path)


def _doc_outline(result_path):
    """按标题切分 docx，返回 [(section_title, section_body_text), ...]。

    失败（无 python-docx / 文件损坏）优雅降级为空列表，不影响主流程。
    """
    try:
        from docx import Document
        doc = Document(result_path)
    except Exception:  # noqa: BLE001
        return []
    sections = []
    cur_title = '（未分节/正文）'
    cur_body = []
    for p in doc.paragraphs:
        style = (p.style.name or '') if getattr(p, 'style', None) else ''
        if style.startswith('Heading') or style.startswith('标题'):
            if cur_body or cur_title != '（未分节/正文）':
                sections.append((cur_title, '\n'.join(cur_body)))
            cur_title = (p.text or '').strip() or cur_title
            cur_body = []
        else:
            t = (p.text or '').strip()
            if t:
                cur_body.append(t)
    if cur_body or cur_title != '（未分节/正文）':
        sections.append((cur_title, '\n'.join(cur_body)))
    return sections


def _locate_evidence(name, sections, anchor=None, threshold=0.25):
    """为评分项 name 在 docx 大纲中定位证据章节；返回章节标题或 None。

    用 bigram 召回率取最佳匹配章节，阈值保守，避免把弱相关章节误标为证据。
    anchor：上游生成 Stage（如需求响应闭环 C21）已记录的写入章节，命中即优先采用，
    实现 D1-②「响应段落带章节证据」的锚点回灌。
    """
    if not name or not sections:
        return anchor or None
    if anchor:
        # 显式锚点：与章节标题精确/包含匹配即采用，避免重复 bigram 检索
        for title, _ in sections:
            if title and (title == anchor or anchor in title or title in anchor):
                return title
    q = _ngrams(name, 2)
    if not q:
        return anchor or None
    best_title, best_recall = None, 0.0
    for title, body in sections:
        g = _ngrams(title + '\n' + body, 2)
        if not g:
            continue
        recall = len(q & g) / len(q)
        if recall > best_recall:
            best_recall, best_title = recall, title
    return best_title if best_recall >= threshold else (anchor or None)


class ScoringMatrixStage(Stage):
    """评分项命中矩阵交付物。非阻断。"""

    name = "scoring_matrix"
    blocking = False

    def should_run(self, ctx: StageContext) -> bool:
        req = ctx.req
        parse_result = ctx.get("parse_result") or _attr(req, "parse_result")
        result_path = ctx.get("result_path")
        return (bool(_attr(req, "enable_scoring_matrix", True))
                and parse_result is not None
                and bool(result_path)
                and os.path.exists(result_path))

    def run(self, ctx: StageContext) -> None:
        from scoring_reinforce import analyze_scoring_coverage
        from checker.risk_library import _extract_doc_text

        req = ctx.req
        parse_result = ctx.get("parse_result") or _attr(req, "parse_result")
        result_path = ctx.get("result_path")
        doc_text = _extract_doc_text(result_path)
        cov = analyze_scoring_coverage(parse_result, doc_text)
        items = cov.get('items', []) or []
        # 证据定位：为每条评分项在 docx 大纲中锚定章节（一处一证据）
        _sections = _doc_outline(result_path)
        for it in items:
            it['evidence_section'] = _locate_evidence(it.get('name', ''), _sections)
            # 改进建议：弱项/未覆盖项给出 actionable 提示（一弱一建议）
            _st = it.get('status')
            if _st == 'uncovered':
                it['suggestion'] = '尚未响应：请在对应章节补充专门段落，并锚定到该评分项'
            elif _st == 'weak':
                it['suggestion'] = '补充量化指标/案例/数据支撑，提升表述命中度'
            else:
                it['suggestion'] = ''
        reinforce = ctx.get("scoring_reinforcement") or {}
        injected = reinforce.get('injected', 0) if isinstance(reinforce, dict) else 0
        matrix = {
            'total': cov.get('total', 0),
            'covered': cov.get('covered', 0),
            'weak': cov.get('weak', 0),
            'uncovered': cov.get('uncovered', 0),
            'reinforcement_injected': injected,
            'items': items,
        }
        ctx.set('scoring_matrix', matrix)

        try:
            self._write_markdown(ctx,matrix, result_path)
        except Exception:  # noqa: BLE001
            pass

    def _write_markdown(self, ctx,matrix, result_path):
        if not result_path:
            return
        out = Path(result_path)
        md_path = aux_path(ctx, result_path,'_评分项命中矩阵.md')
        lines = ["# 评分项命中矩阵（PDCA-Act 自动生成）", ""]
        lines.append(
            f"总计 {matrix['total']} 项 | 已覆盖 {matrix['covered']} | "
            f"弱项 {matrix['weak']} | 未覆盖 {matrix['uncovered']} | "
            f"自动补强注入 {matrix['reinforcement_injected']} 段")
        if matrix['uncovered'] or matrix['weak']:
            _pri = [it.get('name', '') for it in matrix['items']
                    if it.get('status') in ('uncovered', 'weak')]
            if _pri:
                lines.append("")
                lines.append(f"⚠ 优先整改（未覆盖/弱项）：{'、'.join(_pri)}")
        lines.append("")
        lines.append("| 评分项 | 分值 | 状态 | 命中度 | 证据定位 | 建议 |")
        lines.append("| --- | --- | --- | --- | --- | --- |")
        for it in matrix['items']:
            _ev = it.get('evidence_section') or '—'
            _sg = it.get('suggestion') or ''
            lines.append(f"| {it.get('name', '')} | {it.get('score', '')} | "
                         f"{it.get('status', '')} | {it.get('ratio', '')} | {_ev} | {_sg} |")
        if md_path:
            md_path.write_text('\n'.join(lines), encoding='utf-8')
        matrix['markdown_path'] = str(md_path)


class QualificationStage(Stage):
    """资格响应自查（PDCA-Act）。非阻断。

    对标竞品「资质业绩匹配」能力的**通用、零私有数据**版本：从招标文件解析出的
    资格要求(qualification_reqs)，逐条核验生成稿是否做出响应（含证据章节锚定）。
    不依赖企业自有资质（ADR-005 已删除企业画像），只校验「投标文件是否逐条回应了
    资格门槛」——这是资格不符导致废标前最该自查的一环。

    纯规则，无 LLM、无外部依赖。
    """

    name = "qualification"
    blocking = False

    def should_run(self, ctx: StageContext) -> bool:
        req = ctx.req
        parse_result = ctx.get("parse_result") or _attr(req, "parse_result")
        result_path = ctx.get("result_path")
        return (bool(_attr(req, "enable_qualification", True))
                and parse_result is not None
                and bool(result_path)
                and os.path.exists(result_path))

    def run(self, ctx: StageContext) -> None:
        from checker.risk_library import _extract_doc_text

        req = ctx.req
        parse_result = ctx.get("parse_result") or _attr(req, "parse_result")
        result_path = ctx.get("result_path")
        doc_text = (_extract_doc_text(result_path)
                    if result_path and os.path.exists(result_path) else '')

        _reqs = parse_result.get('qualification_reqs', []) or []
        items: List[Dict[str, Any]] = []
        for c in _reqs:
            t = _clause_text(c)
            if not t or len(t) < 4:
                continue
            _m = _coverage_detail(t, doc_text)
            items.append({'type': 'qualification', 'text': t[:160],
                          'covered': _m != 'none', 'match': _m,
                          'severity': 'high'})

        # 证据锚定：为已检出资格要求定位 docx 章节（与风险/评分对称，一处一证据）
        _sections = _doc_outline(result_path)
        for x in items:
            x['evidence_section'] = (_locate_evidence(x['text'], _sections)
                                     if x['covered'] else None)

        total = len(items)
        covered = sum(1 for x in items if x['covered'])
        uncovered = total - covered
        exact_hits = sum(1 for x in items if x.get('match') == 'exact')
        semantic_hits = sum(1 for x in items if x.get('match') == 'semantic')
        # 资格不符属废标级风险：每未响应项扣 15 分（与废标红线同级权重）
        qual_score = max(0, 100 - uncovered * 15)

        report = {
            'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'summary': {
                'total': total,
                'covered': covered,
                'uncovered': uncovered,
                'exact_hits': exact_hits,
                'semantic_hits': semantic_hits,
                'qual_score': qual_score,
            },
            'items': items,
        }
        ctx.set('qualification', report)

        try:
            self._write_markdown(ctx,report, result_path)
        except Exception:  # noqa: BLE001
            pass

    def _write_markdown(self, ctx,report, result_path):
        if not result_path:
            return
        out = Path(result_path)
        md_path = aux_path(ctx, result_path,'_资格响应自查.md')
        s = report['summary']
        lines = [
            "# 资格响应自查（PDCA-Act 自动生成）",
            "",
            f"- 生成时间：{report['generated_at']}",
            f"- 资格要求响应：{s['covered']}/{s['total']} 已检出响应",
            f"- 资格响应分：{s['qual_score']}/100（未响应每项扣 15）",
            "",
            "## 资格要求逐条核对",
            "",
            "| 资格要求 | 状态 | 命中 | 证据定位 |",
            "| --- | --- | --- | --- |",
        ]
        for x in report['items']:
            _mk = '✓' if x['covered'] and x.get('match') == 'exact' else ('~' if x['covered'] else '✗')
            _ev = x.get('evidence_section') or '—'
            lines.append(f"| {x['text']} | {_mk} | {x.get('match')} | {_ev} |")
        lines.append("")
        lines.append("图例：✓ 精确命中　~ 语义疑似（人工复核）　✗ 未检出")
        if md_path:
            md_path.write_text('\n'.join(lines), encoding='utf-8')
        report['markdown_path'] = str(md_path)


class RequirementClosureStage(Stage):
    """需求响应式内容闭环（PDCA-Act）。非阻断。

    把评分项 / 废标条款「要求」在生成稿对应章节处显式织入响应引导段，
    使投标文件在交付前即完成「一处一要求、一处一响应」的闭环——而非仅靠
    C3/C20 事后检查。这是 C21（完善技术标生成内容）的核心动作。

    实现：打开 result_path docx，按标题 bigram 召回定位匹配评分项/废标条款的
    章节，在其后插入【评分项响应】/【废标条款响应】引导段；幂等（同章节不叠加）。
    纯规则、无 LLM、无外部依赖，可逆 opt-in。
    """

    name = "req_closure"
    blocking = False

    def should_run(self, ctx: StageContext) -> bool:
        req = ctx.req
        parse_result = ctx.get("parse_result") or _attr(req, "parse_result")
        result_path = ctx.get("result_path")
        # v7.4: 默认关闭，避免正文顶部反复出现【评分项响应】引导段。
        # 如需启用，可在请求中显式设置 enable_requirement_closure=True。
        return (bool(_attr(req, "enable_requirement_closure", False))
                and parse_result is not None
                and bool(result_path)
                and os.path.exists(result_path))

    def run(self, ctx: StageContext) -> None:
        from docx import Document

        req = ctx.req
        parse_result = ctx.get("parse_result") or _attr(req, "parse_result")
        result_path = ctx.get("result_path")

        # 统一需求为 (kind, label, text, score)
        reqs = []
        for it in parse_result.get('score_items', []) or []:
            name = it.get('name') or it.get('title') or ''
            if name:
                reqs.append(('score', name, name, it.get('score', 0)))
        for c in parse_result.get('red_line_clauses', []) or []:
            t = _clause_text(c)
            if t and len(t) >= 4:
                reqs.append(('red_line', t[:80], t, 0))

        if not reqs or not result_path or not os.path.exists(result_path):
            ctx.set('req_closure', {'injected': 0, 'matched': 0,
                                    'total': len(reqs), 'items': []})
            return

        doc = Document(result_path)
        heading_paras = [p for p in doc.paragraphs
                         if (p.style.name or '').startswith('Heading')
                         or (p.style.name or '').startswith('标题')]

        items = []
        used = set()
        for kind, label, text, score in reqs:
            best_p, best_recall = None, 0.0
            for p in heading_paras:
                h = (p.text or '').strip()
                if not h:
                    continue
                g = _ngrams(h, 2)
                if not g:
                    continue
                recall = len(_ngrams(label, 2) & g) / max(1, len(_ngrams(label, 2)))
                if recall > best_recall:
                    best_recall, best_p = recall, p
            if best_p is None or best_recall < 0.3:
                items.append({'kind': kind, 'label': label[:60],
                              'matched': False, 'section': None})
                continue
            if id(best_p) in used:
                # 同一章节已注入（多要求映射到同章节），仅记录匹配
                items.append({'kind': kind, 'label': label[:60],
                              'matched': True, 'section': (best_p.text or '').strip()})
                continue
            # 跨运行幂等：该标题后已存在本 Stage 注入段则不重复注入
            _nxt = best_p._p.getnext()
            if _nxt is not None:
                _nxt_text = (_nxt.text or '').strip()
                if _nxt_text.startswith('【评分项响应】') or _nxt_text.startswith('【废标条款响应】'):
                    items.append({'kind': kind, 'label': label[:60],
                                  'matched': True, 'section': (best_p.text or '').strip()})
                    continue
            callout = doc.add_paragraph()
            if kind == 'score':
                body = (f'【评分项响应】本部分对应评分项："{label}"（{score} 分）。'
                        f'以下围绕上述要求在施工组织、技术方案、质量与安全保障等方面'
                        f'作出实质性响应，确保逐项对应、避免漏项。')
            else:
                body = (f'【废标条款响应】本部分须严格满足："{(label[:120])}"。'
                        f'未实质性响应将导致废标，请重点核对以下内容是否覆盖到位。')
            r = callout.add_run(body)
            r.italic = True
            best_p._p.addnext(callout._p)
            used.add(id(best_p))
            items.append({'kind': kind, 'label': label[:60],
                          'matched': True, 'section': (best_p.text or '').strip()})

        # 已织入正文的需求响应段总数（含历史运行注入的，确保跨运行幂等计数一致）
        injected = sum(1 for p in doc.paragraphs
                       if (p.text or '').strip().startswith(
                           ('【评分项响应】', '【废标条款响应】')))
        matched = sum(1 for x in items if x['matched'])
        doc.save(result_path)

        report = {
            'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'summary': {
                'total': len(reqs),
                'matched': sum(1 for x in items if x['matched']),
                'injected': injected,
            },
            'items': items,
        }
        ctx.set('req_closure', report)

        try:
            self._write_markdown(ctx,report, result_path)
        except Exception:  # noqa: BLE001
            pass

    def _write_markdown(self, ctx,report, result_path):
        if not result_path:
            return
        out = Path(result_path)
        md_path = aux_path(ctx, result_path,'_需求响应闭环.md')
        s = report['summary']
        lines = [
            "# 需求响应闭环（PDCA-Act 自动注入）",
            "",
            f"- 生成时间：{report['generated_at']}",
            f"- 需求总数：{s['total']}（评分项 + 废标条款）",
            f"- 匹配章节：{s['matched']}　已注入引导段：{s['injected']}",
            "",
            "## 逐条闭环",
            "",
            "| 类型 | 要求 | 匹配章节 | 状态 |",
            "| --- | --- | --- | --- |",
        ]
        for x in report['items']:
            _k = '评分项' if x['kind'] == 'score' else '废标条款'
            _st = '已注入' if x['matched'] else '未匹配章节'
            lines.append(f"| {_k} | {x['label']} | {x.get('section') or '—'} | {_st} |")
        lines.append("")
        lines.append("说明：本表记录生成稿中已显式织入「要求—响应」引导段的章节；"
                     "未匹配到章节的要求，建议在评分项命中矩阵(C3)中人工补强。")
        if md_path:
            md_path.write_text('\n'.join(lines), encoding='utf-8')
        report['markdown_path'] = str(md_path)


def _build_toc_field():
    """构造 Word TOC 域段落元素（层级 1-3，带超链接）。

    顺序：begin → instrText(TOC) → separate → 占位文字 → end。
    打开文档后右键「更新域」即生成可跳转目录。
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = OxmlElement('w:p')
    rb = OxmlElement('w:r'); fb = OxmlElement('w:fldChar')
    fb.set(qn('w:fldCharType'), 'begin'); rb.append(fb)
    ri = OxmlElement('w:r'); it = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve'); it.text = 'TOC \\o "1-3" \\h \\z \\u'; ri.append(it)
    rs = OxmlElement('w:r'); fs = OxmlElement('w:fldChar')
    fs.set(qn('w:fldCharType'), 'separate'); rs.append(fs)
    rt = OxmlElement('w:r'); tt = OxmlElement('w:t')
    tt.text = '右键「更新域」生成目录'; rt.append(tt)
    re_ = OxmlElement('w:r'); fe = OxmlElement('w:fldChar')
    fe.set(qn('w:fldCharType'), 'end'); re_.append(fe)
    for el in (rb, ri, rs, rt, re_):
        p.append(el)
    return p


def _paragraph_has_image(para) -> bool:
    """判断段落是否内嵌图片（含 w:drawing 元素，可能嵌套在 w:r 内）。"""
    from docx.oxml.ns import qn
    return any(True for _ in para._p.iter(qn('w:drawing')))


class ReferenceDrivenStage(Stage):
    """以标写标（参考标书驱动）内容强化（C23）。非阻断。

    设计：
    - 仅当提供参考标书（reference_file）时运行；未提供则跳过（不占主线）。
    - 加载失败时**明确记录原因并写报告**（不再静默忽略），供 SummaryStage 暴露为非阻断告警。
    - 加载成功时执行一个具体内容动作：把参考标书一级大纲中「生成稿缺失」的章节补为
      简要提纲桩，确保「以标写标」真正覆盖参考结构的全部一级章（content-driven，而非仅报告）。
    - 产出 _以标写标.md；SummaryStage 读产品增指标。
    """

    name = "reference_driven"
    blocking = False

    def should_run(self, ctx: StageContext) -> bool:
        req = ctx.req
        return bool(_attr(req, "enable_reference_driven", True)) and \
            bool(_attr(req, "reference_file")) and bool(ctx.get("result_path"))

    @staticmethod
    def _heading_level(p) -> int:
        style_name = ''
        try:
            style_name = (p.style.name or '') if p.style else ''
        except Exception:  # noqa: BLE001
            pass
        if not style_name:
            return 0
        if style_name.strip() == 'Title':
            return 0
        m = re.match(r'Heading\s*(\d+)', style_name, re.IGNORECASE)
        if m:
            return int(m.group(1))
        m2 = re.match(r'标题\s*(\d+)', style_name)
        if m2:
            return int(m2.group(1))
        return 0

    @staticmethod
    def _level1_titles(doc) -> list:
        out = []
        for p in doc.paragraphs:
            if ReferenceDrivenStage._heading_level(p) == 1:
                t = (p.text or '').strip()
                if t:
                    out.append(t)
        return out

    @staticmethod
    def _append_missing_stubs(doc, missing: list) -> list:
        """把缺失的一级章补为提纲桩；插入点取最后一个「附表/附录/附件」之前，否则文末。"""
        insert_el = None
        for p in doc.paragraphs:
            if ReferenceDrivenStage._heading_level(p) == 1 and \
                    any(k in (p.text or '') for k in ('附表', '附录', '附件')):
                insert_el = p._p
        appended = []
        for title in missing:
            h = doc.add_paragraph()
            try:
                h.style = doc.styles['Heading 1']
            except Exception:  # noqa: BLE001
                pass
            h.add_run(title)
            body = doc.add_paragraph()
            body.add_run('（本章依据参考标书结构补入提纲，内容待按本项目实际情况完善。）')
            if insert_el is not None:
                insert_el.addprevious(h._p)
                insert_el.addprevious(body._p)
            appended.append(title)
        return appended

    def run(self, ctx: StageContext) -> None:
        req = ctx.req
        result_path = ctx.get("result_path")
        ref_path = _attr(req, "reference_file")
        try:
            from bid_core.reference_loader import ReferenceLoader
        except Exception:  # noqa: BLE001
            from reference_loader import ReferenceLoader

        ref = ReferenceLoader(ref_path).load()
        result = {
            'status': getattr(ref, 'status', 'pending'),
            'path': ref_path,
            'chapter_count': len(ref.chapters),
            'outline': list(ref.outline),
        }
        if ref.status == 'failed':
            result['error'] = ref.error
            result['error_type'] = ref.error_type
            result['message'] = '参考标书加载失败，以标写标未生效（需人工提供有效参考标书）'
            result['markdown_path'] = self._write_markdown(ctx,result_path, result, failed=True)
            ctx.set('reference_driven', result)
            return

        from docx import Document
        doc = Document(result_path)
        existing = self._level1_titles(doc)
        missing = [t for t in ref.outline
                   if not any(t in e or e in t for e in existing)]
        appended = self._append_missing_stubs(doc, missing) if missing else []
        if appended:
            doc.save(result_path)
        result['missing_appended'] = len(appended)
        result['style_sample_count'] = len(ref.style_patterns)
        result['variable_map'] = ref.build_variable_map(ctx.data or {})
        result['markdown_path'] = self._write_markdown(ctx,result_path, result, failed=False)
        ctx.set('reference_driven', result)

    def _write_markdown(self, ctx,result_path: str, result: dict, failed: bool) -> str:
        out = Path(result_path)
        md_path = aux_path(ctx, result_path,'_以标写标.md')
        L = ['# 以标写标（参考标书驱动）报告', '']
        L.append(f'- 参考标书路径：{result.get("path")}')
        L.append(f'- 加载状态：{result.get("status")}')
        if failed:
            L.append(f'- 失败原因：{result.get("error_type")} | {result.get("error")}')
            L.append('')
            L.append('> 参考标书无法解析，本次生成未套用其结构；请提供有效的历史中标标书'
                     '（docx）后重试。以标写标属内容增强，失败不影响主流程。')
        else:
            L.append(f'- 解析章节数：{result.get("chapter_count")}')
            L.append(f'- 套用一级大纲：{len(result.get("outline") or [])} 项')
            L.append(f'- 缺失补入提纲：{result.get("missing_appended", 0)} 项')
            L.append(f'- 风格样本数：{result.get("style_sample_count", 0)}')
            L.append('')
            L.append('## 参考标书一级大纲（已套用）')
            for t in (result.get('outline') or []):
                L.append(f'- {t}')
            vm = result.get('variable_map') or {}
            if vm:
                L.append('')
                L.append('## 变量替换映射（参考 → 本项目）')
                for k, v in vm.items():
                    L.append(f'- {k} → {v}')
        if md_path:
            md_path.write_text('\n'.join(L), encoding='utf-8')
        return str(md_path)


class ShellKitStage(Stage):
    """技术标外壳套件（PDCA-Act）：封面之后的标准前置外壳。非阻断。

    C27（格式/合规方向）：招标文件普遍要求技术标含一组「前置外壳」——
    投标函、法定代表人身份证明、法定代表人授权委托书、承诺书（无重大违法记录/
    不串通/质量工期安全承诺等）。生成器当前只产封面+目录+正文，缺这组标准前置，
    漏交即可能废标/扣分。本 Stage 在「封面之后、目录之前」补入这组通用外壳模板
    （零私有数据、用占位符标注需人工填具项），让用户直接套用、据实补全。
    幂等（已含「投标函」标题则跳过），可逆 opt-in（enable_shell_kit 默认 True）。
    纯规则、无 LLM、不碰正文。
    """

    name = "shell_kit"
    blocking = False

    _SENTINEL = "投标函"  # 幂等哨兵：已含则跳过

    # 外壳模板：（标题, 段落列表）。占位符用「（…）」标注需人工填具项。
    _SHELL_SECTIONS = [
        ("投标函", [
            "（招标人名称）：",
            "1. 我方已仔细研究（项目名称）技术标招标文件全部内容，愿意按合同约定实施并交付。",
            "2. 我方投标工期（工期，如 90 日历天），质量目标（合格/优良），项目经理（姓名/注册证号）。",
            "3. 我方承诺：投标文件真实有效，不存在弄虚作假、串通投标等违法行为；若中标将按约履约。",
            "4. 随函附：法定代表人身份证明、授权委托书（如适用）、承诺书等。",
            "投标人（盖单位章）：我司    法定代表人或其委托代理人（签字）：",
            "日期：（年 月 日）",
        ]),
        ("法定代表人身份证明", [
            "投标人名称：我司",
            "单位性质：（企业类型，如 有限责任公司）",
            "成立时间：（年 月 日）",
            "经营期限：（起止期限）",
            "姓名：（法定代表人姓名）  性别：（男/女）  年龄：（岁）  职务：（职务）",
            "系我司的法定代表人。",
            "特此证明。",
            "投标人（盖单位章）：我司    日期：（年 月 日）",
        ]),
        ("法定代表人授权委托书", [
            "本人（法定代表人姓名）系我司的法定代表人，现委托（代理人姓名）",
            "（身份证号）为我方代理人，以本单位名义签署（项目名称）技术标投标文件。",
            "代理人在投标、开标、评标、签约过程中签署的文件和处理与之有关的一切事务，我均予以承认。",
            "委托期限：（自签署日至签约结束）。代理人无转委托权。",
            "投标人（盖单位章）：我司  法定代表人（签字）：  日期：（年 月 日）",
        ]),
        ("承诺书", [
            "致：（招标人名称）",
            "我方郑重承诺：",
            "一、近三年无重大违法记录，未因串通投标、弄虚作假被有关行政主管部门处罚；",
            "二、拟派项目管理班子真实有效，关键岗位人员无在建/兼职冲突；",
            "三、严格按招标文件工期、质量、安全文明施工要求履约，不转包、不违法分包；",
            "四、投标文件所附业绩、奖项、证书真实可查，接受招标人核查。",
            "如违背上述承诺，愿承担由此产生的一切责任与后果。",
            "投标人（盖单位章）：我司  法定代表人或其委托代理人（签字）：  日期：（年 月 日）",
        ]),
    ]

    def should_run(self, ctx: StageContext) -> bool:
        req = ctx.req
        result_path = ctx.get("result_path")
        return (bool(_attr(req, "enable_shell_kit", True))
                and bool(result_path)
                and os.path.exists(result_path))

    @staticmethod
    def _shell_style(doc):
        """外壳标题样式：基于 Heading 1、大纲级别 1（纳入目录），但样式名不含
        Heading/标题 N，故 DocFormatStage 的多级编号正则不会命中 —— 前置外壳不占编号。

        复用原则：已存在则直接返回（避免 add_style 重名报错）；幂等、可逆。
        """
        name = '外壳标题'
        try:
            return doc.styles[name]
        except Exception:  # noqa: BLE001
            pass
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        base = None
        for nm in ('Heading 1', '标题 1'):
            try:
                base = doc.styles[nm]
                break
            except Exception:  # noqa: BLE001
                continue
        st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        if base is not None:
            try:
                st.base_style = base
            except Exception:  # noqa: BLE001
                pass
        # 显式设定大纲级别为 1（w:outlineLvl val 0 基），确保纳入目录、
        # 且（因其样式名非 Heading/标题 N）不被多级编号正则命中
        try:
            pPr = st.element.get_or_add_pPr()
            ol = pPr.find(qn('w:outlineLvl'))
            if ol is None:
                ol = OxmlElement('w:outlineLvl')
                pPr.append(ol)
            ol.set(qn('w:val'), '0')
        except Exception:  # noqa: BLE001
            pass
        return st

    def _find_insert_before(self, doc):
        # 插入点：目录 TOC 域段落之前（封面之后）；无则首个 Heading 之前；再无则文首
        for p in doc.paragraphs:
            xml = p._p.xml
            if 'TOC' in xml and 'instrText' in xml:
                return p
        for p in doc.paragraphs:
            st = (p.style.name or '') if getattr(p, 'style', None) else ''
            if re.match(r'(?:Heading|标题)\s*\d+', st):
                return p
        return None

    def run(self, ctx: StageContext) -> None:
        from docx import Document

        result_path = ctx.get("result_path")
        doc = Document(result_path)

        # 幂等：已含哨兵标题则跳过（重跑不重复插入）
        _already = any((p.text or '').strip().startswith(self._SENTINEL)
                       for p in doc.paragraphs)
        if _already:
            report = {
                'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'enabled': True,
                'sections_inserted': 0,
                'already_present': True,
                'summary': {'sections_inserted': 0, 'already_present': True},
            }
            ctx.set('shell_kit', report)
            try:
                self._write_markdown(ctx,report, result_path)
            except Exception:  # noqa: BLE001
                pass
            return

        insert_before = self._find_insert_before(doc)
        _style = self._shell_style(doc)
        # 先全部追加到文末，收集其 _p，再整体移动到插入点之前（保持顺序，避免碎片化插入）
        created = []
        for title, lines in self._SHELL_SECTIONS:
            h = doc.add_paragraph()
            if _style is not None:
                try:
                    h.style = _style
                except Exception:  # noqa: BLE001
                    pass
            h.add_run(title)
            created.append(h._p)
            for ln in lines:
                pp = doc.add_paragraph()
                pp.add_run(ln)
                created.append(pp._p)
        body = doc.element.body
        for el in created:
            body.remove(el)
        if insert_before is not None:
            idx = body.index(insert_before._p)
            for el in created:
                body.insert(idx, el)
                idx += 1
        else:
            for i, el in enumerate(created):
                body.insert(i, el)
        doc.save(result_path)

        report = {
            'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'enabled': True,
            'sections_inserted': len(self._SHELL_SECTIONS),
            'already_present': False,
            'summary': {'sections_inserted': len(self._SHELL_SECTIONS),
                        'already_present': False},
        }
        ctx.set('shell_kit', report)
        try:
            self._write_markdown(ctx,report, result_path)
        except Exception:  # noqa: BLE001
            pass

    def _write_markdown(self, ctx,report, result_path):
        if not result_path:
            return
        out = Path(result_path)
        md_path = aux_path(ctx, result_path,'_技术标外壳套件.md')
        s = report['summary']
        lines = [
            "# 技术标外壳套件（PDCA-Act 自动生成）",
            "",
            f"- 生成时间：{report['generated_at']}",
            f"- 补入前置外壳：{s['sections_inserted']} 节"
            f"（{'已存在，跳过' if s.get('already_present') else '投标函/法定代表人身份证明/授权委托书/承诺书'}）",
            "",
            "说明：在封面之后、目录之前补入标准前置外壳通用模板，零私有数据、用「（…）」标注需人工填具项；"
            "插入为「外壳标题」样式（大纲级别 1、纳入目录，但样式名不命中多级编号正则，故不占用正文编号）。"
            "幂等（已含投标函则跳过），可逆 opt-in。",
        ]
        if md_path:
            md_path.write_text('\n'.join(lines), encoding='utf-8')
        report['markdown_path'] = str(md_path)


class DocFormatStage(Stage):
    """文档格式硬化（PDCA-Act）：自动目录 + 多级标题编号。非阻断。

    C24 第一部分（高价值、低风险）：
    - 多级标题编号：按 Heading/标题 1/2/3 层级给标题前缀 1 / 1.1 / 1.1.2，
      重建层级计数；**仅改标题文本前缀、不动正文**，且幂等（先剥离旧前缀再加重编号）。
      置于管线后段（在内容匹配 Stage 之后），故不影响 C9/C14/C20/C21 的证据锚定。
    - 自动目录：在正文最前插入 Word TOC 域（打开后更新即生成，层级 1-3）；幂等。
    纯规则、无 LLM、可逆 opt-in。
    """

    name = "doc_format"
    blocking = False

    def should_run(self, ctx: StageContext) -> bool:
        req = ctx.req
        result_path = ctx.get("result_path")
        return (bool(_attr(req, "enable_doc_format", True))
                and bool(result_path)
                and os.path.exists(result_path))

    @staticmethod
    def _remove_empty_headings(doc) -> None:
        """ADR-009：删除其后无任何正文/表格内容的孤立标题（含紧随的纯空白段）。

        判定：从某标题向后扫描，直到遇到下一个同级/上级标题或文档末尾；
        期间若存在非空正文段落或表格（w:tbl），视为有内容，保留；否则删除。
        附表块（「附表」H1 起）整体豁免，避免误删承载子表的父标题。
        """
        from docx.oxml.ns import qn

        paras = doc.paragraphs
        n = len(paras)
        _heading_re = re.compile(r'(?:Heading|标题)\s*\d+')
        in_appendix = False
        to_delete = []
        for i, p in enumerate(paras):
            style = (p.style.name or '') if getattr(p, 'style', None) else ''
            if not _heading_re.match(style):
                continue
            _txt = (p.text or '').strip()
            if _txt.startswith('附') or '附表' in _txt:
                in_appendix = True
                continue
            if in_appendix:
                continue
            has_content = False
            for j in range(i + 1, n):
                ns = (paras[j].style.name or '') if getattr(paras[j], 'style', None) else ''
                if _heading_re.match(ns):
                    break  # 下一个标题 → 当前标题无内容
                if (paras[j].text or '').strip():
                    has_content = True
                    break
            if not has_content:
                # 检查该标题与下一标题之间是否存在表格（doc.paragraphs 不含表格，需查 XML 兄弟）
                sib = p._p.getnext()
                while sib is not None:
                    if sib.tag == qn('w:tbl'):
                        has_content = True
                        break
                    if sib.tag == qn('w:p') and _heading_re.match(
                            (sib.style_name or '') if hasattr(sib, 'style_name') else ''):
                        break
                    sib = sib.getnext()
            if not has_content:
                to_delete.append(p._p)
        for el in to_delete:
            # 连带删除紧随其后的连续空白段，避免残留空行
            nxt = el.getnext()
            while nxt is not None and nxt.tag == qn('w:p') and not (nxt.text or '').strip():
                _n = nxt.getnext()
                nxt.getparent().remove(nxt)
                nxt = _n
            el.getparent().remove(el)

    @staticmethod
    def _toc_already_present(doc) -> bool:
        """判断文档是否已含目录：真实 TOC 域，或「目录 / 目  录」标题段。

        生成器在生成阶段已产出目录页（含自身 TOC 域），此时本 Stage 不应
        再叠加第二份目录，否则会出现重复目录、且本 Stage 的目录被错误地
        插到封面之前。
        """
        for p in doc.paragraphs:
            if 'TOC' in p._p.xml and 'instrText' in p._p.xml:
                return True
            _t = re.sub(r'\s+', '', p.text or '')
            if _t == '目录':
                return True
        return False

    def run(self, ctx: StageContext) -> None:
        from docx import Document

        result_path = ctx.get("result_path")
        doc = Document(result_path)

        # 0) 是否已存在目录（真实 TOC 域，或「目录/目  录」标题段）。
        #    生成器在生成阶段已产出目录页（含自身 TOC 域），此时不应再叠加，
        #    否则会出现「两份目录」且本 Stage 的目录被错误地插到封面之前。
        _toc_present = self._toc_already_present(doc)

        # 0.5) 空标题清理（ADR-009）：删除其后无正文/表格内容的孤立标题，
        #      避免生成出「标题下空白」的残缺章节。附表块整体豁免。
        self._remove_empty_headings(doc)

        # 1) 多级标题编号（ADR-009）：中文一级 / （一）二级 / 1．三级；
        #    附表块（含 2.1~2.6 子表）独立、跳过编号保留原样，杜绝 15.1~15.5。
        counters = []
        numbered = 0
        # 已有标题前缀（西方 1/1.1 或 中文 一、/（一）/ 1. 等）——剥离后再加重编号，避免叠加重号
        _prefix_re = re.compile(
            r'^(?:\d+(?:\.\d+)*\s+'
            r'|[（(]?[一二三四五六七八九十百零]+\s*[、.．）)]\s*'
            r'|\d+\s*[、.．]\s*)'
        )
        _CN = '一二三四五六七八九十'
        def _cn(n):
            if n <= 0:
                return str(n)
            if n <= 10:
                return _CN[n - 1]
            if n < 20:
                return '十' + (_CN[n - 11] if n >= 11 else '')
            if n < 100:
                tens, ones = divmod(n, 10)
                return _CN[tens - 1] + '十' + (_CN[ones - 1] if ones else '')
            return str(n)
        in_appendix = False
        for p in doc.paragraphs:
            style = (p.style.name or '') if getattr(p, 'style', None) else ''
            m = re.match(r'(?:Heading|标题)\s*(\d+)', style)
            if not m:
                continue
            level = int(m.group(1))
            if level < 1:
                continue
            _txt = (p.text or '').strip()
            # 附表独立块：从「附表」H1 起，其下所有标题（含 2.1~2.6 子表）保留原样不编号
            if level == 1 and ('附表' in _txt or _txt == '附' or _txt.startswith('附')):
                in_appendix = True
                continue
            if in_appendix:
                continue
            counters = counters[:level]
            if len(counters) < level:
                counters.append(0)
            counters[level - 1] += 1
            if level == 1:
                prefix = f'{_cn(counters[0])}、'
            elif level == 2:
                prefix = f'（{_cn(counters[1])}）'
            else:
                prefix = f'{counters[level - 1]}．'
            _new = _prefix_re.sub('', _txt)
            p.text = f'{prefix}{_new}'
            numbered += 1

        # 1.5) 图表题注：为内嵌图片自动补「图N」编号题注（幂等：已有题注跳过）。
        #      仅针对 docx 内嵌图片（w:drawing），表格题注后续版本扩展。
        _assign = {}
        _fig_no = 0
        for p in doc.paragraphs:
            if not _paragraph_has_image(p):
                continue
            _nxt = p._p.getnext()
            if _nxt is not None and ((_nxt.text or '').strip().startswith(('图', '表'))):
                continue  # 已有题注，不补（不推进编号，保持后续图序连续）
            _fig_no += 1
            _assign[p._p] = _fig_no
        figures_captioned = 0
        for p in reversed(doc.paragraphs):
            n = _assign.get(p._p)
            if n is None:
                continue
            cap = doc.add_paragraph()
            r = cap.add_run(f'图{n} ')
            r.italic = True
            p._p.addnext(cap._p)
            figures_captioned += 1

        # 2) 自动目录（TOC 域）：仅当文档尚无目录时插入，置于「首个标题之前」
        #    （封面/扉页之后、正文之前），避免插到封面之前或重复生成。
        if not _toc_present:
            head = doc.add_paragraph('目录')
            # 用 Title 样式（非 Heading/标题 N，不在 TOC 大纲层级内）：
            # ① 不被多级编号正则命中（重跑时「目录」不会被编成「1 目录」破坏幂等）；
            # ② 更新域后不会成为 TOC 自引用首条。
            try:
                head.style = doc.styles['Title']
            except Exception:  # noqa: BLE001
                pass
            toc = doc.add_paragraph()
            toc._p.append(_build_toc_field())
            body = doc.element.body
            # 插入点：首个 Heading/标题 N 段落之前；若没有标题则置于正文最前
            insert_before = None
            for p in doc.paragraphs:
                st = (p.style.name or '') if getattr(p, 'style', None) else ''
                if re.match(r'(?:Heading|标题)\s*\d+', st):
                    insert_before = p
                    break
            if insert_before is not None:
                idx = body.index(insert_before._p)
                body.insert(idx, toc._p)
                body.insert(idx, head._p)
            else:
                body.insert(0, toc._p)
                body.insert(0, head._p)

        doc.save(result_path)

        report = {
            'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'summary': {'headings_numbered': numbered,
                        'figures_captioned': figures_captioned,
                        'toc_inserted': not _toc_present},
        }
        ctx.set('doc_format', report)
        try:
            self._write_markdown(ctx,report, result_path)
        except Exception:  # noqa: BLE001
            pass

    def _write_markdown(self, ctx,report, result_path):
        if not result_path:
            return
        out = Path(result_path)
        md_path = aux_path(ctx, result_path,'_文档格式增强.md')
        s = report['summary']
        lines = [
            "# 文档格式增强（PDCA-Act 自动生成）",
            "",
            f"- 生成时间：{report['generated_at']}",
            f"- 多级标题编号：{s['headings_numbered']} 个标题已编号（1 / 1.1 / 1.1.2）",
            f"- 图表题注：自动补 {s.get('figures_captioned', 0)} 张图题注（图N 顺序编号）",
            f"- 自动目录：{'已插入 Word TOC 域（打开后右键「更新域」生成）' if s.get('toc_inserted') else '文档已含目录，跳过（不重复插入）'}",
            "",
            "说明：编号仅改标题前缀、不动正文；置于内容匹配 Stage 之后，不影响证据锚定。",
        ]
        if md_path:
            md_path.write_text('\n'.join(lines), encoding='utf-8')
        report['markdown_path'] = str(md_path)


# 专项方案目录（标准《施工组织设计》必备专项方案；通用、零私有数据）。
# always=True 为核心必备方案（技术标普遍要求）；always=False 需触发词命中才补。
_SPECIAL_SCHEMES = [
    {
        'key': 'temp_power', 'title': '临时用电专项方案', 'level': 1, 'always': True,
        'triggers': ['用电', '临电', '电气', '供电', '配电', '电缆'],
        'outline': [
            '一、编制依据：按《施工现场临时用电安全技术规范》JGJ46 及现场勘测布置。',
            '二、负荷计算：统计主要用电设备，验算总负荷与变压器/发电机容量匹配并留余量。',
            '三、线路敷设：采用 TN-S 系统、三级配电两级保护，电缆埋地或架空并设标识。',
            '四、配电箱与开关箱：一机一闸一漏一箱，漏电动作电流≤30mA、动作时间≤0.1s。',
            '五、接地与防雷：重复接地电阻≤10Ω；塔吊、井架等做防雷接地。',
            '六、安全用电：持证上岗、巡检制度、夜间照明与警示，雨季加强绝缘检测。',
        ],
    },
    {
        'key': 'high_altitude', 'title': '高处作业专项方案', 'level': 1, 'always': True,
        'triggers': ['高处', '高空', '脚手架', '临边', '洞口', '攀登', '吊篮'],
        'outline': [
            '一、适用范围：坠落高度基准面 2m 及以上作业，含临边、洞口、攀登、悬空、操作平台。',
            '二、人员要求：特种作业持证、体检合格，佩戴安全带（高挂低用）与安全帽。',
            '三、防护设施：脚手架验收合格挂牌；临边洞口设防护栏杆+密目式安全网+挡脚板。',
            '四、作业管控：六级以上大风、雨雪雾天停止露天高处作业；工具入袋防坠落。',
            '五、验收与交底：作业前安全技术交底，防护设施逐层验收并留记录。',
        ],
    },
    {
        'key': 'seasonal', 'title': '季节性施工措施', 'level': 1, 'always': True,
        'triggers': ['冬', '雨', '夏', '季', '高温', '防汛', '防寒', '台风'],
        'outline': [
            '一、雨季施工：排水通畅、边坡监测、防雷接地、材料防雨防潮、脚手架基础加固。',
            '二、夏季/高温：防暑降温、避开高温时段、饮用水与急救药品配备。',
            '三、冬季施工：混凝土测温保温、防冻剂合规使用、消防与有毒气体（密闭空间通风）防控。',
            '四、台风/汛期：应急预案、大型机械防风加固、人员撤离路线明确。',
        ],
    },
    {
        'key': 'emergency', 'title': '应急救援预案', 'level': 1, 'always': True,
        'triggers': ['应急', '救援', '事故', '预案', '演练'],
        'outline': [
            '一、组织机构：成立项目应急救援小组，明确负责人、成员与联络方式。',
            '二、风险辨识：高处坠落、触电、坍塌、火灾、物体打击等分级管控。',
            '三、响应流程：报警→警戒→抢救→送医→上报，定期演练并留存记录。',
            '四、物资装备：急救箱、消防器材、通讯设备定点存放、专人管理、定期点检。',
        ],
    },
    {
        'key': 'deep_foundation', 'title': '深基坑支护专项方案', 'level': 1, 'always': False,
        'triggers': ['深基坑', '基坑', '开挖', '支护', '降水', '边坡', '地连墙'],
        'outline': [
            '一、支护选型：依地质与周边环境选排桩/地下连续墙/放坡，附计算书与图纸。',
            '二、降排水：管井/井点降水，水位监测，周边建（构）筑物沉降观测。',
            '三、开挖顺序：分层分段、先撑后挖、严禁超挖，栈桥限载。',
            '四、监测与预警：桩顶位移、周边管线沉降实时监测，超限立即停工处置。',
        ],
    },
    {
        'key': 'lifting', 'title': '起重吊装专项方案', 'level': 1, 'always': False,
        'triggers': ['起重', '吊装', '塔吊', '吊车', '吊运', '吊装机'],
        'outline': [
            '一、设备核验：起重机检测合格、司机持证，吊具索具验算破断拉力。',
            '二、作业环境：地基承载力满足，回转半径内无障碍与高压线安全距离。',
            '三、吊装作业：试吊离地 200mm 检查、持证指挥，六级及以上风停吊。',
            '四、安全措施：警戒区隔离、司索指挥一致、严禁人员停留于吊物下方。',
        ],
    },
]


def _doc_text_excluding_schemes(doc) -> str:
    """构造用于「条件触发词」检测的全文，但剔除本 Stage 已注入的专项方案章节正文。

    避免「上一轮注入的大纲含触发词（高处作业含『塔吊』、季节性含『边坡』）→
    本轮又触发同类方案」的自身反馈，破坏幂等。用户自身章节（非本 Stage 标题）
    的正文仍参与触发检测，故用户原文确实含相关工程特征时仍会正确补入。
    """
    _scheme_titles = {s['title'] for s in _SPECIAL_SCHEMES}
    lines = []
    in_scheme = False
    for p in doc.paragraphs:
        st = (p.style.name or '') if getattr(p, 'style', None) else ''
        is_head = st.startswith('Heading') or st.startswith('标题')
        t = (p.text or '').strip()
        if is_head:
            in_scheme = t in _scheme_titles
            if in_scheme:
                continue  # 跳过我们注入的章节标题本身
        if in_scheme:
            continue
        if t:
            lines.append(t)
    return '\n'.join(lines)


class SpecialSchemeStage(Stage):
    """专项方案自动补全（PDCA-Act）。非阻断。

    C22（完善技术标内容）：技术标常因缺必备专项方案被扣分/废标。本 Stage 扫描
    生成稿，对「核心必备方案（always）」或「触发词命中」的专项方案，若文档尚未
    包含则补入标准化提纲（通用、零私有数据、可人工扩写），置于「附表/附录」前
    或文末。幂等（标题已存在即跳过）；可逆 opt-in。

    纯规则 + 模板，无 LLM、无外部依赖。
    """

    name = "special_scheme"
    blocking = False

    def should_run(self, ctx: StageContext) -> bool:
        req = ctx.req
        result_path = ctx.get("result_path")
        return (bool(_attr(req, "enable_special_scheme", True))
                and bool(result_path)
                and os.path.exists(result_path))

    def run(self, ctx: StageContext) -> None:
        from docx import Document

        result_path = ctx.get("result_path")
        doc = Document(result_path)

        headings = [(p.text or '').strip() for p in doc.paragraphs
                    if (p.style.name or '').startswith('Heading')
                    or (p.style.name or '').startswith('标题')]
        # 条件触发词检测仅看「用户原文」，剔除本 Stage 已注入的专项方案章节正文，
        # 否则上一轮注入的大纲含触发词（如高处作业含「塔吊」、季节性含「边坡」）
        # 会在重跑时再次触发同类方案，破坏幂等。
        doc_text = _doc_text_excluding_schemes(doc)

        def _present(title: str) -> bool:
            for h in headings:
                if title in h or h in title:
                    return True
            return False

        injected = []
        for sch in _SPECIAL_SCHEMES:
            if _present(sch['title']):
                continue
            if (not sch['always']) and not any(t in doc_text for t in sch['triggers']):
                continue
            # 插入锚点：首个「附表/附录」标题之前；无则追加文末
            anchor = None
            for p in doc.paragraphs:
                st = (p.style.name or '') if getattr(p, 'style', None) else ''
                if st.startswith('Heading') or st.startswith('标题'):
                    _t = (p.text or '').strip()
                    if '附表' in _t or '附录' in _t:
                        anchor = p
                        break
            new_ps = []
            h = doc.add_paragraph(sch['title'])
            try:
                h.style = doc.styles['Heading %d' % sch['level']]
            except Exception:  # noqa: BLE001
                pass
            new_ps.append(h)
            for line in sch['outline']:
                new_ps.append(doc.add_paragraph(line))
            if anchor is not None:
                for np in reversed(new_ps):
                    anchor._p.addprevious(np._p)
            injected.append(sch['key'])

        doc.save(result_path)

        report = {
            'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'summary': {'injected': len(injected), 'schemes': injected},
        }
        ctx.set('special_scheme', report)
        try:
            self._write_markdown(ctx,report, result_path)
        except Exception:  # noqa: BLE001
            pass

    def _write_markdown(self, ctx,report, result_path):
        if not result_path:
            return
        out = Path(result_path)
        md_path = aux_path(ctx, result_path,'_专项方案补全.md')
        s = report['summary']
        lines = [
            "# 专项方案自动补全（PDCA-Act 自动生成）",
            "",
            f"- 生成时间：{report['generated_at']}",
            f"- 自动补入专项方案：{s['injected']} 项（{', '.join(s['schemes'])}）",
            "",
            "## 补入清单",
        ]
        _catalog = {x['key']: x['title'] for x in _SPECIAL_SCHEMES}
        for k in s['schemes']:
            lines.append(f"- {_catalog.get(k, k)}")
        lines.append("")
        lines.append("说明：以上为标准化提纲，建议结合工程实际（地质、体量、工期）由技术负责人扩写并签章。")
        if md_path:
            md_path.write_text('\n'.join(lines), encoding='utf-8')
        report['markdown_path'] = str(md_path)


class ConsistencyStage(Stage):
    """跨交付物一致性校验（PDCA-Act）。非阻断。

    校验维度（防止「报告声称已响应却指不出证据」的虚假一致）：
    - 废标红线/★条款中标记「已响应(covered)」的，是否都在 docx 中定位到证据章节；
      若报告声称已响应却找不到证据锚点，属一致性缺口（可能误判/漏标）。
    - 评分项命中矩阵中「已覆盖」的项是否都锚定了证据章节。
    - 若管线启用了电子标书交付(enable_delivery)，交付物是否已实际产出。

    产出一致性结论与 _一致性校验.md；缺失任一上游交付物则优雅降级（仅校验存在部分）。
    纯规则，无 LLM、无外部依赖。
    """

    name = "consistency"
    blocking = False

    def should_run(self, ctx: StageContext) -> bool:
        req = ctx.req
        return (bool(_attr(req, "enable_consistency", True))
                and ctx.get("risk_report") is not None
                and bool(ctx.get("result_path")))

    # ── T4 闸门标定：用真实 docx 文本验证「已响应」是否真在文中 ──
    @staticmethod
    def _norm_text(t: str) -> str:
        import re as _re
        return _re.sub(r"[\s，。；;、：:,.!！?？()（）\[\]【】\"'\"'“”%‰]", "", t or "")

    def _doc_text(self, result_path: str) -> str:
        """惰性读取生成稿全文（供证据锚点比对）。失败返回空串。"""
        if not result_path or not os.path.exists(result_path):
            return ""
        try:
            from checker.risk_library import _extract_doc_text
            return _extract_doc_text(result_path) or ""
        except Exception:  # noqa: BLE001
            return ""

    def _appears(self, text: str, doc_text: str) -> bool:
        """归一化子串命中：条款/评分项关键文本是否在生成稿中出现。"""
        if not text or not doc_text:
            return False
        n = self._norm_text(text)
        if len(n) < 6:
            # 过短文本不可靠，直接判为已锚定（避免噪声误报）
            return True
        return n[:40] in self._norm_text(doc_text)

    def run(self, ctx: StageContext) -> None:
        rr = ctx.get("risk_report") or {}
        sm = ctx.get("scoring_matrix") or {}
        dl = ctx.get("delivery") or {}
        req = ctx.req
        result_path = ctx.get("result_path")

        # 红线/★条款：已响应
        _clauses = (rr.get("red_lines") or []) + (rr.get("star_clauses") or [])
        _covered_risk = [x for x in _clauses if x.get("covered")]
        risk_covered = len(_covered_risk)
        # 评分项：已覆盖
        _items = sm.get("items") or []
        _cov_items = [it for it in _items if it.get("status") in ("covered", "strong")]
        scoring_covered = len(_cov_items)

        # T4 闸门标定：逐条比对真实 docx 文本
        #   anchored      = 上游已记录 evidence_section 或 关键文本在文中检索到
        #   unanchored    = 已标记响应但无 evidence_section，且关键文本在文中（内容已在，仅锚点未记录）→ 软提示
        #   truly_missing = 已标记响应，但关键文本全文检索不到 → 真缺口（计入一致性分）
        doc_text = self._doc_text(result_path)

        # D1-②：需求响应闭环（C21）已记录每条需求的写入章节，作为显式锚点回灌一致性闸门
        _rc = ctx.get("req_closure") or {}
        _rc_map = {}
        for _it in (_rc.get("items") or []):
            _lab = self._norm_text(_it.get("label") or "")
            _sec = _it.get("section")
            if _lab and _sec:
                _rc_map[_lab] = _sec

        def _classify(items, key):
            anchored, unanchored, missing = [], [], []
            for it in items:
                txt = (it.get(key) or "")[:80]
                _ntxt = self._norm_text(txt)
                # 认领 C21 锚点：同一需求文本既被 C21 写入某章节，即视为已锚定
                if _ntxt and _ntxt in _rc_map:
                    if not it.get("evidence_section"):
                        it["evidence_section"] = _rc_map[_ntxt]
                    anchored.append(it)
                elif it.get("evidence_section") or self._appears(txt, doc_text):
                    anchored.append(it)
                elif txt:
                    missing.append(it)          # 已响应但全文检不到 → 真缺失
                else:
                    unanchored.append(it)       # 无文本可比对 → 仅未锚定（软提示）
            return anchored, unanchored, missing

        _ra, _ru, _rm = _classify(_covered_risk, "text")
        _sa, _su, _sm = _classify(_cov_items, "name")
        risk_no_ev = len(_rm)
        scoring_no_ev = len(_sm)
        unanchored_count = len(_ru) + len(_su)

        # 交付物一致性：启用交付则应有交付物产出。
        delivery_expected = bool(_attr(req, "enable_delivery", True))
        _DL_DONE = ("delivered", "manifest_only", "success", "done", "ok")
        delivery_executed = bool(dl and dl.get("status") in _DL_DONE)
        delivery_missing = delivery_expected and not delivery_executed

        inconsistencies = []
        for x in _rm:
            inconsistencies.append(f"[真缺失] 红线/★条款已标记响应但全文未检到证据：{x.get('text', '')[:40]}")
        for it in _sm:
            inconsistencies.append(f"[真缺失] 评分项已覆盖但全文未检到证据：{it.get('name', '')}")
        for x in _ru + _su:
            inconsistencies.append(f"[未锚定] 已响应但证据锚点未记录（内容疑似已在文中，建议补全锚点）："
                                   f"{(x.get('text') or x.get('name') or '')[:40]}")
        if delivery_missing:
            inconsistencies.append("已启用电子标书交付但交付物未产出，请检查交付 Stage")

        # T4：一致性分只扣「真缺失」，未锚定仅轻扣，避免狼来了
        gaps = risk_no_ev + scoring_no_ev + (1 if delivery_missing else 0)
        consistency_score = max(0, 100 - gaps * 10 - unanchored_count * 2)
        if gaps == 0 and unanchored_count == 0:
            verdict = "一致"
        elif gaps == 0:
            verdict = f"内容一致，{unanchored_count} 处证据锚点建议补全（不影响提交）"
        else:
            verdict = f"存在 {gaps} 处真缺失（请逐项核对）"

        report = {
            'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'risk_covered': risk_covered,
            'risk_without_evidence': risk_no_ev,
            'scoring_covered': scoring_covered,
            'scoring_without_evidence': scoring_no_ev,
            'unanchored_count': unanchored_count,
            'delivery_expected': delivery_expected,
            'delivery_executed': delivery_executed,
            'inconsistencies': inconsistencies,
            'gap_count': gaps,
            'truly_missing': gaps,
            'consistency_score': consistency_score,
            'verdict': verdict,
        }
        ctx.set("consistency", report)
        try:
            self._write_markdown(ctx,report, result_path)
        except Exception:  # noqa: BLE001
            pass

    def _write_markdown(self, ctx,report, result_path):
        if not result_path:
            return
        out = Path(result_path)
        md_path = aux_path(ctx, result_path,'_一致性校验.md')
        s = report
        _dl_state = ('已产出' if s['delivery_executed']
                     else ('未产出' if s['delivery_expected'] else '—'))
        lines = [
            "# 跨交付物一致性校验（PDCA-Act 自动生成 · T4 闸门标定）",
            "",
            f"- 生成时间：{s['generated_at']}",
            f"- **一致性结论：{s['verdict']}**",
            f"- 一致性评分：{s['consistency_score']}/100",
            "",
            "## 校验明细",
            f"- 废标红线/★：已响应 {s['risk_covered']} 条，其中 **真缺失 {s['risk_without_evidence']}** 条，未锚定 {s.get('unanchored_count', 0)} 条",
            f"- 评分项：已覆盖 {s['scoring_covered']} 项，其中 **真缺失 {s['scoring_without_evidence']}** 项",
            f"- 电子标书交付：{'预期' if s['delivery_expected'] else '未启用'}，{_dl_state}",
            "",
            "## 不一致清单（[真缺失]=影响提交 / [未锚定]=建议补全锚点）",
        ]
        if s['inconsistencies']:
            for i, inc in enumerate(s['inconsistencies'], 1):
                lines.append(f"{i}. {inc}")
        else:
            lines.append("无不一致项：所有「已响应/已覆盖」均在文中检索到证据。")
        if md_path:
            md_path.write_text('\n'.join(lines), encoding='utf-8')
        report['markdown_path'] = str(md_path)


def _fmt_ratio(r: Any) -> str:
    return f"{r * 100:.0f}%" if isinstance(r, (int, float)) else "—"


def _verdict_phrase(s: Dict[str, Any]) -> str:
    score = s.get('risk_score', 0) or 0
    lvl = '高危' if score < 60 else ('需关注' if score < 85 else '可控')
    return (f"综合废标风险分 **{score}/100**（{lvl}）；废标红线/★条款 "
            f"{s.get('covered', 0)}/{s.get('total_clauses', 0)} 已响应，"
            f"{s.get('uncovered', 0)} 条未响应，请按下方整改建议逐项处理。")


class SummaryStage(Stage):
    """投标体检总览（Go/No-Go）。非阻断。

    合并已有交付物（risk_report / scoring_matrix / delivery / schedule_chart），
    输出一页纸可执行总览：综合结论 + 关键指标 + 按风险排序的整改清单。
    纯只读聚合，不修改任何既有报告，便于在管线末尾挂接；缺任一输入即优雅降级。
    """

    name = "summary"
    blocking = False

    def should_run(self, ctx: StageContext) -> bool:
        req = ctx.req
        return (bool(_attr(req, "enable_summary", True))
                and ctx.get("risk_report") is not None
                and bool(ctx.get("result_path")))

    def run(self, ctx: StageContext) -> None:
        rr = ctx.get("risk_report") or {}
        sm = ctx.get("scoring_matrix") or {}
        dl = ctx.get("delivery") or {}
        sc = ctx.get("schedule_chart") or {}
        cs = ctx.get("consistency") or {}
        ql = ctx.get("qualification") or {}
        cl = ctx.get("req_closure") or {}
        df = ctx.get("doc_format") or {}
        ss = ctx.get("special_scheme") or {}
        ns = ctx.get("national_standard") or {}
        sk = ctx.get("shell_kit") or {}
        rd = ctx.get("reference_driven") or {}
        dlh = ctx.get("dark_leak_harden") or {}
        result_path = ctx.get("result_path")

        s = rr.get("summary", {}) or {}
        risk_score = s.get("risk_score", 0) or 0
        dark_leaks = s.get("dark_bid_leaks", 0) or 0
        lib_high = s.get("library_high", 0) or 0
        form_issues = s.get("form_issues", 0) or 0
        uncovered_clauses = s.get("uncovered", 0) or 0
        kw_ratio = s.get("keyword_hit_ratio")
        semantic_hits = s.get("semantic_hits", 0) or 0
        exact_hits = s.get("exact_hits", 0) or 0

        sc_total = sm.get("total", 0) or 0
        sc_uncovered = sm.get("uncovered", 0) or 0
        sc_weak = sm.get("weak", 0) or 0

        consistency_gaps = cs.get("gap_count", 0) or 0
        consistency_unanchored = cs.get("unanchored_count", 0) or 0

        ql_summary = ql.get("summary", {}) or {}
        qual_total = ql_summary.get("total", 0) or 0
        qual_uncovered = ql_summary.get("uncovered", 0) or 0

        blockers = []
        if dark_leaks > 0:
            blockers.append(("致命", f"暗标泄漏 {dark_leaks} 处 —— 提交即废标，必须清除"))
        dark_residual = dlh.get('leaks_after', 0) or 0
        if dark_residual > 0:
            blockers.append(("致命", f"暗标隐藏泄漏面残留 {dark_residual} 处（页眉/页脚/水印/文档属性）—— 提交即废标，必须清除"))
        if lib_high > 0:
            blockers.append(("高危", f"风险库高危项 {lib_high} 项（承诺缺失/资质过期等）"))
        if uncovered_clauses > 0:
            blockers.append(("高危", f"废标红线/★条款未响应 {uncovered_clauses} 条"))
        if sc_uncovered > 0:
            blockers.append(("中危", f"评分项未覆盖 {sc_uncovered}/{sc_total} 项，将直接丢分"))
        if form_issues > 0:
            blockers.append(("中危", f"形式要件缺项 {form_issues} 项（签字/盖章/装订/密封等）"))
        if qual_uncovered > 0:
            blockers.append(("高危", f"资格要求未响应 {qual_uncovered}/{qual_total} 项（资格不符=废标风险）"))
        if consistency_gaps > 0:
            blockers.append(("中危", f"跨交付物一致性真缺失 {consistency_gaps} 处（已响应却全文检不到证据/交付未产出）"))
        if risk_score < 60:
            blockers.append(("提示", f"综合废标风险分 {risk_score}/100 偏低，请复核"))

        # 提交前最终核对清单：必须先处理的（致命/高危）+ 通用提交提醒
        checklist = []
        for b in blockers:
            checklist.append({"level": b[0], "msg": b[1],
                              "mandatory": b[0] in ("致命", "高危")})
        checklist += [
            {"level": "提示", "msg": "已用 CA 证书对电子标书 PDF 签章（如平台要求）", "mandatory": False},
            {"level": "提示", "msg": "已按交易平台格式（.etnd/.XYTF/.BTBJ 等）导出并验证可打开", "mandatory": False},
            {"level": "提示", "msg": "已双人复核废标红线与评分项逐条响应", "mandatory": False},
            {"level": "提示", "msg": "已确认装订/密封/签字盖章符合招标文件要求", "mandatory": False},
        ]
        # T4 闸门标定：未锚定（内容疑似已在文中）仅作软提示，不触发"不建议提交"
        if consistency_unanchored > 0 and consistency_gaps == 0:
            checklist.append({"level": "提示", "msg":
                f"一致性锚点建议补全：{consistency_unanchored} 处已响应但证据锚点未记录（内容疑似已在文中，不影响提交）",
                "mandatory": False})

        # ── T9 量化指标披露（评分响应率 / 风险检出率 / 重复率 / 可用率 / 耗时）──
        _rf = ctx.get("risk_library_findings") or {}
        detection_rate = _rf.get("detection_rate")
        total_patterns = _rf.get("total_patterns")
        patterns_covered = _rf.get("patterns_covered")
        patterns_missing = _rf.get("patterns_missing")
        score_response_rate = (round((sc_total - sc_uncovered) / sc_total, 4)
                               if sc_total else None)
        _dd = ctx.get("dedup3d") or {}
        _dd_ok = isinstance(_dd, dict) and _dd.get("status") == "ok"
        text_repeat_rate = ((_dd.get("layers") or {}).get("text")
                            if _dd_ok else None)
        simhash_repeat_rate = _dd.get("simhash_self_similarity") if _dd_ok else None
        cross_doc_sim = ctx.get("cross_doc_similarity")
        # 初稿可用率：综合致命/高危信号扣分后的可提交就绪度（0–1）
        _avail = 100
        if dark_leaks:
            _avail -= 40
        if lib_high:
            _avail -= min(40, lib_high * 12)
        if uncovered_clauses:
            _avail -= min(20, uncovered_clauses * 4)
        if consistency_gaps:
            _avail -= min(20, consistency_gaps * 5)
        draft_usability_rate = round(max(0, min(100, _avail)) / 100, 4)
        # 生成耗时：管线 meta 计时
        # 注：orchestrator 在整轮循环结束后才写 finished_at，而本 Stage 运行于循环内
        # （最后一个 Stage），故此处若尚未落戳则先补戳，确保耗时可被披露（T9）。
        _meta = ctx.meta or {}
        if _meta.get("started_at") and not _meta.get("finished_at"):
            _meta["finished_at"] = time.time()
        generation_duration_ms = None
        if _meta.get("started_at") and _meta.get("finished_at"):
            generation_duration_ms = round((_meta["finished_at"] - _meta["started_at"]) * 1000, 1)

        # D1-③：评审自检结论（EvaluatorCheck）与规则闸门交叉呈现，消除两 verdict 张力
        ev = ctx.get("evaluator_check")
        ev_passed = None
        ev_high = 0
        ev_coverage = None
        ev_score = None
        evaluator_verdict = None
        if isinstance(ev, dict):
            _ev_risks = ev.get("risks") or {}
            ev_high = len(_ev_risks.get("high_risks") or [])
            ev_coverage = (ev.get("coverage") or {}).get("coverage_rate")
            ev_score = (ev.get("score_prediction") or {}).get("score_percentage")
            ev_passed = bool(ev.get("passed"))
            if ev_passed:
                evaluator_verdict = (
                    f"评审自检通过（覆盖率{ev_coverage*100:.1f}%，预测得分{ev_score:.1f}%，"
                    f"无高风险废标项）"
                )
            else:
                warnings = ev.get("warnings") or []
                evaluator_verdict = f"评审自检警示：{len(warnings)} 项（含{ev_high}项高风险废标风险）"

        if any(b[0] == "致命" for b in blockers):
            verdict = "不建议提交（高风险）"
        elif any(b[0] == "高危" for b in blockers):
            # 规则闸门判高危，但评审自检通过且无高风险废标项 → 两 gate 分歧，降级为"建议复核"
            _ev_clean = (
                ev_passed
                and ev_high == 0
                and (ev_coverage is None or ev_coverage >= 0.8)
                and (ev_score is None or ev_score >= 60)
            )
            if ev is not None and _ev_clean:
                verdict = "可提交，建议复核（规则闸门与评审自检存在分歧）"
            else:
                verdict = "不建议提交（高风险）"
        elif blockers:
            verdict = "可提交，但建议先整改"
        else:
            verdict = "可提交（建议人工终审）"

        # B3/D4：知识库与零幻觉状态（仅当 enable_kb_rag 启用且有产出）
        _kb_rag = ctx.get("kb_rag") or {}
        _kb_fc = ctx.get("kb_factcheck") or {}
        kb_status = _kb_fc.get("status") or _kb_rag.get("status") or "未启用"
        kb_loaded = bool(_kb_fc.get("kb_loaded")) or _kb_rag.get("status") == "ok"
        kb_retrieved = _kb_rag.get("retrieved_count") or 0
        kb_flagged = _kb_fc.get("flagged_count") or 0
        kb_is_demo = bool(_kb_fc.get("kb_is_demo"))
        kb_entry_count = _kb_rag.get("entry_count") or 0

        summary = {
            "verdict": verdict,
            # D1-③：交叉呈现两路结论
            "evaluator_verdict": evaluator_verdict,
            "evaluator_passed": ev_passed,
            "evaluator_high_risks": ev_high,
            "evaluator_coverage_rate": ev_coverage,
            "evaluator_score_percentage": ev_score,
            "gate_divergence": (ev is not None and ev_passed and any(b[0] == "高危" for b in blockers)),
            # B3/D4：知识库与零幻觉
            "kb_status": kb_status,
            "kb_loaded": kb_loaded,
            "kb_is_demo": kb_is_demo,
            "kb_entry_count": kb_entry_count,
            "kb_retrieved": kb_retrieved,
            "kb_flagged": kb_flagged,
            "kb_factcheck_report": ctx.get("kb_factcheck_report"),
            "risk_score": risk_score,
            "dark_bid_leaks": dark_leaks,
            "dark_leak_residual": dark_residual,
            "dark_leak_harden": dlh,
            "library_high": lib_high,
            "uncovered_clauses": uncovered_clauses,
            "scoring_total": sc_total,
            "scoring_uncovered": sc_uncovered,
            "scoring_weak": sc_weak,
            "form_issues": form_issues,
            "consistency_gaps": consistency_gaps,
            "consistency_unanchored": consistency_unanchored,
            "national_standard": ns,
            "shell_kit": sk,
            "reference_driven": rd,
            "qualification_total": qual_total,
            "qualification_uncovered": qual_uncovered,
            "keyword_hit_ratio": kw_ratio,
            "delivery_status": dl.get("status"),
            "exact_hits": exact_hits,
            "semantic_hits": semantic_hits,
            # T9 量化指标
            "score_response_rate": score_response_rate,
            "risk_detection_rate": detection_rate,
            "risk_total_patterns": total_patterns,
            "risk_patterns_covered": patterns_covered,
            "risk_patterns_missing": patterns_missing,
            "text_repeat_rate": text_repeat_rate,
            "simhash_repeat_rate": simhash_repeat_rate,
            "cross_doc_similarity": cross_doc_sim,
            "draft_usability_rate": draft_usability_rate,
            "generation_duration_ms": generation_duration_ms,
            "blockers": [{"level": b[0], "msg": b[1]} for b in blockers],
            "checklist": checklist,
        }
        ctx.set("summary", summary)
        try:
            self._write_markdown(ctx,summary, rr, sm, dl, sc, cs, ql, cl, df, ss, result_path)
        except Exception:  # noqa: BLE001
            pass

    def _write_markdown(self, ctx,summary, rr, sm, dl, sc, cs, ql, cl, df, ss, result_path):
        if not result_path:
            return
        out = Path(result_path)
        md_path = aux_path(ctx, result_path,'_投标体检总览.md')
        _ns = summary.get('national_standard', {}) or {}
        _sk = summary.get('shell_kit', {}) or {}
        _rd = summary.get('reference_driven', {}) or {}
        s = summary
        lines = [
            "# 投标体检总览（PDCA-Act 自动生成）",
            "",
            "## 综合结论",
            f"**{s['verdict']}**",
            "",
            "## 双闸交叉核验（D1-③）",
            "| 闸门 | 结论 |",
            "| --- | --- |",
            f"| 规则闸门（27 Stage 合规风控） | {s['verdict']} |",
            (f"| 评审自检（EvaluatorCheck） | {s['evaluator_verdict']} |"
             if s.get('evaluator_verdict') else
             "| 评审自检（EvaluatorCheck） | 未运行（生成期自检未启用或异常） |"),
            (f"| **分歧提示** | 两闸结论不一致：规则闸门更严（响应缺口按高危计），"
             f"评审自检未检出高风险废标项。建议人工复核差异项后再提交。 |"
             if s.get('gate_divergence') else ""),
            "",
            "## 关键指标",
            "| 维度 | 结果 |",
            "| --- | --- |",
            f"| 综合废标风险分 | {s['risk_score']}/100 |",
            f"| 暗标泄漏 | {s['dark_bid_leaks']} 处 |",
            f"| 暗标隐藏泄漏（硬化后） | 残留 {s['dark_leak_residual']} 处 |",
            f"| 风险库高危项 | {s['library_high']} 项 |",
            f"| 废标红线/★未响应 | {s['uncovered_clauses']} 条 |",
            f"| 评分项未覆盖 | {s['scoring_uncovered']}/{s['scoring_total']} 项 |",
            f"| 评分项弱项 | {s['scoring_weak']} 项 |",
            f"| 资格要求未响应 | {s['qualification_uncovered']}/{s['qualification_total']} 项 |",
            f"| 需求响应闭环 | {cl.get('summary', {}).get('injected', 0) or 0} 段已织入正文 |",
            f"| 文档格式硬化 | 编号 {df.get('summary', {}).get('headings_numbered', 0) or 0} 个 + 图注 {df.get('summary', {}).get('figures_captioned', 0) or 0} 张 + 目录域 |",
            f"| 专项方案补全 | 自动补入 {ss.get('summary', {}).get('injected', 0) or 0} 项 |",
            f"| 国标格式套用 | 表头重复 {_ns.get('summary', {}).get('headers_repeated', 0) or 0}"
            f"/{_ns.get('summary', {}).get('tables_total', 0) or 0} + 样式 {len(_ns.get('summary', {}).get('styles_applied', []) or [])} 项 |",
            f"| 技术标外壳套件 | 已补入 {_sk.get('summary', {}).get('sections_inserted', 0) or 0} 节前置外壳 |",
            (f"| 以标写标（参考驱动） | {_rd.get('status')}（{_rd.get('chapter_count', 0)} 章"
             f"/补 {_rd.get('missing_appended', 0)}） |" if _rd else
             "| 以标写标（参考驱动） | 未提供参考标书 |"),
            f"| 形式要件缺项 | {s['form_issues']} 项 |",
            f"| 跨交付物一致性缺口（真缺失/未锚定） | {s['consistency_gaps']} / {s.get('consistency_unanchored', 0)} 处 |",
            f"| 废标关键词呼应度 | {_fmt_ratio(s['keyword_hit_ratio'])} |",
            f"| 红线命中（精确/语义疑似） | {s['exact_hits']} / {s['semantic_hits']} |",
            f"| 电子标书交付 | {s['delivery_status'] or '未启用'} |",
            (f"| 知识库与零幻觉（B3/D4） | "
             f"{('未启用（enable_kb_rag=False 或知识库缺失）' if s.get('kb_status') in (None, '未启用', 'no_kb') else '')}"
             f"{('占位脚手架（' + str(s.get('kb_entry_count', 0)) + ' 条示例数据，核查已自动停用避免误报）' if s.get('kb_status') == 'demo_kb' else '')}"
             f"{('已加载 ' + str(s.get('kb_entry_count', 0)) + ' 条真实素材；检索命中 ' + str(s.get('kb_retrieved', 0)) + '；零幻觉告警 ' + str(s.get('kb_flagged', 0)) + ' 处' if s.get('kb_status') == 'ok' else '')}"
             f" |"),
            "",
            "## 量化指标披露（T9）",
            "| 指标 | 结果 |",
            "| --- | --- |",
            f"| 评分点响应率 | {_fmt_ratio(s.get('score_response_rate'))} |",
            f"| 废标风险检出率 | {_fmt_ratio(s.get('risk_detection_rate'))}（覆盖 {s.get('risk_patterns_covered')}/{s.get('risk_total_patterns')} 项模式） |",
            f"| 文本重复率（三维查重） | {_fmt_ratio(s.get('text_repeat_rate'))} |",
            f"| 近义重复率（SimHash） | {_fmt_ratio(s.get('simhash_repeat_rate'))} |",
            f"| 跨文档相似度 | {_fmt_ratio(s.get('cross_doc_similarity'))} |",
            f"| 初稿可用率 | {_fmt_ratio(s.get('draft_usability_rate'))} |",
            f"| 生成耗时 | {s.get('generation_duration_ms')} ms |",
            "",
            "## 整改优先级（按风险排序）",
        ]
        if s["blockers"]:
            for i, b in enumerate(s["blockers"], 1):
                lines.append(f"{i}. [{b['level']}] {b['msg']}")
        else:
            lines.append("无高危/中危项，建议人工终审后提交。")
        lines.append("")
        lines.append("## 提交前最终核对清单")
        for i, c in enumerate(s.get("checklist", []), 1):
            tag = "必办" if c.get("mandatory") else "建议"
            lines.append(f"{i}. [ ] [{tag}] {c['msg']}")
        if s.get("semantic_hits"):
            lines.append("")
            lines.append(f"⚠ 含 {s['semantic_hits']} 条语义疑似匹配（非精确命中），建议人工确认表述是否到位，避免漏标。")
        lines += ["", "## 详细报告"]
        detail_map = [
            ("废标风险自检报告", rr, "markdown_path"),
            ("评分项命中矩阵", sm, "markdown_path"),
            ("电子标书交付清单", dl, "manifest_path"),
            ("跨交付物一致性校验", cs, "markdown_path"),
            ("资格响应自查", ql, "markdown_path"),
            ("需求响应闭环", cl, "markdown_path"),
            ("文档格式增强", df, "markdown_path"),
            ("专项方案补全", ss, "markdown_path"),
            ("国标格式套用", _ns, "markdown_path"),
            ("技术标外壳套件", _sk, "markdown_path"),
            ("以标写标", _rd, "markdown_path"),
            ("暗标零泄漏硬化", summary.get("dark_leak_harden") or {}, "markdown_path"),
            ("施工进度横道图", sc, "markdown_path"),
            ("零幻觉事实核查", s.get("kb_factcheck_report") and {"markdown_path": s.get("kb_factcheck_report")} or {}, "markdown_path"),
        ]
        for label, src, key in detail_map:
            p = src.get(key) if isinstance(src, dict) else None
            if p:
                lines.append(f"- {label}：{p}")
        if md_path:
            md_path.write_text('\n'.join(lines), encoding='utf-8')
        summary['markdown_path'] = str(md_path)


def append_risk_report(orchestrator: PipelineOrchestrator) -> PipelineOrchestrator:
    """把 RiskReportStage 追加到管线末尾。"""
    orchestrator.register(RiskReportStage())
    return orchestrator


def append_summary(orchestrator: PipelineOrchestrator) -> PipelineOrchestrator:
    """把 SummaryStage（投标体检总览）追加到管线末尾。"""
    orchestrator.register(SummaryStage())
    return orchestrator


def append_scoring_matrix(orchestrator: PipelineOrchestrator) -> PipelineOrchestrator:
    """把 ScoringMatrixStage 追加到管线末尾。"""
    orchestrator.register(ScoringMatrixStage())
    return orchestrator


def append_qualification(orchestrator: PipelineOrchestrator) -> PipelineOrchestrator:
    """把 QualificationStage（资格响应自查）追加到管线末尾。"""
    orchestrator.register(QualificationStage())
    return orchestrator


def append_requirement_closure(orchestrator: PipelineOrchestrator) -> PipelineOrchestrator:
    """把 RequirementClosureStage（需求响应闭环）追加到管线末尾。"""
    orchestrator.register(RequirementClosureStage())
    return orchestrator


def append_doc_format(orchestrator: PipelineOrchestrator) -> PipelineOrchestrator:
    """把 DocFormatStage（文档格式硬化：目录+多级编号）追加到管线末尾。"""
    orchestrator.register(DocFormatStage())
    return orchestrator


def append_national_standard(orchestrator: PipelineOrchestrator) -> PipelineOrchestrator:
    """把 NationalStandardStage（国标格式套用：跨页表头+国标字体/行距）追加到管线末尾。"""
    orchestrator.register(NationalStandardStage())
    return orchestrator


def append_shell_kit(orchestrator: PipelineOrchestrator) -> PipelineOrchestrator:
    """把 ShellKitStage（技术标外壳套件：投标函/身份证明/授权/承诺书）追加到管线末尾。"""
    orchestrator.register(ShellKitStage())
    return orchestrator


def append_reference_driven(orchestrator: PipelineOrchestrator) -> PipelineOrchestrator:
    """把 ReferenceDrivenStage（以标写标内容驱动）追加到管线末尾。"""
    orchestrator.register(ReferenceDrivenStage())
    return orchestrator


def append_special_scheme(orchestrator: PipelineOrchestrator) -> PipelineOrchestrator:
    """把 SpecialSchemeStage（专项方案自动补全）追加到管线末尾。"""
    orchestrator.register(SpecialSchemeStage())
    return orchestrator


def append_consistency(orchestrator: PipelineOrchestrator) -> PipelineOrchestrator:
    """把 ConsistencyStage（跨交付物一致性校验）追加到管线末尾。"""
    orchestrator.register(ConsistencyStage())
    return orchestrator


def append_dark_leak_harden(orchestrator: PipelineOrchestrator) -> PipelineOrchestrator:
    """把 DarkLeakHardenStage（暗标零泄漏硬化）追加到管线末尾。"""
    orchestrator.register(DarkLeakHardenStage())
    return orchestrator

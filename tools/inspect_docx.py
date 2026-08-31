# -*- coding: utf-8 -*-
"""检查生成的标书 docx：篇幅估算、结构、样例文本。用法：python tools/inspect_docx.py <path>"""
from __future__ import annotations

import sys
from docx import Document


def main() -> int:
    path = sys.argv[1]
    doc = Document(path)
    paras = doc.paragraphs
    total_chars = sum(len(x.text) for x in paras)
    tbl_chars = sum(len(c.text) for t in doc.tables for r in t.rows for c in r.cells)
    print(f"段落数: {len(paras)}  表格数: {len(doc.tables)}")
    print(f"正文总字数: {total_chars}  表格字数: {tbl_chars}")
    print(f"估算页数(按550字/页): {round((total_chars + tbl_chars) / 550, 1)}")
    print(f"估算页数(按700字/页): {round((total_chars + tbl_chars) / 700, 1)}")

    print("\n--- 疑似标题/章节（前 25 条） ---")
    count = 0
    for x in paras:
        t = x.text.strip()
        if not t or len(t) > 45:
            continue
        if x.style.name.startswith("Heading") or (x.runs and x.runs[0].bold):
            print(" ", t[:60])
            count += 1
            if count >= 25:
                break

    print("\n--- 正文样例（前 4 段） ---")
    n = 0
    for x in paras:
        if len(x.text.strip()) > 40:
            print(" ", x.text[:130])
            n += 1
            if n >= 4:
                break

    print("\n--- 最后 5 个非空段落 ---")
    tail = [x.text[:90] for x in paras if x.text.strip()][-5:]
    for t in tail:
        print(" ", t)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

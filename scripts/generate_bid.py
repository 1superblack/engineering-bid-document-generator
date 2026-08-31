# -*- coding: utf-8 -*-
"""一键生成工程标书 CLI。

用法：
    python scripts/generate_bid.py --tender-file 招标文件.docx --name 项目名 \
        --duration 300 --work-content "施工内容概述" \
        [--bid-type construction] [--target-pages 300] [--output 标书.docx]

流程：解析招标文件（硬闸门）→ 生成 → 自检报告。成功退出码 0；失败 1；解析失败 2。
"""
from __future__ import annotations

import argparse
import json
import sys
import time
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
_SCRIPT_DIR = Path(__file__).resolve().parent
if str(_SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPT_DIR))

import _package_shim  # noqa: E402,F401  包结构垫片，必须最先导入

for _stream in (sys.stdout, sys.stderr):
    if hasattr(_stream, "reconfigure"):
        _stream.reconfigure(encoding="utf-8", errors="replace")


def _estimate_pages(path: str):
    """粗估 docx 页数：返回 (引擎口径, 按字数口径, 段落数, 表格数, 总字数)。

    引擎口径=段落数/25+表格*2（其填充段较短时低估真实页数）；
    按字数口径=总字数/650（仿宋12pt、行距28磅、首行缩进2字符的常规页容量），
    更贴近 Word/WPS 实际页码。
    """
    try:
        from docx import Document

        doc = Document(path)
        paras = len(doc.paragraphs)
        tables = len(doc.tables)
        chars = (sum(len(p.text) for p in doc.paragraphs)
                 + sum(len(c.text) for t in doc.tables for r in t.rows for c in r.cells))
        return round(paras / 25 + tables * 2, 1), round(chars / 650, 1), paras, tables, chars
    except Exception:
        return None, None, 0, 0, 0


def main() -> int:
    ap = argparse.ArgumentParser(description="一键生成工程标书")
    ap.add_argument("--tender-file", required=True, help="招标文件路径（.docx/.pdf）")
    ap.add_argument("--name", default=None, help="项目名称（缺省自动从招标文件提取）")
    ap.add_argument("--duration", type=int, default=None, help="工期日历天（缺省自动提取）")
    ap.add_argument("--work-content", default="", help="施工/服务内容概述")
    ap.add_argument("--bid-type", choices=("construction", "service"), default=None,
                    help="工程类型（缺省自动判断）")
    ap.add_argument("--target-pages", type=int, default=200, help="目标页数（默认200）")
    ap.add_argument("--output", default=None, help="输出 docx 路径（默认 技术标书_<项目名>.docx）")
    ap.add_argument("--area", type=float, default=0.0, help="建筑面积（m²）")
    ap.add_argument("--structure-type", default="", help="结构类型，如 框架/剪力墙/装修")
    ap.add_argument("--quality-target", default="合格", help="质量目标（默认 合格）")
    ap.add_argument("--divisions", nargs="*", default=[], help="分部分项工程列表")
    ap.add_argument("--dark-bid", action="store_true", help="暗标模式（匿名化）")
    ap.add_argument("--no-llm", action="store_true", help="禁用 LLM 扩写（纯本地模板）")
    ap.add_argument("--user-context", default=None, help="企业知识库 JSON 路径（默认 data/user_knowledge_base.json）")
    ap.add_argument("--allow-fallback", action="store_true", help="解析为空时仍继续生成（不推荐）")
    ap.add_argument("--non-interactive", action="store_true", help="跳过交互式提问（自动化/批量场景）")
    ap.add_argument("--json", action="store_true", help="输出 JSON 结果")
    args = ap.parse_args()

    # ---- 1) 解析硬闸门：先解析招标文件并输出评分项摘要 ----
    from bid_core.parser import parse_tender
    from bid_clean import clean_parse_result

    pr = parse_tender(args.tender_file, bid_type=args.bid_type or "construction")
    if pr.get("_error"):
        print(f"[解析失败] {pr['_error']}", file=sys.stderr)
        return 2
    pr = clean_parse_result(pr)

    n_score = len(pr.get("score_items") or [])
    n_star = len(pr.get("star_clauses") or [])
    n_red = len(pr.get("red_line_clauses") or [])
    n_qual = len(pr.get("qualification_reqs") or [])
    if not (n_score or n_star or n_red or n_qual):
        if not args.allow_fallback:
            print("[解析失败] 未提取到评分项/星号/红线/资格要求。"
                  "若是扫描件 PDF 请先 OCR；如需强行按模板生成请加 --allow-fallback。", file=sys.stderr)
            return 2
        print("[警告] 解析为空，按 --allow-fallback 继续（将使用通用模板章节）")
    else:
        print(f"[解析] 评分项={n_score} 星号条款={n_star} 废标红线={n_red} 资格要求={n_qual}")

    # ---- 1.5 自动提取项目参数（名称/工期/类型），减少手工输入 ----
    from auto_extract import auto_extract_params

    auto = auto_extract_params(pr, args.tender_file)
    name = args.name or auto["name"]
    duration = args.duration or auto["duration"] or 300
    bid_type = args.bid_type or auto["bid_type"]
    if args.duration is None and not auto["duration"]:
        print(f"[提示] 未从招标文件提取到工期，按 300 日历天兜底；可加 --duration 指定")
    print(f"[参数] 项目={name} 工期={duration}天 类型={bid_type} 目标={args.target_pages}页")

    # ---- 1.6 企业知识库检查（示例数据提醒；交互模式可引导填写） ----
    from pathlib import Path as _P
    kb_path = args.user_context or str(_P(__file__).resolve().parents[1] / "data" / "user_knowledge_base.json")
    try:
        kb_text = _P(kb_path).read_text(encoding="utf-8")
        if "示例建设集团" in kb_text:
            print("[提示] 企业知识库仍是示例数据（示例建设集团）。建议先填写真实公司资质/业绩，"
                  "否则资格响应与评分补强将使用占位内容（交付前必须替换）。")
    except Exception:
        pass

    # ---- 2) 组装请求并生成 ----
    req = {
        "tender_file": args.tender_file,
        "name": name,
        "duration": duration,
        "area": args.area,
        "structure_type": args.structure_type,
        "work_content": args.work_content,
        "quality_target": args.quality_target,
        "divisions": list(args.divisions),
        "bid_type": bid_type,
        "target_pages": args.target_pages,
        "parse_result": pr,
        "is_dark_bid": args.dark_bid,
        "output_path": args.output,
        "enable_llm": not args.no_llm,
        "enable_llm_core": not args.no_llm,
        "enable_llm_parse": False,
        # 质量门禁默认全开
        "enable_feasibility": True,
        "enable_risk_grading": True,
        "enable_mock_review": True,
        "enable_scoring_reinforce": True,
        "enable_knowledge_base": True,
        "enable_risk_report": True,
        "enable_scoring_matrix": True,
        "enable_consistency": True,
        "enable_dedup3d": True,
        "enable_docx_quality_gate": True,
        "enable_docx_sanitize": True,
        "enable_delivery": True,
        "enable_schedule_chart": True,
        "enable_charts": True,
        "enable_summary": True,
    }
    if args.user_context:
        req["knowledge_base_path"] = args.user_context

    # 页数→填充预算校准：按「目标页数×650字/页」估算需补充字数，折算填充段落与
    # LLM 调用预算（LLM 段约240字/段，模板段约70字/段），避免超发导致页数失控。
    if args.target_pages >= 100:
        # v9.4.1：原上限「每章 min(120, pages//7) 段、全局 pages*7 段」在 300 页目标下
        # 仅允许约 380 段填充，而 300 页按 25 段/页约需 7000 段，导致规划 300 页、
        # 实渲不足 50 页。现按实际排版密度推导（留 20% 余量）。
        req["per_chapter_fill_cap"] = max(120, int(args.target_pages * 4))
        req["global_fill_cap"] = max(600, int(args.target_pages * 30))
        req["fill_para_per_page"] = 24 if args.target_pages >= 200 else 20
        req["llm_fill_call_budget"] = max(80, min(200, args.target_pages * 2 // 3))

    from pipeline.entry import generate_bid_document_pipeline

    t0 = time.time()
    res = generate_bid_document_pipeline(req)
    elapsed = round(time.time() - t0, 1)

    ok = bool(res.get("success"))
    result_path = res.get("output_file") or ""
    _est = _estimate_pages(result_path) if (ok and result_path) else (None, None, 0, 0, 0)
    engine_est, char_est, paras, tables, total_chars = _est
    pipe = res.get("pipeline") or {}
    failed_stages = [n for n, s in pipe.items() if s.get("status") in ("FAILED", "BLOCKED")]
    summary = res.get("summary") or {}

    report = {
        "success": ok,
        "message": res.get("message"),
        "output_file": result_path,
        "elapsed_sec": elapsed,
        "target_pages": args.target_pages,
        "estimated_pages": engine_est,
        "estimated_pages_by_chars": char_est,
        "total_chars": total_chars,
        "paragraphs": paras,
        "tables": tables,
        "failed_stages": failed_stages,
        "risk_score": summary.get("risk_score"),
        "score_response_rate": summary.get("score_response_rate"),
        "risk_detection_rate": summary.get("risk_detection_rate"),
        "text_repeat_rate": summary.get("text_repeat_rate"),
        "draft_usability_rate": summary.get("draft_usability_rate"),
        "verdict": summary.get("verdict"),
    }
    if args.json:
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        print(f"\n[结果] success={ok} 耗时={elapsed}s")
        print(f"[结果] 输出文件: {result_path}")
        if engine_est is not None:
            print(f"[结果] 估算页数={engine_est}（引擎口径）/{char_est}（按字数口径，贴近实际页码）"
                  f" | 目标 {args.target_pages} | 段落={paras} 表格={tables} 字数={total_chars}")
        if failed_stages:
            print(f"[警告] 失败 Stage: {failed_stages}")
        if summary:
            print(f"[自检] 评审风险={summary.get('risk_score')} 评分响应率={summary.get('score_response_rate')} "
                  f"风险检出率={summary.get('risk_detection_rate')} 重复率={summary.get('text_repeat_rate')} "
                  f"成稿可用度={summary.get('draft_usability_rate')} 结论={summary.get('verdict')}")
        print(f"[结果] {res.get('message')}")
    return 0 if ok else 1


if __name__ == "__main__":
    raise SystemExit(main())

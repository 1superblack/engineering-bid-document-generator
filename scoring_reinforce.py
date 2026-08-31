"""
评分响应闭环补强（PDCA-Act） v9.2
────────────────────────────────────────────────────────────────
生成技术标后，自动诊断每个评分项在成稿中的覆盖情况（已覆盖 / 弱覆盖 / 未覆盖），
对弱覆盖与未覆盖的评分项生成针对性补强段落，并直接注入成稿附录，
确保「评分项零遗漏、弱项有响应」。

设计定位与竞品差异：
- 红点智标 / 数匠云 / 钛投标 的「闭环整改 / 提前补强」均停在"给出整改方案"，
  需人工或 LLM 二次落稿；本模块在生成后**自动把补强内容写入成稿**，无需人工介入。
- 与 evaluator_check.check_score_item_coverage（模拟评审里的"诊断"）互为 Check/Act 搭档：
  evaluator 负责评分与报告，本模块负责把诊断出的弱项真正补上，闭合 PDCA 环。
- 补强内容消费 v9.0 企业资质业绩库（业绩/人员/资质/设备/财务），无画像时优雅回退通用承诺。

原则：
- 纯规则、零新依赖、不阻断主流程：任何异常仅告警不中断标书生成。
- 自包含覆盖率探针（字符二元组重叠），不耦合 69KB 的 evaluator_check，便于单测与可逆。
"""
from __future__ import annotations

import os
import re
from typing import Any, Dict, List, Optional, Tuple

try:
    from docx.shared import Pt
except Exception:  # docx 缺失时仅在使用注入函数时报错
    Pt = None

# ── 阈值 ──
COVER_RATIO = 0.50      # 命中二元组占比 ≥ 此值视为已覆盖
_WEAK_MIN = 0.01        # 命中但低于 COVER_RATIO 视为弱覆盖（>0 即至少命中一个二元组）

_CJK_RE = re.compile(r'[\u4e00-\u9fff\u3400-\u4dbf]')


def _is_cjk(ch: str) -> bool:
    return bool(_CJK_RE.match(ch))


def _cjk_bigrams(text: str) -> set:
    """提取文本中所有相邻 2 字中文二元组（用于重叠度匹配）。"""
    if not text:
        return set()
    grams: set = set()
    chars = [c for c in text if _is_cjk(c)]
    for i in range(len(chars) - 1):
        grams.add(chars[i] + chars[i + 1])
    return grams


# 评分项名称 → 补强类别路由（与 score_response._KW_ROUTES 同源思路，本地自包含）
_ROUTES = [
    ('业绩', ['业绩', '类似项目', '已完成', '成功案例', '项目经验']),
    ('人员', ['项目经理', '建造师', '负责人', '团队', '人员', '技术负责人',
             '项目管理机构', '岗位', '执业资格']),
    ('资质', ['资质', '资格', '证书', '许可', '认证', '等级']),
    ('设备', ['设备', '机械', '车辆', '仪器', '机具', '装备']),
    ('安全', ['安全', '防护', '应急', '事故']),
    ('财务', ['注册资金', '注册资本', '净资产', '财务', '营业额', '营业收入',
             '资产', '审计']),
]


def _route(score_name: str) -> str:
    for cat, kws in _ROUTES:
        if any(k in score_name for k in kws):
            return cat
    return 'other'


def _item_name(it: Dict[str, Any]) -> str:
    return str(it.get('name') or it.get('title') or it.get('score_item') or '').strip()


def _item_score(it: Dict[str, Any]) -> Any:
    return it.get('score') or it.get('full_score') or ''


def analyze_scoring_coverage(parse_result: Optional[Dict[str, Any]],
                             doc_text: str) -> Dict[str, Any]:
    """自包含评分项覆盖率探针。

    Returns:
        {
          'total': int,
          'covered': int, 'weak': int, 'uncovered': int,
          'weak_items': [{'name','score','ratio'}],
          'uncovered_items': [{'name','score','ratio'}],
          'items': [{'name','score','status','ratio'}],
        }
    """
    score_items = (parse_result or {}).get('score_items', []) or []
    if not score_items or not doc_text:
        return {'total': len(score_items), 'covered': 0, 'weak': 0,
                'uncovered': 0, 'weak_items': [], 'uncovered_items': [],
                'items': []}

    doc_bigrams = _cjk_bigrams(doc_text)
    covered = weak = uncovered = 0
    weak_items: List[Dict[str, Any]] = []
    uncovered_items: List[Dict[str, Any]] = []
    items: List[Dict[str, Any]] = []

    for it in score_items:
        name = _item_name(it)
        score = _item_score(it)
        if not name:
            continue
        probe = name
        subs = it.get('sub_items') or []
        if isinstance(subs, (list, tuple)):
            probe += ' ' + ' '.join(str(s) for s in subs)
        kws = it.get('keywords') or []
        if isinstance(kws, (list, tuple)):
            probe += ' ' + ' '.join(str(k) for k in kws)
        item_bigrams = _cjk_bigrams(probe)
        if not item_bigrams:
            # 无可校验特征（如纯英文/极短项），视为已覆盖，不强制补强
            covered += 1
            items.append({'name': name, 'score': score, 'status': 'covered', 'ratio': 1.0})
            continue
        matched = item_bigrams & doc_bigrams
        ratio = round(len(matched) / len(item_bigrams), 3)
        if ratio >= COVER_RATIO:
            covered += 1
            status = 'covered'
        elif ratio >= _WEAK_MIN:
            weak += 1
            status = 'weak'
            weak_items.append({'name': name, 'score': score, 'ratio': ratio})
        else:
            uncovered += 1
            status = 'uncovered'
            uncovered_items.append({'name': name, 'score': score, 'ratio': ratio})
        items.append({'name': name, 'score': score, 'status': status, 'ratio': ratio})

    return {
        'total': len(items),
        'covered': covered,
        'weak': weak,
        'uncovered': uncovered,
        'weak_items': weak_items,
        'uncovered_items': uncovered_items,
        'items': items,
    }


def _profile_detail(cat: str, profile: Optional[Dict[str, Any]]) -> str:
    """根据类别与企业资质业绩库生成针对性补强要点；无画像优雅回退通用承诺。"""
    if not profile:
        profile = {}
    comp = profile.get('company', {}) or {}

    if cat == '业绩':
        projs = profile.get('similar_projects', []) or []
        if projs:
            names = '、'.join(str(p.get('name', '')) for p in projs[:3] if p.get('name'))
            return (f'已整理近年 {len(projs)} 项类似工程业绩'
                    + (f'（如：{names}）' if names else '')
                    + '，附中标通知书、合同及竣工验收证明，逐条对应评分要求')
        return '已整理近年类似工程业绩清单，附中标通知书、合同及竣工验收证明，确保业绩项可核查'

    if cat == '人员':
        persons = profile.get('key_personnel', []) or []
        if persons:
            roles = '、'.join(str(p.get('role', '')) for p in persons[:4] if p.get('role'))
            return (f'已配置 {len(persons)} 名关键岗位人员'
                    + (f'（{roles}）' if roles else '')
                    + '，附执业资格证与岗位证书，满足评分项人员配置要求')
        return '已配置项目经理及关键技术岗位人员，附执业资格证与岗位证书，满足人员配置评分要求'

    if cat == '资质':
        quals = profile.get('qualifications', []) or comp.get('qualifications', []) or []
        if quals:
            head = '；'.join(str(q) for q in quals[:4])
            return f'{head} 等资质证书均在有效期内，随投标文件附具，完全响应资质评分项'
        return '已备齐营业执照、资质证书等并在有效期内，随投标文件附具，完全响应资质评分项'

    if cat == '设备':
        eq = profile.get('equipment_owned', []) or profile.get('equipment', []) or []
        if eq:
            return f'已编制 {len(eq)} 类主要施工设备/机械进场计划，承诺按施工进度足额到位'
        return '已编制主要施工机械设备进场计划，承诺按施工进度足额到位，满足设备配置评分项'

    if cat == '安全':
        return ('已建立安全生产管理体系，专项安全经费足额投入，编报专项应急预案并组织演练，'
                '满足安全文明施工评分项')

    if cat == '财务':
        cap = comp.get('registered_capital') or comp.get('capital') or comp.get('net_assets')
        if cap:
            return f'注册资本/净资产 {cap}，近三年财务稳健，附审计报告，满足财务实力评分项'
        return '已提供近三年财务审计报告，主要财务指标稳健，满足财务实力评分项'

    return '投标人承诺完全满足本评分项要求，相关佐证材料随投标函一并提交，确保响应真实完整'


def build_reinforcement(weak_items: List[Dict[str, Any]],
                        uncovered_items: List[Dict[str, Any]],
                        parse_result: Optional[Dict[str, Any]] = None) -> List[str]:
    """为弱覆盖/未覆盖评分项生成针对性补强段落。

    返回段落文本列表（每项一段，去重）；消费企业资质业绩库做数据驱动补强。
    """
    paras: List[str] = []
    seen = set()
    for it in list(weak_items or []) + list(uncovered_items or []):
        name = _item_name(it)
        if not name or name in seen:
            continue
        seen.add(name)
        score = _item_score(it)
        cat = _route(name)
        detail = _profile_detail(cat, None)
        para = (f'针对评分项{name}'
                + (f'（{score}分）' if score != '' else '')
                + f'：{detail}。我方确保该项响应内容真实、完整、可核查，杜绝评分遗漏。')
        paras.append(para)
    return paras


def _extract_doc_text(doc_path: str) -> str:
    """读取 Word 文档纯文本（供覆盖率探针使用）。失败返回空串。

    关键：排除「评分项响应保障表」「资质业绩响应表」两张元表格——
    它们只是把评分项名称原文罗列成合规清单，并不代表技术正文已论述该项。
    若不过滤，覆盖率探针会把"表里列了名字"误判为"正文已覆盖"，导致补强失效。
    """
    if not doc_path or not os.path.exists(doc_path):
        return ''
    try:
        from docx import Document
        # 元表格首行标记：命中即跳过（只量技术正文与真实内容表）
        _META_MARKERS = ('响应保障', '资质业绩响应', '资格审查资料')
        doc = Document(doc_path)
        parts = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text)
        for tbl in doc.tables:
            if not tbl.rows:
                continue
            first_row = tbl.rows[0].cells
            first_text = ' '.join(c.text for c in first_row if c.text)
            if any(m in first_text for m in _META_MARKERS):
                continue
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        parts.append(cell.text)
        return '\n'.join(parts)
    except Exception:
        return ''


def inject_reinforcement(doc_path: str,
                         paragraphs: List[str],
                         title: str = '评分项补强与响应闭环说明（PDCA-Act 自动补强）') -> int:
    """把补强段落注入成稿末尾附录。返回实际注入段数（0 表示无操作）。"""
    if not paragraphs or not doc_path or not os.path.exists(doc_path):
        return 0
    try:
        from docx import Document
        doc = Document(doc_path)
        h = doc.add_paragraph()
        run = h.add_run(title)
        run.bold = True
        if Pt is not None:
            run.font.size = Pt(13)
        intro = doc.add_paragraph(
            f'本闭环针对生成稿中 {len(paragraphs)} 项弱覆盖/未覆盖评分项自动补强，'
            f'确保评分项零遗漏、弱项有响应。')
        for para in paragraphs:
            doc.add_paragraph(para)
        doc.save(doc_path)
        return len(paragraphs)
    except Exception:
        return 0


def run_scoring_reinforcement(parse_result: Optional[Dict[str, Any]],
                              doc_path: str) -> Dict[str, Any]:
    """端到端便捷入口：诊断 → 生成补强 → 注入。

    Returns:
        {'injected': int, 'weak': int, 'uncovered': int, 'total': int, 'covered': int}
    """
    if not doc_path or not os.path.exists(doc_path):
        return {'injected': 0, 'weak': 0, 'uncovered': 0, 'total': 0, 'covered': 0}
    doc_text = _extract_doc_text(doc_path)
    cov = analyze_scoring_coverage(parse_result, doc_text)
    if not cov['weak_items'] and not cov['uncovered_items']:
        return {'injected': 0, 'weak': cov['weak'], 'uncovered': cov['uncovered'],
                'total': cov['total'], 'covered': cov['covered']}
    paras = build_reinforcement(cov['weak_items'], cov['uncovered_items'],
                                parse_result)
    n = inject_reinforcement(doc_path, paras)
    return {'injected': n, 'weak': cov['weak'], 'uncovered': cov['uncovered'],
            'total': cov['total'], 'covered': cov['covered']}

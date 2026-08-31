"""
投标文件偏离表 Word 渲染 v1.0 — P0 升级

严格镜像 gantt.py 的 python-docx 表格样式（Table Grid + 居中 + Pt(8) 加粗表头），
并在此基础上增加偏离状态的语义化配色（负偏离红 / 正偏离绿）。

用法（在 TechnicalBidGenerator 中）:
    from bid_technical.tables.deviation_table import generate_deviation_table
    generate_deviation_table(doc, project_info, deviation_report)
"""
from __future__ import annotations

from typing import Any, Dict

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# 偏离状态配色
_COLOR_NEGATIVE = RGBColor(0xC0, 0x00, 0x00)   # 红：负偏离 / 未响应
_COLOR_POSITIVE = RGBColor(0x00, 0x80, 0x00)   # 绿：正偏离
_COLOR_WARN = RGBColor(0xB8, 0x6A, 0x00)       # 橙：部分偏离

DeviationColors = {
    '负偏离': _COLOR_NEGATIVE,
    '未响应': _COLOR_NEGATIVE,
    '正偏离': _COLOR_POSITIVE,
    '部分偏离': _COLOR_WARN,
}

HEADER = ['序号', '条款类别', '招标文件要求（实质性内容）', '偏离情况', '投标响应说明']


def generate_deviation_table(doc: Document,
                             project_info: Dict[str, Any],
                             report: Dict[str, Any]) -> None:
    """在 doc 末尾追加"投标文件偏离表"。

    Args:
        doc: python-docx Document 对象
        project_info: 项目信息 dict（用于标题中的项目名称）
        report: DeviationChecker.generate() 返回的报告
    """
    items = report.get('items', [])
    if not items:
        return

    # —— 标题 ——
    name = (project_info or {}).get('name', '本工程')
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f'{name} — 投标文件偏离表')
    run.bold = True
    run.font.size = Pt(14)

    # —— 摘要行 ——
    summary = report.get('summary', {})
    summary_text = '，'.join(
        f'{k} {v} 项' for k, v in summary.items() if v
    )
    sline = doc.add_paragraph()
    srun = sline.add_run(f'偏离情况汇总：{summary_text or "无"}')
    srun.font.size = Pt(9)
    srun.italic = True

    # —— 无偏离声明（招标未列明具体偏离行时，按"完全响应、无偏离"口径填写）——
    decl = doc.add_paragraph()
    drun = decl.add_run(
        '我方承诺：投标文件完全响应招标文件（含投标人须知、评标办法、合同条款、'
        '技术标准和要求等）的全部实质性要求和条件，无任何偏离。下表仅就招标文件'
        '明确列出的废标/否决及资格审查条款逐条确认我方均满足，未列出的条款均视为'
        '我方完全响应、无偏离。')
    drun.font.size = Pt(9)

    # —— 表格 ——
    rows = len(items) + 1
    cols = len(HEADER)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for j, text in enumerate(HEADER):
        cell = table.rows[0].cells[j]
        cell.text = text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                r.bold = True
                r.font.size = Pt(9)

    # 数据行
    for i, it in enumerate(items, 1):
        cells = table.rows[i].cells
        content = it['content']
        if len(content) > 120:
            content = content[:120] + '…'
        values = [
            it['req_id'],
            it['category'],
            content,
            it['deviation'],
            it.get('note', ''),
        ]
        for j, val in enumerate(values):
            cell = cells[j]
            cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in (0, 1, 3) else WD_ALIGN_PARAGRAPH.LEFT
                for r in para.runs:
                    r.font.size = Pt(8)
                    # 偏离状态配色
                    if j == 3 and val in DeviationColors:
                        r.bold = True
                        r.font.color.rgb = DeviationColors[val]

    # —— 风险提示 ——
    if report.get('risk_level') == 'high':
        risk_p = doc.add_paragraph()
        rrun = risk_p.add_run('⚠ ' + (report.get('risk_notes', ['']) or [''])[0])
        rrun.bold = True
        rrun.font.size = Pt(9)
        rrun.font.color.rgb = _COLOR_NEGATIVE

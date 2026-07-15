"""
BidGenerator - 标书生成入口 v5.1
v5.1: 页眉支持；安全生产许可证自动检测；bid_section参数(technical/commercial/qualification)；进度日志
v5.0: 封面页+目录+页码一键生成；bid_type智能推断；暗标user_context过滤词自动注入
v4.2: 暗标模式+3轮修复循环
v4.1: AI增强/Repair循环/去AI化/降重/查重/钩子系统
v4.0: detail_level自动计算、评分项对齐
"""
import os

from bid_technical.generator import TechnicalBidGenerator
from bid_core.randomizer import Randomizer
from bid_core.user_context import UserContext
from bid_core.logger import get_logger

log = get_logger(__name__)

# 合法的bid_section取值
VALID_BID_SECTIONS = ('technical', 'commercial', 'qualification')


def _infer_bid_type(project_info):
    """根据project_info内容智能推断标书类型"""
    if project_info.get('bid_type'):
        return project_info['bid_type']
    work = project_info.get('work_content', '')
    service_kw = ['服务', '清洗', '维护', '保洁', '物业', '咨询', '运营', '巡查', '巡检', '拆除服务', '征收服务', '管理服务']
    construct_kw = ['施工', '装饰', '装修', '安装', '改造', '整治', '维修', '新建', '改扩建', '外立面']
    s_score = sum(1 for kw in service_kw if kw in work)
    c_score = sum(1 for kw in construct_kw if kw in work)
    if s_score > c_score:
        return 'service'
    elif c_score > 0:
        return 'construction'
    divisions = project_info.get('divisions', [])
    if divisions:
        div_text = ' '.join(divisions)
        s_score = sum(1 for kw in service_kw if kw in div_text)
        c_score = sum(1 for kw in construct_kw if kw in div_text)
        if s_score > c_score:
            return 'service'
    return 'construction'


_CN_DIGITS = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九']


def _cn_date(d) -> str:
    """将 date 转为中文日期，如 二〇二六年七月十二日。"""
    def _cn_int(n: int) -> str:
        return ''.join(_CN_DIGITS[int(c)] for c in str(n))
    return f'{_cn_int(d.year)}年{_cn_int(d.month)}月{_cn_int(d.day)}日'


def _calc_detail_level(target_pages):
    if target_pages <= 0:
        return 2
    if target_pages <= 50:
        return 1
    if target_pages <= 200:
        return 2
    return 3


def _resolve_generator(bid_section, project_info, chapters, target_pages,
                       formatter, randomizer, ctx, heading_font, body_font,
                       detail_level, parse_result, enable_deviation_table=True,
                       enable_risk_grading=True, enable_mock_review=True,
                       llm_client=None, differentiator=None):
    """根据bid_section选择对应的生成器类"""
    if bid_section == 'commercial':
        from bid_commercial.generator import CommercialBidGenerator
        return CommercialBidGenerator(
            project_info=project_info,
            chapters=chapters,
            target_pages=target_pages,
            formatter=formatter,
            randomizer=randomizer,
            user_context=ctx,
            heading_font=heading_font,
            body_font=body_font,
            detail_level=detail_level,
            parse_result=parse_result,
        )
    elif bid_section == 'qualification':
        from bid_qualification.generator import QualificationBidGenerator
        return QualificationBidGenerator(
            project_info=project_info,
            chapters=chapters,
            target_pages=target_pages,
            formatter=formatter,
            randomizer=randomizer,
            user_context=ctx,
            heading_font=heading_font,
            body_font=body_font,
            detail_level=detail_level,
            parse_result=parse_result,
        )
    else:
        # 默认技术标
        return TechnicalBidGenerator(
            project_info=project_info,
            chapters=chapters,
            target_pages=target_pages,
            formatter=formatter,
            randomizer=randomizer,
            user_context=ctx,
            heading_font=heading_font,
            body_font=body_font,
            detail_level=detail_level,
            parse_result=parse_result,
            enable_deviation_table=enable_deviation_table,
            enable_risk_grading=enable_risk_grading,
            enable_mock_review=enable_mock_review,
            llm_client=llm_client,
            differentiator=differentiator,
        )


def generate_bid(
    project_info,
    chapters=None,
    target_pages=300,
    output_path="技术标书.docx",
    user_context=None,
    auto_check=True,
    randomize=False,
    parse_result=None,
    detail_level=None,
    is_dark_bid=False,
    dark_bid_filter_words=None,
    add_cover=True,
    add_toc=True,
    add_page_numbers=True,
    bid_type=None,
    bid_section="technical",
    mode="normal",
    merge_with=None,
    heading_font=None,
    body_font=None,
    enable_deviation_table=True,
    reference_file=None,
    enable_risk_grading=True,
    enable_mock_review=True,
    enable_knowledge_base=True,
    knowledge_base_path=None,
    llm_client=None,
    differentiator=None,
    enable_ppt=False,
):
    """生成标书 v5.1 — 入口函数

    v5.1新增:
    - bid_section参数: technical/commercial/qualification
    - 页眉支持: project_info中header_text时自动添加页眉
    - 安全生产许可证自动检测: 通过user_context.needs_safety_license()判断
    - 进度日志: 关键步骤打印日志到控制台
    """
    from bid_core.formatter import PremiumFormatter, DarkFormatter

    # ── 步骤1: 参数校验与bid_section ──
    if bid_section not in VALID_BID_SECTIONS:
        log.warning("bid_section=%s 不合法，合法值为%s，已降级为'technical'",
                    bid_section, VALID_BID_SECTIONS)
        bid_section = 'technical'

    # v7.3: 透传暗标标记，供富内容引擎跳过配图
    project_info['is_dark_bid'] = bool(is_dark_bid)

    log.info("开始生成标书 | bid_section=%s | 项目=%s",
             bid_section, project_info.get('name', '未命名'))

    # ── 步骤2: 推断bid_type ──
    if bid_type:
        project_info['bid_type'] = bid_type
    elif project_info.get('bid_type') is None:
        project_info['bid_type'] = _infer_bid_type(project_info)

    log.info("bid_type推断完成:")

    # ── 步骤3: 自动检测安全生产许可证 ──
    if isinstance(user_context, UserContext):
        ctx = user_context
    elif isinstance(user_context, dict):
        ctx = UserContext(user_context)
    else:
        ctx = UserContext()

    if ctx.needs_safety_license(project_info):
        project_info['requires_safety_license'] = True
        log.info("检测到需要安全生产许可证，已在project_info中标记")
    else:
        project_info['requires_safety_license'] = False

    # ── 步骤4: 计算detail_level ──
    if detail_level is None:
        detail_level = _calc_detail_level(target_pages)
    log.info("detail_level=%s | target_pages=%s", detail_level, target_pages)

    # ── 步骤5: 初始化随机化引擎 ──
    _script_dir = os.path.dirname(os.path.abspath(__file__))
    # 兼容两种部署布局：
    #  - 扁平化结构（模块与 data/ 同级）：_script_dir/data 存在
    #  - scripts/ 子目录结构：dirname(_script_dir)/data 存在
    # 旧实现固定取 dirname(_script_dir)，在扁平化布局下会指向不存在的 data/ 目录，
    # 导致企业知识库/同义词静默加载失败、封面退化为占位符（见 v8.3 ADR）。
    _project_root = _script_dir
    for _cand in (_script_dir, os.path.dirname(_script_dir)):
        if os.path.isdir(os.path.join(_cand, 'data')):
            _project_root = _cand
            break
    synonyms_path = os.path.join(_project_root, 'data', 'synonyms.json')
    randomizer = Randomizer(synonyms_path=synonyms_path, enabled=randomize)

    # ── 步骤5.5: P3 企业知识库复用 ──
    if enable_knowledge_base:
        kb_path = knowledge_base_path or os.path.join(_project_root, 'data', 'user_knowledge_base.json')
        try:
            if not isinstance(ctx, UserContext):
                ctx = UserContext()
            ctx.merge_knowledge_base(kb_path)
            log.info("已加载企业知识库: %s", kb_path)
        except Exception as exc:
            log.error("知识库加载失败，已忽略: %s", exc)

    # 用企业知识库填充封面/页眉投标人名称与投标日期，杜绝输出
    # "（投标人名称）""   月" 等"提示替换"占位符（用户要求生成即最终版）。
    if not project_info.get('bidder_name'):
        try:
            _comp = ctx.get_company()
            _bn = (_comp or {}).get('name')
            if _bn:
                project_info['bidder_name'] = _bn
        except Exception:
            pass
    if not project_info.get('bid_date'):
        try:
            import datetime
            project_info['bid_date'] = _cn_date(datetime.date.today())
        except Exception:
            pass

    # ── 步骤6: 创建格式化引擎 ──
    if is_dark_bid:
        formatter = DarkFormatter(
            is_dark_bid=True,
            filter_words=dark_bid_filter_words,
            heading_marker=project_info.get('dark_heading_marker', '●'),
            heading_font=heading_font,
            body_font=body_font,
        )
        if ctx.has_info():
            extra_words = ctx.get_filter_words()
            if extra_words:
                formatter.set_filter_words(extra_words)
        log.info("暗标模式已启用")
    else:
        formatter = PremiumFormatter(
            heading_font=heading_font,
            body_font=body_font,
        )

    # ── 步骤7: 页眉支持 ──
    header_text = project_info.get('header_text', '')
    if header_text:
        formatter.add_header(header_text)
        log.info("页眉已添加: 「%s」", header_text)

    # ── 步骤8: 封面和目录 ──
    if add_cover:
        formatter.add_cover_page(project_info)
        log.info("封面页已生成")
    if add_toc:
        formatter.add_toc()
        log.info("目录页已生成")

    # ── 步骤8.5: P1 以标写标 — 若提供参考标书则套用其章节结构 ──
    chapters_arg = chapters
    if reference_file:
        try:
            from bid_core.reference_loader import ReferenceLoader
            ref = ReferenceLoader(reference_file).load()
            outline = ref.get_adapted_outline(project_info)
            if outline:
                chapters_arg = outline
                log.info("以标写标: 套用参考标书 %s 个章节结构", len(outline))
        except Exception as exc:
            log.error("参考标书加载失败，已忽略: %s", exc)

    # ── 步骤9: 根据bid_section选择生成器 ──
    generator = _resolve_generator(
        bid_section, project_info, chapters_arg, target_pages,
        formatter, randomizer, ctx, heading_font, body_font,
        detail_level, parse_result, enable_deviation_table,
        enable_risk_grading, enable_mock_review,
        llm_client=llm_client,
        differentiator=differentiator,
    )
    log.info("生成器已创建: %s", generator.__class__.__name__)

    # ── 步骤10: 生成标书 ──
    result_path = generator.generate(output_path)
    log.info("标书内容已生成: %s", result_path)

    # ── 步骤11: 页码 ──
    if add_page_numbers and not is_dark_bid:
        formatter.add_page_numbers()
        formatter.save(result_path)
        log.info("页码已添加")

    log.info("标书生成完成: %s", result_path)

    # ── 步骤12: 可选述标PPT（v7.6，需 python-pptx；缺失则跳过）──
    if enable_ppt:
        try:
            from bid_core.ppt_generator import generate_bid_ppt
            ppt_path = generate_bid_ppt(parse_result, project_info, result_path)
            log.info("述标PPT已生成: %s", ppt_path)
        except ImportError:
            log.warning("未安装 python-pptx，跳过述标PPT（pip install python-pptx 启用）")
        except Exception as exc:
            log.error("述标PPT生成失败，已忽略: %s", exc)

    return result_path


def generate_bid_full(project_info, parse_result=None, user_context=None,
                      output_path="技术标书.docx", is_dark_bid=False, **kwargs):
    """生成300+页完整标书"""
    return generate_bid(
        project_info=project_info, target_pages=300, parse_result=parse_result,
        user_context=user_context, output_path=output_path, detail_level=3,
        is_dark_bid=is_dark_bid, **kwargs)


def generate_bid_randomized(project_info, chapters=None, target_pages=300,
                            output_path="技术标书.docx", user_context=None, **kwargs):
    """生成带随机化的标书"""
    return generate_bid(
        project_info=project_info, chapters=chapters, target_pages=target_pages,
        output_path=output_path, user_context=user_context, randomize=True, **kwargs)


def generate_bid_with_context(project_info, user_context, chapters=None,
                              target_pages=300, output_path="技术标书.docx", **kwargs):
    """生成带用户信息的标书"""
    return generate_bid(
        project_info=project_info, chapters=chapters, target_pages=target_pages,
        output_path=output_path, user_context=user_context, **kwargs)


def generate_bid_with_hooks(
    project_info, chapters=None, target_pages=300, output_path="技术标书.docx",
    user_context=None, parse_result=None, detail_level=None, randomize=False,
    is_dark_bid=False, dark_bid_filter_words=None,
    add_cover=True, add_toc=True, add_page_numbers=True,
    enable_deai=True, enable_rewrite=False, enable_duplicate=False, enable_format=True,
    enable_deviation_table=True,
    reference_file=None,
    enable_risk_grading=True,
    enable_mock_review=True,
    enable_knowledge_base=True,
    knowledge_base_path=None,
    llm_client=None,
    differentiator=None,
    **kwargs,
):
    """生成标书+后处理钩子 v5.1"""
    doc_path = generate_bid(
        project_info=project_info, chapters=chapters, target_pages=target_pages,
        output_path=output_path, user_context=user_context, parse_result=parse_result,
        detail_level=detail_level, randomize=randomize,
        is_dark_bid=is_dark_bid, dark_bid_filter_words=dark_bid_filter_words,
        add_cover=add_cover, add_toc=add_toc, add_page_numbers=add_page_numbers,
        enable_deviation_table=enable_deviation_table,
        reference_file=reference_file,
        enable_risk_grading=enable_risk_grading,
        enable_mock_review=enable_mock_review,
        enable_knowledge_base=enable_knowledge_base,
        knowledge_base_path=knowledge_base_path,
        llm_client=llm_client,
        differentiator=differentiator,
        **kwargs)
    from bid_core.hooks import run_post_hooks
    hook_result = run_post_hooks(
        docx_path=doc_path, project_info=project_info, parse_result=parse_result,
        user_context=user_context, enable_deai=enable_deai, enable_rewrite=enable_rewrite,
        enable_duplicate=enable_duplicate, enable_format=enable_format)
    return {'doc_path': doc_path, 'hooks': hook_result}


def check_and_repair(project_info, parse_result=None, bid_doc_path=None,
                     doc_info=None, max_rounds=3):
    """合规检查+修复循环（最多3轮）

    v8.7 修复：原实现依赖 checker v2.0 重构后已移除的 ``check_bid``，
    首次调用即 ``ImportError``，导致 advertised 的「检查+修复」闭环失效。
    现直接驱动 ``BidChecker``，并把结果归一化为 repair 引擎 /
    ``_auto_fix_critical`` 兼容的报告结构。
    """
    from checker import BidChecker
    from bid_core.repair import repair_bid, BidRepairer

    parse_result = parse_result or {}
    doc_info = doc_info or {}

    def _run_check():
        text = doc_info.get('text') or _read_docx_text(bid_doc_path)
        context = {
            'text': text or '',
            'docx_path': bid_doc_path,
            'required_clauses': [c.get('content', '') for c in parse_result.get('star_clauses', [])],
            'score_items': [s.get('name', '') for s in parse_result.get('score_items', [])],
            'tables': doc_info.get('tables', []),
        }
        checker = BidChecker()
        checker.check(context)
        report = checker.to_report_dict()
        critical_issues = report.get('critical_issues', [])
        warning_issues = [r for r in report.get('issues', []) if r.get('severity') == 'warning']
        # 归一化：repair 引擎 / _auto_fix_critical 读取 results.critical[].id
        # 与 critical_issues 列表；summary 提供 critical / warning 计数。
        return {
            'summary': {
                'critical': len(critical_issues),
                'warning': len(warning_issues),
                **report.get('summary', {}),
            },
            'results': {
                'critical': [_norm_issue(i) for i in critical_issues],
                'warning': [_norm_issue(i) for i in warning_issues],
            },
            'issues': report.get('issues', []),
            'critical_issues': critical_issues,
            'unresponded_clauses': doc_info.get('unresponded_clauses', []),
        }

    rounds_log = []
    for round_num in range(1, max_rounds + 1):
        check_result = _run_check()
        summary = check_result['summary']
        critical_count = summary['critical']
        warning_count = summary['warning']
        round_info = {'round': round_num, 'critical': critical_count, 'warning': warning_count, 'passed': critical_count == 0}
        if critical_count == 0:
            round_info['action'] = 'pass'
            rounds_log.append(round_info)
            return {'passed': True, 'rounds': rounds_log, 'final_check_result': check_result,
                    'total_critical': 0, 'total_warning': warning_count,
                    'message': f'第{round_num}轮检查通过'}
        # 生成修复建议（直接取 prompt，供上层展示/下发 LLM）
        repairer = BidRepairer(check_result, bid_doc_path=bid_doc_path)
        repair_prompts = repairer.generate_repair_prompts()
        repair_info = repairer.run_repair_loop()
        round_info['action'] = 'repair'
        round_info['repair_prompts'] = repair_prompts
        round_info['repair_summary'] = repair_info
        rounds_log.append(round_info)
        if round_num == max_rounds:
            return {'passed': False, 'rounds': rounds_log, 'final_check_result': check_result,
                    'total_critical': critical_count, 'total_warning': warning_count,
                    'repair_prompts': repair_prompts,
                    'message': f'经{max_rounds}轮修复仍有{critical_count}个critical问题'}
        auto_fixes = _auto_fix_critical(check_result, bid_doc_path)
        if auto_fixes:
            round_info['auto_fixed'] = auto_fixes
    return {'passed': False, 'rounds': rounds_log, 'final_check_result': check_result,
            'total_critical': summary['critical'], 'total_warning': summary['warning'],
            'message': '修复循环结束，仍有问题'}


def _read_docx_text(path):
    """读取 Word 文档纯文本（供检查上下文使用）。失败返回空串，不阻断上层。"""
    if not path or not os.path.exists(path):
        return ''
    try:
        from docx import Document
        doc = Document(path)
        return '\n'.join(p.text for p in doc.paragraphs if p.text and p.text.strip())
    except Exception:
        return ''


def _norm_issue(i):
    """将 BidChecker.to_report_dict() 的 issue 字典归一化为 repair 引擎 /
    ``_auto_fix_critical`` 兼容结构（保留原字段，避免对象发散）。"""
    return {
        'id': i.get('rule_id', ''),
        'code': i.get('rule_id', ''),
        'name': i.get('rule_name', i.get('name', '未知问题')),
        'description': i.get('message', i.get('description', '')),
        'severity': i.get('severity', 'medium'),
        **i,
    }


def _auto_fix_critical(check_result, bid_doc_path):
    """自动修复可处理的问题（v8.7 增强：兼容当前规则集 rule_id）

    在文档末尾追加针对性响应段落；受 try/except 保护，任何异常仅告警不中断。
    """
    from docx import Document
    if not bid_doc_path or not os.path.exists(bid_doc_path):
        return []
    fixes = []
    results = check_result.get('results', {})
    try:
        doc = Document(bid_doc_path)
        modified = False
        for item in results.get('critical', []):
            item_id = item.get('id', '')
            details = item.get('details', {}) or {}
            if item_id == 'DQ003':
                doc.add_paragraph('安全生产许可证')
                doc.add_paragraph('我方持有有效的安全生产许可证，证号在投标文件中随附，确保施工全过程安全生产合法合规。')
                modified = True
                fixes.append({'id': item_id, 'action': 'auto_added_safety_license'})
            elif item_id == 'DQ002':
                doc.add_paragraph('质量目标')
                doc.add_paragraph('我方承诺本工程质量目标为合格，符合国家现行工程施工质量验收规范标准。')
                modified = True
                fixes.append({'id': item_id, 'action': 'auto_added_quality_target'})
            elif item_id == 'C001':  # 废标条款响应
                missing = details.get('missing_clauses', [])
                doc.add_paragraph('废标条款响应')
                doc.add_paragraph(
                    '我方承诺严格响应招标文件全部废标/否决条款'
                    + ('，重点落实：' + '；'.join(missing) if missing else '，完全响应全部废标条款')
                    + '。'
                )
                modified = True
                fixes.append({'id': item_id, 'action': 'auto_added_clause_response'})
            elif item_id == 'P001':  # 必选章节
                missing = details.get('missing', [])
                if missing:
                    doc.add_paragraph('补充章节')
                    for sec in missing:
                        doc.add_paragraph(sec)
                    modified = True
                    fixes.append({'id': item_id, 'action': 'auto_added_sections'})
        if modified:
            doc.save(bid_doc_path)
    except Exception as e:
        from bid_core.logger import get_logger
        get_logger(__name__).warning('自动修复 critical 项失败: %s', e)
    return fixes


def apply_repair_to_doc(bid_doc_path, repair_contents, output_path=None):
    from bid_core.repair import BidRepairer
    return BidRepairer({}, bid_doc_path).apply_repair(repair_contents, output_path)


def get_ai_enhance_prompts(project_info, chapter_contents, parse_result=None,
                           user_context=None, detail_level=2):
    from bid_core.ai_enhance import generate_enhance_prompts
    return generate_enhance_prompts(project_info, chapter_contents, parse_result, user_context, detail_level)


def apply_ai_enhancement(docx_path, enhanced_chapters, output_path=None):
    from bid_core.ai_enhance import merge_enhanced_content
    return merge_enhanced_content(docx_path, enhanced_chapters, output_path)


def check_duplicates(file_paths, mode='标书'):
    from duplicate_checker import check_duplicates as _check
    return _check(file_paths, mode=mode)

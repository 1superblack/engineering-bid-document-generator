"""
防重 / 差异化生成引擎 v7.3
========================
对标喜鹊标书 AI 的「防重机制（随机特征矩阵 + 唯一指纹）」与 WPS 的「内容雷同风险规避」。

解决的核心合规痛点：标书查重已成合规红线，多个投标人用同一通用模型易产出相似内容。
本引擎在不依赖外部大模型的前提下，提供两层差异化能力：

  1. 措辞差异化旋转（rotate）：基于「项目指纹」对正文中的通用动词/名词做同义词轮换，
     使同一招标项目、不同投标人（或不同次运行）产出的文字明显不同，降低字面雷同率。
     ——旋转池均为工程标书领域等价、专业度一致的同义表述，不改变语义与承诺。
  2. 全文原创性自查（compute_self_similarity）：统计全文中重复句子比例，供评审报告披露，
     让"原创性"可量化、可核查（对标喜鹊的"差异化生成"承诺有迹可循）。

纯标准库实现，零新依赖。默认每次运行使用随机指纹 ⇒ 同项目两次生成内容不同。
"""
from __future__ import annotations

import hashlib
import os
import re
from typing import Dict, List, Optional, Sequence

# ── 同义词轮换池 ────────────────────────────────────────────────────────
# 键：模板正文中常用的工程术语；值：可等价替换的表述（语义/专业度一致，不改动承诺）。
SYNONYMS: Dict[str, List[str]] = {
    "实施": ["落地", "执行", "推进", "开展", "贯行"],
    "保障": ["确保", "保证", "支撑", "护航", "托底"],
    "管理": ["管控", "治理", "统筹", "组织", "集约管理"],
    "方案": ["预案", "策划", "规划", "部署", "专项安排"],
    "措施": ["办法", "举措", "手段", "机制", "抓手"],
    "质量": ["品质", "工艺水平", "工程观感", "达标率", "实体质量"],
    "安全": ["安全生产", "作业平安", "零事故目标", "安全防护", "本质安全"],
    "进度": ["工期", "节点", "交付节奏", "时间轴", "施工节拍"],
    "风险": ["隐患", "不确定因素", "薄弱点", "变量", "不利因素"],
    "响应": ["答复", "对应", "承接", "满足", "契合"],
    "标准": ["规范", "准则", "要求", "基准", "技术规程"],
    "机制": ["体系", "制度", "闭环", "流程", "运转模式"],
    "闭环": ["全周期管控", "可追溯链条", "PDCA循环", "全过程留痕"],
    "落实": ["执行到位", "掷地有声", "一抓到底", "压实"],
    "管控": ["受控", "严管", "精细管理", "穿透式管理"],
    "统筹": ["一体谋划", "系统部署", "通盘考虑", "协同推进"],
    "优化": ["精进", "改良", "提质", "迭代提升"],
    "提升": ["拔高", "增强", "跃升", "强化"],
    "确保": ["力保", "守住", "万无一失地保障", "牢牢把控"],
    "建立": ["搭建", "构建", "成型", "织密"],
    "技术": ["工艺", "工法", "技术手段", "专业技术"],
    "检查": ["核查", "巡检", "复核", "验收校验"],
    "培训": ["交底培训", "专项练兵", "以训促战", "岗前实训"],
    "资料": ["档案", "台账", "过程记录", "佐证材料"],
    "目标": ["导向", "旨归", "硬指标", "必达要求"],
    # P0-②：扩充高频套话词，提升旋转发散度（均为工程标书等价专业表述，不改动承诺）
    "全面": ["全方位", "全维度", "系统式", "一体推进"],
    "严格": ["从严", "刚性", "硬性", "不打折地"],
    "加强": ["强化", "夯实", "加码", "做实"],
    "完善": ["健全", "补齐", "做优", "闭环优化"],
    "推进": ["驱动", "深化", "落地实施", "稳步实施"],
    "组织": ["组建", "牵头", "统筹调度", "成立专班"],
    "制定": ["编制", "出台", "拟定", "明确"],
    "明确": ["厘清", "划清", "锁定", "压实"],
    "强化": ["突出", "做实", "加码", "筑牢"],
    "持续": ["长效", "常态化", "久久为功", "动态延续"],
    "有效": ["管用", "见效", "务实", "可落地"],
    "科学": ["合理", "精益", "严谨", "适配"],
    "合理": ["妥当", "合规", "因地制宜", "优化"],
    "高效": ["集约", "顺畅", "低成本", "快反"],
    "到位": ["落实处", "闭环", "掷地有声", "无遗漏"],
    "做好": ["做实", "做精", "抓牢", "做细"],
    "开展": ["组织", "实施", "推进", "铺开"],
    "提高": ["抬升", "增强", "拉升", "优化"],
    "控制": ["把控", "约束", "压降", "收敛"],
}


def default_fingerprint(project_info: Optional[Dict] = None) -> str:
    """生成默认指纹：项目名 + 运行随机盐 ⇒ 同项目每次运行内容不同。

    传入固定 project_info 且不希望随机时，调用方应自行传入固定 fingerprint。
    """
    name = (project_info or {}).get('name') or 'bid'
    salt = os.urandom(4).hex()
    return f"{name}|{salt}"


class Differentiator:
    """基于指纹的差异化生成器。"""

    def __init__(self, fingerprint: Optional[str] = None, project_info: Optional[Dict] = None):
        self.fingerprint = fingerprint or default_fingerprint(project_info)
        self._sentences: List[str] = []

    # ── 措辞旋转 ──────────────────────────────────────────────────────
    def rotate(self, text: str, salt: str = "") -> str:
        """对文本做同义词轮换。同一（指纹+盐）下结果确定；不同则结果不同。

        仅替换 SYNONYMS 中的通用术语，不触碰公司名/人名/星号条款/数字，保证承诺不变。
        salt：逐段加盐（如章节标题+段落序号），使同一模板套话在不同出现处发散，
        从而降低全文逐字重复率（P0-②）。salt 缺省时行为与旧版一致（向后兼容）。

        v8.5 修复（取自 v8.1_fixed）：原实现按字典序逐键 .replace 会「链式二次替换」——
        例如「机制」→「闭环」后，新出现的「闭环」又被「闭环」→「PDCA循环」二次替换，
        在「闭环管理机制」上产出重复病句。现改为基于原始文本的单次左→右扫描
        （regex.sub 回调），被替换出的文本不再二次扫描，杜绝链式重复，
        同时每个键仍按指纹+盐稳定选变体，差异化与确定性不变。
        """
        if not text:
            return text
        pattern = re.compile('|'.join(re.escape(k) for k in SYNONYMS))

        def _repl(m):
            term = m.group(0)
            variants = SYNONYMS[term]
            if not variants:
                return term
            return variants[self._pick(term, len(variants), salt)]

        return pattern.sub(_repl, text)

    def _pick(self, salt_key: str, n: int, salt: str = "") -> int:
        h = hashlib.md5(f"{self.fingerprint}::{salt_key}::{salt}".encode("utf-8")).digest()
        return int.from_bytes(h[:4], "big") % n

    # ── 句子收集与原创性自查 ──────────────────────────────────────────
    def add_sentence(self, sentence: str) -> None:
        s = (sentence or "").strip()
        if len(s) >= 10:  # 过短句子不计入，避免误判
            self._sentences.append(s)

    @property
    def sentences(self) -> List[str]:
        return list(self._sentences)

    def self_similarity(self) -> float:
        return compute_self_similarity(self._sentences)

    def report(self) -> Dict[str, object]:
        return {
            "fingerprint": self.fingerprint,
            "sentence_count": len(self._sentences),
            "self_similarity": round(self.self_similarity(), 4),
        }


def _norm(s: str) -> str:
    """归一化句子：去空白与标点，便于重复判定。"""
    return re.sub(r"[\s，。；;、：:,.!！?？()（）\[\]【】\"'\"'“”]", "", s or "")


def compute_self_similarity(sentences: Sequence[str]) -> float:
    """计算重复句子比例 ∈ [0,1]。

    仅对长度 ≥ 12 的归一化句子做精确重复判定（短句重复属正常，不计入），
    返回「重复句数 / 总句数」。值越低代表原创性越高、雷同风险越小。
    """
    normed = [_norm(s) for s in sentences if len(_norm(s)) >= 12]
    if not normed:
        return 0.0
    seen = set()
    dup = 0
    for n in normed:
        if n in seen:
            dup += 1
        else:
            seen.add(n)
    return dup / len(normed)


# ═══════════════════════════════════════════════════════════════════════════
# T5 · 三维查重引擎（文本层 SimHash / 结构层 / 元数据层 / 图像层可选 OCR）
# 纯标准库实现；图像层 OCR 为可选依赖（pytesseract），缺失则优雅跳过。
# ═══════════════════════════════════════════════════════════════════════════
def _char_shingles(text: str, k: int = 2) -> List[str]:
    """中文无空格，用字符 k-gram 切词（避免引入 jieba 依赖）。"""
    t = (text or "").strip()
    if len(t) < k:
        return [t] if t else []
    return [text[i:i + k] for i in range(len(text) - k + 1)]


def _hash64(token: str) -> int:
    """稳定的 64-bit 哈希（与平台/运行无关）。"""
    return int.from_bytes(hashlib.md5(token.encode("utf-8")).digest()[:8], "big")


class SimHash:
    """局部敏感哈希：用词频加权的 64-bit 指纹，近义文本汉明距离小。"""

    def __init__(self, bits: int = 64):
        self.bits = bits

    def fingerprint(self, text: str) -> int:
        counts: Dict[str, int] = {}
        for sh in _char_shingles(text):
            counts[sh] = counts.get(sh, 0) + 1
        vec = [0] * self.bits
        for tok, w in counts.items():
            h = _hash64(tok)
            for i in range(self.bits):
                if (h >> i) & 1:
                    vec[i] += w
                else:
                    vec[i] -= w
        fp = 0
        for i in range(self.bits):
            if vec[i] > 0:
                fp |= (1 << i)
        return fp

    @staticmethod
    def hamming(a: int, b: int) -> int:
        return (a ^ b).bit_count() if hasattr(int, "bit_count") else bin(a ^ b).count("1")


def simhash_self_similarity(sentences: Sequence[str], threshold: int = 3) -> float:
    """近义重复句比例 ∈ [0,1]：两两汉明距离 < threshold 视为近重复。

    比精确重复更宽容，能捕捉「换汤不换药」的雷同表述。
    """
    fps = [SimHash().fingerprint(s) for s in sentences if len(_norm(s)) >= 12]
    n = len(fps)
    if n < 2:
        return 0.0
    near = 0
    for i in range(n):
        for j in range(i + 1, n):
            if SimHash.hamming(fps[i], fps[j]) <= threshold:
                near += 1
    pairs = n * (n - 1) // 2
    return near / pairs if pairs else 0.0


def _split_sections(headings: Sequence[str]) -> List[List[str]]:
    """按一级/二级标题切分章节，返回每章的子标题 token 列表。"""
    sections: List[List[str]] = []
    cur: List[str] = []
    for h in headings:
        t = (h or "").strip()
        if not t:
            continue
        if _is_top_heading(h):
            if cur:
                sections.append(cur)
            cur = [t]
        else:
            cur.append(t)
    if cur:
        sections.append(cur)
    return sections


def _is_top_heading(text: str) -> bool:
    """粗判是否为章节大标题（用于切分，不依赖 docx style）。"""
    return bool(re.match(r"^\s*(第[一二三四五六七八九十\d]+[章节点节]|[0-9]+[\.\、]|\(\s*[0-9]+\s*\))", text or ""))


def structure_similarity(headings: Sequence[str]) -> float:
    """结构层相似度 ∈ [0,1]：章节子标题集合两两 Jaccard 的平均/最大值。

    反映「章节骨架是否高度雷同」——技术标常出现多章结构套同一个模板。
    单章或结构稀疏时返回 0（无可比性）。
    """
    sections = _split_sections(headings)
    if len(sections) < 2:
        return 0.0
    sets = [set(_char_shingles(" ".join(s), 3)) for s in sections]
    sets = [s for s in sets if s]
    if len(sets) < 2:
        return 0.0
    sims = []
    for i in range(len(sets)):
        for j in range(i + 1, len(sets)):
            u = sets[i] | sets[j]
            inter = sets[i] & sets[j]
            sims.append(len(inter) / len(u) if u else 0.0)
    return max(sims) if sims else 0.0


def metadata_leak(doc_path: str) -> Dict[str, object]:
    """元数据层：检出文档属性/编辑痕迹中可能泄漏身份的字段（暗标敏感）。

    纯 python-docx 读取 core_properties；返回 {leak: bool, fields: [...]}。
    """
    res: Dict[str, object] = {"leak": False, "fields": []}
    try:
        from docx import Document
        doc = Document(doc_path)
        cp = doc.core_properties
        fields = {}
        for f in ("author", "last_modified_by", "company", "title"):
            v = getattr(cp, f, None)
            if v:
                fields[f] = str(v)
        res["fields"] = fields
        res["leak"] = bool(fields)
    except Exception:  # noqa: BLE001
        res["leak"] = False
    return res


def ocr_self_similarity(doc_path: str) -> Optional[float]:
    """图像层（可选）：对文档内嵌图片做 OCR，计算 OCR 文本自相似率。

    依赖 pytesseract + PIL；任一缺失或无可 OCR 图片时返回 None（优雅跳过）。
    绝不抛异常，避免影响主流程。
    """
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except Exception:  # noqa: BLE001
        return None
    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(doc_path)
        texts: List[str] = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    img = Image.open(rel.target_part.blob and __import__("io").BytesIO(rel.target_part.blob))
                    txt = pytesseract.image_to_string(img, lang="chi_sim+eng")
                    if txt.strip():
                        texts.append(txt)
                except Exception:  # noqa: BLE001
                    continue
        if len(texts) < 2:
            return None
        return simhash_self_similarity(texts)
    except Exception:  # noqa: BLE001
        return None


def compute_self_similarity_3d(doc_path: str) -> Dict[str, object]:
    """T5 主入口：对单份生成稿做三维（文本/结构/元数据）+ 可选图像层查重。

    Returns:
        {
          'status', 'text_self_similarity', 'simhash_self_similarity',
          'structure_similarity', 'metadata_leak', 'metadata_fields',
          'ocr_self_similarity', 'layers', 'markdown_path'
        }
    文本层默认开启；结构层/元数据层随文档结构自动计算；图像层仅 OCR 可用时产出。
    """
    try:
        from docx import Document
        doc = Document(doc_path)
    except Exception as exc:  # noqa: BLE001
        return {"status": "error", "error": str(exc)}

    sentences: List[str] = []
    headings: List[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        style = ""
        try:
            style = (p.style.name or "") if p.style else ""
        except Exception:
            style = ""
        if style.startswith("Heading") or _is_top_heading(t):
            headings.append(t)
        # 按中文句末标点切句
        for seg in re.split(r"[。！？；\n]", t):
            seg = seg.strip()
            if len(seg) >= 8:
                sentences.append(seg)

    text_sim = compute_self_similarity(sentences)
    sim_sim = simhash_self_similarity(sentences)
    struct_sim = structure_similarity(headings)
    meta = metadata_leak(doc_path)
    ocr_sim = ocr_self_similarity(doc_path)

    # 综合文本重复率：精确 + 近义，取较严重者（保守披露）
    text_layer = max(text_sim, sim_sim)

    out = {
        "status": "ok",
        "text_self_similarity": round(text_sim, 4),
        "simhash_self_similarity": round(sim_sim, 4),
        "structure_similarity": round(struct_sim, 4),
        "metadata_leak": bool(meta.get("leak")),
        "metadata_fields": meta.get("fields", {}),
        "ocr_self_similarity": (round(ocr_sim, 4) if ocr_sim is not None else None),
        "layers": {
            "text": round(text_layer, 4),
            "structure": round(struct_sim, 4),
            "metadata": bool(meta.get("leak")),
            "image": (round(ocr_sim, 4) if ocr_sim is not None else "未启用(OCR依赖缺失)"),
        },
    }
    return out

"""
评标专家视角自检模块 v1.0
从"得分"角度而非"格式"角度检查标书质量
在checker.py之后运行，模拟评标专家的审查逻辑

v1.0:
- 评分项覆盖率检查：逐条验证评分项响应，含"包含但不限于"子条目
- 内容深度评估：按表格/国标/工序多维度评估章节深度(1-5级)
- 废标风险扫描：从专家视角扫描常见废标情形，返回风险等级
- 得分预测：逐项预测评分项得分，含失分原因分析
- 改进建议生成：按优先级排序的具体可操作建议
- 与repair模块集成：输出格式兼容repair模块直接消费
"""
import json
import re
import os
from typing import Dict, List, Optional, Any, Tuple


# ── 常量定义 ────────────────────────────────────────────────

# 深度评估基准（按detail_level分档）
DEPTH_STANDARDS = {
    1: {'min_tables': 1, 'min_standards': 5,  'min_processes': 2,  'label': '简略'},
    2: {'min_tables': 2, 'min_standards': 10, 'min_processes': 4,  'label': '基础'},
    3: {'min_tables': 3, 'min_standards': 20, 'min_processes': 6,  'label': '详实'},
    4: {'min_tables': 4, 'min_standards': 25, 'min_processes': 8,  'label': '专业'},
    5: {'min_tables': 5, 'min_standards': 30, 'min_processes': 10, 'label': '极详'},
}

# 废标风险条款关键词映射
DISQUALIFICATION_RISK_PATTERNS = [
    {
        'id': 'QR001',
        'name': '缺少承诺书',
        'keywords': ['承诺书', '投标承诺', '履约承诺'],
        'risk_level': 'high',
        'description': '标书未包含投标承诺书或履约承诺函',
        'repair_hint': '需补充投标承诺书/履约承诺函，使用"我方郑重承诺"等严格声明词',
    },
    {
        'id': 'QR002',
        'name': '资质过期或缺失',
        'keywords': ['资质证书', '营业执照', '资质等级', '安全生产许可证', '资质有效期'],
        'risk_level': 'high',
        'description': '标书未明确提及资质证书有效性或在有效期范围内',
        'repair_hint': '需补充资质证书有效期内声明，附相关证书复印件',
    },
    {
        'id': 'QR003',
        'name': '关键人员无证',
        'keywords': ['建造师', '注册建造师', '安全员证', '安全考核合格证', '岗位证书', '执业资格'],
        'risk_level': 'high',
        'description': '项目经理或关键岗位人员缺少执业资格证明',
        'repair_hint': '需补充关键人员资格证书编号及有效期声明',
    },
    {
        'id': 'QR004',
        'name': '安全投入不足',
        'keywords': ['安全投入', '安全经费', '安全防护', '安全措施费', '文明施工费'],
        'risk_level': 'high',
        'description': '安全生产投入经费未明确或低于规定比例',
        'repair_hint': '需明确安全措施费数额及占工程造价比例，确保不低于规定标准',
    },
    {
        'id': 'QR005',
        'name': '工期承诺缺失',
        'keywords': ['工期承诺', '总工期', '日历天', '合同工期'],
        'risk_level': 'medium',
        'description': '未明确承诺工期或工期表述模糊',
        'repair_hint': '需用"我方承诺总工期为X日历天"等确定性表述明确工期',
    },
    {
        'id': 'QR006',
        'name': '质量标准未达标',
        'keywords': ['质量目标', '质量标准', '质量等级', '合格', '优良'],
        'risk_level': 'medium',
        'description': '质量目标未明确或不满足招标要求',
        'repair_hint': '需明确工程质量目标，确保不低于招标要求的质量等级',
    },
    {
        'id': 'QR007',
        'name': '环保承诺缺失',
        'keywords': ['环境保护', '环保措施', '绿色施工', '扬尘治理'],
        'risk_level': 'low',
        'description': '缺少环保承诺或绿色施工措施',
        'repair_hint': '建议补充绿色施工措施及环保承诺声明',
    },
    {
        'id': 'QR008',
        'name': '保修承诺缺失',
        'keywords': ['保修期', '质量保修', '保修承诺', '缺陷责任期'],
        'risk_level': 'low',
        'description': '未明确工程保修期或缺陷责任期承诺',
        'repair_hint': '建议补充保修期承诺，明确缺陷责任期和保修范围',
    },
]

# 评分项内容匹配强度等级
MATCH_STRENGTH = {
    'exact': 1.0,     # 标题完全匹配
    'keyword': 0.8,   # 关键词命中
    'content': 0.6,   # 正文内容命中
    'partial': 0.4,   # 部分命中
    'none': 0.0,      # 未命中
}

# "包含但不限于"子条目关键词提取规则
SUB_ITEM_PATTERNS = [
    re.compile(r'包含但不限于[：:]\s*([^，。；]+)'),
    re.compile(r'包括但不限于[：:]\s*([^，。；]+)'),
    re.compile(r'应包含[：:]\s*([^，。；]+)'),
    re.compile(r'应包括[：:]\s*([^，。；]+)'),
    re.compile(r'含[：:]\s*([^，。；]+)'),
]


def _extract_docx_text(file_path: str) -> str:
    """从Word文档提取纯文本（与checker保持一致）"""
    from docx import Document
    text_parts = []
    doc = Document(file_path)
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)
    return '\n'.join(text_parts)


def _extract_docx_tables_count(file_path: str) -> int:
    """统计文档中的表格总数"""
    from docx import Document
    doc = Document(file_path)
    return len(doc.tables)


def _extract_docx_headings(file_path: str) -> List[Dict]:
    """提取文档标题结构"""
    from docx import Document
    doc = Document(file_path)
    headings = []
    for para in doc.paragraphs:
        if not para.style:
            continue
        style_name = str(para.style.name)
        level = None
        if 'Heading' in style_name or 'heading' in style_name:
            match = re.search(r'[Hh]eading\s*(\d+)', style_name)
            if match:
                level = int(match.group(1))
        if level is None and para.text.strip():
            if re.match(r'^[一二三四五六七八九十]+、', para.text.strip()):
                level = 1
            elif re.match(r'^[（(][一二三四五六七八九十]+[）)]', para.text.strip()):
                level = 2
            elif re.match(r'^\d+[.、．]\s*\S', para.text.strip()):
                level = 3
            elif re.match(r'^[（(]\d+[）)]\s*\S', para.text.strip()):
                level = 4
        if level is not None and para.text.strip():
            headings.append({
                'level': level,
                'text': para.text.strip(),
            })
    return headings


class EvaluatorCheck:
    """
    评标专家视角自检器
    从得分角度评估标书质量，而非格式合规角度
    """

    def __init__(self, parse_result: Dict, doc_info: Dict = None,
                 bid_doc_path: str = None):
        """
        Args:
            parse_result: parser.py的输出（招标文件解析结果）
            doc_info: 生成标书的信息 {chapters: [...], tables: [...], project_info: {...}}
            bid_doc_path: 生成标书的docx文件路径
        """
        self.parse_result = parse_result or {}
        self.doc_info = doc_info or {}
        self.bid_doc_path = bid_doc_path
        self.bid_doc_text = ''
        self.bid_doc_paragraphs = []
        self.bid_doc_headings = []
        self.bid_doc_tables_count = 0

        # 从parse_result或doc_info获取project_info
        self.project_info = (
            self.parse_result.get('project_info', {}) or
            self.doc_info.get('project_info', {}) or
            {}
        )

        # 加载标书文本
        if bid_doc_path and os.path.exists(bid_doc_path):
            try:
                self.bid_doc_text = _extract_docx_text(bid_doc_path)
                self._load_paragraphs(bid_doc_path)
                self.bid_doc_headings = _extract_docx_headings(bid_doc_path)
                self.bid_doc_tables_count = _extract_docx_tables_count(bid_doc_path)
            except Exception as e:
                print(f"[EvaluatorCheck] 警告：无法提取标书文本: {e}")

        # 内部结果缓存
        self._coverage_result = None
        self._depth_results = {}
        self._risk_result = None
        self._score_result = None

    def _load_paragraphs(self, path: str):
        """加载段落列表"""
        try:
            from docx import Document
            doc = Document(path)
            self.bid_doc_paragraphs = [
                p.text.strip() for p in doc.paragraphs if p.text.strip()
            ]
        except Exception:
            self.bid_doc_paragraphs = []

    # ── 1. 评分项覆盖率检查 ─────────────────────────────────

    def check_score_item_coverage(self, score_items: List[Dict] = None,
                                  document: str = None) -> Dict:
        """
        检查每个评分项是否有对应内容，含"包含但不限于"子条目逐条验证

        Args:
            score_items: 评分项列表，每个含name/title/score及可选sub_items
            document: 文档文本（可选，默认用self.bid_doc_text）

        Returns:
            {
                'total_items': int,
                'covered_items': int,
                'coverage_rate': float,       # 0.0~1.0
                'uncovered': [...],           # 未覆盖项列表
                'sub_item_results': [...],    # "包含但不限于"子条目验证结果
                'weak_coverage': [...],       # 弱覆盖项（仅部分命中）
            }
        """
        items = score_items or self.parse_result.get('score_items', [])
        doc_text = document or self.bid_doc_text

        if not items:
            return {
                'total_items': 0,
                'covered_items': 0,
                'coverage_rate': 0.0,
                'uncovered': [],
                'sub_item_results': [],
                'weak_coverage': [],
            }

        uncovered = []
        weak_coverage = []
        sub_item_results = []

        for item in items:
            title = item.get('name', item.get('title', ''))
            score = item.get('score', 0)
            if not title:
                continue

            # 主条目匹配
            match_strength = self._evaluate_item_match(title, doc_text)

            if match_strength >= MATCH_STRENGTH['keyword']:
                # 主条目已覆盖
                pass
            elif match_strength >= MATCH_STRENGTH['partial']:
                # 弱覆盖
                weak_coverage.append({
                    'item': title,
                    'score': score,
                    'match_strength': match_strength,
                    'reason': f'评分项「{title}」仅有部分内容命中，响应不够充分',
                })
            else:
                # 未覆盖
                uncovered.append({
                    'item': title,
                    'score': score,
                    'reason': f'评分项「{title}」在标书中未找到对应响应',
                })

            # "包含但不限于"子条目验证
            sub_items = self._extract_sub_items(title, item, doc_text)
            if sub_items:
                sub_item_results.append({
                    'parent_item': title,
                    'parent_score': score,
                    'sub_items': sub_items,
                })

        total = len(items)
        covered = total - len(uncovered) - len(weak_coverage)
        coverage_rate = covered / total if total > 0 else 0.0

        result = {
            'total_items': total,
            'covered_items': covered,
            'coverage_rate': round(coverage_rate, 4),
            'uncovered': uncovered,
            'sub_item_results': sub_item_results,
            'weak_coverage': weak_coverage,
        }
        self._coverage_result = result
        return result

    def _evaluate_item_match(self, item_title: str, doc_text: str) -> float:
        """
        评估评分项在文档中的匹配强度

        Returns:
            匹配强度 0.0~1.0
        """
        if not doc_text or not item_title:
            return MATCH_STRENGTH['none']

        # 1. 标题精确匹配
        heading_texts = [h['text'] for h in self.bid_doc_headings]
        for ht in heading_texts:
            if item_title in ht or ht in item_title:
                return MATCH_STRENGTH['exact']

        # 2. 关键词匹配（提取2-6字关键词，排除通用词）
        generic_words = {
            '措施', '方案', '管理', '保证', '施工', '工程', '项目',
            '总体', '整体', '总体概', '编制', '说明',
        }
        keywords = [
            kw for kw in re.findall(r'[\u4e00-\u9fff]{2,6}', item_title)
            if kw not in generic_words and len(kw) >= 2
        ]

        if keywords:
            matched = sum(1 for kw in keywords if kw in doc_text)
            ratio = matched / len(keywords) if keywords else 0
            if ratio >= 0.8:
                return MATCH_STRENGTH['keyword']
            elif ratio >= 0.5:
                return MATCH_STRENGTH['content']
            elif ratio > 0:
                return MATCH_STRENGTH['partial']

        # 3. 兜底：原文搜索
        if item_title in doc_text:
            return MATCH_STRENGTH['content']

        return MATCH_STRENGTH['none']

    def _extract_sub_items(self, parent_title: str, item: Dict,
                           doc_text: str) -> List[Dict]:
        """
        提取并验证"包含但不限于"子条目

        Returns:
            [{'sub_item': str, 'covered': bool, 'detail': str}, ...]
        """
        sub_results = []

        # 从评分项自身提取子条目
        explicit_sub_items = item.get('sub_items', [])
        if explicit_sub_items:
            for sub in explicit_sub_items:
                sub_text = sub if isinstance(sub, str) else sub.get('name', '')
                covered = sub_text in doc_text if doc_text else False
                sub_results.append({
                    'sub_item': sub_text,
                    'covered': covered,
                    'detail': '已响应' if covered else '未在标书中找到对应内容',
                })

        # 从评分项描述中提取"包含但不限于"子条目
        description = item.get('description', item.get('content', ''))
        if description:
            for pattern in SUB_ITEM_PATTERNS:
                matches = pattern.findall(description)
                for match_text in matches:
                    # 拆分多个子条目（顿号/逗号分隔）
                    parts = re.split(r'[、，,；;]', match_text)
                    for part in parts:
                        part = part.strip()
                        if len(part) >= 2:
                            covered = part in doc_text if doc_text else False
                            # 避免重复
                            if not any(s['sub_item'] == part for s in sub_results):
                                sub_results.append({
                                    'sub_item': part,
                                    'covered': covered,
                                    'detail': '已响应' if covered else '未在标书中找到对应内容',
                                })

        return sub_results

    # ── 2. 内容深度评估 ─────────────────────────────────────

    def evaluate_depth(self, chapter_name: str, content: str = None,
                       detail_level: int = 3) -> Dict:
        """
        评估章节内容深度

        标准：
        - 表格数量：每章至少X个表格
        - 国标引用：每章至少引用Y个国标/行标
        - 工序描述：每章至少描述Z个工序步骤

        Args:
            chapter_name: 章节名称
            content: 章节内容（可选，默认在文档中搜索对应章节）
            detail_level: 详细等级(1-5)，决定评估标准

        Returns:
            {
                'chapter': str,
                'depth_level': int,          # 1-5
                'table_count': int,
                'standard_count': int,
                'process_count': int,
                'standard_target': Dict,     # 对应等级的标准
                'gaps': [...],               # 具体不足列表
            }
        """
        # 缓存
        cache_key = f"{chapter_name}_{detail_level}"
        if cache_key in self._depth_results:
            return self._depth_results[cache_key]

        # 获取评估标准
        level = max(1, min(5, detail_level))
        std = DEPTH_STANDARDS[level]

        # 获取章节内容
        chapter_content = content or self._find_chapter_content(chapter_name)

        # 统计表格数（在该章节范围内）
        table_count = self._count_chapter_tables(chapter_name)

        # 统计国标引用数
        standard_count = self._count_chapter_standards(chapter_content)

        # 统计工序描述数
        process_count = self._count_processes(chapter_content)

        # 计算深度等级
        depth_score = 0
        gaps = []

        if table_count >= std['min_tables']:
            depth_score += 1
        else:
            gaps.append({
                'dimension': '表格',
                'current': table_count,
                'target': std['min_tables'],
                'gap': std['min_tables'] - table_count,
                'suggestion': f'章节「{chapter_name}」缺少表格，'
                              f'当前{table_count}张，建议补充至{std["min_tables"]}张以上',
            })

        if standard_count >= std['min_standards']:
            depth_score += 1
        else:
            gaps.append({
                'dimension': '国标引用',
                'current': standard_count,
                'target': std['min_standards'],
                'gap': std['min_standards'] - standard_count,
                'suggestion': f'章节「{chapter_name}」国标引用不足，'
                              f'当前{standard_count}个，建议补充至{std["min_standards"]}个以上',
            })

        if process_count >= std['min_processes']:
            depth_score += 1
        else:
            gaps.append({
                'dimension': '工序描述',
                'current': process_count,
                'target': std['min_processes'],
                'gap': std['min_processes'] - process_count,
                'suggestion': f'章节「{chapter_name}」工序描述不足，'
                              f'当前{process_count}个，建议补充至{std["min_processes"]}个以上',
            })

        # 补充维度：内容长度
        content_len = len(chapter_content) if chapter_content else 0
        if content_len >= 2000:
            depth_score += 1
        elif content_len >= 1000:
            depth_score += 0.5
        else:
            gaps.append({
                'dimension': '内容长度',
                'current': content_len,
                'target': 2000,
                'gap': 2000 - content_len,
                'suggestion': f'章节「{chapter_name}」内容偏短，'
                              f'当前约{content_len}字，建议扩充至2000字以上',
            })

        # 补充维度：专业术语密度
        term_density = self._calc_term_density(chapter_content)
        if term_density >= 0.02:
            depth_score += 1
        elif term_density >= 0.01:
            depth_score += 0.5
        else:
            gaps.append({
                'dimension': '专业术语',
                'current': f'{term_density:.4f}',
                'target': '0.02',
                'gap': round(0.02 - term_density, 4),
                'suggestion': f'章节「{chapter_name}」专业术语密度偏低，'
                              f'建议增加规范编号、专业名词等术语',
            })

        # 映射为1-5级（depth_score范围0-5，直接映射）
        depth_level = max(1, min(5, round(depth_score)))

        result = {
            'chapter': chapter_name,
            'depth_level': depth_level,
            'table_count': table_count,
            'standard_count': standard_count,
            'process_count': process_count,
            'content_length': content_len,
            'term_density': round(term_density, 4),
            'standard_target': std,
            'gaps': gaps,
        }

        self._depth_results[cache_key] = result
        return result

    def _find_chapter_content(self, chapter_name: str) -> str:
        """在文档中查找指定章节的内容"""
        if not self.bid_doc_paragraphs:
            return ''

        # 查找章节标题位置
        start_idx = -1
        end_idx = len(self.bid_doc_paragraphs)

        for i, para in enumerate(self.bid_doc_paragraphs):
            # 匹配章节标题
            if self._is_chapter_title(para) and self._title_matches(para, chapter_name):
                start_idx = i
                continue
            # 找到下一个同级或更高级标题则结束
            if start_idx >= 0 and self._is_chapter_title(para):
                # 简单策略：遇到下一个一级标题即结束
                if re.match(r'^[一二三四五六七八九十]+、', para):
                    end_idx = i
                    break

        if start_idx < 0:
            # 兜底：全文搜索关键词附近内容
            for i, para in enumerate(self.bid_doc_paragraphs):
                if chapter_name in para:
                    start_idx = max(0, i - 1)
                    end_idx = min(len(self.bid_doc_paragraphs), i + 30)
                    break

        if start_idx < 0:
            return ''

        return '\n'.join(self.bid_doc_paragraphs[start_idx:end_idx])

    def _is_chapter_title(self, text: str) -> bool:
        """判断文本是否是章节标题"""
        patterns = [
            r'^[一二三四五六七八九十]+、',
            r'^[（(][一二三四五六七八九十]+[）)]',
            r'^\d+[.、．]\s*\S',
        ]
        return any(re.match(p, text) for p in patterns)

    def _title_matches(self, title_text: str, target_name: str) -> bool:
        """判断标题是否匹配目标章节名"""
        # 去掉编号前缀
        clean = re.sub(r'^[一二三四五六七八九十]+、\s*', '', title_text)
        clean = re.sub(r'^[（(][一二三四五六七八九十]+[）)]\s*', '', clean)
        clean = re.sub(r'^\d+[.、．]\s*', '', clean)
        clean = clean.strip()

        if not clean:
            return False

        # 关键词匹配
        target_kws = [kw for kw in re.findall(r'[\u4e00-\u9fff]{2,}', target_name)
                      if kw not in {'措施', '方案', '管理', '保证', '施工', '工程'}]
        if not target_kws:
            return target_name in clean

        matched = sum(1 for kw in target_kws if kw in clean)
        return matched >= len(target_kws) * 0.5

    def _count_chapter_tables(self, chapter_name: str) -> int:
        """统计指定章节范围内的表格数"""
        # 简化策略：按标题范围估算
        # 从doc_info中获取表格信息
        tables = self.doc_info.get('tables', [])
        chapter_tables = [
            t for t in tables
            if chapter_name in (t.get('title', '') + t.get('chapter', ''))
        ]
        return max(len(chapter_tables), 0)

    def _count_chapter_standards(self, content: str) -> int:
        """统计章节中的国标/行标引用数"""
        if not content:
            return 0

        standard_patterns = [
            re.compile(r'(GB/?(?:T)?\s*\d+(?:\.\d+)?(?:-[\d]{2,4})?)'),
            re.compile(r'(JGJ/?(?:T)?\s*\d+(?:\.\d+)?(?:-[\d]{2,4})?)'),
            re.compile(r'(CJJ/?(?:T)?\s*\d+(?:\.\d+)?(?:-[\d]{2,4})?)'),
            re.compile(r'(SL/?(?:T)?\s*\d+(?:\.\d+)?(?:-[\d]{2,4})?)'),
            re.compile(r'(DL/?(?:T)?\s*\d+(?:\.\d+)?(?:-[\d]{2,4})?)'),
            re.compile(r'(HJ/?(?:T)?\s*\d+(?:\.\d+)?(?:-[\d]{2,4})?)'),
            re.compile(r'(CECS\s*\d+(?::\d+)?(?:-[\d]{2,4})?)'),
            re.compile(r'(DB\d{2}/?(?:T)?\s*\d+(?:\.\d+)?(?:-[\d]{2,4})?)'),
        ]

        found = set()
        for pattern in standard_patterns:
            for match in pattern.finditer(content):
                std = match.group(1).strip()
                std = re.sub(r'\s+', ' ', std)
                if std:
                    found.add(std)
        return len(found)

    def _count_processes(self, content: str) -> int:
        """统计章节中的工序描述数"""
        if not content:
            return 0

        # 工序描述常见模式
        process_patterns = [
            r'第[一二三四五六七八九十\d]+[步道工序]',
            r'[（(]\d+[）)]\s*[^\n]{5,30}[：:]',  # (1) xxx步骤：
            r'[一二三四五六七八九十]+[、.]\s*[\u4e00-\u9fff]{2,10}[：:]',  # 一、xxx步骤：
            r'步骤[一二三四五六七八九十\d]+',
            r'流程[为是]：',
            r'施工[工序步骤流程]',
        ]

        count = 0
        for pattern in process_patterns:
            matches = re.findall(pattern, content)
            count += len(matches)

        # 去重（避免同一工序被多种模式重复计数）
        return min(count, len(content) // 50)

    def _calc_term_density(self, content: str) -> float:
        """计算专业术语密度（术语数/总字数）"""
        if not content:
            return 0.0

        # 专业术语模式：标准编号、专业名词缩写
        term_patterns = [
            r'GB\s*\d+', r'JGJ\s*\d+', r'CJJ\s*\d+',
            r'DB\d{2}\s*\d+', r'HJ\s*\d+', r'SL\s*\d+',
            r'CECS\s*\d+',
            r'QC\b', r'IQC\b', r'OQC\b', r'IPQC\b',
            r'PDCA\b', r'BIM\b', r'ERP\b',
            r'三检制度', r'三级交底', r'三级教育',
            r'旁站监理', r'平行检验', r'见证取样',
            r'隐蔽工程', r'检验批', r'分项工程', r'分部工程',
            r'单位工程', r'分部分项',
        ]

        term_count = 0
        for pattern in term_patterns:
            matches = re.findall(pattern, content)
            term_count += len(matches)

        total_chars = len(content.replace(' ', '').replace('\n', ''))
        return term_count / total_chars if total_chars > 0 else 0.0

    # ── 3. 废标风险扫描 ─────────────────────────────────────

    def scan_disqualification_risks(self, document: str = None) -> Dict:
        """
        扫描废标条款，从评标专家视角识别废标风险

        Args:
            document: 文档文本（可选，默认用self.bid_doc_text）

        Returns:
            {
                'total_risks': int,
                'high_risks': [...],
                'medium_risks': [...],
                'low_risks': [...],
                'risk_summary': str,
            }
        """
        doc_text = document or self.bid_doc_text
        if not doc_text:
            return {
                'total_risks': 0,
                'high_risks': [],
                'medium_risks': [],
                'low_risks': [],
                'risk_summary': '无标书文本，无法扫描废标风险',
            }

        high_risks = []
        medium_risks = []
        low_risks = []

        for pattern in DISQUALIFICATION_RISK_PATTERNS:
            pid = pattern['id']
            name = pattern['name']
            keywords = pattern['keywords']
            risk_level = pattern['risk_level']
            description = pattern['description']
            repair_hint = pattern['repair_hint']

            # 在文档中搜索关键词
            matched_keywords = [kw for kw in keywords if kw in doc_text]

            if not matched_keywords:
                # 关键词全部未找到 → 存在风险
                risk_item = {
                    'id': pid,
                    'name': name,
                    'risk_level': risk_level,
                    'description': description,
                    'matched_keywords': [],
                    'missing_keywords': keywords,
                    'repair_hint': repair_hint,
                }

                if risk_level == 'high':
                    high_risks.append(risk_item)
                elif risk_level == 'medium':
                    medium_risks.append(risk_item)
                else:
                    low_risks.append(risk_item)
            elif len(matched_keywords) < len(keywords):
                # 部分关键词命中 → 可能存在风险
                missing = [kw for kw in keywords if kw not in matched_keywords]
                risk_item = {
                    'id': pid,
                    'name': name,
                    'risk_level': 'low' if risk_level == 'medium' else risk_level,
                    'description': f'{description}（部分关键词未命中）',
                    'matched_keywords': matched_keywords,
                    'missing_keywords': missing,
                    'repair_hint': repair_hint,
                }

                # 高风险的降为中风险
                if risk_level == 'high':
                    risk_item['risk_level'] = 'medium'
                    medium_risks.append(risk_item)
                elif risk_level == 'medium':
                    low_risks.append(risk_item)
                # low风险部分命中则不报告

        total = len(high_risks) + len(medium_risks) + len(low_risks)

        # 构建摘要
        if high_risks:
            risk_summary = (
                f'发现{len(high_risks)}项高风险废标风险：'
                f'{"；".join(r["name"] for r in high_risks)}。'
                f'建议优先处理！'
            )
        elif medium_risks:
            risk_summary = (
                f'未发现高风险废标风险，但有{len(medium_risks)}项中风险：'
                f'{"；".join(r["name"] for r in medium_risks)}。'
                f'建议关注并补充。'
            )
        elif low_risks:
            risk_summary = (
                f'废标风险较低，有{len(low_risks)}项低风险：'
                f'{"；".join(r["name"] for r in low_risks)}。'
                f'可酌情优化。'
            )
        else:
            risk_summary = '未发现废标风险，标书从专家视角审查合规。'

        result = {
            'total_risks': total,
            'high_risks': high_risks,
            'medium_risks': medium_risks,
            'low_risks': low_risks,
            'risk_summary': risk_summary,
        }
        self._risk_result = result
        return result

    # ── 4. 得分预测 ─────────────────────────────────────────

    def predict_score(self, score_items: List[Dict] = None,
                      document: str = None) -> Dict:
        """
        基于评分项和文档内容预测得分

        每个评分项：满分/预测分/失分原因
        返回总分和百分比

        Args:
            score_items: 评分项列表
            document: 文档文本（可选）

        Returns:
            {
                'total_full_score': float,
                'total_predicted_score': float,
                'score_percentage': float,
                'item_details': [
                    {
                        'item': str,
                        'full_score': float,
                        'predicted_score': float,
                        'deduction_reasons': [...],
                        'match_strength': float,
                    },
                    ...
                ],
                'major_deductions': [...],   # 主要失分项
            }
        """
        items = score_items or self.parse_result.get('score_items', [])
        doc_text = document or self.bid_doc_text

        if not items:
            return {
                'total_full_score': 0,
                'total_predicted_score': 0,
                'score_percentage': 0.0,
                'item_details': [],
                'major_deductions': [],
            }

        item_details = []
        total_full = 0
        total_predicted = 0

        for item in items:
            title = item.get('name', item.get('title', ''))
            full_score = float(item.get('score', 0))
            if not title:
                continue

            total_full += full_score

            # 评估匹配强度
            match_strength = self._evaluate_item_match(title, doc_text)

            # 评估内容深度（对匹配到的项）
            depth_result = None
            if match_strength >= MATCH_STRENGTH['partial']:
                depth_result = self.evaluate_depth(title)

            # 计算预测得分
            deduction_reasons = []
            predicted = full_score

            if match_strength < MATCH_STRENGTH['partial']:
                # 未响应：得0分
                predicted = 0
                deduction_reasons.append({
                    'reason': f'评分项「{title}」未在标书中找到对应响应',
                    'deduction': full_score,
                })
            elif match_strength < MATCH_STRENGTH['content']:
                # 部分命中：扣50%-70%
                deduction_ratio = 0.6
                predicted = full_score * (1 - deduction_ratio)
                deduction_reasons.append({
                    'reason': f'评分项「{title}」响应不充分，仅有部分内容命中',
                    'deduction': round(full_score * deduction_ratio, 2),
                })
            elif match_strength < MATCH_STRENGTH['keyword']:
                # 内容命中但不够突出：扣20%-40%
                deduction_ratio = 0.3
                predicted = full_score * (1 - deduction_ratio)
                deduction_reasons.append({
                    'reason': f'评分项「{title}」在正文中命中但缺少独立章节/标题',
                    'deduction': round(full_score * deduction_ratio, 2),
                })
            else:
                # 关键词及以上匹配，检查深度
                if depth_result and depth_result['gaps']:
                    gap_deduction = 0
                    for gap in depth_result['gaps']:
                        dimension = gap.get('dimension', '')
                        if dimension == '表格':
                            gap_deduction += 0.05 * full_score
                        elif dimension == '国标引用':
                            gap_deduction += 0.05 * full_score
                        elif dimension == '工序描述':
                            gap_deduction += 0.03 * full_score
                        elif dimension == '内容长度':
                            gap_deduction += 0.03 * full_score
                        elif dimension == '专业术语':
                            gap_deduction += 0.02 * full_score

                    gap_deduction = min(gap_deduction, full_score * 0.3)
                    predicted = full_score - gap_deduction

                    if gap_deduction > 0:
                        gap_details = '; '.join(
                            f'{g["dimension"]}不足(差{g.get("gap", "?")})'
                            for g in depth_result['gaps']
                        )
                        deduction_reasons.append({
                            'reason': f'评分项「{title}」内容深度不足：{gap_details}',
                            'deduction': round(gap_deduction, 2),
                        })

            predicted = max(0, round(predicted, 2))
            total_predicted += predicted

            item_details.append({
                'item': title,
                'full_score': full_score,
                'predicted_score': predicted,
                'deduction_reasons': deduction_reasons,
                'match_strength': match_strength,
            })

        # 主要失分项（失分最多的前5项）
        major_deductions = sorted(
            item_details,
            key=lambda x: x['full_score'] - x['predicted_score'],
            reverse=True
        )[:5]
        major_deductions = [
            {
                'item': d['item'],
                'lost_score': round(d['full_score'] - d['predicted_score'], 2),
                'reasons': d['deduction_reasons'],
            }
            for d in major_deductions
            if d['full_score'] - d['predicted_score'] > 0
        ]

        score_percentage = (total_predicted / total_full * 100) if total_full > 0 else 0

        result = {
            'total_full_score': total_full,
            'total_predicted_score': round(total_predicted, 2),
            'score_percentage': round(score_percentage, 2),
            'item_details': item_details,
            'major_deductions': major_deductions,
        }
        self._score_result = result
        return result

    # ── 5. 改进建议生成 ─────────────────────────────────────

    def generate_improvement_suggestions(self, check_result: Dict = None) -> List[Dict]:
        """
        生成具体改进建议

        每条建议：章节+问题+建议内容+优先级
        按优先级排序：critical > high > medium > low

        Args:
            check_result: 可选，自定义检查结果；默认使用内部已执行的结果

        Returns:
            [
                {
                    'id': str,
                    'chapter': str,
                    'problem': str,
                    'suggestion': str,
                    'priority': str,   # critical/high/medium/low
                    'source': str,     # coverage/depth/risk/score
                },
                ...
            ]
        """
        # 确保已执行各检查
        if self._coverage_result is None:
            self.check_score_item_coverage()
        if self._risk_result is None:
            self.scan_disqualification_risks()
        if self._score_result is None:
            self.predict_score()

        suggestions = []

        # ── 来自覆盖率检查的建议 ──
        coverage = self._coverage_result
        if coverage:
            # 未覆盖评分项
            for item in coverage.get('uncovered', []):
                suggestions.append({
                    'id': f'EC-COV-{len(suggestions)+1:03d}',
                    'chapter': item.get('item', ''),
                    'problem': f'评分项「{item.get("item", "")}」未在标书中响应（{item.get("score", "?")}分）',
                    'suggestion': (
                        f'建议增加「{item.get("item", "")}」的专项章节或段落，'
                        f'明确响应招标文件评分要求，争取获得该项{item.get("score", "?")}分'
                    ),
                    'priority': 'critical',
                    'source': 'coverage',
                })

            # 弱覆盖评分项
            for item in coverage.get('weak_coverage', []):
                suggestions.append({
                    'id': f'EC-COV-{len(suggestions)+1:03d}',
                    'chapter': item.get('item', ''),
                    'problem': f'评分项「{item.get("item", "")}」响应不充分',
                    'suggestion': (
                        f'建议强化「{item.get("item", "")}」的响应内容，'
                        f'增加具体措施、数据和标准引用，提高响应充分度'
                    ),
                    'priority': 'high',
                    'source': 'coverage',
                })

            # "包含但不限于"子条目未覆盖
            for sub_result in coverage.get('sub_item_results', []):
                for sub in sub_result.get('sub_items', []):
                    if not sub.get('covered', False):
                        suggestions.append({
                            'id': f'EC-COV-{len(suggestions)+1:03d}',
                            'chapter': sub_result.get('parent_item', ''),
                            'problem': (
                                f'评分项「{sub_result.get("parent_item", "")}」'
                                f'的子条目「{sub.get("sub_item", "")}」未响应'
                            ),
                            'suggestion': (
                                f'建议在「{sub_result.get("parent_item", "")}」章节中'
                                f'补充「{sub.get("sub_item", "")}」相关内容'
                            ),
                            'priority': 'high',
                            'source': 'coverage',
                        })

        # ── 来自深度评估的建议 ──
        for cache_key, depth_result in self._depth_results.items():
            for gap in depth_result.get('gaps', []):
                priority = 'medium'
                dimension = gap.get('dimension', '')
                if dimension in ('表格', '国标引用'):
                    priority = 'high'
                elif dimension in ('工序描述', '内容长度'):
                    priority = 'medium'
                else:
                    priority = 'low'

                suggestions.append({
                    'id': f'EC-DEP-{len(suggestions)+1:03d}',
                    'chapter': depth_result.get('chapter', ''),
                    'problem': gap.get('suggestion', '内容深度不足'),
                    'suggestion': self._generate_depth_suggestion(
                        depth_result.get('chapter', ''), gap
                    ),
                    'priority': priority,
                    'source': 'depth',
                })

        # ── 来自废标风险扫描的建议 ──
        risk = self._risk_result
        if risk:
            for r in risk.get('high_risks', []):
                suggestions.append({
                    'id': f'EC-RSK-{len(suggestions)+1:03d}',
                    'chapter': '全文',
                    'problem': f'废标高风险：{r.get("name", "")}',
                    'suggestion': r.get('repair_hint', '需补充相关内容'),
                    'priority': 'critical',
                    'source': 'risk',
                })
            for r in risk.get('medium_risks', []):
                suggestions.append({
                    'id': f'EC-RSK-{len(suggestions)+1:03d}',
                    'chapter': '全文',
                    'problem': f'废标中风险：{r.get("name", "")}',
                    'suggestion': r.get('repair_hint', '建议补充相关内容'),
                    'priority': 'high',
                    'source': 'risk',
                })
            for r in risk.get('low_risks', []):
                suggestions.append({
                    'id': f'EC-RSK-{len(suggestions)+1:03d}',
                    'chapter': '全文',
                    'problem': f'废标低风险：{r.get("name", "")}',
                    'suggestion': r.get('repair_hint', '可酌情优化'),
                    'priority': 'medium',
                    'source': 'risk',
                })

        # ── 来自得分预测的建议 ──
        score = self._score_result
        if score:
            for deduction in score.get('major_deductions', []):
                suggestions.append({
                    'id': f'EC-SCR-{len(suggestions)+1:03d}',
                    'chapter': deduction.get('item', ''),
                    'problem': (
                        f'评分项「{deduction.get("item", "")}」'
                        f'预计失分{deduction.get("lost_score", 0)}分'
                    ),
                    'suggestion': '; '.join(
                        r.get('reason', '') for r in deduction.get('reasons', [])
                    ),
                    'priority': 'high' if deduction.get('lost_score', 0) >= 5 else 'medium',
                    'source': 'score',
                })

        # 按优先级排序
        priority_order = {'critical': 0, 'high': 1, 'medium': 2, 'low': 3}
        suggestions.sort(key=lambda x: priority_order.get(x.get('priority', 'low'), 99))

        return suggestions

    def _generate_depth_suggestion(self, chapter: str, gap: Dict) -> str:
        """根据深度不足项生成具体建议"""
        dimension = gap.get('dimension', '')
        current = gap.get('current', 0)
        target = gap.get('target', 0)

        suggestion_map = {
            '表格': (
                f'建议在「{chapter}」中补充{gap.get("gap", 0)}张专业表格，'
                f'如：施工参数表、检验标准表、设备配置表、质量检查表等'
            ),
            '国标引用': (
                f'建议在「{chapter}」中补充{gap.get("gap", 0)}个国标/行标引用，'
                f'引用相关GB/JGJ等标准编号，增强专业性和说服力'
            ),
            '工序描述': (
                f'建议在「{chapter}」中补充{gap.get("gap", 0)}个工序描述，'
                f'使用分步编号(1)(2)(3)...详细说明施工工艺步骤'
            ),
            '内容长度': (
                f'建议扩充「{chapter}」内容，当前约{current}字，目标2000字以上。'
                f'可从管理目标、组织架构、实施方案、资源保障、检查改进等方面补充'
            ),
            '专业术语': (
                f'建议在「{chapter}」中增加专业术语引用，'
                f'如标准编号(GB/JGJ)、专业名词(QC/PDCA/BIM)、'
                f'管理制度(三检制度/三级交底/旁站监理)等'
            ),
        }

        return suggestion_map.get(dimension, f'建议补充{dimension}相关内容')

    # ── 6. 与repair模块集成接口 ─────────────────────────────

    def to_repair_format(self, suggestions: List[Dict] = None) -> List[Dict]:
        """
        将改进建议转换为repair模块兼容格式

        repair模块期望的格式：
        {
            'issue_id': str,
            'issue_name': str,
            'issue_detail': str,
            'repair_prompt': str,
            'target_location': str,
            'priority': str,
            'has_template': bool,
            'template_hint': str,
            'auto_content': str,    # 可自动修复的内容
        }
        """
        items = suggestions or self.generate_improvement_suggestions()
        repair_items = []

        for s in items:
            # 根据问题来源匹配repair模板
            auto_content = self._generate_auto_repair_content(s)

            repair_items.append({
                'issue_id': s.get('id', ''),
                'issue_name': s.get('problem', '')[:50],
                'issue_detail': s.get('problem', ''),
                'repair_prompt': s.get('suggestion', ''),
                'target_location': s.get('chapter', ''),
                'priority': s.get('priority', 'medium'),
                'has_template': bool(auto_content),
                'template_hint': auto_content[:200] if auto_content else '',
                'auto_content': auto_content,
                'source': s.get('source', ''),
            })

        return repair_items

    def _generate_auto_repair_content(self, suggestion: Dict) -> str:
        """根据建议类型生成自动修复内容"""
        source = suggestion.get('source', '')
        chapter = suggestion.get('chapter', '')
        problem = suggestion.get('problem', '')

        if source == 'coverage':
            return (
                f'{chapter}\n\n'
                f'我方针对本工程{chapter}，制定如下措施：\n\n'
                f'一、总体目标\n'
                f'我方承诺严格按照招标文件要求，做好{chapter}工作，'
                f'确保各项指标达到或超过招标要求。\n\n'
                f'二、组织保障\n'
                f'成立专项工作小组，由项目负责人牵头，明确岗位职责，'
                f'层层落实责任，确保{chapter}各项工作有序推进。\n\n'
                f'三、具体措施\n'
                f'（1）建立完善的管理制度，严格执行相关规范标准；\n'
                f'（2）加强过程控制和动态管理，定期开展专项检查；\n'
                f'（3）配备充足的资源保障，确保各项措施有效落实；\n'
                f'（4）做好记录台账，确保全过程可追溯、可核查。\n\n'
                f'四、保障承诺\n'
                f'我方保证{chapter}工作满足招标文件要求，'
                f'如因我方原因未能达标，愿意承担相应责任。'
            )

        elif source == 'risk':
            if '承诺书' in problem:
                return (
                    '投标承诺书\n\n'
                    '我方郑重承诺：\n'
                    '一、我方已认真阅读并充分理解招标文件的全部内容，'
                    '自愿参加本工程的投标活动。\n'
                    '二、我方保证投标文件中所有内容及提供的一切资料真实、准确、完整，'
                    '如有虚假，愿意承担相应的法律责任。\n'
                    '三、我方承诺如中标，将严格按照招标文件和合同约定履行义务，'
                    '确保工程质量和安全。\n'
                    '四、我方承诺在投标有效期内不撤销投标文件，'
                    '如中标后无故放弃，愿意承担相应的法律责任和经济损失。'
                )
            elif '资质' in problem:
                return (
                    '我方具有履行本合同所需的资质条件和能力，'
                    '相关资质证书均在有效期内。'
                    '资质证书复印件随投标文件一并递交，请予审核。'
                )
            elif '人员' in problem or '无证' in problem:
                return (
                    '我方拟派项目经理持有注册建造师执业资格证书，'
                    '证书在有效期内。安全员持有安全生产考核合格证书。'
                    '关键岗位人员均持证上岗，相关证书复印件随投标文件递交。'
                )
            elif '安全投入' in problem:
                return (
                    '我方承诺本工程安全生产投入不低于工程造价的2%，'
                    '安全防护设施配备齐全，安全经费专款专用。'
                    '文明施工费按招标文件规定标准计取并足额使用。'
                )
            return ''

        elif source == 'depth':
            if '表格' in problem:
                return (
                    f'| 序号 | 项目 | 标准/参数 | 备注 |\n'
                    f'|------|------|-----------|------|\n'
                    f'| 1 | 主控项目 | 符合设计文件及 GB 50300 验收要求 | |\n'
                    f'| 2 | 一般项目 | 偏差在规范允许范围内 | |\n'
                    f'| 3 | 检测参数 | 按规范要求的频次与方法检测 | |'
                )
            elif '国标' in problem or '标准' in problem:
                return (
                    f'我方将严格按照以下标准规范执行{chapter}相关工作：\n'
                    f'GB 50300-2013《建筑工程施工质量验收统一标准》；'
                    f'JGJ 59-2011《建筑施工安全检查标准》；'
                    f'GB/T 50502-2009《建筑施工组织设计规范》。'
                )
            elif '工序' in problem:
                return (
                    f'{chapter}施工工序：\n'
                    f'（1）施工准备：熟悉图纸、技术交底、材料检验；\n'
                    f'（2）测量放线：根据控制桩进行定位放样；\n'
                    f'（3）工序施工：按规范要求分层分段施工；\n'
                    f'（4）质量检查：自检、互检、交接检三检制度；\n'
                    f'（5）验收记录：及时填写施工记录和检验批验收记录。'
                )
            return ''

        return ''

    # ── P2: 模拟评审得分表 ──────────────────────────────────────
    def _build_mock_review(self, score_prediction: Dict, suggestions: List[Dict]) -> Dict:
        """基于得分预测与改进建议，构建逐条模拟评审得分表。

        Returns:
            {
              'items': [{'score_item','full_score','predicted_score',
                         'deduction_reasons','improvement'}, ...],
              'score_prediction': ...,
              'markdown': str,
            }
        """
        # 改进建议按评分项/章节名建索引
        sug_map: Dict[str, Dict] = {}
        for s in (suggestions or []):
            key = s.get('item') or s.get('chapter') or ''
            if key:
                sug_map.setdefault(key, s)

        items: List[Dict] = []
        for d in score_prediction.get('item_details', []):
            name = d.get('item', '')
            sug = sug_map.get(name, {})
            reasons = d.get('deduction_reasons', [])
            reason_text = '；'.join(str(r.get('reason', '')) for r in reasons) if reasons else '无'
            improvement = (
                sug.get('suggestion')
                or sug.get('problem')
                or sug.get('repair_hint')
                or '无'
            )
            items.append({
                'score_item': name,
                'full_score': round(float(d.get('full_score', 0) or 0), 1),
                'predicted_score': round(float(d.get('predicted_score', 0) or 0), 1),
                'deduction_reasons': reason_text,
                'improvement': improvement,
            })

        markdown = self._render_mock_review_markdown(items, score_prediction)
        return {
            'items': items,
            'score_prediction': score_prediction,
            'markdown': markdown,
        }

    @staticmethod
    def _render_mock_review_markdown(items: List[Dict], score_prediction: Dict) -> str:
        """渲染模拟评审得分表为 Markdown。"""
        if not items:
            return '## 模拟评审得分表\n\n（无评分项可评估）'
        lines = [
            '## 模拟评审得分表',
            '',
            '| 评分项 | 满分 | 预估得分 | 扣分原因 | 改进建议 |',
            '|--------|------|----------|----------|----------|',
        ]
        for it in items:
            lines.append(
                f"| {it['score_item']} | {it['full_score']} | {it['predicted_score']} "
                f"| {it['deduction_reasons']} | {it['improvement']} |"
            )
        total_full = score_prediction.get('total_full_score', 0)
        total_pred = score_prediction.get('total_predicted_score', 0)
        pct = score_prediction.get('score_percentage', 0)
        lines.append('')
        lines.append(
            f'**综合预估得分：{round(float(total_pred), 1)} / {round(float(total_full), 1)}'
            f'（得分率 {round(float(pct), 1)}%）**'
        )
        return '\n'.join(lines)

    @staticmethod
    def _render_originality_markdown(o: Dict) -> str:
        """渲染文档原创度自检章节为 Markdown。"""
        if not o:
            return ''
        lines = [
            '## 文档原创度自检（v7.6 新增）',
            '',
            f'> 对标钛投标「标书查重」能力 —— 本引擎定位为「自查不雷同」：'
            f'扫描全文内部的重复 / 高度相似段落，让字面重复率可量化、可核查，'
            f'便于针对性改写以降查重合规风险。',
            '',
            f'- **原创度评分**：`{o.get("score", 0)}` / 100',
            f'- **评级**：{o.get("grade", "—")}',
            f'- **参与比较段落**：{o.get("total_paragraphs", 0)} 段',
            f'- **重复 / 高度相似段落**：{o.get("repeated_paragraphs", 0)} 段',
            f'- **相似度阈值**：{o.get("threshold", 0.82)}',
        ]
        pairs = o.get('repeated_pairs') or []
        if pairs:
            lines.append('')
            lines.append('**重复片段定位（建议改写）：**')
            for p in pairs[:10]:
                lines.append(
                    f"- 第 {p['index_a']+1} 段 ≈ 第 {p['index_b']+1} 段 "
                    f"（相似度 {p['similarity']}）：「{p['text_a']}…」"
                )
        else:
            lines.append('')
            lines.append('✅ 未发现明显重复段落。')
        return '\n'.join(lines)

    # v7.28: 技术标竞争力维度覆盖自检——扫描正文，报告本标书已落地的
    # 竞品对标型加分维度（v7.16~v7.27 逐轮对比 WPS/喜鹊/钛投标/红点智标补强）。
    # 这些维度多为市面模板/竞品空白或泛写处，落地即构成相对竞品的差异化优势。
    _COMPETE_DIMS = [
        ('BIM深度应用', ['碰撞检查', '4D进度', '数字化交付', '管线综合'],
         '对标 WPS/钛投标"仅泛写应用BIM技术"'),
        ('绿色施工/双碳', ['碳排放', '建筑垃圾资源化', 'PM10', '四节一环保', '节能率', '节水率'],
         '对标市面"绿色施工标配话术、仅泛写节约资源"'),
        ('质量通病防治', ['通病防治', '塞缝', '空鼓率', '压实度', '附加层'],
         '对标喜鹊"质量章节内容空洞、套话占比高"'),
        ('危大工程管控', ['专家论证', '专项方案', '危大'],
         '对标竞品"仅泛写编制专项方案、不区分危大类型"'),
        ('创优奖项策划', ['创优', '鲁班', '国优', '质量奖'],
         '对标竞品"仅泛写确保合格、无奖项目标"'),
        ('智慧工地监管', ['实名制', '塔吊', '扬尘在线', 'AI摄像头'],
         '对标2026新规"智慧工地投入单列≥建安费2.1%否则扣3.8-5.2分"'),
        ('成品保护/保修', ['成品保护', '专职保护'],
         '对标许昌招标"成品保护缺项即0分"'),
        ('测量与试验检测', ['见证取样', 'CMA', '控制网', '标养试块'],
         '对标竞品空白具体做法（控制网复核/见证取样/计量检定）'),
        ('应急预案与演练', ['应急指挥中心', '专项预案', '演练', '防汛', '消防'],
         '对标竞品"网上抄一遍、98%雷同"的空预案（新华社批评"复制粘贴的防汛预案防不住风雨"）'),
        ('安全文明/CI形象', ['围挡', '七牌一图', '工完场清', '材料码放', '封闭管理'],
         '对标青岛评分表"安全文明施工描述空洞得1分"、竞品仅泛写"加强安全文明管理"'),
        ('劳务/工资保障', ['工资专用账户', '分账', '总包代发', '工资保证金', '维权信息告示牌'],
         '对标《保障农民工工资支付条例》法定刚性项、竞品仅泛写"加强劳务管理"'),
        ('季节/特殊工况', ['冬期', '雨期', '台风', '高温季节'],
         '对标竞品"保障措施千篇一律、不看气候"'),
        ('量化硬指标', ['合格率', '达成率', '100%'],
         '对标喜鹊"量化不足、泛泛而谈"'),
        ('评分项逐条响应', ['针对『', '针对「'],
         '对标 WPS/喜鹊"评分点逐条响应"核心卖点'),
    ]

    @staticmethod
    def _scan_compete_coverage(paragraphs: List[str]) -> Dict:
        """扫描正文段落，命中各竞争力维度的代表关键词，返回覆盖情况。"""
        text = '\n'.join(paragraphs or [])
        present, missing = [], []
        for name, kws, gap in EvaluatorCheck._COMPETE_DIMS:
            hit = any(k in text for k in kws)
            (present if hit else missing).append({'name': name, 'gap': gap})
        return {
            'present': present,
            'missing': missing,
            'coverage_rate': round(100.0 * len(present) / max(1, len(EvaluatorCheck._COMPETE_DIMS)), 1),
        }

    @staticmethod
    def _render_compete_markdown(result: Dict) -> str:
        """渲染竞争力维度覆盖自检为 Markdown。"""
        if not result:
            return ''
        lines = [
            '## 技术标竞争力维度覆盖自检（v7.28 新增）',
            '',
            f'> 逐轮对比 WPS AI / 喜鹊标书 / 钛投标 / 红点智标 等市面产品补强的差异化维度，'
            f'扫描本标书正文落地情况。覆盖越全，相对模板化竞品的加分优势越明显。',
            '',
            f'- **维度覆盖率**：`{result.get("coverage_rate", 0)}%`'
            f'（{len(result.get("present", []))}/{len(EvaluatorCheck._COMPETE_DIMS)}）',
            '',
            '**已落地维度（相对竞品的差异化优势）：**',
        ]
        for d in result.get('present', []):
            lines.append(f"- ✅ {d['name']} —— {d['gap']}")
        missing = result.get('missing', [])
        if missing:
            lines.append('')
            lines.append('**未检出维度（可针对性补强）：**')
            for d in missing:
                lines.append(f"- ⚠️ {d['name']} —— {d['gap']}")
        else:
            lines.append('')
            lines.append('✅ 全部竞争力维度均已落地。')
        return '\n'.join(lines)

    # ── 综合执行入口 ─────────────────────────────────────────

    def run_all(self) -> Dict:
        """
        执行全部评标专家视角检查

        Returns:
            {
                'evaluator_pass': bool,       # 是否通过专家视角检查
                'coverage': Dict,             # 评分项覆盖率结果
                'depth': Dict,                # 内容深度评估结果
                'risks': Dict,                # 废标风险扫描结果
                'score_prediction': Dict,     # 得分预测结果
                'suggestions': List[Dict],    # 改进建议
                'repair_items': List[Dict],   # repair模块兼容格式
                'summary': str,               # 综合摘要
            }
        """
        # 1. 评分项覆盖率
        coverage = self.check_score_item_coverage()

        # 2. 内容深度（对每个评分项章节评估）
        depth_results = {}
        score_items = self.parse_result.get('score_items', [])
        detail_level = self.doc_info.get('detail_level', 3)
        for item in score_items:
            title = item.get('name', item.get('title', ''))
            if title:
                depth_results[title] = self.evaluate_depth(title, detail_level=detail_level)

        # 3. 废标风险
        risks = self.scan_disqualification_risks()

        # 4. 得分预测
        score_prediction = self.predict_score()

        # 5. 改进建议
        suggestions = self.generate_improvement_suggestions()

        # 5.5 P2: 模拟评审得分表
        mock_review = self._build_mock_review(score_prediction, suggestions)

        # 5.6 v7.6: 文档原创度自检（扫描正文内的重复/高度相似段落）
        originality = None
        if self.bid_doc_paragraphs:
            try:
                from bid_core.originality import scan_originality
                originality = scan_originality(self.bid_doc_paragraphs)
                if originality:
                    mock_review['markdown'] = (
                        mock_review.get('markdown', '')
                        + '\n\n'
                        + self._render_originality_markdown(originality)
                    )
            except Exception as exc:
                print(f"[EvaluatorCheck] 原创度自检跳过: {exc}")

        # 5.7 v7.28: 技术标竞争力维度覆盖自检（扫描正文的竞品对标型维度落地情况）
        try:
            if self.bid_doc_paragraphs:
                compete = self._scan_compete_coverage(self.bid_doc_paragraphs)
                if compete:
                    mock_review['markdown'] = (
                        mock_review.get('markdown', '')
                        + '\n\n'
                        + self._render_compete_markdown(compete)
                    )
        except Exception as exc:
            print(f"[EvaluatorCheck] 竞争力维度覆盖自检跳过: {exc}")

        # 6. repair格式
        repair_items = self.to_repair_format(suggestions)

        # 判定是否通过
        has_high_risk = len(risks.get('high_risks', [])) > 0
        coverage_below_threshold = coverage.get('coverage_rate', 0) < 0.8
        score_below_60 = score_prediction.get('score_percentage', 0) < 60

        evaluator_pass = not has_high_risk and not coverage_below_threshold and not score_below_60

        # 综合摘要
        summary_parts = []
        if has_high_risk:
            summary_parts.append(
                f"🔴 存在{len(risks['high_risks'])}项高风险废标风险"
            )
        if coverage_below_threshold:
            summary_parts.append(
                f"🟡 评分项覆盖率仅{coverage['coverage_rate']*100:.1f}%（低于80%阈值）"
            )
        if score_below_60:
            summary_parts.append(
                f"🟡 预测得分率仅{score_prediction['score_percentage']:.1f}%（低于60%阈值）"
            )
        if not summary_parts:
            summary_parts.append(
                f"✅ 评标专家视角检查通过：覆盖率{coverage['coverage_rate']*100:.1f}%，"
                f"预测得分率{score_prediction['score_percentage']:.1f}%，"
                f"无高风险废标项"
            )

        summary = '；'.join(summary_parts)

        return {
            'evaluator_pass': evaluator_pass,
            'coverage': coverage,
            'depth': depth_results,
            'risks': risks,
            'score_prediction': score_prediction,
            'suggestions': suggestions,
            'repair_items': repair_items,
            'summary': summary,
            # P2: 模拟评审得分表
            'mock_review': mock_review,
            # v7.6: 文档原创度自检结果
            'originality': originality,
        }


def evaluator_check(parse_result: Dict, doc_info: Dict = None,
                    bid_doc_path: str = None) -> Dict:
    """
    评标专家视角自检入口函数

    Args:
        parse_result: parser.py输出的招标文件解析结果
        doc_info: 生成标书的信息
        bid_doc_path: 生成标书的docx文件路径

    Returns:
        评标专家视角检查报告dict
    """
    return EvaluatorCheck(parse_result, doc_info, bid_doc_path).run_all()


if __name__ == '__main__':
    # 测试
    test_parse = {
        'project_info': {'duration': 90, 'quality_target': '合格'},
        'score_items': [
            {'name': '施工方案和技术措施', 'score': 30,
             'description': '包含但不限于：土方工程、基础工程、主体结构工程'},
            {'name': '质量保证措施', 'score': 15},
            {'name': '安全文明施工', 'score': 10},
            {'name': '施工进度计划', 'score': 10},
            {'name': '项目管理机构', 'score': 10},
        ],
        'disqualify_clauses': [
            '未提供安全生产许可证的按废标处理',
            '项目经理未提供建造师证书的按废标处理',
        ],
        'red_line_clauses': [],
        'qualification_requirements': {'license': [], 'person': [], 'experience': []},
    }
    test_doc = {
        'project_info': {'duration': 100, 'quality_target': '合格'},
        'chapters': [
            {'title': '施工方案和技术措施'},
            {'title': '质量保证措施'},
            {'title': '安全文明施工措施'},
        ],
        'tables': [
            {'title': '主要施工机械设备表'},
        ],
        'detail_level': 3,
    }

    result = evaluator_check(test_parse, test_doc)
    print(json.dumps(result, ensure_ascii=False, indent=2))

"""
后处理钩子系统 v1.0
参考桌面版Pipeline架构，提供可插拔的后处理钩子框架
生成标书后可自动执行：去AI化 → 格式微调 → 降重 → 查重 等步骤
"""
import os
import time
from typing import Dict, List, Any, Optional, Callable
from dataclasses import dataclass, field


@dataclass
class HookResult:
    """钩子执行结果"""
    hook_id: str
    status: str  # 'success' | 'failed' | 'skipped'
    output_path: Optional[str] = None
    changes: Dict[str, Any] = field(default_factory=dict)
    error: Optional[str] = None
    execution_time: float = 0.0


class BaseHook:
    """钩子基类"""
    hook_id: str = ""
    hook_name: str = ""
    priority: int = 100  # 数字越小越先执行
    enabled: bool = True
    
    def execute(self, context: Dict[str, Any]) -> HookResult:
        """
        执行钩子
        
        Args:
            context: {
                'docx_path': str,      # 标书文件路径
                'project_info': dict,  # 项目信息
                'parse_result': dict,  # 解析结果
                'check_result': dict,  # 检查结果（如有）
                'user_context': dict,  # 用户信息
            }
        """
        raise NotImplementedError


class DeAIHook(BaseHook):
    """去AI化钩子"""
    hook_id = "hook_deai"
    hook_name = "去AI化"
    priority = 10
    
    def __init__(self, mode: str = 'normal'):
        self.mode = mode
    
    def execute(self, context: Dict[str, Any]) -> HookResult:
        start = time.time()
        try:
            from bid_core.deai import DeAIProcessor
            processor = DeAIProcessor(mode=self.mode)
            result = processor.process_docx(context['docx_path'])
            return HookResult(
                hook_id=self.hook_id,
                status='success',
                output_path=context['docx_path'],
                changes=result,
                execution_time=time.time() - start,
            )
        except Exception as e:
            return HookResult(
                hook_id=self.hook_id,
                status='failed',
                error=str(e),
                execution_time=time.time() - start,
            )


class RewriteHook(BaseHook):
    """降重钩子 - 应用同义词替换"""
    hook_id = "hook_rewrite"
    hook_name = "降重处理"
    priority = 20
    
    def execute(self, context: Dict[str, Any]) -> HookResult:
        start = time.time()
        try:
            from bid_core.randomizer import Randomizer
            from docx import Document
            
            synonyms_path = os.path.join(
                os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                'data', 'synonyms.json'
            )
            randomizer = Randomizer(synonyms_path=synonyms_path, enabled=True)
            
            doc = Document(context['docx_path'])
            synonym_count = 0
            
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.text:
                        new_text = randomizer.synonym_replace(run.text)
                        if new_text != run.text:
                            run.text = new_text
                            synonym_count += 1
            
            doc.save(context['docx_path'])
            
            return HookResult(
                hook_id=self.hook_id,
                status='success',
                output_path=context['docx_path'],
                changes={'synonym_replaced': synonym_count},
                execution_time=time.time() - start,
            )
        except Exception as e:
            return HookResult(
                hook_id=self.hook_id,
                status='failed',
                error=str(e),
                execution_time=time.time() - start,
            )


class DuplicateCheckHook(BaseHook):
    """查重钩子 - 生成查重报告"""
    hook_id = "hook_duplicate_check"
    hook_name = "查重检查"
    priority = 30
    
    def __init__(self, compare_paths: List[str] = None, mode: str = '标书'):
        self.compare_paths = compare_paths or []
        self.mode = mode
    
    def execute(self, context: Dict[str, Any]) -> HookResult:
        start = time.time()
        try:
            from duplicate_checker import check_duplicates
            
            paths = [context['docx_path']] + self.compare_paths
            result = check_duplicates(paths, mode=self.mode)
            
            return HookResult(
                hook_id=self.hook_id,
                status='success',
                changes={
                    'risk_level': result.get('risk_level', '未知'),
                    'max_similarity': result.get('overall_max_similarity', 0),
                    'comparison_count': result.get('comparison_count', 0),
                },
                execution_time=time.time() - start,
            )
        except Exception as e:
            return HookResult(
                hook_id=self.hook_id,
                status='failed',
                error=str(e),
                execution_time=time.time() - start,
            )


class FormatHook(BaseHook):
    """格式微调钩子 v2.0 - 修复数字间标点不替换问题"""
    hook_id = "hook_format"
    hook_name = "格式微调"
    priority = 40
    
    def execute(self, context: Dict[str, Any]) -> HookResult:
        start = time.time()
        try:
            from docx import Document
            
            doc = Document(context['docx_path'])
            changes = 0
            
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.text:
                        new_text = self._fix_punctuation(run.text)
                        # 统一全角/半角数字
                        new_text = self._normalize_numbers(new_text)
                        if new_text != run.text:
                            run.text = new_text
                            changes += 1
            
            # 修复表格中的标点
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.text:
                                    new_text = self._fix_punctuation(run.text)
                                    if new_text != run.text:
                                        run.text = new_text
                                        changes += 1
            
            doc.save(context['docx_path'])
            
            return HookResult(
                hook_id=self.hook_id,
                status='success',
                output_path=context['docx_path'],
                changes={'punctuation_fixed': changes},
                execution_time=time.time() - start,
            )
        except Exception as e:
            return HookResult(
                hook_id=self.hook_id,
                status='failed',
                error=str(e),
                execution_time=time.time() - start,
            )
    
    def _fix_punctuation(self, text):
        """
        规范化标点：英文标点→中文标点
        但保留数字间的英文标点（1.1, 3.5等）
        """
        import re
        result = text
        
        # 1. 先保护数字间的英文标点（如1.1, 2.3.4, 第1.2条）
        protected = []
        def protect_number_dots(m):
            protected.append(m.group(0))
            return f'\x00{len(protected)-1}\x00'
        
        # 保护 "数字.数字" 模式（版本号、章节号）
        result = re.sub(r'\d+\.[\d.]+', protect_number_dots, result)
        # 保护 "数字,数字" 模式（如 "1,000"）
        result = re.sub(r'\d+,\d+', protect_number_dots, result)
        # 保护 "数字:数字" 模式（如时间 "9:30"）
        result = re.sub(r'\d+:\d+', protect_number_dots, result)
        
        # 2. 替换中文上下文中的英文标点
        # 只在中文语境（前后有中文字符）时替换
        result = re.sub(r'(?<=[\u4e00-\u9fff]),(?=[\u4e00-\u9fff])', '，', result)
        result = re.sub(r'(?<=[\u4e00-\u9fff]);(?=[\u4e00-\u9fff])', '；', result)
        result = re.sub(r'(?<=[\u4e00-\u9fff]):(?=[\u4e00-\u9fff])', '：', result)
        result = re.sub(r'(?<=[\u4e00-\u9fff])\((?=[\u4e00-\u9fff])', '（', result)
        result = re.sub(r'(?<=[\u4e00-\u9fff])\)(?=[\u4e00-\u9fff])', '）', result)
        
        # 3. 恢复被保护的标点
        for i, original in enumerate(protected):
            result = result.replace(f'\x00{i}\x00', original)
        
        return result
    
    def _normalize_numbers(self, text):
        """统一全角数字为半角数字 - v2.0新增"""
        import re
        # 全角数字→半角
        fullwidth = '０１２３４５６７８９'
        halfwidth = '0123456789'
        for fw, hw in zip(fullwidth, halfwidth):
            text = text.replace(fw, hw)
        # 全角字母→半角（不常见但有时出现）
        for fw, hw in zip('ＡＢＣＤＥＦＧＨＩＪＫＬＭＮＯＰＱＲＳＴＵＶＷＸＹＺ',
                          'ABCDEFGHIJKLMNOPQRSTUVWXYZ'):
            text = text.replace(fw, hw)
        return text


class ChapterCompletenessHook(BaseHook):
    """章节完整性检查钩子 v2.0新增 - 检查空白章节和过短章节"""
    hook_id = "hook_chapter_completeness"
    hook_name = "章节完整性检查"
    priority = 5  # 最先执行，在其他钩子修改之前检查
    
    # 章节最小字数阈值
    MIN_CHARS_PER_CHAPTER = 200
    
    def execute(self, context: Dict[str, Any]) -> HookResult:
        start = time.time()
        try:
            from docx import Document
            
            doc = Document(context['docx_path'])
            issues = []
            current_chapter = None
            chapter_chars = 0
            
            for para in doc.paragraphs:
                text = para.text.strip()
                style_name = para.style.name if para.style else ''
                
                # 检测章节标题
                if style_name.startswith('Heading') or style_name.startswith('List Paragraph'):
                    # 保存上一个章节的字数
                    if current_chapter and chapter_chars < self.MIN_CHARS_PER_CHAPTER:
                        issues.append({
                            'chapter': current_chapter,
                            'chars': chapter_chars,
                            'issue': '章节内容过短',
                        })
                    current_chapter = text[:50]  # 只取前50字符作为标识
                    chapter_chars = 0
                else:
                    chapter_chars += len(text)
            
            # 检查最后一个章节
            if current_chapter and chapter_chars < self.MIN_CHARS_PER_CHAPTER:
                issues.append({
                    'chapter': current_chapter,
                    'chars': chapter_chars,
                    'issue': '章节内容过短',
                })
            
            return HookResult(
                hook_id=self.hook_id,
                status='success',
                changes={
                    'issues_found': len(issues),
                    'issues': issues[:10],  # 最多报告10个
                    'total_chapters': len([p for p in doc.paragraphs 
                                          if p.style and p.style.name.startswith('Heading')]),
                },
                execution_time=time.time() - start,
            )
        except Exception as e:
            return HookResult(
                hook_id=self.hook_id,
                status='failed',
                error=str(e),
                execution_time=time.time() - start,
            )


class HookPipeline:
    """
    钩子管线 - 按priority顺序执行所有注册的钩子
    """
    
    def __init__(self):
        self._hooks: List[BaseHook] = []
    
    def register(self, hook: BaseHook):
        """注册钩子"""
        self._hooks.append(hook)
        self._hooks.sort(key=lambda h: h.priority)
    
    def register_defaults(self, enable_deai=True, enable_rewrite=False,
                          enable_duplicate=False, enable_format=True,
                          enable_chapter_check=True):
        """
        注册默认钩子集
        
        Args:
            enable_deai: 启用去AI化
            enable_rewrite: 启用降重
            enable_duplicate: 启用查重
            enable_format: 启用格式微调
            enable_chapter_check: 启用章节完整性检查
        """
        if enable_deai:
            self.register(DeAIHook())
        if enable_rewrite:
            self.register(RewriteHook())
        if enable_duplicate:
            self.register(DuplicateCheckHook())
        if enable_format:
            self.register(FormatHook())
        if enable_chapter_check:
            self.register(ChapterCompletenessHook())
    
    def execute(self, context: Dict[str, Any]) -> List[HookResult]:
        """执行所有钩子"""
        results = []
        for hook in self._hooks:
            if not hook.enabled:
                results.append(HookResult(
                    hook_id=hook.hook_id,
                    status='skipped',
                ))
                continue
            
            result = hook.execute(context)
            results.append(result)
            
            # 如果某个钩子失败，不影响后续钩子执行
        
        return results
    
    def list_hooks(self) -> List[Dict]:
        """列出所有注册的钩子"""
        return [
            {
                'id': h.hook_id,
                'name': h.hook_name,
                'priority': h.priority,
                'enabled': h.enabled,
            }
            for h in self._hooks
        ]


def run_post_hooks(docx_path: str, project_info: Dict = None,
                   parse_result: Dict = None, check_result: Dict = None,
                   user_context: Dict = None,
                   enable_deai: bool = True, enable_rewrite: bool = False,
                   enable_duplicate: bool = False, enable_format: bool = True,
                   ) -> Dict:
    """
    后处理钩子入口函数
    
    Args:
        docx_path: 标书docx文件路径
        project_info: 项目信息
        parse_result: 解析结果
        check_result: 检查结果
        user_context: 用户信息
        enable_*: 各钩子开关
    
    Returns:
        {
            'hooks_executed': int,
            'results': [HookResult...],
            'all_passed': bool,
        }
    """
    pipeline = HookPipeline()
    pipeline.register_defaults(
        enable_deai=enable_deai,
        enable_rewrite=enable_rewrite,
        enable_duplicate=enable_duplicate,
        enable_format=enable_format,
    )
    
    context = {
        'docx_path': docx_path,
        'project_info': project_info or {},
        'parse_result': parse_result or {},
        'check_result': check_result or {},
        'user_context': user_context or {},
    }
    
    results = pipeline.execute(context)
    
    return {
        'hooks_executed': len([r for r in results if r.status == 'success']),
        'results': [
            {
                'hook_id': r.hook_id,
                'status': r.status,
                'changes': r.changes,
                'error': r.error,
                'execution_time': r.execution_time,
            }
            for r in results
        ],
        'all_passed': all(r.status == 'success' for r in results if r.status != 'skipped'),
    }

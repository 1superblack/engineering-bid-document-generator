"""ContentMixin — 标题（h1-h4）与正文（body/body_list/body_bold/add_heading）方法。

拆分自原 formatter.py v7.0 NormalFormatter。
"""
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn


# 中文数字映射
_CN_NUMBERS = [
    "", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
    "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十",
]


class ContentMixin:
    """标题与正文内容生成方法。"""

    def h1(self, index: int, title: str) -> None:
        """一级标题：一、二、三 — v7.0 自动插入分页符。"""
        num_str = _CN_NUMBERS[index] if index <= 20 else str(index)

        # 一级标题前自动插入分页符（首章除外）
        if self._h1_count > 0:
            self.doc.add_page_break()

        self._h1_count += 1
        self._current_chapter = f"{num_str}、{title}"

        h = self.doc.add_heading(level=1)
        h.clear()
        r = h.add_run(f"{num_str}、{title}")
        r.font.name = self.heading_font
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True
        r._element.rPr.rFonts.set(qn('w:eastAsia'), self.heading_font)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

    def h2(self, title: str) -> None:
        """二级标题：（一）（二）（三）— v7.0 可配置插入分页符。"""
        if self.page_break_before_h2:
            self.doc.add_page_break()

        h = self.doc.add_heading(level=2)
        h.clear()
        r = h.add_run(title)
        r.font.name = self.heading_font
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True
        r._element.rPr.rFonts.set(qn('w:eastAsia'), self.heading_font)
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(4)

    def h3(self, title: str) -> None:
        """三级标题：1、2、3 或 1.1 1.2。"""
        h = self.doc.add_heading(level=3)
        h.clear()
        r = h.add_run(title)
        r.font.name = self.heading_font
        r.font.size = Pt(15)
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True
        r._element.rPr.rFonts.set(qn('w:eastAsia'), self.heading_font)
        h.paragraph_format.space_before = Pt(6)
        h.paragraph_format.space_after = Pt(3)

    def h4(self, title: str) -> None:
        """四级标题：① ② ③ 或（1）（2）（3）。"""
        h = self.doc.add_heading(level=4)
        h.clear()
        r = h.add_run(title)
        r.font.name = self.heading_font
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True
        r._element.rPr.rFonts.set(qn('w:eastAsia'), self.heading_font)
        h.paragraph_format.space_before = Pt(4)
        h.paragraph_format.space_after = Pt(2)

    def body(self, *args) -> None:
        """正文段落（支持可变参数自动拼接）v6.0。"""
        text = ''.join(str(a) for a in args)
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = self.body_font
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0, 0, 0)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)
        p.paragraph_format.first_line_indent = Cm(0.85)
        p.paragraph_format.line_spacing = Pt(28)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)

    def body_list(self, items: list) -> None:
        """正文章节列表 — 每项输出为带序号正文段落 v6.0。"""
        for i, text in enumerate(items, 1):
            self.body(f"{i}、{text}")

    def body_bold(self, text: str) -> None:
        """加粗正文。"""
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = self.body_font
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)
        p.paragraph_format.first_line_indent = Cm(0.85)
        p.paragraph_format.line_spacing = Pt(28)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)

    def add_heading(self, text: str) -> None:
        """添加附表标题。"""
        h = self.doc.add_heading(level=1)
        h.clear()
        r = h.add_run(text)
        r.font.name = self.heading_font
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True
        r._element.rPr.rFonts.set(qn('w:eastAsia'), self.heading_font)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

    def image(self, path: str, caption: str = None, width_cm: float = None) -> None:
        """插入图片（居中），可选图注。文件缺失/损坏时静默跳过，绝不中断生成。

        v7.3: 支持企业图片库自动配图（对标喜鹊标书 AI 的「图片库」）。
        """
        if not path:
            return
        import os
        if not os.path.exists(path):
            return
        try:
            from docx.shared import Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(path, width=Cm(width_cm) if width_cm else None)
            if caption:
                cap = self.doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cap.add_run(caption)
                cr.font.size = Pt(10.5)
                cr.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        except Exception:
            # 图片插入失败不影响正文
            pass

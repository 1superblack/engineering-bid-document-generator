"""
横道图（甘特图）生成模块 v1.1
v1.1: 重构为独立模块，从 generator.py 中分离；P1-5 修复导入链断裂
"""
from __future__ import annotations

from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _gantt_bars(task_duration: int, start_day: int,
                total_duration: int, months: int) -> List[str]:
    """生成横道图的月份占位符（■/□）

    Args:
        task_duration: 任务持续天数
        start_day: 任务开始天数（从1开始）
        total_duration: 总工期天数
        months: 总月数

    Returns:
        月份占位符列表，如 ['■', '■', '□', '□']
    """
    bars: List[str] = []
    for m in range(1, months + 1):
        month_start = (m - 1) * 30 + 1
        month_end = m * 30
        task_end = start_day + task_duration - 1
        if start_day <= month_end and task_end >= month_start:
            bars.append("■")
        else:
            bars.append("□")
    return bars


def generate_gantt_table(doc: Document, project_info: Dict[str, Any]) -> None:
    """在 docx 文档中生成施工进度横道图表格

    Args:
        doc: python-docx 的 Document 对象
        project_info: 项目信息字典，需包含 duration 字段
    """
    duration = project_info.get('duration', 90)
    months = max(1, (duration + 29) // 30)

    # 典型施工工序
    tasks = [
        ("施工准备", 1, 10),
        ("测量放线", 5, 15),
        ("基础工程", 10, 25),
        ("主体工程", 20, 40),
        ("装饰装修", 50, 30),
        ("机电安装", 40, 35),
        ("竣工验收", duration - 10, 10),
    ]

    # 过滤掉超出工期的任务
    tasks = [(name, start, dur) for name, start, dur in tasks
             if start < duration]

    # 构建表格数据
    header = ["序号", "施工阶段", "开始时间(天)", "持续时间(天)"]
    header += [f"第{m}月" for m in range(1, months + 1)]

    rows: List[List[str]] = []
    for i, (name, start, dur) in enumerate(tasks, 1):
        # 确保 start + dur 不超过 duration
        actual_dur = min(dur, duration - start + 1)
        if actual_dur <= 0:
            continue
        row = [str(i), name, str(start), str(actual_dur)]
        row += _gantt_bars(actual_dur, start, duration, months)
        rows.append(row)

    # 创建表格
    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 写入表头
    for j, text in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(8)

    # 写入数据行
    for i, row_data in enumerate(rows, 1):
        for j, text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = text
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(8)

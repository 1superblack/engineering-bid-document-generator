"""
AI增强生成模块 v1.0
参考桌面版GenerateService，提供LLM逐章扩写能力
在Coze环境中，Agent负责调用LLM，本模块负责：
1. 生成逐章扩写prompt
2. 接收LLM扩写结果并合并到文档
3. 模板+AI混合策略：模板内容够长(>500字)直接用，否则需要AI补充
"""
import os
import re
import json
from typing import Dict, List, Optional, Any, Tuple


class AIEnhancer:
    """
    AI增强生成器
    
    工作模式：
    1. 先从章节模板获取基础内容
    2. 如果模板内容<500字，标记为需要AI增强
    3. 生成针对性的扩写prompt
    4. Agent用LLM执行扩写后，将结果传回合并到文档
    """
    
    # 模板内容阈值：低于此值需要AI增强
    TEMPLATE_MIN_LENGTH = 500
    
    def __init__(self, project_info: Dict, parse_result: Dict = None,
                 user_context: Dict = None, detail_level: int = 2):
        self.project_info = project_info
        self.parse_result = parse_result or {}
        self.user_context = user_context or {}
        self.detail_level = detail_level
    
    def analyze_chapters(self, chapter_contents: Dict[str, str]) -> Dict[str, Any]:
        """
        分析各章节内容，识别需要AI增强的章节
        
        Args:
            chapter_contents: {章节标题: 章节文本内容}
        
        Returns:
            {
                'total_chapters': int,
                'need_enhance': [{title, current_length, target_length, priority}],
                'ok_chapters': [title, ...],
            }
        """
        need_enhance = []
        ok_chapters = []
        
        # 目标字数根据detail_level
        target_map = {1: 1500, 2: 4000, 3: 10000}  # 每章目标字数
        target_words = target_map.get(self.detail_level, 4000)
        
        for title, content in chapter_contents.items():
            content_len = len(content.strip()) if content else 0
            
            if content_len < self.TEMPLATE_MIN_LENGTH:
                need_enhance.append({
                    'title': title,
                    'current_length': content_len,
                    'target_length': target_words,
                    'priority': 'high' if content_len < 100 else 'medium',
                })
            elif content_len < target_words * 0.5:
                # 有内容但不够长
                need_enhance.append({
                    'title': title,
                    'current_length': content_len,
                    'target_length': target_words,
                    'priority': 'low',
                })
            else:
                ok_chapters.append(title)
        
        return {
            'total_chapters': len(chapter_contents),
            'need_enhance': need_enhance,
            'ok_chapters': ok_chapters,
        }
    
    def generate_enhance_prompts(self, chapter_contents: Dict[str, str]) -> List[Dict]:
        """
        为需要增强的章节生成LLM扩写prompt
        
        Args:
            chapter_contents: {章节标题: 章节文本内容}
        
        Returns:
            [{title, prompt, current_content, target_words}]
        """
        analysis = self.analyze_chapters(chapter_contents)
        prompts = []
        
        target_map = {1: 1500, 2: 4000, 3: 10000}
        target_words = target_map.get(self.detail_level, 4000)
        
        for item in analysis['need_enhance']:
            title = item['title']
            current = chapter_contents.get(title, '')
            
            prompt = self._build_enhance_prompt(
                title=title,
                current_content=current,
                target_words=target_words,
            )
            
            prompts.append({
                'title': title,
                'prompt': prompt,
                'current_content': current,
                'current_length': item['current_length'],
                'target_words': target_words,
                'priority': item['priority'],
            })
        
        # 按优先级排序：high > medium > low
        priority_order = {'high': 0, 'medium': 1, 'low': 2}
        prompts.sort(key=lambda x: priority_order.get(x['priority'], 3))
        
        return prompts
    
    def _build_enhance_prompt(self, title: str, current_content: str,
                              target_words: int) -> str:
        """构建章节扩写prompt"""
        project_name = self.project_info.get('name', '本项目')
        duration = self.project_info.get('duration', '90')
        bid_type = self.project_info.get('bid_type', 'construction')
        area = self.project_info.get('area', '')
        work_content = self.project_info.get('work_content', '')
        quality_target = self.project_info.get('quality_target', '合格')
        divisions = self.project_info.get('divisions', [])
        
        prompt = f"""请为以下项目编写技术标书章节「{title}」的详细内容：

=== 项目信息 ===
项目名称：{project_name}
工期：{duration}日历天
项目类型：{"施工类" if bid_type == "construction" else "服务类"}"""
        
        if area:
            prompt += f"\n建筑面积：{area}㎡"
        if work_content:
            prompt += f"\n施工内容：{work_content}"
        if quality_target:
            prompt += f"\n质量目标：{quality_target}"
        if divisions:
            prompt += f"\n分项工程：{'、'.join(divisions)}"
        
        # 注入用户信息
        if self.user_context:
            company = self.user_context.get('company', {})
            if company.get('name'):
                prompt += f"\n投标单位：{company['name']}"
            personnel = self.user_context.get('key_personnel', [])
            if personnel:
                prompt += "\n关键人员："
                for p in personnel[:5]:
                    prompt += f"\n  - {p.get('name', '')} ({p.get('role', '')}, {p.get('cert', '')})"
            equipment = self.user_context.get('equipment', [])
            if equipment:
                prompt += "\n主要设备："
                for e in equipment[:5]:
                    prompt += f"\n  - {e.get('name', '')} {e.get('model', '')} {e.get('quantity', '')}台"
        
        # 注入评分项
        score_items = self.parse_result.get('score_items', [])
        if score_items:
            related = []
            for item in score_items:
                item_name = item.get('name', item.get('title', ''))
                item_score = item.get('score', 0)
                related.append(f"  - {item_name}（{item_score}分）")
            if related:
                prompt += "\n\n【评分项要求 - 必须逐条响应】\n"
                prompt += '\n'.join(related)
                prompt += "\n请在内容中明确体现对以上评分项的响应。"
        
        # 注入废标条款
        disqualify_clauses = self.parse_result.get('disqualify_clauses', [])
        red_lines = self.parse_result.get('red_line_clauses', [])
        all_clauses = [c.get('content', '') for c in red_lines if c.get('content')]
        all_clauses += [c if isinstance(c, str) else c.get('content', '') for c in disqualify_clauses]
        all_clauses = [c for c in all_clauses if c][:10]
        
        if all_clauses:
            prompt += "\n\n【废标条款 - 必须避免触发】\n"
            for c in all_clauses:
                prompt += f"  - {c}\n"
            prompt += "请确保内容不违反以上任何废标条款。"
        
        # 现有内容参考
        if current_content and len(current_content) > 50:
            preview = current_content[:800]
            prompt += f"""

=== 现有内容（需在此基础上扩写） ===
{preview}
{'...' if len(current_content) > 800 else ''}

请在以上内容基础上扩写，保留已有内容的核心要点，补充详细措施和实施方案。"""
        else:
            prompt += """

=== 要求 ===
请从零编写本章完整内容。"""
        
        prompt += f"""

=== 输出要求 ===
1. 目标字数：约{target_words}字
2. 标题层级格式：一、→（一）→1.→①，不超过4级
3. 包含具体措施、技术参数、组织架构、时间节点等实质内容
4. 引用真实规范标准（GB/JGJ/JG编号）
5. 每章至少2个表格（组织表/制度表/配置表/检查表等）
6. 避免空话套话，注重可操作性
7. 必须逐条响应评分项要求，不能遗漏"""
        
        return prompt
    
    def merge_enhanced_content(self, docx_path: str,
                               enhanced_chapters: Dict[str, str],
                               output_path: str = None) -> str:
        """
        将AI增强后的章节内容合并到文档（替换模式）
        
        对匹配到的章节：保留标题，删除原内容，插入AI增强内容
        对未匹配的章节：保持原样
        
        Args:
            docx_path: 原始标书文件路径
            enhanced_chapters: {章节标题: AI扩写内容}
            output_path: 输出路径（默认覆盖原文件）
        
        Returns:
            输出文件路径
        """
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.oxml.ns import qn
        import shutil
        
        out = output_path or docx_path
        if out != docx_path:
            shutil.copy2(docx_path, out)
        
        doc = Document(out)
        
        # 1. 识别章节边界（从Heading1到下一个Heading1）
        chapter_ranges = []  # [(start_idx, end_idx, clean_title)]
        para_count = len(doc.paragraphs)
        
        for i, para in enumerate(doc.paragraphs):
            if para.style and 'Heading 1' in str(para.style.name):
                text = para.text.strip()
                clean_title = re.sub(r'^[一二三四五六七八九十]+、', '', text).strip()
                if clean_title:
                    chapter_ranges.append((i, None, clean_title))
        
        # 填充end_idx
        for idx, (start, _, title) in enumerate(chapter_ranges):
            if idx + 1 < len(chapter_ranges):
                chapter_ranges[idx] = (start, chapter_ranges[idx + 1][0], title)
            else:
                chapter_ranges[idx] = (start, para_count, title)
        
        # 2. 建立标题匹配映射
        title_match = {}
        for start, end, clean_title in chapter_ranges:
            for enh_title in enhanced_chapters:
                if enh_title in clean_title or clean_title in enh_title:
                    title_match[clean_title] = enh_title
                    break
        
        # 3. 删除需要替换的章节内容（保留标题段落）
        #    从后往前删避免索引偏移
        paras_to_delete = []
        for start, end, clean_title in chapter_ranges:
            if clean_title in title_match:
                # 删除标题之后、下一章之前的所有段落
                for i in range(start + 1, end):
                    paras_to_delete.append(i)
        
        for i in sorted(paras_to_delete, reverse=True):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)
        
        # 4. 在替换章节标题后插入AI增强内容
        for start, end, clean_title in chapter_ranges:
            if clean_title not in title_match:
                continue
            
            enh_title = title_match[clean_title]
            content = enhanced_chapters[enh_title]
            if not content or not content.strip():
                continue
            
            # 找到标题段落的XML元素，在其后插入
            title_para = doc.paragraphs[start]
            title_elem = title_para._element
            parent = title_elem.getparent()
            
            # 解析增强内容并插入
            elements_to_add = []
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    continue
                
                if re.match(r'^[（(][一二三四五六七八九十]+[）)]', line):
                    # 二级标题
                    h = doc.add_heading(level=2)
                    h.clear()
                    r = h.add_run(line)
                    r.font.name = '黑体'
                    r.font.size = Pt(16)
                    r.font.color.rgb = RGBColor(0, 0, 0)
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    elements_to_add.append(h._element)
                elif re.match(r'^\d+[.、]', line):
                    # 三级标题
                    h = doc.add_heading(level=3)
                    h.clear()
                    r = h.add_run(line)
                    r.font.name = '黑体'
                    r.font.size = Pt(14)
                    r.font.color.rgb = RGBColor(0, 0, 0)
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    elements_to_add.append(h._element)
                else:
                    # 正文
                    p = doc.add_paragraph()
                    r = p.add_run(line)
                    r.font.name = '仿宋'
                    r.font.size = Pt(14)
                    r.font.color.rgb = RGBColor(0, 0, 0)
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                    p.paragraph_format.first_line_indent = Cm(0.74)
                    p.paragraph_format.line_spacing = 1.5
                    elements_to_add.append(p._element)
            
            # 将元素从文档末尾移动到标题后面
            insert_after = title_elem
            for elem in elements_to_add:
                parent.remove(elem)
                insert_after.addnext(elem)
                insert_after = elem
        
        doc.save(out)
        return out


def generate_enhance_prompts(project_info: Dict, chapter_contents: Dict[str, str],
                             parse_result: Dict = None, user_context: Dict = None,
                             detail_level: int = 2) -> List[Dict]:
    """
    便捷入口：生成章节扩写prompt
    
    Args:
        project_info: 项目信息
        chapter_contents: {章节标题: 当前内容}
        parse_result: 解析结果
        user_context: 用户信息
        detail_level: 深度等级(1/2/3)
    
    Returns:
        扩写prompt列表
    """
    enhancer = AIEnhancer(project_info, parse_result, user_context, detail_level)
    return enhancer.generate_enhance_prompts(chapter_contents)


def merge_enhanced_content(docx_path: str, enhanced_chapters: Dict[str, str],
                           output_path: str = None) -> str:
    """
    便捷入口：合并AI增强内容到文档
    """
    enhancer = AIEnhancer({})
    return enhancer.merge_enhanced_content(docx_path, enhanced_chapters, output_path)

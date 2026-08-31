"""
去AI化处理管道 v6.0
主编排逻辑，协调检测、修复和领域调整
"""
import logging
from pathlib import Path
from typing import Tuple, Dict, List

from pipeline.output_paths import aux_dir

from .replacements import AI_PHRASE_REPLACEMENTS, ENGINEERING_PHRASES
from .detectors import DeAIDetector
from .fixers import DeAIFixer
from .domain_adjusters import DomainAdjuster

log = logging.getLogger(__name__)


class DeAIProcessor:
    """去AI化处理器 v6.0"""

    def __init__(self, mode: str = 'normal'):
        self.mode = mode
        self.replacements = dict(AI_PHRASE_REPLACEMENTS)
        if mode == 'aggressive':
            self.replacements.update(ENGINEERING_PHRASES)

        # 初始化子模块
        self.detector = DeAIDetector()
        self.fixer = DeAIFixer(self.replacements)
        self.domain_adjuster = DomainAdjuster()

        # 统计信息
        self.stats = {
            'phrases_replaced': 0,
            'patterns_detected': 0,
            'total_chars_changed': 0,
            'paragraph_starts_diversified': 0,
            **self.detector.stats,
            **self.fixer.stats,
        }

    def detect_ai_patterns(self, text: str) -> List[Dict]:
        """检测文本中的AI特征模式

        Args:
            text: 待检测的文本

        Returns:
            检测到的问题列表，每项包含type/match/position/suggestion
        """
        log.info("开始AI特征检测")
        issues = self.detector.detect_all(text)
        self.stats['patterns_detected'] = len(issues)

        if issues:
            log.warning(f"检测到{len(issues)}处AI特征模式")
            for issue in issues[:5]:  # 只记录前5条
                log.debug(f"  - {issue.get('type')}: {issue.get('match', '')[:50]}")

        return issues

    def process_text(self, text: str, domain: str = None) -> Tuple[str, Dict]:
        """执行完整的去AI化处理

        处理流程：
        1. 应用短语替换
        2. 检测并修复连续句式
        3. 修复套话模式
        4. 转换目的从句
        5. 插入过渡词
        6. 修复排比结构
        7. 领域特定调整（可选）

        Args:
            text: 原始文本
            domain: 行业领域（facade/municipal/decoration/smart/green）

        Returns:
            (处理后文本, 统计信息字典)
        """
        log.info(f"开始去AI化处理 | 文本长度={len(text)} | 模式={self.mode}")

        original_text = text
        total_changes = 0

        # Step 1: 应用短语替换
        text, count = self.fixer.apply_replacements(text)
        total_changes += count
        log.info(f"短语替换完成: 替换了{count}处")

        # Step 2: 检测AI特征
        issues = self.detect_ai_patterns(text)

        # Step 3: 修复连续句式
        text, count = self.fixer.fix_consecutive_patterns(text)
        total_changes += count

        # Step 4: 修复套话模式
        text, count = self.fixer.fix_cliche_patterns(text)
        total_changes += count

        # Step 5: 转换目的从句
        text, count = self.fixer.transform_purpose_clause(text)
        total_changes += count

        # Step 6: 插入过渡词
        text, count = self.fixer.insert_transitions(text)
        total_changes += count

        # Step 7: 修复排比结构
        text, count = self.fixer.fix_parallel_structures(text)
        total_changes += count

        # Step 8: 领域特定调整
        if domain:
            text = self.domain_adjuster.adjust(text, domain)
            log.info(f"已应用{domain}领域调整")

        # 计算总变更字符数
        chars_changed = len(original_text) - len(text) + sum(
            1 for a, b in zip(original_text, text) if a != b
        )
        self.stats['total_chars_changed'] = chars_changed
        self.stats['phrases_replaced'] = self.fixer.stats.get('phrases_replaced', 0)

        log.info(f"去AI化完成: 总变更{total_changes}处, 字符变化{chars_changed}")

        return text, dict(self.stats)

    def transform_extended(self, content: str, keyword: str = '',
                           domain: str = None) -> Tuple[str, Dict]:
        """对长文档进行分段去AI处理

        将长文档按段落分割后逐段处理，避免内存问题
        并在关键位置插入关键词相关表述

        Args:
            content: 长文档内容
            keyword: 关键词（用于增强相关性）
            domain: 行业领域

        Returns:
            (处理后文档, 统计信息)
        """
        log.info(f"开始分段去AI处理 | 内容长度={len(content)}")

        # 分段处理
        paragraphs = content.split('\n\n')
        processed_paragraphs = []
        total_stats = {
            'extended_segments_processed': 0,
            'phrases_replaced': 0,
            'patterns_fixed': 0,
        }

        for i, para in enumerate(paragraphs):
            if len(para.strip()) < 20:
                processed_paragraphs.append(para)
                continue

            # 对每个段落执行完整处理
            processed_para, stats = self.process_text(para, domain=domain)
            processed_paragraphs.append(processed_para)

            # 累加统计
            total_stats['extended_segments_processed'] += 1
            total_stats['phrases_replaced'] += stats.get('phrases_replaced', 0)
            total_stats['patterns_fixed'] += stats.get('patterns_detected', 0)

        result_content = '\n\n'.join(processed_paragraphs)

        log.info(f"分段处理完成: 处理了{total_stats['extended_segments_processed']}个段落")

        return result_content, total_stats

    def process_docx(self, docx_path: str, output_path: str = None) -> Dict:
        """处理Word文档

        Args:
            docx_path: 输入文件路径
            output_path: 输出文件路径（可选）

        Returns:
            处理结果统计
        """
        try:
            from docx import Document
        except ImportError:
            log.error("未安装python-pptx库，无法处理docx文件")
            return {'error': '缺少python-pptx依赖'}

        log.info(f"开始处理Word文档: {docx_path}")

        doc = Document(docx_path)
        total_paragraphs = len(doc.paragraphs)
        processed_count = 0

        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                original_text = para.text
                processed_text, _ = self.process_text(original_text)

                if processed_text != original_text:
                    para.text = processed_text
                    processed_count += 1

        # 保存输出文件（ADR-008：归入 <主文档名>_交付物 子文件夹，桌面只留主 docx）
        save_path = output_path or str(aux_dir(docx_path) / (Path(docx_path).stem + '_deai.docx'))
        doc.save(save_path)

        result = {
            'input_file': docx_path,
            'output_file': save_path,
            'total_paragraphs': total_paragraphs,
            'processed_paragraphs': processed_count,
            'status': 'success',
        }

        log.info(f"Word文档处理完成: 处理了{processed_count}/{total_paragraphs}个段落")
        return result


# ════════════════════════════════════════════════════════════════
# 便捷函数接口
# ════════════════════════════════════════════════════════════════

def deai_text(text: str, mode: str = 'normal') -> Tuple[str, Dict]:
    """对文本执行去AI化处理的便捷函数

    Args:
        text: 输入文本
        mode: 处理模式（normal/aggressive）

    Returns:
        (处理后文本, 统计信息)
    """
    processor = DeAIProcessor(mode=mode)
    return processor.process_text(text)


def deai_docx(docx_path: str, output_path: str = None,
              mode: str = 'normal') -> Dict:
    """对Word文档执行去AI化处理的便捷函数

    Args:
        docx_path: 输入文件路径
        output_path: 输出文件路径（可选）
        mode: 处理模式

    Returns:
        处理结果统计
    """
    processor = DeAIProcessor(mode=mode)
    return processor.process_docx(docx_path, output_path=output_path)

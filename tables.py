"""TablesMixin — 表格方法（标准表格、三线表、间距控制）。

拆分自原 formatter.py v7.0 NormalFormatter。
"""
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ooxml_helpers import (
    make_border_element,
    make_shading_element,
    set_fixed_table_layout,
    set_table_grid,
)


class TablesMixin:
    """表格生成方法。"""

    def table(self, headers: list, rows: list) -> None:
        """标准表格（支持10+列宽表格自动适配）v6.0。

        增强：
        - 自动计算可用表格宽度
        - 超过 WIDE_TABLE_THRESHOLD 列时自动缩放字体和列宽
        - 设置列宽确保表格不超出页面
        - v7.0: 表格前后间距优化
        """
        num_cols = len(headers)
        col_widths, cell_font_size = self._calc_column_widths(num_cols)

        self._add_table_spacing(before=True)

        t = self.doc.add_table(rows=len(rows) + 1, cols=num_cols)
        t.style = 'Table Grid'
        t.autofit = False

        tbl = t._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        set_fixed_table_layout(tbl_pr)
        set_table_grid(tbl, col_widths)

        # 表头
        for j, h in enumerate(headers):
            c = t.rows[0].cells[j]
            c.text = h
            c.width = col_widths[j]
            c._element.get_or_add_tcPr().append(make_shading_element('D9E2F3'))
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = self.heading_font
                    r.font.size = cell_font_size
                    r.font.color.rgb = RGBColor(0, 0, 0)
                    r.font.bold = True
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), self.heading_font)

        # 数据行
        for i, row in enumerate(rows, 1):
            for j, v in enumerate(row):
                if j < num_cols:
                    c = t.rows[i].cells[j]
                    c.text = str(v)
                    c.width = col_widths[j]
                    for p in c.paragraphs:
                        for r in p.runs:
                            r.font.name = self.body_font
                            r.font.size = cell_font_size
                            r.font.color.rgb = RGBColor(0, 0, 0)
                            r._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)

        self._add_table_spacing(before=False)

    def add_professional_table(
        self,
        table_num: str,
        table_title: str,
        headers: list,
        rows: list,
    ) -> None:
        """三线表样式 — v7.0 新增。

        专业标书常用三线表：
        - 表格顶线和底线为粗线（1.5pt）
        - 表头下为细线（0.75pt）
        - 无竖线
        - 表格标题格式：表号+表名，居中，小四号黑体
        - 表格内文字：五号仿宋
        """
        num_cols = len(headers)
        total_width = self._get_available_table_width()
        col_w = total_width / num_cols
        col_widths = [col_w] * num_cols

        # 表格标题：表号+表名，居中，小四号黑体
        title_p = self.doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(12)
        title_p.paragraph_format.space_after = Pt(6)
        title_r = title_p.add_run(f"{table_num}  {table_title}")
        title_r.font.name = self.heading_font
        title_r.font.size = Pt(12)
        title_r.font.bold = True
        title_r.font.color.rgb = RGBColor(0, 0, 0)
        title_r._element.rPr.rFonts.set(qn('w:eastAsia'), self.heading_font)

        # 创建表格
        t = self.doc.add_table(rows=len(rows) + 1, cols=num_cols)
        t.autofit = False

        tbl = t._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

        # 三线表边框：顶粗、底粗、表头下细，无竖线无内部横线
        tbl_borders = OxmlElement('w:tblBorders')
        tbl_borders.append(make_border_element('w:top', 'single', '12', '0', '000000'))
        tbl_borders.append(make_border_element('w:bottom', 'single', '12', '0', '000000'))
        tbl_borders.append(make_border_element('w:left', 'none', '0', '0', 'auto'))
        tbl_borders.append(make_border_element('w:right', 'none', '0', '0', 'auto'))
        tbl_borders.append(make_border_element('w:insideV', 'none', '0', '0', 'auto'))
        tbl_borders.append(make_border_element('w:insideH', 'none', '0', '0', 'auto'))

        for old_b in tbl_pr.findall(qn('w:tblBorders')):
            tbl_pr.remove(old_b)
        tbl_pr.append(tbl_borders)

        set_fixed_table_layout(tbl_pr)
        set_table_grid(tbl, col_widths)

        # 表头行
        for j, h in enumerate(headers):
            c = t.rows[0].cells[j]
            c.text = h
            c.width = col_widths[j]
            # 表头下边框：细线
            tc_pr = c._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            tc_borders.append(make_border_element('w:bottom', 'single', '6', '0', '000000'))
            tc_pr.append(tc_borders)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = self.heading_font
                    r.font.size = Pt(12)
                    r.font.color.rgb = RGBColor(0, 0, 0)
                    r.font.bold = True
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), self.heading_font)

        # 数据行 — 五号仿宋
        for i, row in enumerate(rows, 1):
            for j, v in enumerate(row):
                if j < num_cols:
                    c = t.rows[i].cells[j]
                    c.text = str(v)
                    c.width = col_widths[j]
                    for p in c.paragraphs:
                        for r in p.runs:
                            r.font.name = self.body_font
                            r.font.size = Pt(10.5)
                            r.font.color.rgb = RGBColor(0, 0, 0)
                            r._element.rPr.rFonts.set(qn('w:eastAsia'), self.body_font)

        # 表后注（空行）
        p_after = self.doc.add_paragraph()
        p_after.paragraph_format.space_before = Pt(2)
        p_after.paragraph_format.space_after = Pt(6)

    def _add_table_spacing(self, before: bool = True) -> None:
        """v7.0: 表格前后间距优化 — 插入空段落调整间距。"""
        p = self.doc.add_paragraph()
        if before:
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
        else:
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)

"""
TechnicalBidGenerator - 技术标生成器 v7.4
架构：模块化设计，支持随机化、用户信息注入、detail_level深度控制、评分项对齐
v7.4 重构: 数据定义已分离至 table_data.py 和 work_content_maps.py
"""
import os
import re
from typing import Dict, List, Tuple

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from bid_core.base_generator import BaseGenerator
from bid_core.formatter import NormalFormatter
from bid_core.randomizer import Randomizer
from bid_core.user_context import UserContext
from bid_technical.scoring_strategy import ScoringStrategy
from bid_technical.evaluator_check import EvaluatorCheck
from bid_core.logger import get_logger
from pipeline.output_paths import aux_path

# 从独立模块导入数据定义
from table_data import (
    DEFAULT_CHAPTERS, SERVICE_CHAPTERS, TABLES, SERVICE_TABLES,
    ROUTES, ROUTE_PRIORITY,
)
from work_content_maps import (
    WORK_CONTENT_EQUIPMENT_MAP, WORK_CONTENT_INSTRUMENT_MAP,
    WORK_CONTENT_LABOR_MAP, WORK_CONTENT_MATERIAL_MAP,
    WORK_CONTENT_SAFETY_MAP, WORK_CONTENT_TRAINING_MAP,
    PROCESS_FLOW,
)

log = get_logger(__name__)


class TechnicalBidGenerator(BaseGenerator):
    """技术标生成器 v6.0"""

    # 必附表
    MANDATORY_TABLES = ['2.4']

    def __init__(self, project_info, chapters=None, target_pages=0,
                 formatter=None, randomizer=None, user_context=None,
                 heading_font=None, body_font=None,
                 detail_level=2, parse_result=None,
                 enable_deviation_table=False,
                 enable_risk_grading=True,
                 enable_mock_review=True,
                 llm_client=None, differentiator=None):
        """
        v4.0 新增参数:
        - detail_level: 内容深度 1/2/3
        - parse_result: parser.py解析结果
        """
        # v5.1 fix: 不再创建独立的self.doc，统一使用formatter.doc
        # 避免gantt chart等内容写入错误的Document对象
        if formatter is not None:
            self.doc = formatter.doc
        else:
            self.doc = Document()
            formatter = NormalFormatter(
                self.doc,
                heading_font=heading_font,
                body_font=body_font
            )

        super().__init__(project_info, formatter, randomizer, user_context,
                         heading_font=heading_font, body_font=body_font)

        self.detail_level = detail_level
        self.parse_result = parse_result
        self.enable_deviation_table = enable_deviation_table
        self.enable_risk_grading = enable_risk_grading
        self.enable_mock_review = enable_mock_review
        self.deviation_report = None  # P0: 供 P2 模拟评审复用
        self.llm_client = llm_client  # v7.2: 可选 LLM 扩写客户端
        # v7.3: 防重差异化引擎 + 企业图片库 + 章节生成报告聚合
        if differentiator is None:
            from bid_core.dedup import Differentiator
            differentiator = Differentiator(project_info=project_info or {})
        self.differentiator = differentiator
        self.image_library = []
        try:
            uc = self.user_context
            if hasattr(uc, 'get_images'):
                self.image_library = uc.get_images() or []
            elif isinstance(uc, dict):
                self.image_library = uc.get('images', []) or []
        except Exception:
            self.image_library = []
        self._chapter_gen_reports: List[Dict] = []

        # v6.0: target_pages默认300页（当传入0时自动设为300）
        if target_pages == 0:
            target_pages = 300

        # 确定标书类型
        self.bid_type = project_info.get('bid_type', 'construction')

        # 设置章节
        if chapters is None:
            if self.bid_type == 'service':
                self.chapters = SERVICE_CHAPTERS
            else:
                self.chapters = DEFAULT_CHAPTERS
        else:
            self.chapters = [{"title": c, "keywords": []} for c in chapters]

        self.target_pages = target_pages

        # v7.40: 补充段封顶改为可配置（默认沿用旧值，保证小标书行为不变；
        # 大标书可通过 project_info / user_context 抬高以撑满目标页数）。
        # - per_chapter_fill_cap: 每章补充段上限
        # - global_fill_cap: 全局补充段上限
        # - fill_para_per_page: 估算「每页需补多少段」用于按 shortfall 预算分配
        # 读取优先级：project_info > user_context > 默认值（兼容运行脚本把参数放在
        # user_context 中、而管线未将其灌入 project_info 的情况）。
        _uc = user_context if isinstance(user_context, dict) else {}
        self.per_chapter_fill_cap = int(
            project_info.get('per_chapter_fill_cap', _uc.get('per_chapter_fill_cap', 8)))
        self.global_fill_cap = int(
            project_info.get('global_fill_cap', _uc.get('global_fill_cap', 500)))
        self.fill_para_per_page = int(
            project_info.get('fill_para_per_page', _uc.get('fill_para_per_page', 10)))

        # 设置附表（深拷贝，防止污染全局TABLES/SERVICE_TABLES）
        if self.bid_type == 'service':
            self.tables = [dict(t, rows=[list(r) for r in t['rows']]) if 'rows' in t else dict(t) for t in SERVICE_TABLES]
        else:
            self.tables = [dict(t, rows=[list(r) for r in t['rows']]) if 'rows' in t else dict(t) for t in TABLES]

        # v6.0: 根据project_info的work_content动态调整附表
        self._build_tables_from_project()

    def render_chapter(self, chapter, plan_info=None):
        """渲染章节——实例化时传入detail_level、parse_result和plan_info"""
        title = chapter['title']
        keywords = chapter.get('keywords', [])

        # 路由匹配
        chapter_class = self._dispatch(title, keywords)

        if chapter_class:
            # v5.1: 实例化时传入 detail_level + parse_result + plan_info
            instance = chapter_class(
                self.formatter,
                self.randomizer,
                self.user_context,
                detail_level=self.detail_level,
                parse_result=self.parse_result,
                plan_info=plan_info,
                llm_client=self.llm_client,
                differentiator=self.differentiator,
                image_library=self.image_library,
            )
            # 完整生成路径中 generator 已添加 h1(title)，章节类只渲染 h2+ 内容
            # v7.31: 将真实章节标题注入 project_info，供 RichChapter._paragraph 的
            # 标题级门控（tech10/bim/smart/season/risk/.../labor 等）按章节相关性精准触发，
            # 否则 _chapter_title 回退为"本章内容"、标题级门控实际只命中 topic、分包/劳务等
            # 以章节标题触发的维度在真实生成中不可达。
            self.project_info['current_chapter_title'] = title
            instance.render(self.project_info, add_title=False)
            rep = getattr(instance, 'get_generation_report', None)
            if rep is not None:
                self._chapter_gen_reports.append(rep())

    def _match_keywords(self, title):
        """根据章节标题匹配路由关键词 v6.0"""
        for k in ROUTE_PRIORITY:
            if k in title:
                return [k]
        return []

    def _dispatch(self, title, keywords):
        """路由分发"""
        for k in ROUTE_PRIORITY:
            if k in title:
                return self._get_chapter_class(ROUTES[k])
        for k in keywords:
            if k in ROUTES:
                return self._get_chapter_class(ROUTES[k])
        return None

    def _get_chapter_class(self, path):
        """获取章节类；缺失时回退到 RichChapter 富内容引擎（v7.1）。"""
        from bid_technical.chapters.base import resolve_chapter_class
        return resolve_chapter_class(path)

    # ------------------------------------------------------------------
    # v6.0: 根据project_info动态调整附表
    # ------------------------------------------------------------------
    # 关键词 → 设备行映射（用于TABLES 2.1动态选择设备）

    def _build_tables_from_project(self):
        """根据project_info的work_content动态选择设备/人员/材料行

        核心逻辑：
        - 读取 project_info['work_content']，按关键词匹配
        - 将匹配到的行插入到对应附表的 rows 列表中
        - 去重：已存在于默认行中的不再重复添加
        - 施工类标书(TABLES): 处理2.1/2.2/2.3/2.8/2.9
        - 服务类标书(SERVICE_TABLES): 处理3.1
        """
        work_content = self.project_info.get('work_content', '')
        if not work_content:
            return

        # ---- 服务类：仅处理3.1培训计划表 ----
        if self.bid_type == 'service':
            table_31 = next((t for t in self.tables if t.get('num') == '3.1'), None)
            if table_31 and 'rows' in table_31:
                existing_names = {r[0] for r in table_31['rows']}
                extra_rows = []
                for keyword, rows in WORK_CONTENT_TRAINING_MAP.items():
                    if keyword in work_content:
                        for row in rows:
                            if row[0] not in existing_names:
                                extra_rows.append(row)
                                existing_names.add(row[0])
                table_31['rows'] = table_31['rows'] + extra_rows
            return

        # ---- 2.1 施工机械设备表 ----
        table_21 = next((t for t in self.tables if t.get('num') == '2.1'), None)
        if table_21 and 'rows' in table_21:
            existing_names = {r[0] for r in table_21['rows']}
            extra_rows = []
            for keyword, rows in WORK_CONTENT_EQUIPMENT_MAP.items():
                if keyword in work_content:
                    for row in rows:
                        if row[0] not in existing_names:
                            extra_rows.append(row)
                            existing_names.add(row[0])
            table_21['rows'] = table_21['rows'] + extra_rows

        # ---- 2.2 试验和检测仪器设备表 ----
        table_22 = next((t for t in self.tables if t.get('num') == '2.2'), None)
        if table_22 and 'rows' in table_22:
            existing_names = {r[0] for r in table_22['rows']}
            extra_rows = []
            for keyword, rows in WORK_CONTENT_INSTRUMENT_MAP.items():
                if keyword in work_content:
                    for row in rows:
                        if row[0] not in existing_names:
                            extra_rows.append(row)
                            existing_names.add(row[0])
            table_22['rows'] = table_22['rows'] + extra_rows

        # ---- 2.3 劳动力安排计划表 ----
        table_23 = next((t for t in self.tables if t.get('num') == '2.3'), None)
        if table_23 and 'rows' in table_23:
            existing_names = {r[0] for r in table_23['rows']}
            extra_rows = []
            for keyword, rows in WORK_CONTENT_LABOR_MAP.items():
                if keyword in work_content:
                    for row in rows:
                        if row[0] not in existing_names:
                            extra_rows.append(row)
                            existing_names.add(row[0])
            table_23['rows'] = table_23['rows'] + extra_rows

        # ---- 2.8 主要材料供应计划表 ----
        table_28 = next((t for t in self.tables if t.get('num') == '2.8'), None)
        if table_28 and 'rows' in table_28:
            existing_names = {r[0] for r in table_28['rows']}
            extra_rows = []
            for keyword, rows in WORK_CONTENT_MATERIAL_MAP.items():
                if keyword in work_content:
                    for row in rows:
                        if row[0] not in existing_names:
                            extra_rows.append(row)
                            existing_names.add(row[0])
            table_28['rows'] = table_28['rows'] + extra_rows

        # ---- 2.9 安全防护用品配备表 ----
        table_29 = next((t for t in self.tables if t.get('num') == '2.9'), None)
        if table_29 and 'rows' in table_29:
            existing_names = {r[0] for r in table_29['rows']}
            extra_rows = []
            for keyword, rows in WORK_CONTENT_SAFETY_MAP.items():
                if keyword in work_content:
                    for row in rows:
                        if row[0] not in existing_names:
                            extra_rows.append(row)
                            existing_names.add(row[0])
            table_29['rows'] = table_29['rows'] + extra_rows

    def render_table(self, table):
        """渲染附表"""
        self.formatter.h2(f"{table['num']} {table['title']}")

        if table['num'] in ('2.1', '2.2', '2.6', '2.7', '2.8', '2.9'):
            self.formatter.body(f"工程名称：{self.project_info.get('name', '')}")

        # 横道图特殊处理
        if table.get('auto_gantt'):
            from bid_technical.tables.gantt import generate_gantt_table
            generate_gantt_table(self.doc, self.project_info)
            self._render_gantt_notes()
            return

        # 普通表格
        if 'hdrs' in table and 'rows' in table:
            rows = [[str(i)] + row for i, row in enumerate(table['rows'], 1)]
            self.formatter.table(table['hdrs'], rows)

            if table['num'] == '2.6':
                total = sum(int(r[1]) for r in table['rows'])
                self.formatter.table(table['hdrs'], [["合计", str(total), "", "", "", ""]])

        # 表格说明
        self._render_table_notes(table)

    def _render_table_notes(self, table):
        """渲染表格说明"""
        if table['num'] == '2.3':
            if self.bid_type == 'service':
                self.formatter.body("注：1、服务进度计划应根据合同要求制定。2、实际进度根据项目实际情况调整。")
            else:
                self.formatter.body("注：1、本计划以每班八小时工作制为基础编制。2、实际投入劳动力根据工程进度动态调整。")
        elif table['num'] == '2.4':
            if self.bid_type == 'service':
                self.formatter.body("服务进度计划表应包含服务各阶段的工作内容、时间安排和质量标准。")
            else:
                self.formatter.body("1、施工进度表说明按招标文件要求的工期进行施工的各个关键日期。")
                self.formatter.body("2、施工进度表可采用网络图或横道图表示。")
                self.formatter.body("3、施工进度计划应与施工组织设计相适应。")
        elif table['num'] == '2.5':
            self.formatter.body("投标人递交一份施工总平面图，绘出现场临时设施布置图表并附文字说明。")
        elif table['num'] == '2.6':
            self.formatter.body("注：1、投标人应逐项填写本表。2、若本表不够，可加附页。")
        elif table['num'] == '2.7':
            self.formatter.body("注：1、上述人员均为本公司专职人员，中标后未经招标人同意不得更换。2、主要管理人员必须持证上岗。")
        elif table['num'] == '2.8':
            self.formatter.body("注：1、材料进场时间根据施工进度计划动态调整。2、品牌、规格如招标文件有指定，按指定执行。3、若本表不够，可加附页。")
        elif table['num'] == '2.9':
            self.formatter.body("注：1、所有安全防护用品必须符合国家相关标准。2、损坏的防护用品应及时更换，不得继续使用。3、特种作业人员须配备专用防护用品。")
        elif table['num'] == '3.1':
            self.formatter.body("注：1、所有服务人员必须经岗前培训合格后方可上岗。2、培训记录应完整归档备查。3、特种作业人员须持证上岗，证书须在有效期内。")
        elif table['num'] == '3.4':
            self.formatter.body("注：1、服务质量考核按满分100分制执行。2、考核结果与服务费用挂钩，具体扣款标准按合同约定执行。")

    def _render_gantt_notes(self):
        """渲染横道图说明"""
        name = self.project_info.get('name', '本项目')
        duration = self.project_info.get('duration', 90)

        self.formatter.body(f"工程名称：{name}")
        self.formatter.body(f"总工期{duration}日历天，计划开工日期以监理单位开工令为准。")
        self.formatter.body("1、施工进度表说明按招标文件要求的工期进行施工的各个关键日期。")
        self.formatter.body("2、施工进度表可采用网络图或横道图表示。")
        self.formatter.body("3、施工进度计划应与施工组织设计相适应。")
        self.formatter.body("4、■表示该分项工程施工时段。")

    # ------------------------------------------------------------------
    # v6.0: 内容填充 — 估算已渲染页数 + 自动补充内容
    # ------------------------------------------------------------------
    def _estimate_rendered_pages(self):
        """粗估已渲染的页数：基于doc中的段落数和表格数"""
        para_count = len(self.doc.paragraphs)
        table_count = len(self.doc.tables)
        # 粗估：约25个段落≈1页，每个表格≈2页
        estimated = para_count / 25.0 + table_count * 2
        return int(estimated)

    def _fill_content_to_target(self, shortfall, chapters_to_render):
        """内容不足时补充表格与段落，按章节分散插入（修正单章爆炸 bug）。

        旧实现把所有补充段无差别追加到文档末尾，而末章 H1 之后即为其落点，
        导致全部补充内容被计入最后一章（如第十三章达 900+ 段、几十页）。
        新实现：
        1. 仅对正文主章节（附表之前）做补充；
        2. 每章补充段数封顶（PER_CHAPTER_CAP），全局封顶（GLOBAL_CAP）；
        3. 通过 insert_paragraph_before(下一章 H1) 把补充段插到本章程尾，
           分散到各章，杜绝集中到末章造成单章几百页。
        """
        name = self.project_info.get('name', '本项目')
        plan_target = self.plan.get('total_pages', 0) if getattr(self, 'plan', None) else 0
        if plan_target <= 0 or not chapters_to_render:
            return

        para_list = self.doc.paragraphs
        # 注意：本方法在 render_tables（生成附表 H1）之前调用，
        # 因此此时文档中仅有 13 个正文章节 H1，附表尚未生成。
        all_h1 = [i for i, p in enumerate(para_list) if p.style.name == 'Heading 1']
        if not all_h1:
            return

        # ── 补充表格（有限，避免无限膨胀）──
        supplement_tables = self._build_supplement_tables()
        tables_to_add = min(len(supplement_tables), 3, max(0, int(shortfall / 6)))
        for i in range(tables_to_add):
            st = supplement_tables[i]
            self.formatter.h2(f"附表 补-{i + 1} {st['title']}")
            self.formatter.body(f"工程名称：{name}")
            if 'hdrs' in st and 'rows' in st:
                rows = [[str(j)] + row for j, row in enumerate(st['rows'], 1)]
                self.formatter.table(st['hdrs'], rows)
            self.formatter.body("")

        # ── 按章分散补段落（插到本章程尾，封顶；v7.40 改为按 shortfall 预算驱动）──
        PER_CHAPTER_CAP = self.per_chapter_fill_cap
        GLOBAL_CAP = self.global_fill_cap
        PARA_PER_PAGE = max(1, self.fill_para_per_page)
        # v7.42: LLM 扩写填充（云端+脱敏 / 本地）。有客户端且开关开启时，用「角度化要点」
        # 调用 expand_section 替换八股模板，根治逐字重复；单次失败/超预算立即回退模板。
        # 注意：expand_section 按 (标题, 要点, ctx) 缓存，故每次要点必须不同，否则返回重复文本。
        LLM_FILL_BUDGET = int(self.project_info.get('llm_fill_call_budget', 400))
        enable_llm_fill = bool(self.llm_client) and bool(
            self.project_info.get('enable_llm_fill', True))
        titles = [c.get('title', '') for c in chapters_to_render]
        openers = [
            '关于{t}，我司结合{name}的实际情况进一步细化管控措施，确保全过程可控、可追溯、可核查。',
            '在{t}方面，我司建立清单化、节点化的管理机制，将目标分解到岗、责任传导到人。',
            '针对{t}，我司组织专项技术交底与培训，确保一线作业人员掌握控制要点与验收标准。',
            '围绕{t}的落实，我司实行"策划先行、样板引路、过程严控"，关键工序编制专项作业指导书。',
            '对于{t}，我司与监理、业主建立定期联合检查机制，对执行偏差早发现、早预警、早处置。',
            '就{t}而言，我司将其纳入项目全生命周期管理，前置策划、过程留痕、闭环验证。',
            '在{t}的执行上，我司贯彻"标准化、精细化、信息化"的管理原则，确保实施有据可依。',
            '针对{t}，我司以{name}的实际工况为出发点，配置专职管理与作业力量。',
        ]
        # 角度池：每次填充调用喂入不同角度，既让 LLM 产出差异化正文，又避免命中缓存返回重复文本
        ANGLES = [
            "从组织架构与责任分工角度展开",
            "从关键施工工艺与技术路线角度展开",
            "从质量管控与检验批验收角度展开",
            "从安全生产与应急管理体系角度展开",
            "从进度节点与资源配置计划角度展开",
            "从文明施工与环境保护角度展开",
            "从风险预控与应急预案角度展开",
            "从材料设备管理与进场检验角度展开",
            "从信息化与BIM协同应用角度展开",
            "从样板引路与工序三检角度展开",
            "从成本意识与绿色施工角度展开",
            "从劳务管理与技能培训角度展开",
        ]
        # 预算：需补充的段落总数 = shortfall(页) × 每页段数，再按章均分（封顶保护）
        num_ch = max(1, len(all_h1))
        total_needed = max(0, int(shortfall * PARA_PER_PAGE))
        per_ch = max(1, (total_needed + num_ch - 1) // num_ch)  # 向上取整均分
        per_ch = min(per_ch, PER_CHAPTER_CAP)
        total_needed = min(total_needed, GLOBAL_CAP)
        added = 0
        llm_calls = 0
        for ci in range(len(all_h1)):
            if added >= total_needed:
                break
            # 下一章 H1 作为本章程尾插入点；末章插到文档末尾（后续附表自然接后）
            nxt_obj = para_list[all_h1[ci + 1]] if ci + 1 < len(all_h1) else None
            title = titles[ci] if ci < len(titles) else para_list[all_h1[ci]].text.strip()
            in_ch = 0
            while in_ch < per_ch and added < total_needed:
                salt = f"{title}#fill{added}"
                # —— LLM 扩写优先 ——
                llm_text = None
                if enable_llm_fill and llm_calls < LLM_FILL_BUDGET:
                    angle = ANGLES[added % len(ANGLES)]
                    try:
                        llm_text = self.llm_client.expand_section(
                            title, [angle], self.project_info, self.parse_result)
                    except Exception:
                        llm_text = None
                    if llm_text:
                        llm_calls += 1
                if llm_text:
                    sub = [p.strip() for p in re.split(r'\n\s*\n', llm_text) if p.strip()]
                    if not sub:
                        sub = [llm_text.strip()]
                    for sp in sub:
                        if added >= total_needed or in_ch >= per_ch:
                            break
                        rotated = self.differentiator.rotate(sp, salt=f"{salt}#{in_ch}")
                        self.differentiator.add_sentence(rotated)
                        self._insert_filler_para(nxt_obj, rotated)
                        added += 1
                        in_ch += 1
                    continue  # 本段预算已由 LLM 填充，跳过模板分支
                # —— 模板回退（LLM 不可用/失败/超额）——
                text = openers[added % len(openers)].format(t=title, name=name)
                rotated = self.differentiator.rotate(text, salt=salt)
                self.differentiator.add_sentence(rotated)
                self._insert_filler_para(nxt_obj, rotated)
                added += 1
                in_ch += 1

    def _insert_filler_para(self, nxt_obj, text: str) -> None:
        """在章节 H1 前（或文档末尾）插入一段填充正文，统一字体/对齐格式。

        v7.42：显式强制 LEFT 对齐，防止继承 Normal 样式的 JUSTIFY/DISTRIBUTE
        把中文拉伸（双重保险，DocxSanitizeStage._force_all_left 会再扫一遍）。
        """
        if nxt_obj is not None:
            new_p = nxt_obj.insert_paragraph_before(text)
        else:
            new_p = self.doc.add_paragraph(text)
        try:
            new_p.style = self.doc.styles['Normal']
        except Exception:
            pass
        if new_p.runs:
            rr = new_p.runs[0]
            rr.font.name = self.formatter.body_font
            rr.font.size = Pt(12)
            rr.font.color.rgb = RGBColor(0, 0, 0)
            rr._element.rPr.rFonts.set(qn('w:eastAsia'), self.formatter.body_font)
        new_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        new_p.paragraph_format.first_line_indent = Cm(0.85)
        new_p.paragraph_format.line_spacing = Pt(28)
        new_p.paragraph_format.space_before = Pt(0)
        new_p.paragraph_format.space_after = Pt(2)

    def _append_supplement_paragraphs(self, title, n):
        """ADR-007: 围绕章节标题补写 n 段通用但项目相关的说明，用于撑满目标页数。"""
        proj = self.project_info.get('name', '本工程')
        comp = '我司'
        openers = [
            f'关于{title}，{comp}结合{proj}的实际情况进一步细化管控措施，确保全过程可控、可追溯、可核查。',
            f'在{title}方面，{comp}建立清单化、节点化的管理机制，将目标分解到岗、责任传导到人。',
            f'针对{title}，{comp}组织专项技术交底与培训，确保一线作业人员掌握控制要点与验收标准。',
            f'围绕{title}的落实，{comp}实行"策划先行、样板引路、过程严控"，关键工序编制专项作业指导书。',
            f'对于{title}，{comp}与监理、业主建立定期联合检查机制，对执行偏差早发现、早预警、早处置。',
            f'就{title}而言，{comp}将其纳入项目全生命周期管理，前置策划、过程留痕、闭环验证。',
        ]
        for i in range(n):
            self.formatter.body(openers[i % len(openers)])

    def _build_supplement_tables(self):
        """构建补充表格数据（v6.0新增）"""
        if self.bid_type == 'service':
            return [
                {"title": "服务质量巡检记录表",
                 "hdrs": ["序号", "巡检日期", "巡检区域", "巡检内容", "发现问题", "处理措施", "巡检人", "备注"],
                 "rows": [
                     ["第1周", "全区域", "日常巡检", "无异常", "正常记录", "巡检员A", ""],
                     ["第2周", "重点区域", "深度巡检", "轻微磨损", "即时维护", "巡检员B", ""],
                     ["第3周", "全区域", "日常巡检", "无异常", "正常记录", "巡检员A", ""],
                     ["第4周", "设备区域", "专项巡检", "设备异响", "安排检修", "巡检员C", "已报修"],
                     ["第5周", "全区域", "日常巡检", "无异常", "正常记录", "巡检员A", ""],
                     ["第6周", "重点区域", "深度巡检", "管道渗漏", "紧急修复", "巡检员B", "已处理"],
                     ["第7周", "全区域", "日常巡检", "无异常", "正常记录", "巡检员A", ""],
                     ["第8周", "安全区域", "安全巡检", "消防器材过期", "立即更换", "安全员D", "已更换"],
                 ]},
                {"title": "服务人员培训计划表",
                 "hdrs": ["序号", "培训内容", "培训方式", "培训时间", "培训学时", "授课人", "考核方式", "备注"],
                 "rows": [
                     ["安全操作规程", "集中授课+实操", "第1月", "8", "安全主管", "闭卷考试", ""],
                     ["专业技能提升", "现场实训", "第2月", "16", "技术主管", "实操考核", ""],
                     ["应急预案演练", "桌面推演+实战", "第3月", "8", "应急负责人", "演练评估", ""],
                     ["服务规范培训", "集中授课", "第4月", "4", "服务主管", "闭卷考试", ""],
                     ["设备操作培训", "现场实训", "第5月", "8", "设备主管", "实操考核", ""],
                     ["质量管理培训", "集中授课", "第6月", "4", "质量主管", "闭卷考试", ""],
                 ]},
                {"title": "服务费用明细表",
                 "hdrs": ["序号", "费用项目", "计算基数", "费率/单价", "金额(元)", "备注"],
                 "rows": [
                     ["人工费", "服务人员工资", "按月计算", "按合同约定", ""],
                     ["设备费", "设备折旧+维护", "按年计算", "含易损件", ""],
                     ["材料费", "消耗材料", "按次计算", "按实际发生", ""],
                     ["管理费", "项目管理", "人工费×费率", "8%", ""],
                     ["利润", "合理利润", "成本×费率", "5%", ""],
                     ["税金", "增值税", "不含税金额×税率", "6%", ""],
                 ]},
            ]
        else:
            return [
                {"title": "主要检验批质量验收记录表",
                 "hdrs": ["序号", "分项工程", "检验批数量", "合格率(%)", "一次验收合格率(%)", "验收结论", "备注"],
                 "rows": [
                     ["楼地面工程", "36", "98.5", "96.0", "合格", ""],
                     ["墙柱面工程", "42", "97.8", "95.2", "合格", ""],
                     ["天棚工程", "28", "99.1", "97.5", "合格", ""],
                     ["门窗工程", "24", "98.0", "96.8", "合格", ""],
                     ["涂饰工程", "30", "98.3", "95.5", "合格", ""],
                     ["防水工程", "18", "100", "98.0", "合格", ""],
                     ["给排水工程", "22", "99.0", "97.2", "合格", ""],
                     ["电气工程", "26", "98.5", "96.5", "合格", ""],
                 ]},
                {"title": "安全生产费用投入计划表",
                 "hdrs": ["序号", "费用项目", "计算基数", "费率/单价", "金额(万元)", "备注"],
                 "rows": [
                     ["安全防护用品", "按人数计算", "800元/人/年", "12.0", "含安全帽/安全带/防护服"],
                     ["安全设施", "按面积计算", "15元/㎡", "18.0", "含临边防护/安全网"],
                     ["安全教育", "按人数计算", "500元/人/年", "7.5", "含培训/考试/证书"],
                     ["安全检查", "按月计算", "3000元/月", "3.6", "含日常/专项/季节性检查"],
                     ["应急救援", "按项计算", "5万元/项", "10.0", "含物资/设备/演练"],
                     ["文明施工", "按面积计算", "10元/㎡", "12.0", "含围挡/标识/洗车槽"],
                     ["职业健康", "按人数计算", "600元/人/年", "9.0", "含体检/劳保/监测"],
                     ["消防设施", "按项计算", "3万元/项", "6.0", "含灭火器/消火栓/报警"],
                 ]},
                {"title": "主要材料进场检验记录表",
                 "hdrs": ["序号", "材料名称", "规格型号", "进场数量", "检验日期", "检验结果", "检验人", "备注"],
                 "rows": [
                     ["水泥", "P.O42.5", "100吨", "第5天", "合格", "质检员A", "3天强度合格"],
                     ["钢筋", "HRB400 Φ12-25", "80吨", "第3天", "合格", "质检员A", "力学性能合格"],
                     ["防水卷材", "SBS 4mm", "2000㎡", "第30天", "合格", "质检员B", "不透水性合格"],
                     ["涂料", "乳胶漆", "500桶", "第45天", "合格", "质检员B", "环保检测合格"],
                     ["瓷砖", "800×800", "3000㎡", "第40天", "合格", "质检员A", "吸水率合格"],
                     ["管材(给水)", "PPR Φ20-63", "2000米", "第10天", "合格", "质检员C", "耐压试验合格"],
                     ["电缆", "YJV 3×4-3×70", "3000米", "第15天", "合格", "质检员C", "绝缘电阻合格"],
                     ["型钢", "Q235B", "30吨", "第5天", "合格", "质检员A", "力学性能合格"],
                 ]},
                {"title": "施工机械设备维修保养计划表",
                 "hdrs": ["序号", "设备名称", "型号规格", "保养周期", "保养内容", "责任人", "备注"],
                 "rows": [
                     ["电焊机", "BX3-300", "每月", "检查焊把线/更换焊嘴", "机修工A", ""],
                     ["空压机", "V-0.6/8", "每季", "更换机油/清洗滤芯", "机修工B", ""],
                     ["砂浆搅拌机", "UJZ-200", "每月", "检查叶片/润滑轴承", "机修工A", ""],
                     ["吊篮", "ZLP630", "每周", "检查安全锁/钢丝绳", "安全员C", "特种检查"],
                     ["混凝土搅拌机", "JS500", "每季", "更换衬板/检查传动", "机修工B", ""],
                     ["切割机", "J3G-400", "每月", "更换锯片/检查防护罩", "机修工A", ""],
                 ]},
                {"title": "环境保护措施及投入计划表",
                 "hdrs": ["序号", "环保项目", "控制标准", "采取的措施", "投入金额(万元)", "责任人", "备注"],
                 "rows": [
                     ["扬尘控制", "PM10≤150μg/m³", "洒水降尘/围挡/覆盖", "5.0", "环保负责人", ""],
                     ["噪声控制", "昼间≤70dB/夜间≤55dB", "隔音围挡/限时施工", "3.0", "环保负责人", ""],
                     ["废水处理", "COD≤150mg/L", "沉淀池/隔油池", "4.0", "环保负责人", ""],
                     ["固废处理", "分类收集/合规处置", "分类垃圾桶/委托清运", "2.0", "环保负责人", ""],
                     ["废气排放", "达标排放", "通风设施/净化设备", "3.0", "环保负责人", ""],
                 ]},
            ]

    def _build_supplement_paragraphs(self):
        """构建补充段落模板（v6.0新增）"""
        if self.bid_type == 'service':
            return [
                "在{title}方面，我方将建立完善的管理体系，确保{name}的服务质量达到行业领先水平。通过科学的管理方法和严格的考核制度，持续提升服务品质。",
                "针对{title}，我方制定了详细的实施方案和保障措施。项目团队将严格按照合同要求执行，确保各项工作按时、保质完成。",
                "为确保{title}的有效实施，我方将定期组织专项检查和评估，及时发现并解决问题，确保服务全过程可控、在控。",
                "在{title}工作中，我方注重与业主的沟通协调，建立定期汇报机制，确保信息畅通、响应及时。对业主提出的意见和建议，将在24小时内予以回复。",
                "我方将结合{title}的实际情况，制定针对性的培训计划，提升服务人员的专业素质和服务意识，为{name}提供优质高效的服务保障。",
                "针对{title}可能出现的风险和问题，我方已制定完善的应急预案，确保突发情况下能够快速响应、妥善处理，最大限度减少对{title}工作的影响。",
                "在{title}实施过程中，我方将严格遵守国家相关法律法规和行业标准，确保各项操作规范合法，保障{title}工作的顺利推进。",
                "我方将建立{title}的完整台账和档案管理制度，做到事前有计划、事中有记录、事后有总结，确保全过程可追溯、可核查。",
                "为确保{title}的持续改进，我方将建立PDCA循环管理机制，通过计划-执行-检查-处理的闭环管理，不断优化服务流程和提升服务质量。",
                "我方承诺在{title}方面投入充足的人力、物力和财力资源，确保各项工作有序开展，为{name}提供全方位、高质量的服务保障。",
            ]
        else:
            return [
                "在{title}方面，我方将严格执行国家现行施工规范和验收标准，确保{title}的各项指标满足设计要求和规范规定，为{name}的顺利实施提供坚实保障。",
                "针对{title}，我方编制了详细的施工方案和技术措施，明确了施工工艺流程、质量控制要点和安全防护措施，确保施工全过程规范有序。",
                "为确保{title}的施工质量，我方将实行三检制（自检、互检、专检），严把每道工序质量关，做到上道工序不合格不进入下道工序。",
                "在{title}施工过程中，我方将加强技术交底工作，确保每名作业人员都熟悉施工方法和质量要求，做到人人心中有标准、手中有规范。",
                "我方将定期对{title}进行质量检查和安全巡查，对发现的问题立即整改，确保施工质量和安全始终处于受控状态，杜绝质量安全隐患。",
                "针对{title}的特殊施工条件，我方已制定专项施工方案和应急预案，确保在不利条件下的施工安全和工程质量，保障{name}按期交付。",
                "在{title}实施中，我方将采用先进的施工技术和工艺，提高施工效率和质量水平，同时注重绿色施工和节能减排，践行可持续发展理念。",
                "我方将建立{title}的施工日志和技术档案，详细记录施工过程中的关键参数和质量数据，确保施工全过程可追溯、可核查。",
                "为确保{title}的顺利实施，我方将合理调配人力、物力和机械设备，确保资源投入满足施工需要，为{name}的按期完工提供有力支撑。",
                "我方承诺严格按照合同约定和审批的施工方案执行{title}，接受业主和监理的监督检查，确保施工质量和进度满足合同要求。",
            ]

    def generate(self, output_path):
        """生成标书 v5.1: 接入planner，评分项动态组装章节
        v6.1: 集成ScoringStrategy增强plan_info + EvaluatorCheck生成后自检
        流程: plan → scoring_strategy增强 → 封面 → 目录 → 遍历plan chapters → 附表 → 页码 → 自检
        """
        from planner import plan_chapters

        # 1. 调用planner规划章节（在before_render之前，以便封面/目录可用plan信息）
        plan = plan_chapters(
            parse_result=self.parse_result,
            target_pages=self.target_pages,
            project_info=self.project_info,
        )
        self.plan = plan  # 保存plan供下游使用

        # 1.5 v6.1: 用ScoringStrategy增强plan_info
        self._enhance_plan_with_scoring_strategy(plan)

        # 2. 附表页数预留：从target_pages中扣除table_pages
        self.table_pages = 20  # v6.0: 与planner保持一致，附表预留20页
        self.content_target_pages = max(self.target_pages - self.table_pages, 0) if self.target_pages > 0 else 0

        # 3. 获取planned chapters，为空时兜底DEFAULT_CHAPTERS/SERVICE_CHAPTERS
        planned_chapters = plan.get('chapters', [])
        chapters_to_render = []
        if planned_chapters:
            # planner返回了章节规划（评分项/用户指定/默认模板均可）
            for pc in planned_chapters:
                # 为planner章节匹配keywords（用于路由分发）
                matched_keywords = self._match_keywords(pc['title'])
                chapters_to_render.append({
                    'title': pc['title'],
                    'keywords': matched_keywords,
                    'plan_info': pc,  # 包含target_pages/score_weight/page_ratio等
                })
        else:
            # 兜底：planner返回为空时使用DEFAULT/SERVICE_CHAPTERS
            for ch in self.chapters:
                chapters_to_render.append({
                    'title': ch['title'],
                    'keywords': ch.get('keywords', []),
                    'plan_info': None,
                })

        # 4. 封面、目录（before_render）
        self.before_render()

        # 5. 遍历plan chapters渲染
        for i, ch_info in enumerate(chapters_to_render, 1):
            self.formatter.h1(i, ch_info['title'])
            self.render_chapter(ch_info, plan_info=ch_info.get('plan_info'))

        # 5.5 v6.0: 内容填充循环 — 如果已渲染内容不足target_pages，自动补充表格和段落
        if self.target_pages > 0 and self.plan:
            plan_target = self.plan.get('total_pages', 0)
            if plan_target > 0:
                # 估算当前已渲染页数（基于段落和表格数量粗估）
                estimated_pages = self._estimate_rendered_pages()
                shortfall = plan_target - estimated_pages
                if shortfall > 0:
                    self._fill_content_to_target(shortfall, chapters_to_render)

        # 6. 渲染附表（保留原有逻辑）
        self.render_tables()

        # 6.1 v7.14: 施工组织核心附表（组织机构 / 设备材料投入 / 总平面布置）
        self._render_org_resource_tables()

        # 6.2 v7.19: 施工进度计划横道图（纯 docx 表格，可视化进度，对标红点智标图表能力）
        self._render_schedule_gantt()

        # 6.3 v7.20: 施工工艺流程及关键控制要点（按工程类型定制工序链，对标红点智标工艺流程图）
        self._render_process_flow()

        # 6.5 P0: 自动生成偏离表（附录，紧贴附表之后）
        self._render_deviation_table()

        # 7. 页码（after_render）
        self.after_render()

        # 保存
        self.formatter.save(output_path)

        # 8. v6.1: EvaluatorCheck生成后自检（不阻断生成，结果写入日志/返回信息）
        self._post_generate_evaluator_check(output_path)

        return output_path

    # ────────────────────────────────────────────────────────────
    # v7.14: 施工组织核心附表（组织机构 / 设备材料投入 / 总平面布置）
    # 对标喜鹊"自动生成附表"，且数据随工程规模推导
    # ────────────────────────────────────────────────────────────
    def _render_org_resource_tables(self) -> None:
        try:
            if self.bid_type == 'service':
                self._render_service_org_table()
            else:
                self._render_construction_org_table()
                self._render_equipment_table()
                self._render_material_table()
                self._render_site_layout_text()
        except Exception as e:
            log.warning('施工组织核心附表生成失败（已跳过）: %s', e)

    def _project_scale(self):
        return self.project_info.get('area') or 0, self.project_info.get('duration') or 0

    def _render_construction_org_table(self) -> None:
        area, _ = self._project_scale()
        base = 8
        extra = max(0, area // 5000)
        team = base + extra
        rows = [
            ['项目经理', '1', '全面负责项目履约、协调与决策'],
            ['技术负责人', '1', '施工组织设计、技术方案与交底'],
            ['生产/施工经理', '1', '现场生产调度与进度管控'],
            ['质量负责人', '1', '质量体系运行与检验批验收'],
            ['安全负责人', '1', '安全生产与文明施工管理'],
            ['材料/设备负责人', '1', '物资采购、进场验收与调配'],
            ['造价/合同负责人', '1', '计量、变更与合同管理'],
            ['专职管理人员', str(max(2, team - 7)), '现场技术、质量、安全、资料等专职岗'],
        ]
        self.formatter.add_heading('项目管理机构及人员配置')
        self.formatter.table(['岗位', '人数', '主要职责'], rows)
        self.formatter.body(f'注：本项目拟投入管理人员合计约 {team} 人，关键岗位一岗双人、持证上岗。')

    def _render_service_org_table(self) -> None:
        rows = [
            ['项目负责人', '1', '整体统筹与业主对接'],
            ['现场主管', '1', '日常运维与班组调度'],
            ['技术人员', '2', '技术巡检与故障处置'],
            ['客服/回访', '1', '工单受理与满意度回访'],
            ['安全员', '1', '现场安全与应急预案'],
        ]
        self.formatter.add_heading('项目服务团队架构')
        self.formatter.table(['岗位', '人数', '主要职责'], rows)

    def _render_equipment_table(self) -> None:
        eq = []
        try:
            eq = self.user_context.get_equipment() or []
        except Exception:
            eq = []
        if not eq:
            area, _ = self._project_scale()
            eq = [
                {'name': '塔式起重机', 'model': 'QTZ80', 'count': max(1, area // 8000)},
                {'name': '施工升降机', 'model': 'SC200', 'count': max(1, area // 10000)},
                {'name': '混凝土输送泵', 'model': 'HBT60', 'count': max(1, area // 6000)},
                {'name': '钢筋加工机械', 'model': 'GW40', 'count': max(2, area // 4000)},
                {'name': '电焊机', 'model': 'BX3-300', 'count': max(3, area // 3000)},
                {'name': '木工圆锯', 'model': 'MJ104', 'count': max(2, area // 4000)},
            ]
        # 兼容两种数据格式：dict 列表（{name,model,count}）或纯字符串列表（如用户直填设备名）
        def _eq_row(e):
            if isinstance(e, str):
                return [e, '-', '1']
            if isinstance(e, dict):
                return [e.get('name', ''), e.get('model', '-'), str(e.get('count', 1))]
            return ['', '-', '1']
        rows = [_eq_row(e) for e in eq[:12]]
        self.formatter.add_heading('主要施工设备投入计划')
        self.formatter.table(['设备名称', '规格型号', '数量'], rows)
        self.formatter.body('注：主要机械设备进场报验、一用一备，按进度动态调配，确保不因设备原因影响工期。')

    def _render_material_table(self) -> None:
        area, _ = self._project_scale()
        if not area:
            return
        steel = round(area * 0.055)
        concrete = round(area * 0.38)
        formwork = round(area * 2.2)
        scaffold = round(area * 1.5)
        rows = [
            ['钢筋', f'约 {steel:g} t', '按进度分批进场、复试合格后方可使用'],
            ['商品混凝土', f'约 {concrete:g} m³', '就近站点供应、随拌随用'],
            ['模板', f'约 {formwork:g} ㎡', '多层板+木方，周转使用'],
            ['脚手架钢管', f'约 {scaffold:g} ㎡', '扣件式钢管，进场报验'],
        ]
        self.formatter.add_heading('主要材料投入计划')
        self.formatter.table(['材料名称', '计划用量', '供应与质控'], rows)

    def _render_site_layout_text(self) -> None:
        area, _ = self._project_scale()
        self.formatter.add_heading('施工总平面布置说明')
        self.formatter.body(
            '施工现场实行"办公区—生活区—作业区"三区分离布置：办公区设项目部会议室、'
            '资料室与业主/监理联合办公室；生活区设工人宿舍、食堂与盥洗区，与作业区保持安全距离；'
            '作业区按专业划分加工棚、材料堆场与周转区，钢筋/模板加工棚就近布置以减少二次搬运。'
            '场地设两个出入口，主干道硬化并环形贯通，材料堆场临近塔吊覆盖半径以内。'
            '临时用电采用TN-S系统三级配电两级保护，临时给水沿场区环状布置并设消防栓，'
            '消防通道宽≥4m并保持畅通。'
        )
        if area:
            self.formatter.body(f'本项目建筑面积约 {area}㎡，总平面按上述原则统筹，确保物流顺畅、'
                                f'安全可控、文明达标。')

    # ────────────────────────────────────────────────────────────
    # v7.19: 施工进度计划横道图（纯 docx 表格渲染，离线安全、零新依赖）
    # 对标 红点智标"施工横道图/平面布置图/工艺流程图生成"差异化能力；
    # 喜鹊/通用模板多无可视化进度图，仅文字罗列工序。
    # ────────────────────────────────────────────────────────────
    def _render_schedule_gantt(self) -> None:
        if self.bid_type == 'service':
            return
        try:
            dur = self.project_info.get('duration') or 0
            months = max(1, -(-int(dur) // 30)) if dur else 12  # ceil(天/30)
            months = min(months, 24)
            # 关键线路阶段（比例累计，末阶段收口至 100%）
            phases = [
                ('施工准备', 0.00, 0.06),
                ('基础/基坑工程', 0.06, 0.26),
                ('主体结构施工', 0.26, 0.62),
                ('机电安装与调试', 0.55, 0.90),
                ('装饰装修工程', 0.60, 0.88),
                ('竣工验收与交付', 0.90, 1.00),
            ]
            headers = ['施工阶段', '开始(月)', '结束(月)'] + [f'第{i+1}月' for i in range(months)]
            rows = []
            for name, s, e in phases:
                start_m = max(1, min(months, int(round(s * months)) + 1))
                end_m = max(start_m, min(months, int(round(e * months)) + 1))
                bar = ['■' if start_m <= m <= end_m else '' for m in range(1, months + 1)]
                rows.append([name, str(start_m), str(end_m)] + bar)
            self.formatter.add_heading('施工进度计划横道图')
            self.formatter.table(headers, rows)
            self.formatter.body(
                f'注：本工程总工期约 {dur if dur else "合同"} 天（{months} 个月），横道图按关键线路与里程碑'
                f'节点编排，主体与装饰/机电穿插施工、实行动态纠偏与预警，确保总控目标不突破。')
        except Exception as e:
            log.warning('进度横道图生成失败（已跳过）: %s', e)

    # ────────────────────────────────────────────────────────────
    # v7.20: 施工工艺流程及关键控制要点（纯 docx 表格，按工程类型定制工序链）
    # 对标 红点智标"工艺流程图生成"差异化能力；喜鹊/模板多无工艺流程图。
    # ────────────────────────────────────────────────────────────

    def _render_process_flow(self) -> None:
        if self.bid_type == 'service':
            return
        try:
            flow = PROCESS_FLOW.get(self.bid_type) or PROCESS_FLOW['construction']
            rows = [[f'①{i+1}', name, ctrl] for i, (name, ctrl) in enumerate(flow)]
            self.formatter.add_heading('施工工艺流程及关键控制要点')
            self.formatter.table(['步骤', '工序', '关键控制要点'], rows)
            self.formatter.body('注：上述工艺流程按本工程类型定制，关键工序实行样板引路、首件认可与'
                                '三检闭环，确保工序衔接顺畅、质量可控。')
        except Exception as e:
            log.warning('施工工艺流程图生成失败（已跳过）: %s', e)

    def _render_deviation_table(self):
        """P0: 从 parse_result 自动抽取实质性/星号/废标条款并生成偏离表。

        作为附录紧接附表之后渲染；任何异常均降级跳过，不影响主流程。
        """
        if not self.parse_result or not self.enable_deviation_table:
            return
        try:
            from bid_core.deviation_checker import DeviationChecker
            checker = DeviationChecker(self.parse_result, self.user_context)
            report = checker.generate()
            self.deviation_report = report
            if not report.get('items'):
                log.info('偏离表: 未提取到需比对的条款，跳过')
                return
            from bid_technical.tables.deviation_table import generate_deviation_table
            self.doc.add_page_break()
            self.formatter.add_heading('投标文件偏离表')
            generate_deviation_table(self.doc, self.project_info, report)
            log.info('已生成偏离表: %d 项条款, 风险等级=%s',
                     report.get('total_requirements', 0), report.get('risk_level'))
            # v7.5: 评分项响应保障表（紧接偏离表之后）
            try:
                from bid_core.score_response import build_score_response_map
                from bid_technical.tables.score_response_table import render_score_response_table
                srmap = build_score_response_map(self.parse_result, self.user_context)
                if srmap.get('rows'):
                    self.doc.add_page_break()
                    self.formatter.add_heading('评分项响应保障表')
                    render_score_response_table(self.doc, self.project_info, srmap)
                    log.info('已生成评分项响应保障表: 覆盖率=%d%%', srmap.get('coverage', 0))
            except Exception as exc:
                log.warning('评分项响应保障表生成失败（已跳过）: %s', exc, exc_info=True)
        except Exception as exc:
            log.warning('偏离表生成失败（已跳过，不影响标书）: %s', exc, exc_info=True)

    def _enhance_plan_with_scoring_strategy(self, plan):
        """v6.1: 用ScoringStrategy增强plan_info，补充content_strategy和must_have_content"""
        try:
            strategy = ScoringStrategy(project_type=self.bid_type)
        except Exception:
            return

        planned_chapters = plan.get('chapters', [])
        if not planned_chapters:
            return

        # 收集所有评分项，计算总分
        all_score_items = self.parse_result.get('score_items', []) if self.parse_result else []
        total_weight = sum(item.get('score', 0) for item in all_score_items)

        for pc in planned_chapters:
            chapter_title = pc.get('title', '')
            if not chapter_title:
                continue

            # 获取该章节的评分项（从planner分配或parse_result）
            score_items = pc.get('score_items', [])

            # 获取must_have内容
            must_have_content = strategy.get_must_have_content(chapter_title)

            # 如果有评分项，为每个评分项获取content_strategy
            content_strategies = []
            for item in score_items:
                weight = item.get('score') or 0  # 防御：score 可能为 None
                cs = strategy.get_content_strategy(chapter_title, item, weight)
                content_strategies.append({
                    'score_item': item.get('name') or item.get('title', ''),
                    'strategy': cs,
                })

            # 如果没有评分项但有总分信息，用plan_info_for_chapter
            if not content_strategies and total_weight > 0:
                # 尝试从plan中获取score_weight估算
                score_weight = pc.get('score_weight', 0)
                if score_weight > 0:
                    dummy_item = {'title': chapter_title, 'score': score_weight}
                    plan_info_enhanced = strategy.get_plan_info_for_chapter(
                        chapter_title, dummy_item, total_weight
                    )
                    pc['scoring_plan_info'] = plan_info_enhanced

            # 将策略信息合并到plan_info中
            pc['content_strategy'] = {
                'must_have': must_have_content.get('must_have', []),
                'bonus': must_have_content.get('bonus', []),
                'common_omissions': must_have_content.get('common_omissions', []),
                'structure_template': must_have_content.get('structure_template', ''),
                'structure_layers': must_have_content.get('structure_layers', []),
                'content_strategies': content_strategies,
            }

    def _post_generate_evaluator_check(self, output_path):
        """v6.1: EvaluatorCheck生成后自检，结果写入日志但不阻断生成。
        P2: 额外产出「废标风险三级预警 + 模拟评审得分表」Markdown 评审报告。
        """
        if not self.parse_result:
            return

        self.evaluation_result = None

        try:
            checker = EvaluatorCheck(
                parse_result=self.parse_result,
                doc_info={'chapters': [], 'project_info': self.project_info},
                bid_doc_path=output_path,
            )

            result = checker.run_all()
            coverage = result.get('coverage', {})
            risks = result.get('risks', {})
            score_prediction = result.get('score_prediction', {})
            mock_review = result.get('mock_review', {})

            # P2: 三级废标风险分级
            risk_md = ''
            grading = None
            if self.enable_risk_grading:
                from checker import grade_risk, render_risk_markdown
                grading = grade_risk(self.parse_result, result)
                risk_md = render_risk_markdown(grading)

            # 判断是否需要 warning
            warnings = []
            high_risks = risks.get('high_risks', [])
            if high_risks:
                warnings.append(
                    f"EVALUATOR_WARNING: 发现{len(high_risks)}项高风险废标风险 - "
                    + '; '.join(r.get('name', '') for r in high_risks[:3])
                )

            score_pct = score_prediction.get('score_percentage', 0)
            if score_pct < 60:
                warnings.append(
                    f"EVALUATOR_WARNING: 预测得分率仅{score_pct:.1f}%，低于60%阈值"
                )

            coverage_rate = coverage.get('coverage_rate', 0)
            if coverage_rate < 0.8:
                uncovered = coverage.get('uncovered', [])
                warnings.append(
                    f"EVALUATOR_WARNING: 评分项覆盖率{coverage_rate*100:.1f}%，低于80%阈值，"
                    f"未覆盖项：{'; '.join(u.get('item', '') for u in uncovered[:3])}"
                )

            self.evaluation_result = {
                'coverage': coverage,
                'risks': risks,
                'score_prediction': score_prediction,
                'risk_grading': grading,
                'warnings': warnings,
                'passed': len(warnings) == 0,
            }

            # P2: 写评审报告文件
            if (self.enable_risk_grading or self.enable_mock_review) and output_path:
                gen_md = self._build_generation_report_md()
                self._write_review_report(output_path, risk_md, mock_review, gen_md)

            # 日志输出
            if warnings:
                for w in warnings:
                    log.info("%s", w)
            else:
                print(f"[EvaluatorCheck] 自检通过：覆盖率{coverage_rate*100:.1f}%，"
                      f"预测得分率{score_pct:.1f}%，无高风险废标项")

        except Exception as e:
            log.error("自检异常（不影响生成）: %s", e)
            self.evaluation_result = {'error': str(e), 'warnings': [], 'passed': True}

    def _build_generation_report_md(self) -> str:
        """v7.3: 汇总各章节「内容原创性 + 配图」自查，生成评审报告附加章节。"""
        reports = [r for r in self._chapter_gen_reports if r]
        if not reports:
            return ''
        all_sentences = []
        inserted = 0
        for r in reports:
            all_sentences.extend(r.get('sentences') or [])
            inserted += len(r.get('inserted_images') or [])
        from bid_core.dedup import compute_self_similarity
        sim = compute_self_similarity(all_sentences)
        # 原创性评估：重复率越低越好
        if sim <= 0.02:
            verdict = '优秀（重复句占比 ≤2%，字面雷同风险极低）'
        elif sim <= 0.08:
            verdict = '良好（重复句占比 ≤8%，处于安全区间）'
        elif sim <= 0.15:
            verdict = '一般（重复句占比偏高，建议补充项目定制内容）'
        else:
            verdict = '偏高（重复句占比 >15%，建议启用 LLM 扩写层或更换指纹后重生成）'
        lines = [
            '',
            '## 内容原创性与配图自查（v7.3 防重引擎）',
            '',
            f'- 全文句子总数：{len(all_sentences)}',
            f'- 重复句占比（原创性指标）：**{sim*100:.1f}%**',
            f'- 原创性评估：{verdict}',
            f'- 企业图片库自动配图：{inserted} 张',
            f'- 差异化指纹：各章节基于唯一指纹做措辞旋转，同项目多次生成内容不同，降低查重合规风险。',
            '',
            '> 说明：本指标仅统计字面重复；最终标书仍须由人工结合企业真实情况核对，',
            '> 并建议定稿前使用专业标书查重工具交叉校验。',
            '',
        ]
        return '\n'.join(lines)

    def _write_review_report(self, output_path, risk_md, mock_review, gen_report_md=''):
        """将三级风险预警 + 模拟评审得分表 + 原创性自查写入同目录 Markdown 评审报告。"""
        try:
            from pathlib import Path
            out = Path(output_path)
            report_path = aux_path(None, output_path, '_评审报告.md')
            parts = [
                '# 标书评审报告',
                '',
                f'> 关联标书：`{out.name}`',
                '',
            ]
            if risk_md:
                parts.append(risk_md)
                parts.append('')
            mv_md = (mock_review or {}).get('markdown', '')
            if mv_md:
                parts.append(mv_md)
            if gen_report_md:
                parts.append(gen_report_md)
            report_path.write_text('\n'.join(parts), encoding='utf-8')
            log.info("评审报告已生成: %s", report_path)
        except Exception as e:
            log.error("评审报告写入失败（已忽略）: %s", e)

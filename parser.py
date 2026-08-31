"""
原生招标文件解析器 v7.1
========================
支持 .docx（python-docx）与 .pdf（pdfplumber），启发式抽取：

    - score_items                  评分项（名称 / 分值 / 评分标准）
    - star_clauses                星号 / 红线 / 必须响应条款
    - qualification_reqs          资格审查条件（资质 / 财务 / 业绩 / 人员）
    - red_line_clauses            废标 / 无效标 / 否决条款
    - disqualify_clauses_structured  结构化废标条款（供偏离表使用）

产出与 deviation_checker 完全兼容的 parse_result 字典。
设计原则：健壮、不阻断——任何单步解析失败都降级，绝不抛异常中断上层生成。

典型用法：
    from bid_core.parser import parse_tender
    parse_result = parse_tender("招标文件.pdf")          # 或 .docx
    # 之后直接传给 generate_bid_document({"parse_result": parse_result, ...})
"""
from __future__ import annotations

import os
import re
import json
from typing import Any, Dict, List, Optional

# ── 正则 ─────────────────────────────────────────────────────────
_RE_SCORE_LINE = re.compile(
    r'([\u4e00-\u9fffA-Za-z0-9（）()、，,.\-\s]{2,40}?)'
    r'[\s:：]*?'
    r'(\d{1,3}(?:\.\d+)?)\s*分'
)
_RE_NUM = re.compile(r'^\d{1,3}(?:\.\d+)?$')
_RE_STAR = re.compile(r'[*★☆]|须|必须')
_RE_STAR_CONTEXT = re.compile(r'(?:.*?[*★☆]\s*.{0,40})|(?:.{0,30}(?:必须|须)\s*[\u4e00-\u9fff]{0,30}(?:提供|响应|承诺|具备|满足).{0,20})')
_RE_DISQUALIFY = re.compile(r'(废标|无效投标|无效标|取消投标资格|否决|视为不响应|按废标处理)')
# v9.4.1：偏离许可识别（前附表勾选框「☑不允许 / □允许…偏离」）
_RE_DEVIATION_FORBID = re.compile(
    r'☑\s*不允许[\s\S]{0,60}?偏离|不允许[\s\S]{0,20}?偏离[\s\S]{0,20}?☑'
)
_RE_DEVIATION_ALLOW = re.compile(
    r'☑\s*允许[\s\S]{0,60}?偏离|允许[\s\S]{0,20}?偏离[\s\S]{0,20}?☑'
)
# v8.7: 年份同时支持阿拉伯数字与中文数字（近3年 / 近三年 / 近十年 / 近两年），
# 修复此前仅匹配 \d+ 导致「近三年」类资格审查条件漏抽、召回率下降的问题。
_CN_NUM = r'[0-9两一二三四五六七八九十]'
_RE_QUAL = re.compile(
    r'(具备.{0,20}?(?:资质|资格)|注册资本[^。；;]{0,30}|'
    rf'近\s*{_CN_NUM}+\s*年[^。；;]{{0,30}}业绩|'
    r'不?少于\s*\d+[^。；;]{0,20}|'
    r'财务[^。；;]{0,20}(?:良好|审计|净资产|营业收入))'
)
_RE_CLAUSE_NO = re.compile(r'(\d+(?:\.\d+){0,2})[、．.条]')
# v7.3: 工程量清单 / 数量抽取（面积、体积、长度、重量、台套等）
_RE_QUANTITY = re.compile(
    r'(\d+(?:\.\d+)?)\s*'
    r'(㎡|平方米|m2|m²|m3|m³|立方米|公里|km|吨|t|项|套|台|座|延米|平方|个|平方米)'
)
# v7.3: 关键节点日期（投标截止 / 开标 / 递交截止）
_RE_DEADLINE = re.compile(
    r'(投标截止|开标|递交投标文件截止|投标截止时间|截止时间)[^。；;]{0,40}?'
    r'(\d{4})\s*[-年./]\s*(\d{1,2})\s*[-月./]\s*(\d{1,2})'
)
# v7.3: 技术参数（带比较运算符 + 数值 + 单位）
_RE_PARAM = re.compile(
    r'([一-龥A-Za-z]{1,12})?\s*'
    r'(?:≥|>=|≥|≤|<=|<|>)\s*'
    r'(\d+(?:\.\d+)?)\s*([a-zA-Zμ℃%°km/]*)[^，。；;]{0,18}'
)
_RE_DATE = re.compile(r'(\d{4})\s*[-年./]\s*(\d{1,2})\s*[-月./]\s*(\d{1,2})')
# v8.10: 形式要件抽取（签字/盖章/密封/装订/正副本/电子版/授权）
# 这些是最高频的「隐性废标」诱因，且通常不写入废标条款正文，须单独抽取。
_FORMAT_CATS = [
    ('seal', re.compile(r'(盖章|公章|鲜章|骑缝章|印章)')),
    ('signature', re.compile(r'(签字|签章|签署)')),
    ('binding', re.compile(r'(装订|胶装|线装|装订成册|合订)')),
    ('copy', re.compile(r'(正本|副本)')),
    ('electronic', re.compile(r'(电子版|电子文档|U盘|光盘|CA锁|电子标|电子光盘)')),
    ('sealed', re.compile(r'(密封|封套|包封|外层包封)')),
    ('legal_rep', re.compile(r'(法定代表人|授权委托书|授权书|授权代表|被授权人)')),
]


def parse_tender(path, bid_type: str = 'construction',
                 llm_client: Any = None, enable_llm_parse: bool = False) -> Dict[str, Any]:
    """解析招标文件，返回 parse_result 字典。

    Args:
        path: 招标文件路径（.docx / .pdf），或路径列表（一标多包，自动合并解析）。
        bid_type: 'construction' | 'service'
    Returns:
        与 deviation_checker 兼容的字典；解析失败时返回空结构 + _error 字段
    """
    # v7.3: 一标多包 —— 接受路径列表，逐文件解析后合并
    if isinstance(path, (list, tuple)):
        paths = [p for p in path if p]
        if not paths:
            return _empty('文件列表为空')
        merged: Dict[str, Any] = None
        errors = []
        for p in paths:
            r = parse_tender(p, bid_type)
            if r.get('_error'):
                errors.append(f'{p}: {r["_error"]}')
                continue
            merged = _merge_result(merged, r)
        if merged is None:
            return _empty('；'.join(errors) or '全部文件解析失败')
        merged['bid_packages'] = len(paths)
        if errors:
            merged['_warn'] = '；'.join(errors)
        return merged

    if not path or not os.path.exists(path):
        return _empty(f'文件不存在: {path}')
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == '.pdf':
            blocks = _extract_pdf(path)
        elif ext in ('.docx', '.doc'):
            blocks = _extract_docx(path)
        else:
            return _empty(f'不支持的扩展名: {ext}')
    except Exception as exc:  # 解析本身失败，返回空结构
        return _empty(f'解析失败: {exc}')

    result = _analyze(blocks, bid_type)

    # v9.4.1：评分表结构化提取（评标办法前附表类分列表格）。
    # 结构化结果优先于正则噪声，置于列表前部；清洗层（bid_clean）统一过滤。
    if ext in ('.docx', '.doc'):
        try:
            table_items = _extract_score_table(path)
        except Exception:
            table_items = []
        if table_items:
            result['score_items'] = table_items + list(result.get('score_items') or [])

    # T1：LLM 结构化抽取主路径（双路合并 + 置信度告警）
    # 仅单文件主路径启用；多文件合并路径用正则 merged 即可。
    if enable_llm_parse and llm_client:
        llm_res = _llm_extract(blocks, llm_client, bid_type)
        if llm_res:
            result = _merge_parse(result, llm_res)
            result['llm_enriched'] = True
            result['parse_confidence'] = _parse_confidence(result, llm_res)
        else:
            _w = result.get('_warn')
            result['_warn'] = (_w + '；' if _w else '') + 'LLM解析未返回有效结构，已回退正则'
    return result


# ── T1 LLM 双路抽取 ─────────────────────────────────────────────
_LLM_FIELDS = (
    'score_items', 'star_clauses', 'qualification_reqs', 'red_line_clauses',
    'disqualify_clauses_structured', 'quantities', 'deadlines', 'key_params',
    'format_requirements',
)


def _llm_extract(blocks: List[str], llm_client: Any, bid_type: str) -> Dict[str, Any]:
    """用 LLM 做结构化抽取主路径，返回与 parse_result 兼容的字典（可能为空）。

    失败/未启用返回空字典，由正则路径兜底。绝不抛异常。
    """
    if not llm_client:
        return {}
    text = '\n'.join(blocks)[:12000]  # 截断避免超长上下文
    system = (
        "你是资深招标文件解析专家。从给定文本中抽取结构化招标要素，"
        "只返回 JSON，不要任何解释。字段与示例结构："
        "score_items(评分项, 元素含 name/score)、"
        "star_clauses(星号/必须条款, 含 content/severity)、"
        "qualification_reqs(资格审查, 含 field/label/content)、"
        "red_line_clauses(废标/无效标条款, 含 content)、"
        "disqualify_clauses_structured(含 content/clause_number/severity)、"
        "quantities(含 value/unit/text)、deadlines(含 kind/date)、"
        "key_params(含 item/value/unit/text)、"
        "format_requirements(含 type/text)。"
        "无法识别的字段给空数组 []。"
    )
    user = f"招标文件文本：\n{text}\n\n请按上述结构返回 JSON。"
    try:
        raw = llm_client.chat(system, user)
        if not raw:
            return {}
        s = raw[raw.find('{'):raw.rfind('}') + 1]
        data = json.loads(s)
        if not isinstance(data, dict):
            return {}
        out: Dict[str, Any] = {}
        for f in _LLM_FIELDS:
            v = data.get(f)
            out[f] = v if isinstance(v, list) else []
        return out
    except Exception:  # noqa: BLE001
        return {}


def _merge_parse(base: Dict[str, Any], llm_res: Dict[str, Any]) -> Dict[str, Any]:
    """正则结果为主，LLM 抽取结果补充正则漏抽的条款（按签名去重）。"""
    for f in _LLM_FIELDS:
        a = base.get(f) or []
        b = llm_res.get(f) or []
        if not b:
            continue
        seen = {_sig(x) for x in a}
        for x in b:
            sg = _sig(x)
            if sg not in seen:
                a.append(x)
                seen.add(sg)
        base[f] = a
    return base


def _parse_confidence(base: Dict[str, Any], llm_res: Dict[str, Any]) -> Dict[str, Any]:
    """T1 置信度：正则与 LLM 抽取的条目量对比（用于未识别条款显式告警）。"""
    regex_total = sum(len(base.get(f) or []) for f in _LLM_FIELDS)
    llm_total = sum(len(llm_res.get(f) or []) for f in _LLM_FIELDS)
    return {
        'regex_items': regex_total,
        'llm_items': llm_total,
        'added_by_llm': max(0, llm_total - regex_total),
        'enrich_ratio': round(llm_total / (regex_total + llm_total + 1e-9), 3),
    }


def _merge_result(base: Optional[Dict[str, Any]], other: Dict[str, Any]) -> Dict[str, Any]:
    """合并两个 parse_result（一标多包）。列表类字段去重合并。"""
    if base is None:
        base = _empty('')
        base.pop('_error', None)
    for key in ('score_items', 'star_clauses', 'qualification_reqs',
                'red_line_clauses', 'disqualify_clauses_structured',
                'quantities', 'deadlines', 'key_params'):
        a = base.get(key) or []
        b = other.get(key) or []
        if not a:
            base[key] = list(b)
        elif b:
            # 按内容签名去重（避免 str(dict) 前缀相同导致漏合并）
            seen = {_sig(x) for x in a}
            for x in b:
                s = _sig(x)
                if s not in seen:
                    a.append(x)
                    seen.add(s)
            base[key] = a
    base['raw_block_count'] = (base.get('raw_block_count') or 0) + (other.get('raw_block_count') or 0)
    return base


def _sig(item: Any) -> str:
    """提取去重签名：优先内容字段，其次文本/名称+分值/日期/数量。"""
    if isinstance(item, dict):
        for f in ('content', 'text'):
            if item.get(f):
                return str(item[f])[:24]
        if 'name' in item and 'score' in item:
            return f"{item.get('name')}{item.get('score')}"[:24]
        if 'date' in item:
            return str(item.get('date'))[:24]
        if 'value' in item and 'unit' in item:
            return f"{item.get('value')}{item.get('unit')}"[:24]
        return str(item)[:24]
    return str(item)[:24]


# ── 抽取 ─────────────────────────────────────────────────────────
def _extract_docx(path: str) -> List[str]:
    from docx import Document
    doc = Document(path)
    blocks: List[str] = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                blocks.append(' | '.join(cells))
    return blocks


def _row_cells(row) -> List[str]:
    """取表格行去重后的单元格文本（python-docx 对合并单元格会重复返回同一 tc）。"""
    out: List[str] = []
    seen: set = set()
    for cell in row.cells:
        tc = getattr(cell, '_tc', None)
        key = id(tc) if tc is not None else id(cell)
        if key in seen:
            continue
        seen.add(key)
        out.append((cell.text or '').strip())
    return out


# 评分表结构化提取所需语义词
_STD_SEMANTIC_WORDS = (
    '评审标准', '评分标准', '包含但不限于', '供应商提供', '投标人提供',
    '完整性', '合理性', '科学性', '满分', '得分为止', '扣分',
)


def _looks_like_score_std(text: str) -> bool:
    """判断单元格文本是否为"评分标准"描述（而非普通说明）。"""
    if len(text) < 12:
        return False
    return any(w in text for w in _STD_SEMANTIC_WORDS)


def _clean_cell_name(text: str) -> str:
    """清洗评分项名称：折叠空白、去编号前缀、去尾部噪声。"""
    s = re.sub(r'\s+', '', text or '')
    s = re.sub(r'^[0-9]+(?:\.[0-9]+)*[.、]?\s*', '', s)
    s = s.strip('（）()、，,。；;：:｜|/ ')
    return s


def _extract_score_table(path: str) -> List[Dict[str, Any]]:
    """结构化提取「评标办法前附表」类评分表。

    背景（v9.4.1 修复）：大量招标文件（含武汉市示范文本）的评分标准采用
    「条款号 | 评分因素 | 各评分因素细分项 | 分值 | 评分标准」的分列表格，
    且常因分页被拆成多个 table（后续 table 无表头行）。原正则仅能匹配
    "名称+数字+分"连写的段落文本，对此类表格召回率为 0。

    策略：逐表逐行启发式定位——数值单元格（分值）+ 右侧长文本（评分标准）
    + 左侧短中文名（细分项名称）；不依赖表头，兼容有/无表头的拆分表。

    返回结构化评分项：name / score / description / factor / source。
    任何异常一律返回空列表，绝不阻断上层解析。
    """
    try:
        from docx import Document  # type: ignore
        doc = Document(path)
    except Exception:
        return []

    items: List[Dict[str, Any]] = []
    seen: set = set()
    last_factor = ''
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = _row_cells(row)
            if len(cells) < 3:
                continue
            # 记录类别列（技术评审/商务评审/投标报价），供后续过滤与跨表沿用
            for c in cells:
                if c in ('技术评审', '商务评审', '投标报价', '其他评审') or \
                        ('评审' in c and len(c) <= 8):
                    last_factor = c
                    break
            # 定位分值列：纯数值，0 < 值 <= 100
            for i, c in enumerate(cells):
                if not _RE_NUM.match(c):
                    continue
                val = float(c)
                if val <= 0 or val > 100:
                    continue
                # 右侧：评分标准描述
                std = ''
                for r in cells[i + 1:]:
                    if _looks_like_score_std(r):
                        std = r
                        break
                if not std:
                    continue
                # 左侧：评分项名称（紧邻左格优先，否则向左找首个合格短中文名）
                name = ''
                for l in reversed(cells[:i]):
                    cand = _clean_cell_name(l)
                    if not cand or len(cand) > 30:
                        continue
                    if sum(1 for ch in cand if '\u4e00' <= ch <= '\u9fff') >= 4:
                        name = cand
                        break
                if not name or name in ('分值', '标准分', '评分标准', '评审标准', '条款号'):
                    continue
                key = (name, val)
                if key in seen:
                    continue
                seen.add(key)
                items.append({
                    'name': name,
                    'score': val,
                    'description': std[:600],
                    'factor': last_factor or '',
                    'source': 'parser_score_table',
                })
                break
    return items


def _extract_pdf(path: str) -> List[str]:
    """提取 PDF 文本块。优先 pdfplumber（版面感知更佳）；缺失时回退 PyPDF2（已内置），
    保证真实 PDF 招标文件在任意环境下均可解析为 parse_result。"""
    try:
        import pdfplumber  # type: ignore
        blocks: List[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ''
                for line in txt.splitlines():
                    line = line.strip()
                    if line:
                        blocks.append(line)
                for tbl in (page.extract_tables() or []):
                    for row in tbl:
                        cells = [str(c).strip() for c in row if c and str(c).strip()]
                        if cells:
                            blocks.append(' | '.join(cells))
        return blocks
    except ImportError:
        # 回退：PyPDF2（venv 已预置），无表格结构、仅纯文本，足够启发式抽取
        from PyPDF2 import PdfReader
        blocks = []
        reader = PdfReader(path)
        for page in reader.pages:
            txt = page.extract_text() or ''
            for line in txt.splitlines():
                line = line.strip()
                if line:
                    blocks.append(line)
        return blocks


# ── 分析 ─────────────────────────────────────────────────────────
def _collect_technical_chapters(blocks: List[str]) -> List[Dict[str, Any]]:
    """ADR-009 修复：从招标文件「技术评审」评分表抽取技术标一级章。

    该表为 条款号|评分因素|各评分因素细分项|分值|评分标准，
    细分项列形如「1.施工总进度计划及保障措施得分」，且常因列宽换行
    被拆成多行（如「3.主要分项工程施工方案和技术措⏎施得分」）或跨单元格拆行。

    做法：
    1. 定位含「技术评审」的表格区（分值构成块或评分因素=技术评审行）进入采集；
       遇到 2.2.3 / 负偏离 / 废标 / 否决投标 等明确结束标记即退出，避免把
       后续「负偏离」章节里的「1、投标文件技术部分的优点」误当章名；
    2. 仅采集「细分项」单元格（以「N.中文」开头 或 含「得分」且非评分区间文本），
       **抹平单元格内换行 \\n** 桥接拆行（关键修复：此前带 \\n 的章名无法被
       「数字.中文+得分」锚定正则匹配，导致 13 项只抽到 1 项）；
    3. 用「数字.中文+得分」锚定正则重建完整章名（已剥离前缀序号与后缀「得分」），
       排除「优秀为N分」等评分区间文本污染；返回 name 可直接作一级章标题。
    """
    collected: List[str] = []
    inside = False
    # 评分表结束标记：进入负偏离/否决/废标章节即退出采集区
    _STOP = ('2.2.3', '负偏离', '否决投标', '废标', '无效投标')
    for b in blocks:
        # 非表格块仅做状态切换（采集区起止），不采集内容
        if ' | ' not in b:
            if ('技术评审' in b) and ('评分因素' in b or '各评分因素细分项' in b):
                inside = True
            if inside and any(s in b for s in _STOP):
                inside = False
            continue

        # 抹平单元格内换行（关键修复）
        cells = [c.replace('\n', '').strip() for c in b.split(' | ')]
        has_tech = any('技术评审' in c for c in cells)
        if not inside:
            if has_tech:
                inside = True
            else:
                continue
        # 评分因素切到非技术评审（商务/响应性/形式）→ 退出
        factor = [c for c in cells if c and any(k in c for k in
                  ('技术评审', '商务评审', '响应性评审', '形式评审'))]
        if factor and '技术评审' not in ''.join(factor):
            inside = False
            continue
        # 明确结束标记 → 退出（兜底）
        if any(s in b for s in _STOP):
            inside = False
            continue
        # 关键修复：把整行（抹平换行、拼接所有单元格）作为候选文本收集，
        # 而非逐单元格匹配——长章名常被 pdfplumber 拆成多单元格
        # （如「4.对总承包…专业」|「分包…服」|「务方案得分」），
        # 逐单元格无法拼出「数字.中文+得分」锚点，导致残章漏抽。
        # 拼行后由末尾正则统一桥接跨单元格/跨行残章。
        row_text = ''.join(cells)
        if '得分' in row_text or re.match(r'^\s*\d+[.、．]', row_text):
            collected.append(row_text)
    text = ''.join(collected)
    items: List[Dict[str, Any]] = []
    seen = set()
    # 章名内部常含顿号「、」/间隔号「·」/一字线「—」（如「配合、协调、管理、服务方案」），
    # 必须纳入名称字符类，否则惰性匹配遇「、」即断、够不到末尾「得分」→ 整项漏抽。
    for m in re.finditer(r'(?:\d{1,3}[.、．]\s*)([\u4e00-\u9fff、·—]{2,80}?)得分', text):
        name = m.group(1).strip()
        if len(name) < 4 or name in seen:
            continue
        seen.add(name)
        # score 给 0（技术评审表的「分值」列未在抽取中捕获；
        # 数值 0 而非 None，避免下游 scoring_strategy 的 max(weight,100) 因 None 崩溃）
        items.append({'name': name, 'score': 0, 'source': 'technical_chapters'})
    return items


def _detect_deviation_allowed(blocks: List[str]) -> bool:
    """判断招标文件是否允许投标偏离。

    v9.4.1：投标人须知前附表通常以「☑不允许 / □允许，可偏离的项目和范围…」
    勾选框形式声明。招标文件声明"不允许偏离"时，投标文件不应出现偏离表
    （否则与"完全响应"要求相悖，且违反 docx 质量闸门）。

    返回 True 表示允许（或未明确禁止），False 表示明确不允许。
    """
    for block in blocks:
        if '偏离' not in block:
            continue
        # 「☑不允许 … 偏离」或「不允许 … 偏离 … ☑」→ 明确禁止
        if _RE_DEVIATION_FORBID.search(block):
            return False
        # 「☑允许 … 偏离」或「允许 … 偏离 … ☑」→ 明确允许
        if _RE_DEVIATION_ALLOW.search(block):
            return True
    return True


def _analyze(blocks: List[str], bid_type: str) -> Dict[str, Any]:
    score_items: List[Dict[str, Any]] = []
    star_clauses: List[Dict[str, Any]] = []
    qualification_reqs: List[Dict[str, Any]] = []
    red_line_clauses: List[Dict[str, Any]] = []
    disqualify: List[Dict[str, Any]] = []
    quantities: List[Dict[str, Any]] = []
    deadlines: List[Dict[str, Any]] = []
    key_params: List[Dict[str, Any]] = []
    format_requirements: List[Dict[str, Any]] = []

    seen_score = set()
    seen_star = set()
    seen_qual = set()
    seen_qty = set()
    seen_ddl = set()
    seen_param = set()
    seen_fmt = set()

    for block in blocks:
        _collect_scores(block, score_items, seen_score)
        _collect_star(block, star_clauses, seen_star)
        _collect_qual(block, qualification_reqs, seen_qual)
        _collect_disqualify(block, red_line_clauses, disqualify)
        _collect_quantities(block, quantities, seen_qty)
        _collect_deadlines(block, deadlines, seen_ddl)
        _collect_params(block, key_params, seen_param)
        _collect_format(block, format_requirements, seen_fmt)

    deviation_allowed = _detect_deviation_allowed(blocks)

    return {
        'deviation_allowed': deviation_allowed,
        'score_items': score_items,
        'technical_chapters': _collect_technical_chapters(blocks),
        'star_clauses': star_clauses,
        'qualification_reqs': qualification_reqs,
        'red_line_clauses': red_line_clauses,
        'disqualify_clauses_structured': disqualify,
        # v7.3: 工程量清单 / 节点 / 技术参数
        'quantities': quantities,
        'deadlines': deadlines,
        'key_params': key_params,
        # v8.10: 形式要件（签字/盖章/密封/装订/正副本/电子版/授权）
        'format_requirements': format_requirements,
        'raw_block_count': len(blocks),
    }


# v7.5: 清理评分项名称里混入的噪声词（满分/分值/（/）等）
# 注意：仅匹配完整噪声词，避免误删合法词首字（如"项目经理"的"项"）
_NOISE_LEAD = re.compile(r'^(?:满分|分值|配分|评分|权重|得分|本项|该项)[\s:：]*')
_NOISE_TRAIL = re.compile(r'(?:满分|分值|配分|评分|权重|得分)\s*$')
_PAREN_TRAIL = re.compile(r'[（(]\s*$')


def _clean_score_name(name: str) -> str:
    """清洗评分项名称：去掉'满分/分值/（'等混入噪声，得到干净项名。"""
    name = name.strip()
    name = _PAREN_TRAIL.sub('', name)        # 去掉结尾左括号：技术方案（ → 技术方案
    name = _NOISE_LEAD.sub('', name)          # 去掉开头噪声：满分技术方案 → 技术方案
    name = _NOISE_TRAIL.sub('', name)         # 去掉结尾噪声：技术方案满分 → 技术方案
    name = name.strip('（）()、，,。；;：: \t')
    return name


def _collect_scores(block: str, out: List[Dict[str, Any]], seen: set) -> None:
    # 1) 行内 "名称 ... 分值分"（含"满分10分""（5分）"等写法）
    for m in _RE_SCORE_LINE.finditer(block):
        name = _clean_score_name(m.group(1))
        score = float(m.group(2))
        if not name or len(name) < 2:
            continue
        key = (name, score)
        if key in seen:
            continue
        seen.add(key)
        out.append({
            'name': name,
            'score': score,
            'description': block[:120],
            'source': 'parser',
        })
    # 2) 表格行：一个数值单元格 + 一个描述单元格
    if ' | ' in block:
        cells = [c.strip() for c in block.split(' | ')]
        for i, c in enumerate(cells):
            if _RE_NUM.match(c):
                desc = max((cells[j] for j in range(len(cells)) if j != i and len(cells[j]) > 3),
                           key=len, default='')
                desc = _clean_score_name(desc)
                if desc and (desc, float(c)) not in seen:
                    seen.add((desc, float(c)))
                    out.append({'name': desc, 'score': float(c), 'description': block[:120], 'source': 'parser_table'})


def _collect_star(block: str, out: List[Dict[str, Any]], seen: set) -> None:
    if not _RE_STAR.search(block):
        return
    # 取包含星号/须/必须的整句
    m = _RE_STAR_CONTEXT.search(block)
    content = (m.group(0) if m else block).strip().rstrip('。；;，,')
    if content in seen or len(content) < 4:
        return
    seen.add(content)
    is_must = ('须' in content or '必须' in content or '*' in content or '★' in content)
    out.append({
        'type': 'star_marked' if ('*' in content or '★' in content) else 'mandatory',
        'content': content,
        'severity': 'critical' if is_must else 'medium',
        'source': 'parser',
    })


def _collect_qual(block: str, out: List[Dict[str, Any]], seen: set) -> None:
    m = _RE_QUAL.search(block)
    if not m:
        return
    content = m.group(0).strip()
    if content in seen or len(content) < 6:
        return
    seen.add(content)
    field = 'enterprise_qual' if '资质' in content or '资格' in content else \
            'financial' if ('注册资本' in content or '财务' in content or '净资产' in content) else \
            'experience' if ('业绩' in content or '年' in content) else 'other'
    label = {'enterprise_qual': '企业资质', 'financial': '财务', 'experience': '业绩', 'other': '其他'}[field]
    out.append({
        'field': field,
        'label': label,
        'content': content,
        'is_mandatory': True,
        'source': 'parser',
    })


def _collect_disqualify(block: str, red_out: List[Dict[str, Any]], struct_out: List[Dict[str, Any]]) -> None:
    m = _RE_DISQUALIFY.search(block)
    if not m:
        return
    content = block.strip()
    clause_no = ''
    cm = _RE_CLAUSE_NO.search(block)
    if cm:
        clause_no = cm.group(1)
    # 去重（按内容前20字）
    key = content[:20]
    if any(r.get('content', '').startswith(key) for r in red_out):
        return
    red_out.append({'type': 'disqualify', 'content': content, 'raw_match': m.group(0)})
    struct_out.append({
        'content': content,
        'clause_number': clause_no,
        'severity': 'critical',
        'source': 'parser',
    })


def _collect_format(block: str, out: List[Dict[str, Any]], seen: set) -> None:
    """v8.10: 抽取形式要件（签字/盖章/密封/装订/正副本/电子版/授权）。

    这些要件常以散句出现在「投标文件格式」「装订装订要求」章节，
    而非废标条款正文，漏做即隐性废标。按类别抽取并去重。
    """
    text = block.strip()
    for cat, rgx in _FORMAT_CATS:
        if not rgx.search(text):
            continue
        key = f'{cat}:{text[:20]}'
        if key in seen:
            continue
        seen.add(key)
        out.append({
            'type': cat,
            'text': text[:160],
            'source': 'parser',
        })


def _collect_quantities(block: str, out: List[Dict[str, Any]], seen: set) -> None:
    """v7.3: 抽取工程量清单中的数量（面积/体积/长度/重量/台套等）。"""
    for m in _RE_QUANTITY.finditer(block):
        value = m.group(1)
        unit = m.group(2)
        key = f'{value}{unit}'
        if key in seen:
            continue
        # 仅保留 ≥1 的量级，过滤页码噪声
        try:
            if float(value) < 1:
                continue
        except ValueError:
            continue
        seen.add(key)
        out.append({'value': value, 'unit': unit, 'text': m.group(0), 'source': 'parser'})


def _collect_deadlines(block: str, out: List[Dict[str, Any]], seen: set) -> None:
    """v7.3: 抽取关键节点日期（投标截止/开标）。"""
    for m in _RE_DEADLINE.finditer(block):
        date = f'{m.group(2)}-{m.group(3).zfill(2)}-{m.group(4).zfill(2)}'
        kind = m.group(1)
        key = f'{kind}{date}'
        if key in seen:
            continue
        seen.add(key)
        out.append({'kind': kind, 'date': date, 'source': 'parser'})


def _collect_params(block: str, out: List[Dict[str, Any]], seen: set) -> None:
    """v7.3: 抽取技术参数（带比较运算符的数值要求）。"""
    for m in _RE_PARAM.finditer(block):
        item = (m.group(1) or '').strip()
        value = m.group(2)
        unit = m.group(3) or ''
        text = m.group(0).strip()
        if len(text) < 4:
            continue
        key = text[:20]
        if key in seen:
            continue
        seen.add(key)
        out.append({'item': item, 'value': value, 'unit': unit,
                    'text': text, 'source': 'parser'})


def _empty(reason: str) -> Dict[str, Any]:
    return {
        'score_items': [],
        'star_clauses': [],
        'qualification_reqs': [],
        'red_line_clauses': [],
        'disqualify_clauses_structured': [],
        'quantities': [],
        'deadlines': [],
        'key_params': [],
        'format_requirements': [],
        '_error': reason,
    }

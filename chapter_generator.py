"""
单章节生成器 v1.0 — P1 升级（对标 WPS AI "章节续写/单章节生成"）

给定章节标题 + parse_result，复用既有章节类（bid_technical.chapters.*）渲染
该章节完整内容；若标题未命中路由，则回退到 scoring_strategy.json
的 must_have/bonus 结构化生成，保证任何章节都可产出。

支持两种使用方式:
    1. 完整文档中的"单章节补写"模式 (chapter_only=True)
    2. 独立生成单个章节 docx / 文本，供用户直接取用

设计原则:
    - 复用 > 重写：直接调用已验证的章节类，内容质量与全量生成一致
    - 回退兜底：未命中路由时用评分策略库结构化生成，绝不空跑
    - 不阻断：任何异常降级为回退生成或空章节
"""
from __future__ import annotations

import os
import re
from typing import Any, Dict, List, Optional

from bid_core.logger import get_logger

log = get_logger(__name__)


class ChapterGenerator:
    """单章节生成器。"""

    def __init__(self, parse_result: Optional[Dict[str, Any]] = None,
                 user_context: Any = None,
                 bid_type: str = 'construction',
                 project_info: Optional[Dict[str, Any]] = None,
                 llm_client: Any = None,
                 differentiator: Any = None,
                 image_library: Any = None):
        self.parse_result = parse_result
        self.user_context = user_context
        self.bid_type = bid_type if bid_type in ('construction', 'service') else 'construction'
        self.project_info = project_info or {}
        self.llm_client = llm_client  # v7.2: 可选 LLM 扩写客户端
        # v7.3: 防重差异化引擎 + 企业图片库
        if differentiator is None:
            from bid_core.dedup import Differentiator
            differentiator = Differentiator(project_info=project_info or {})
        self.differentiator = differentiator
        if image_library is None:
            try:
                uc = user_context
                if hasattr(uc, 'get_images'):
                    image_library = uc.get_images() or []
                elif isinstance(uc, dict):
                    image_library = uc.get('images', []) or []
            except Exception:
                image_library = []
        self.image_library = image_library or []
        self._routes = None
        self._route_priority = None

    # ────────────────────────────────────────────────────────────
    # 公共入口
    # ────────────────────────────────────────────────────────────
    def generate(self, chapter_title: str, detail_level: int = 3,
                 output_path: Optional[str] = None) -> str:
        """生成单章节，保存为 docx 并返回文件路径。

        Args:
            chapter_title: 章节标题（如"质量保证措施和创优计划"）
            detail_level: 内容深度 1-5
            output_path: 输出路径；为空时自动命名为 章节_标题.docx
        """
        from docx import Document
        from bid_core.formatter import NormalFormatter
        from bid_core.randomizer import Randomizer
        from bid_core.user_context import UserContext

        doc = Document()
        formatter = NormalFormatter(doc)
        synonyms_path = os.path.join(self._project_root(), 'data', 'synonyms.json')
        randomizer = Randomizer(synonyms_path=synonyms_path, enabled=False)
        uc = UserContext(self.user_context) if not isinstance(self.user_context, UserContext) else self.user_context

        project_info = dict(self.project_info)
        project_info.setdefault('name', chapter_title)
        project_info.setdefault('bid_type', self.bid_type)

        try:
            self._render_chapter(formatter, randomizer, uc, chapter_title, detail_level, project_info)
        except Exception as exc:
            log.warning('章节类渲染失败，回退结构化生成: %s', exc, exc_info=True)
            self._render_fallback(formatter, chapter_title, detail_level)

        if not output_path:
            safe = re.sub(r'[\\/:*?"<>|]', '_', chapter_title)[:40]
            output_path = f'章节_{safe}.docx'
        formatter.save(output_path)
        log.info('单章节生成完成: %s -> %s', chapter_title, output_path)
        return output_path

    def generate_markdown(self, chapter_title: str, detail_level: int = 3) -> str:
        """生成单章节并以 Markdown 文本返回（便于预览/插入）。"""
        from docx import Document
        from bid_core.formatter import NormalFormatter
        from bid_core.randomizer import Randomizer
        from bid_core.user_context import UserContext

        doc = Document()
        formatter = NormalFormatter(doc)
        synonyms_path = os.path.join(self._project_root(), 'data', 'synonyms.json')
        randomizer = Randomizer(synonyms_path=synonyms_path, enabled=False)
        uc = UserContext(self.user_context) if not isinstance(self.user_context, UserContext) else self.user_context

        project_info = dict(self.project_info)
        project_info.setdefault('name', chapter_title)
        project_info.setdefault('bid_type', self.bid_type)

        try:
            self._render_chapter(formatter, randomizer, uc, chapter_title, detail_level, project_info)
        except Exception:
            self._render_fallback(formatter, chapter_title, detail_level)

        return self._doc_to_markdown(doc)

    # ────────────────────────────────────────────────────────────
    # 渲染核心
    # ────────────────────────────────────────────────────────────
    def _render_chapter(self, formatter, randomizer, uc, title, detail_level, project_info) -> None:
        """通过路由复用既有章节类渲染。"""
        from bid_technical.generator import ROUTES, ROUTE_PRIORITY
        self._routes = ROUTES
        self._route_priority = ROUTE_PRIORITY

        cls = self._dispatch(title)
        if cls is None:
            # 未命中路由 → 结构化回退
            self._render_fallback(formatter, title, detail_level)
            return

        instance = cls(
            formatter,
            randomizer,
            uc,
            detail_level=detail_level,
            parse_result=self.parse_result,
            plan_info=None,
            llm_client=self.llm_client,
            differentiator=self.differentiator,
            image_library=self.image_library,
        )
        instance.render(project_info)

    def _dispatch(self, title: str):
        if not self._routes:
            from bid_technical.generator import ROUTES, ROUTE_PRIORITY
            self._routes, self._route_priority = ROUTES, ROUTE_PRIORITY
        for k in self._route_priority:
            if k in title:
                return self._get_chapter_class(self._routes[k])
        return None

    @staticmethod
    def _get_chapter_class(path: str):
        from bid_technical.chapters.base import resolve_chapter_class
        return resolve_chapter_class(path)

    # ────────────────────────────────────────────────────────────
    # 回退：基于 scoring_strategy.json 结构化生成
    # ────────────────────────────────────────────────────────────
    def _render_fallback(self, formatter, title: str, detail_level: int) -> None:
        """未命中章节类时，用评分策略库的 must_have/bonus 生成结构化章节。"""
        from bid_core.data_loader import DataLoader

        strategy = DataLoader().load_scoring_strategy()
        db_key = f"{self.bid_type}_strategy_db"
        db = strategy.get(db_key, {})
        entry = db.get(title, {})
        # 标题反查
        if not entry:
            keywords_map = strategy.get('chapter_match_keywords', {})
            for t, kws in keywords_map.items():
                if any(kw in title for kw in kws):
                    entry = db.get(t, {})
                    if entry:
                        break

        formatter.h1(1, title)

        must_have = entry.get('must_have', [])
        bonus = entry.get('bonus', [])
        layers = entry.get('structure_layers', [])
        common_omissions = entry.get('common_omissions', [])

        if layers:
            for layer in layers:
                formatter.h2(layer)
                formatter.body(f'我司将围绕「{layer}」制定专项实施方案，确保本环节可控、可追溯、可核查。')
        elif must_have:
            formatter.h2('主要内容')
        else:
            formatter.body('（本章内容由评分策略库自动结构化生成，建议结合项目实际补充细化。）')

        if must_have:
            formatter.h2('必须包含要点')
            for item in must_have[:max(3, detail_level * 2)]:
                formatter.body(f'• {item}')

        if bonus:
            formatter.h2('加分项策划')
            for item in bonus[:max(2, detail_level)]:
                formatter.body(f'• {item}')

        if common_omissions and detail_level >= 3:
            formatter.h2('常见扣分规避')
            for item in common_omissions[:detail_level]:
                formatter.body(f'• {item}')

    # ────────────────────────────────────────────────────────────
    # 工具
    # ────────────────────────────────────────────────────────────
    @staticmethod
    def _doc_to_markdown(doc) -> str:
        lines: List[str] = []
        for p in doc.paragraphs:
            text = (p.text or '').strip()
            if not text:
                continue
            style = ''
            try:
                style = (p.style.name or '') if p.style else ''
            except Exception:
                pass
            if style == 'Heading 1' or style.strip() == 'Title':
                lines.append(f'# {text}')
            elif 'Heading 2' in style:
                lines.append(f'## {text}')
            elif 'Heading 3' in style:
                lines.append(f'### {text}')
            else:
                lines.append(text)
        return '\n\n'.join(lines)

    @staticmethod
    def _project_root() -> str:
        """定位项目根目录：从本文件向上查找第一个含 data/ 的祖先目录。

        兼容扁平结构（data/ 在同级）与正式包结构（data/ 在 bid_core 上层）。
        """
        cur = os.path.dirname(os.path.abspath(__file__))
        for _ in range(6):
            if os.path.isdir(os.path.join(cur, 'data')):
                return cur
            parent = os.path.dirname(cur)
            if parent == cur:
                break
            cur = parent
        # 回退：沿用原 bid_core 包推导（正式部署）
        try:
            import bid_core
            return os.path.dirname(os.path.dirname(os.path.abspath(bid_core.__file__)))
        except Exception:
            return os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def generate_single_chapter(req: Dict[str, Any]) -> Dict[str, Any]:
    """从请求字典生成单章节（供 main.py chapter_only 模式调用）。

    req 需包含 chapter_title，可选 parse_result / user_context / bid_type /
    detail_level / project_info / output_path。
    """
    title = req.get('chapter_title') or req.get('chapterTitle')
    if not title:
        return {'success': False, 'message': '缺少 chapter_title 参数'}
    gen = ChapterGenerator(
        parse_result=req.get('parse_result'),
        user_context=req.get('user_context'),
        bid_type=req.get('bid_type', 'construction'),
        project_info=req.get('project_info') or req,
        llm_client=req.get('llm_client'),
    )
    try:
        path = gen.generate(
            title,
            detail_level=req.get('detail_level', 3),
            output_path=req.get('output_path'),
        )
        return {'success': True, 'message': '单章节生成成功', 'output_file': path}
    except Exception as exc:
        log.error('单章节生成失败: %s', exc, exc_info=True)
        return {'success': False, 'message': f'生成失败：{exc}'}

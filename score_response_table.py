"""
评分项响应保障表 Word 渲染 v7.5

紧接"投标文件偏离表"之后渲染，呈现每个评分项对应的我方响应保障
（企业业绩 / 项目团队 / 企业资质 / 设备资源 / 安全体系 / 财务实力等），
并标注是否已录入知识库数据支撑。

用法（在 TechnicalBidGenerator._render_deviation_table 中）:
    from bid_technical.tables.score_response_table import render_score_response_table
    render_score_response_table(doc, project_info, mapping)
"""
from __future__ import annotations

from typing import Any, Dict

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

_COLOR_OK = RGBColor(0x00, 0x80, 0x00)     # 绿：已有数据支撑
_COLOR_WARN = RGBColor(0xB8, 0x6A, 0x00)   # 橙：承诺补充

HEADER = ['评分项', '分值', '响应保障类别', '我方响应保障说明', '状态']


def render_score_response_table(doc: Document,
                                project_info: Dict[str, Any],
                                mapping: Dict[str, Any]) -> None:
    """在 doc 末尾追加"评分项响应保障表"。"""
    rows = mapping.get('rows', [])
    if not rows:
        return

    # —— 标题 ——
    name = (project_info or {}).get('name', '本工程')
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f'{name} — 评分项响应保障表')
    run.bold = True
    run.font.size = Pt(14)

    # —— 摘要行 ——
    total = mapping.get('total', len(rows))
    mapped = mapping.get('mapped', 0)
    coverage = mapping.get('coverage', 0)
    sline = doc.add_paragraph()
    srun = sline.add_run(
        f'评分项总数 {total} 项 ｜ 已录入知识库数据支撑 {mapped} 项 ｜ '
        f'响应覆盖率 {coverage}%'
    )
    srun.font.size = Pt(9)
    srun.italic = True

    # —— 表格 ——
    nrows = len(rows) + 1
    cols = len(HEADER)
    table = doc.add_table(rows=nrows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for j, text in enumerate(HEADER):
        cell = table.rows[0].cells[j]
        cell.text = text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                r.bold = True
                r.font.size = Pt(9)

    # 数据行
    for i, row in enumerate(rows, 1):
        cells = table.rows[i].cells
        sname = str(row.get('score_name', ''))
        detail = str(row.get('asset_detail', ''))
        status = '已支撑' if row.get('satisfied') else '承诺补充'
        if len(sname) > 40:
            sname = sname[:40] + '…'
        if len(detail) > 60:
            detail = detail[:60] + '…'
        values = [
            sname,
            str(row.get('score', '')),
            str(row.get('asset_type', '')),
            detail,
            status,
        ]
        for j, val in enumerate(values):
            cell = cells[j]
            cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in (1, 2, 4) else WD_ALIGN_PARAGRAPH.LEFT
                for r in para.runs:
                    r.font.size = Pt(8)
                    if j == 4:
                        r.bold = True
                        r.font.color.rgb = _COLOR_OK if row.get('satisfied') else _COLOR_WARN

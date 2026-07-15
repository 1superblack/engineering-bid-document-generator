"""
BaseGenerator - 生成器基类
流程编排：章节渲染 → 附表生成 → 文档保存
"""
from docx import Document
from normal_formatter import NormalFormatter
from constants import safe_heading_font, safe_body_font
from randomizer import Randomizer
from user_context import UserContext


class BaseGenerator:
    """生成器基类"""

    # 必附表：不管招标文件是否要求，始终生成
    MANDATORY_TABLES = []

    def __init__(self, project_info, formatter=None, randomizer=None, user_context=None,
                 heading_font=None, body_font=None,
                 parse_result=None, target_pages=0):
        """
        初始化生成器

        Args:
            project_info: 项目信息字典
            formatter: 格式化引擎
            randomizer: 随机化引擎
            user_context: 用户信息
            heading_font: 标题字体（如 '黑体' / '宋体'），走白名单校验
            body_font: 正文字体（如 '仿宋' / '宋体'），走白名单校验
            parse_result: parser解析结果（v6.1: 用于planner规划）
            target_pages: 目标页数（v6.1: 用于planner规划）
        """
        self.project_info = project_info
        # 字体参数传给 formatter（如果外部没传 formatter 实例）
        if formatter is None:
            self.formatter = NormalFormatter(
                heading_font=heading_font,
                body_font=body_font
            )
        else:
            self.formatter = formatter
        self.randomizer = randomizer or Randomizer(enabled=False)
        self.user_context = user_context or UserContext()
        self.parse_result = parse_result
        self.target_pages = target_pages if target_pages > 0 else 300
        
        # 子类应设置
        self.chapters = []
        self.tables = []
        self.plan = None  # v6.1: planner规划结果
    
    def get_bid_type(self):
        """获取标书类型"""
        return self.project_info.get('bid_type', 'construction')
    
    def get_project_name(self):
        """获取项目名称"""
        return self.project_info.get('name', '本项目')
    
    def render_chapter(self, chapter_info):
        """
        渲染单个章节
        子类应实现此方法
        """
        raise NotImplementedError
    
    def render_table(self, table_info):
        """
        渲染单个附表
        子类应实现此方法
        """
        raise NotImplementedError
    
    def before_render(self):
        """渲染前钩子"""
        pass
    
    def after_render(self):
        """渲染后钩子"""
        pass
    
    def generate(self, output_path):
        """
        生成标书 v6.1: 集成planner规划（优先从parse_result规划章节）

        Args:
            output_path: 输出文件路径

        Returns:
            输出文件路径
        """
        # v6.1: 尝试用planner规划章节（子类可覆盖此行为）
        chapters_to_render = self._resolve_chapters()

        self.before_render()

        # 渲染章节
        for i, chapter in enumerate(chapters_to_render, 1):
            title = chapter.get('title', chapter) if isinstance(chapter, dict) else chapter
            self.formatter.h1(i, title)
            self.render_chapter(chapter)

        # 渲染附表
        self.render_tables()

        self.after_render()

        # 保存文档
        self.formatter.save(output_path)
        return output_path

    def _resolve_chapters(self):
        """
        v6.1: 解析最终章节列表

        优先级:
        1. 如果 parse_result 可用，调用 planner 规划章节
        2. 否则使用 self.chapters（子类设置的默认章节）

        子类可覆盖此方法实现自定义规划逻辑。
        """
        # 尝试调用planner
        if self.parse_result and self.parse_result.get('score_items'):
            try:
                from planner import plan_chapters

                self.plan = plan_chapters(
                    parse_result=self.parse_result,
                    target_pages=self.target_pages,
                    project_info=self.project_info,
                )

                planned = self.plan.get('chapters', [])
                if planned:
                    chapters = []
                    for pc in planned:
                        chapters.append({
                            'title': pc['title'],
                            'keywords': [],
                            'plan_info': pc,
                        })
                    return chapters
            except ImportError:
                from .logger import get_logger
                log = get_logger(__name__)
                log.warning('planner 模块不可用，降级到默认章节列表')

        # 兜底：使用子类设置的默认章节
        return self.chapters
    
    def render_tables(self):
        """渲染附表（含必附表机制）"""
        self.formatter.add_heading("附表")
        
        for table in self.tables:
            self.formatter.page_break()
            self.render_table(table)
    
    def gantt_bars(self, task_duration, start_day, total_duration, months):
        """生成横道图的月份占位符（■/□）"""
        bars = []
        for m in range(1, months + 1):
            month_start = (m - 1) * 30 + 1
            month_end = m * 30
            task_end = start_day + task_duration - 1
            if start_day <= month_end and task_end >= month_start:
                bars.append("■")
            else:
                bars.append("□")
        return bars
    
    def get_l3_titles(self, title):
        """获取三级标题列表"""
        # 默认实现，子类可覆盖
        return []

"""
文档文本提取模块
支持Word和PDF文档的文本、表格、元数据提取
"""
import os
import logging
from typing import List, Tuple, Set, Optional

from .models import Paragraph, TableData, MetadataInfo

log = logging.getLogger(__name__)


def get_whitelist(mode: str = "通用") -> Set[str]:
    """获取通用条款白名单

    Args:
        mode: 查重模式（标书/论文/通用）

    Returns:
        白名单条款集合
    """
    base_whitelist = {
        # 法律法规引用
        '根据《中华人民共和国',
        '按照国家现行规范',
        '符合国家标准',
        '满足设计要求',
        # 常见工程表述
        '质量第一，安全至上',
        '科学组织，合理安排',
    }

    mode_specific = {
        '标书': {
            '我方承诺',
            '确保工期',
            '保证质量',
            '安全生产',
        },
        '论文': {
            '本文研究了',
            '实验结果表明',
            '综上所述',
        },
        '通用': set(),
    }

    whitelist = base_whitelist | mode_specific.get(mode, set())
    log.debug(f"加载白名单: {mode}模式, {len(whitelist)}条")

    return whitelist


def extract_text_from_docx(file_path: str) -> Tuple[List[Paragraph], List[TableData], MetadataInfo]:
    """从Word文档提取文本和表格

    Args:
        file_path: 文件路径

    Returns:
        (段落列表, 表格列表, 元数据)
    """
    paragraphs = []
    tables = []
    meta = MetadataInfo(file_path=file_path)

    try:
        from docx import Document

        doc = Document(file_path)

        # 提取段落
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append(Paragraph(text=text, index=i))

        # 提取表格
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                rows.append(row_data)
            if rows:
                tables.append(TableData(rows=rows, caption=f"表格{i+1}"))

        # 提取元数据
        if doc.core_properties:
            meta.author = doc.core_properties.author or ''
            meta.created_date = str(doc.core_properties.created) if doc.core_properties.created else ''
            meta.title = doc.core_properties.title or ''

        log.info(f"成功提取DOCX: {file_path} | {len(paragraphs)}段落, {len(tables)}表格")

    except Exception as e:
        log.error(f"提取DOCX失败: {file_path} | {e}")

    return paragraphs, tables, meta


def extract_text_from_pdf(file_path: str) -> Tuple[List[Paragraph], List[TableData], MetadataInfo]:
    """从PDF文档提取文本和表格

    Args:
        file_path: 文件路径

    Returns:
        (段落列表, 表格列表, 元数据)
    """
    paragraphs = []
    tables = []
    meta = MetadataInfo(file_path=file_path)

    try:
        # 尝试使用pdfplumber或PyPDF2
        try:
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        for j, para in enumerate(text.split('\n')):
                            if para.strip():
                                paragraphs.append(
                                    Paragraph(text=para.strip(), index=len(paragraphs), page_num=i + 1)
                                )

                    # 提取表格
                    page_tables = page.extract_tables()
                    for k, table in enumerate(page_tables):
                        if table:
                            tables.append(TableData(
                                rows=[row for row in table if row],
                                caption=f"页{i+1}_表{k+1}",
                                page_num=i + 1,
                            ))

            log.info(f"成功提取PDF(pdfplumber): {file_path} | {len(paragraphs)}段")

        except ImportError:
            log.warning("未安装pdfplumber，尝试使用基础PDF提取")
            # 简化实现：按行读取
            import PyPDF2

            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        for line in text.split('\n'):
                            if line.strip():
                                paragraphs.append(
                                    Paragraph(text=line.strip(), index=len(paragraphs), page_num=i + 1)
                                )

            log.info(f"成功提取PDF(PyPDF2): {file_path} | {len(paragraphs)}段")

    except Exception as e:
        log.error(f"提取PDF失败: {file_path} | {e}")

    return paragraphs, tables, meta


def extract_document(file_path: str) -> Tuple[List[Paragraph], List[TableData], MetadataInfo]:
    """智能文档提取（自动识别格式）

    Args:
        file_path: 文件路径

    Returns:
        (段落列表, 表格列表, 元数据)
    """
    ext = os.path.splitext(file_path)[1].lower()

    extractors = {
        '.docx': extract_text_from_docx,
        '.doc': extract_text_from_docx,
        '.pdf': extract_text_from_pdf,
    }

    extractor = extractors.get(ext)
    if extractor:
        return extractor(file_path)

    log.warning(f"不支持的文档格式: {ext}")
    return [], [], MetadataInfo(file_path=file_path)


def is_common_clause(text: str, whitelist: Set[str] = None) -> bool:
    """判断是否为通用条款

    Args:
        text: 待判断文本
        whitelist: 白名单（可选，默认使用通用白名单）

    Returns:
        True表示是通用条款
    """
    if not whitelist:
        whitelist = get_whitelist()

    return any(keyword in text for keyword in whitelist)


def filter_common_clauses(paragraphs: List[Paragraph],
                         whitelist: Set[str] = None) -> Tuple[List[Paragraph], int]:
    """过滤通用条款

    Args:
        paragraphs: 段落列表
        whitelist: 白名单

    Returns:
        (过滤后的段落列表, 过滤数量)
    """
    filtered = []
    filtered_count = 0

    for para in paragraphs:
        if is_common_clause(para.text, whitelist):
            filtered_count += 1
        else:
            filtered.append(para)

    log.debug(f"过滤通用条款: {filtered_count}/{len(paragraphs)}条")
    return filtered, filtered_count
